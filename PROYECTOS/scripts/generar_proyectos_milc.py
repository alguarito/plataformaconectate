from __future__ import annotations

import csv
import importlib.util
import json
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path("/Users/alvarocardenasorozco/Desktop/PROYECTOS/ACADEMICO")
BASE_SCRIPT = Path("/Users/alvarocardenasorozco/Desktop/ACADEMICO/generar_guias_sexto_milc.py")
PLAN = ROOT / "plan_area_tecnologia_2025.docx"
GUIDES_ROOT = ROOT / "GUIAS_TIC_MILC_NUEVAS"
OUT = ROOT / "PROYECTOS_TIC_MILC_NUEVOS"
AUTHOR = "Dr. Álvaro Cárdenas Orozco"
XELATEX = "xelatex"

GRADE_NAMES = {
    6: "sexto",
    7: "septimo",
    8: "octavo",
    9: "noveno",
    10: "decimo",
    11: "undecimo",
}

GRADE_LABELS = {
    6: "Sexto",
    7: "Séptimo",
    8: "Octavo",
    9: "Noveno",
    10: "Décimo",
    11: "Undécimo",
}

TABLE_START = {6: 0, 7: 3, 8: 6, 9: 9, 10: 12, 11: 15}


@dataclass
class PeriodData:
    number: int
    name: str
    dba: str
    referentes: str
    estandares: str
    competencias: str
    aprendizajes: str
    cognitivos: str
    procedimentales: str
    actitudinales: str


@dataclass
class ProjectData:
    grade: int
    period: PeriodData
    topics: list[str]
    old_title: str
    old_product: str
    central_theme: str
    challenge: str
    product: str
    active_strategy: str
    design_thinking_focus: str


def load_base():
    spec = importlib.util.spec_from_file_location("milc_guides_base", BASE_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"No se pudo cargar {BASE_SCRIPT}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


BASE = load_base()


def clean(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split()).strip()


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in str(text))


def period_name(number: int) -> str:
    return {1: "Primero", 2: "Segundo", 3: "Tercero"}[number]


def extract_milc_method() -> str:
    doc = Document(str(PLAN))
    for p in doc.paragraphs:
        text = clean(p.text)
        if text.startswith("Modelo MILC:"):
            return text
    return (
        "Modelo MILC: Modelo de Investigación Liberadora y Científica. Integra la praxis "
        "investigativa con la liberación del sujeto educativo frente a la colonización tecnológica."
    )


def extract_period_data(grade: int, period: int) -> PeriodData:
    doc = Document(str(PLAN))
    table = doc.tables[TABLE_START[grade] + period - 1]
    row1 = [clean(c.text) for c in table.rows[1].cells]
    row2 = [clean(c.text) for c in table.rows[2].cells]
    row5 = [clean(c.text) for c in table.rows[5].cells]
    return PeriodData(
        number=period,
        name=period_name(period),
        dba=row1[0].replace("DBA:", "").strip(),
        referentes=row2[0].replace("REFERENTES:", "").strip(),
        estandares=row5[0],
        competencias=row5[1],
        aprendizajes=row5[2],
        cognitivos=row5[3],
        procedimentales=row5[4],
        actitudinales=row5[5],
    )


def old_project_path(grade: int, period: int) -> Path:
    grade_word = {
        6: "sexto",
        7: "septimo",
        8: "octavo",
        9: "noveno",
        10: "decimo",
        11: "undecimo",
    }[grade]
    return ROOT / str(grade) / f"PERIODO {period}" / f"proyecto_integrador_grado_{grade_word}_periodo_{period}.docx"


def extract_old_project(grade: int, period: int) -> tuple[str, str]:
    path = old_project_path(grade, period)
    if not path.exists():
        return ("Proyecto integrador de periodo", "")
    doc = Document(str(path))
    texts = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    product = ""
    title = "Proyecto integrador de periodo"
    for text in texts:
        if text.startswith("Nombre sugerido del proyecto:"):
            title = text.replace("Nombre sugerido del proyecto:", "").strip()
            break
    for text in texts:
        if text.startswith("Producto final esperado:"):
            product = text.replace("Producto final esperado:", "").strip()
            break
    return title, product


def load_topics(grade: int, period: int) -> list[str]:
    grade_name = GRADE_NAMES[grade]
    temario_path = GUIDES_ROOT / grade_name / "fuente" / f"temario_{grade_name}.json"
    if not temario_path.exists() and grade == 6:
        temario_path = GUIDES_ROOT / grade_name / "fuente" / "temario_sexto.json"
    data = json.loads(temario_path.read_text(encoding="utf-8"))
    return [item["titulo"] for item in data if item["periodo"] == period]


def central_theme(period: PeriodData) -> str:
    bits = clean(period.aprendizajes).split(".")
    first = bits[0].strip()
    return first if first else period.dba


def infer_strategy(grade: int, period: int, topics: list[str]) -> tuple[str, str, str, str]:
    text = " ".join(topics).lower()
    if grade >= 10 or "proyecto" in text or "emprendimiento" in text or "solución" in text:
        return (
            "Design Thinking + Aprendizaje Basado en Proyectos",
            "empatizar con usuarios reales, definir un reto, idear, prototipar y validar",
            "Prototipo, propuesta técnica o solución tecnológica con validación social.",
            "Diseñar y sustentar una solución tecnológica con impacto social, ético y comunitario.",
        )
    if "programación" in text or "scratch" in text or "algorit" in text:
        return (
            "PRIMM + prototipado por bloques",
            "predecir, ejecutar, investigar, modificar y crear una solución interactiva",
            "Prototipo funcional, simulación o experiencia interactiva documentada.",
            "Crear un producto computacional que resuelva una necesidad o comunique una idea.",
        )
    if "hardware" in text or "mantenimiento" in text or "redes" in text or "computador" in text:
        return (
            "Aprendizaje experiencial + laboratorio técnico",
            "observar, diagnosticar, experimentar, registrar evidencias y proponer mejoras",
            "Diagnóstico técnico, protocolo de cuidado o mapa de sistema con evidencias.",
            "Investigar un problema técnico del entorno y proponer una respuesta verificable.",
        )
    if "fuentes" in text or "internet" in text or "fake news" in text or "seguridad" in text:
        return (
            "Indagación guiada + campaña de alfabetización digital",
            "formular preguntas, validar fuentes, producir evidencias y comunicar recomendaciones",
            "Campaña, guía visual o informe de validación de información y seguridad digital.",
            "Ayudar a la comunidad escolar a decidir mejor frente a información y riesgos digitales.",
        )
    if "excel" in text or "hoja" in text or "datos" in text:
        return (
            "Aprendizaje basado en datos + experimento escolar",
            "recolectar datos, organizarlos, analizarlos y tomar decisiones con evidencia",
            "Informe con tablas, gráficos y recomendación basada en datos.",
            "Usar datos reales para comprender un problema escolar o comunitario.",
        )
    return (
        "Design Thinking + producto comunicativo",
        "observar una necesidad, organizar información, crear un producto y evaluar su utilidad",
        "Producto comunicativo, diagnóstico aplicado o recurso digital con reflexión ética.",
        "Transformar lo aprendido en un producto útil para otros.",
    )


def build_project(grade: int, period: int) -> ProjectData:
    pdata = extract_period_data(grade, period)
    topics = load_topics(grade, period)
    old_title, old_product = extract_old_project(grade, period)
    strategy, dt_focus, product, challenge = infer_strategy(grade, period, topics)
    if old_product:
        product = f"{product} Puede dialogar con el producto antiguo sugerido: {old_product}."
    return ProjectData(
        grade=grade,
        period=pdata,
        topics=topics,
        old_title=old_title,
        old_product=old_product,
        central_theme=central_theme(pdata),
        challenge=challenge,
        product=product,
        active_strategy=strategy,
        design_thinking_focus=dt_focus,
    )


def stem(project: ProjectData) -> str:
    return f"proyecto-{project.period.number}-{project.grade}-TIC"


def topic_items(topics: list[str]) -> str:
    return "\n".join(rf"\item {tex_escape(t)}" for t in topics)


def tex_template(project: ProjectData, milc_method: str) -> str:
    pdata = project.period
    return rf"""
\documentclass[11pt,letterpaper]{{article}}
\usepackage[margin=1.75cm]{{geometry}}
\usepackage[table]{{xcolor}}
\usepackage{{fontspec}}
\usepackage{{tikz}}
\usepackage[most]{{tcolorbox}}
\usepackage{{tabularx}}
\usepackage{{multicol}}
\usepackage{{array}}
\usepackage{{fancyhdr}}
\usepackage{{enumitem}}

\setmainfont{{Helvetica Neue}}
\setsansfont{{Helvetica Neue}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{6pt}}
\renewcommand{{\arraystretch}}{{1.22}}

\definecolor{{milcMagenta}}{{HTML}}{{E6007E}}
\definecolor{{milcVino}}{{HTML}}{{5A0038}}
\definecolor{{milcTurquesa}}{{HTML}}{{0093A5}}
\definecolor{{milcVerde}}{{HTML}}{{58A923}}
\definecolor{{milcMostaza}}{{HTML}}{{E5B400}}
\definecolor{{milcOcre}}{{HTML}}{{B86600}}
\definecolor{{milcNegro}}{{HTML}}{{241816}}
\definecolor{{milcGris}}{{HTML}}{{F7F7F4}}
\definecolor{{milcLinea}}{{HTML}}{{E8E2DD}}
\color{{milcNegro}}

\pagestyle{{fancy}}
\fancyhf{{}}
\lhead{{\small Proyecto integrador · Grado {project.grade}}}
\rhead{{\small Periodo {tex_escape(pdata.name)} · MILC}}
\cfoot{{\small\thepage}}
\renewcommand{{\headrulewidth}}{{0pt}}

\newtcolorbox{{titlebox}}[1]{{enhanced,colback=#1,colframe=#1,arc=7mm,boxrule=0pt,left=7mm,right=7mm,top=3mm,bottom=3mm,before skip=8pt,after skip=8pt,fontupper=\sffamily\bfseries\Large\color{{white}}}}
\newtcolorbox{{softbox}}[3]{{enhanced,colback=#2,colframe=#1,borderline west={{4pt}}{{0pt}}{{#1}},arc=2mm,boxrule=.4pt,left=4mm,right=4mm,top=3mm,bottom=3mm,before skip=6pt,after skip=8pt,title={{#3}},coltitle=white,colbacktitle=#1,fonttitle=\sffamily\bfseries,attach boxed title to top left={{xshift=3mm,yshift=-2mm}},boxed title style={{arc=2mm, boxrule=0pt}}}}
\newtcolorbox{{drawbox}}[2]{{enhanced,colback=white,colframe=#1,arc=4mm,boxrule=1pt,left=5mm,right=5mm,top=4mm,bottom=4mm,before skip=8pt,after skip=8pt,title={{#2}},coltitle=white,colbacktitle=#1,fonttitle=\sffamily\bfseries,attach boxed title to top left={{xshift=4mm,yshift=-2mm}},boxed title style={{arc=3mm, boxrule=0pt}}}}
\newcommand{{\respuesta}}[1]{{\par\vspace{{2mm}}\foreach \n in {{1,...,#1}}{{\noindent\color{{milcLinea}}\rule{{\linewidth}}{{0.5pt}}\par\vspace{{5mm}}}}\color{{milcNegro}}}}
\newcommand{{\checkbox}}{{\textcolor{{milcTurquesa}}{{\fbox{{\phantom{{x}}}}}}\hspace{{5pt}}}}

\begin{{document}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[milcMagenta] (current page.south west) rectangle (current page.north east);
  \foreach \x in {{-1,1,...,23}}{{
    \foreach \y in {{-1,1,...,31}}{{
      \draw[milcVino, opacity=.12, line width=.6pt] (\x,\y) .. controls +(1,.6) and +(-1,-.6) .. +(2,1.4);
      \draw[milcVino, opacity=.10, line width=.6pt] (\x+.6,\y+.4) arc (210:510:.45);
    }}
  }}
\end{{tikzpicture}}

\vspace*{{.85cm}}
\begin{{tcolorbox}}[
  enhanced,colback=milcVino,colframe=milcVino,arc=12mm,width=.86\textwidth,
  boxrule=0pt,left=12mm,right=12mm,top=11mm,bottom=11mm,center,
  overlay={{\draw[milcMagenta,line width=2pt,dash pattern=on 9pt off 8pt,rounded corners=10mm]
  ([xshift=7mm,yshift=-7mm]frame.north west) rectangle ([xshift=-7mm,yshift=7mm]frame.south east);}}]
  \centering
  {{\sffamily\bfseries\color{{white}}\fontsize{{15}}{{17}}\selectfont PROYECTO INTEGRADOR · TECNOLOGÍA E INFORMÁTICA}}\\[1.0cm]
  {{\sffamily\bfseries\color{{white}}\fontsize{{30}}{{33}}\selectfont Proyecto de periodo {pdata.number}}}\\[3mm]
  {{\sffamily\bfseries\color{{milcMostaza}}\fontsize{{25}}{{28}}\selectfont {tex_escape(stem(project))}}}\\[8mm]
  {{\sffamily\color{{white}}\Large Grado {project.grade} · Periodo {tex_escape(pdata.name)}}}\\[1.0cm]
  {{\sffamily\bfseries\color{{white}}\large Institución Educativa Sor María Juliana}}\\[2mm]
  {{\sffamily\color{{white}}Autor: {tex_escape(AUTHOR)}}}\\[4mm]
  {{\sffamily\itshape\color{{white}}Proyecto activo: investigar, diseñar, prototipar, validar y reflexionar}}
\end{{tcolorbox}}
\vfill
\begin{{tcolorbox}}[enhanced,colback=milcVino!72!black,colframe=milcVino!72!black,arc=5mm,boxrule=0pt,width=.86\textwidth,left=6mm,right=6mm,top=4mm,bottom=4mm,center]
{{\color{{white}}\sffamily\small MILC · Design Thinking · Producto · Evidencias · Sustentación}}
\end{{tcolorbox}}

\newpage
\begin{{titlebox}}{{milcVino}}
Sentido del proyecto
\end{{titlebox}}
\begin{{multicols}}{{2}}
Este proyecto cierra el periodo mediante una experiencia activa. No se trata de repetir contenidos, sino de convertir los aprendizajes en una respuesta útil, visible y argumentada ante una necesidad escolar, comunitaria o digital.

\columnbreak
\begin{{softbox}}{{milcVerde}}{{milcGris}}{{Método MILC desde el plan de área}}
{tex_escape(milc_method)}
\end{{softbox}}
\end{{multicols}}

\begin{{softbox}}{{milcMostaza}}{{milcGris}}{{Reto central}}
{tex_escape(project.challenge)}
\end{{softbox}}

\begin{{tabularx}}{{\textwidth}}{{>{{\bfseries}}p{{3.4cm}}X}}
\rowcolor{{milcMagenta!12}}Autor & {tex_escape(AUTHOR)}\\
DBA del periodo & {tex_escape(pdata.dba)}\\
Estrategia activa & {tex_escape(project.active_strategy)}\\
Foco de diseño & {tex_escape(project.design_thinking_focus)}\\
Producto final & {tex_escape(project.product)}\\
\end{{tabularx}}

\begin{{titlebox}}{{milcTurquesa}}
Aprendizajes que integra
\end{{titlebox}}
\begin{{multicols}}{{2}}
\begin{{itemize}}[leftmargin=*]
{topic_items(project.topics)}
\end{{itemize}}
\end{{multicols}}

\newpage
\begin{{titlebox}}{{milcVerde}}
Fase 1 · Escuta y empatía
\end{{titlebox}}
\begin{{softbox}}{{milcVerde}}{{milcGris}}{{Propósito}}
Escuchar el contexto, reconocer una necesidad real y formular un problema que valga la pena trabajar. Esta fase puede incluir entrevistas, observación, encuesta breve o revisión de evidencias.
\end{{softbox}}
\checkbox Identificamos una necesidad escolar, comunitaria o digital.\\
\checkbox Escuchamos al menos una voz distinta a la del equipo.\\
\checkbox Formulamos una pregunta de proyecto clara.

\textbf{{Problema o necesidad detectada:}}
\respuesta{{2}}
\textbf{{Usuarios o personas afectadas:}}
\respuesta{{1}}
\textbf{{Pregunta de proyecto:}}
\respuesta{{2}}

\begin{{drawbox}}{{milcVerde}}{{Mapa de empatía}}
Dibuja o esquematiza qué piensa, siente, necesita y hace la persona o grupo relacionado con el problema.
\vspace{{5.2cm}}
\end{{drawbox}}

\newpage
\begin{{titlebox}}{{milcTurquesa}}
Fase 2 · Sistematización e investigación
\end{{titlebox}}
\begin{{softbox}}{{milcTurquesa}}{{milcGris}}{{Propósito}}
Organizar información, conceptos y evidencias para comprender el problema con rigor. Aquí se conectan los aprendizajes del periodo con datos, ejemplos, fuentes o pruebas.
\end{{softbox}}

\begin{{tabularx}}{{\textwidth}}{{>{{\bfseries}}p{{3.3cm}}X}}
Conceptos clave & {tex_escape(pdata.aprendizajes)}\\
Competencias & {tex_escape(pdata.competencias)}\\
Desempeño cognitivo & {tex_escape(pdata.cognitivos)}\\
Desempeño procedimental & {tex_escape(pdata.procedimentales)}\\
Desempeño actitudinal & {tex_escape(pdata.actitudinales)}\\
\end{{tabularx}}

\textbf{{Tres hallazgos de investigación:}}
\respuesta{{3}}
\textbf{{Criterios que debe cumplir la solución o producto:}}
\respuesta{{2}}

\newpage
\begin{{titlebox}}{{milcMagenta}}
Fase 3 · Praxis, diseño y prototipo
\end{{titlebox}}
\begin{{softbox}}{{milcMagenta}}{{milcGris}}{{Propósito}}
Pasar de la idea a la acción. El equipo diseña, experimenta, prototipa o produce una evidencia concreta que pueda revisarse y mejorarse.
\end{{softbox}}

\begin{{multicols}}{{2}}
\textbf{{Idea de solución o producto:}}
\respuesta{{2}}
\textbf{{Materiales, herramientas o recursos:}}
\respuesta{{2}}
\columnbreak
\textbf{{Experimento, prueba o validación:}}
\respuesta{{2}}
\textbf{{Mejora después de la prueba:}}
\respuesta{{2}}
\end{{multicols}}

\begin{{drawbox}}{{milcMagenta}}{{Boceto del prototipo o producto}}
Dibuja la primera versión del producto, campaña, informe, prototipo, recurso digital o solución que presentará tu equipo.
\vspace{{6cm}}
\end{{drawbox}}

\newpage
\begin{{titlebox}}{{milcMostaza}}
Entregables y sustentación
\end{{titlebox}}
\begin{{softbox}}{{milcMostaza}}{{milcGris}}{{Entregables mínimos}}
\checkbox Ficha de problema y pregunta de proyecto.\\
\checkbox Evidencias de escuta: entrevista, encuesta, observación o registro.\\
\checkbox Sistematización: conceptos, datos, fuentes o criterios.\\
\checkbox Producto final o prototipo.\\
\checkbox Sustentación breve con reflexión ética y social.
\end{{softbox}}

\begin{{tabularx}}{{\textwidth}}{{p{{4.1cm}}XXX}}
\rowcolor{{milcVino}}\color{{white}}\bfseries Criterio & \color{{white}}\bfseries 5 & \color{{white}}\bfseries 3 & \color{{white}}\bfseries 1\\
Problema y escuta & Problema real y bien sustentado. & Problema comprensible con poca evidencia. & Problema confuso o sin contexto.\\
Sistematización & Organiza conceptos y evidencias con rigor. & Presenta información básica. & Desordena o no conecta información.\\
Praxis y producto & Producto útil, probado y mejorado. & Producto funcional con ajustes pendientes. & Producto incompleto o poco pertinente.\\
Comunicación & Sustenta con claridad y evidencia. & Explica parcialmente. & No logra justificar decisiones.\\
Evaluación liberadora & Reflexiona sobre impacto y responsabilidad. & Reflexiona de forma general. & No relaciona el proyecto con la comunidad.\\
\end{{tabularx}}

\newpage
\begin{{titlebox}}{{milcVino}}
Fase 4 · Evaluación liberadora
\end{{titlebox}}
\begin{{softbox}}{{milcVino}}{{milcGris}}{{Cierre}}
La evaluación liberadora pregunta qué aprendimos, a quién sirve lo producido y qué responsabilidad asumimos frente a la tecnología, la comunidad y la infoesfera.
\end{{softbox}}
\textbf{{Lo más importante que aprendimos fue:}}
\respuesta{{2}}
\textbf{{Nuestro producto ayuda porque:}}
\respuesta{{2}}
\textbf{{Una mejora posible sería:}}
\respuesta{{2}}
\begin{{drawbox}}{{milcVino}}{{Símbolo del proyecto}}
Dibuja un símbolo que represente el sentido social, ético o comunitario del proyecto.
\vspace{{4.4cm}}

{{\small\bfseries Nuestro símbolo significa:}} \textcolor{{milcLinea}}{{\rule{{.60\linewidth}}{{0.5pt}}}}
\end{{drawbox}}
\end{{document}}
"""


def setup_dirs() -> dict[str, Path]:
    dirs = {
        "base": OUT,
        "tex": OUT / "tex",
        "pdf": OUT / "pdf",
        "build": OUT / "build",
        "fuente": OUT / "fuente",
        "scripts": OUT / "scripts",
    }
    for d in dirs.values():
        d.mkdir(parents=True, exist_ok=True)
    return dirs


def compile_tex(tex_path: Path, pdf_out: Path, build_dir: Path) -> tuple[bool, str]:
    job_dir = build_dir / tex_path.stem
    if job_dir.exists():
        shutil.rmtree(job_dir)
    job_dir.mkdir(parents=True)
    work_tex = job_dir / tex_path.name
    shutil.copy2(tex_path, work_tex)
    proc = subprocess.run(
        [XELATEX, "-interaction=nonstopmode", "-halt-on-error", work_tex.name],
        cwd=job_dir,
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        return False, (proc.stdout + "\n" + proc.stderr)[-2400:]
    shutil.copy2(job_dir / f"{tex_path.stem}.pdf", pdf_out)
    return True, "ok"


def main() -> None:
    dirs = setup_dirs()
    milc_method = extract_milc_method()
    rows = []
    fuente = []
    for grade in range(6, 12):
        for period in range(1, 4):
            project = build_project(grade, period)
            name = stem(project)
            tex_path = dirs["tex"] / f"{name}.tex"
            pdf_path = dirs["pdf"] / f"{name}.pdf"
            tex_path.write_text(tex_template(project, milc_method), encoding="utf-8")
            ok, message = compile_tex(tex_path, pdf_path, dirs["build"])
            print(f"{name}: {'PDF ok' if ok else 'PDF error'}")
            rows.append(
                {
                    "proyecto": name,
                    "grado": grade,
                    "periodo": period,
                    "tex": str(tex_path),
                    "pdf": str(pdf_path) if ok else "",
                    "estado_pdf": "ok" if ok else "error",
                    "mensaje": message,
                }
            )
            fuente.append(
                {
                    "proyecto": name,
                    "grado": grade,
                    "periodo": period,
                    "temas": project.topics,
                    "estrategia_activa": project.active_strategy,
                    "producto": project.product,
                }
            )
    with (OUT / "manifest_proyectos.csv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
    (dirs["fuente"] / "proyectos_periodo.json").write_text(json.dumps(fuente, ensure_ascii=False, indent=2), encoding="utf-8")
    shutil.copy2(Path(__file__), dirs["scripts"] / "generar_proyectos_milc.py")
    print(f"\nSalida: {OUT}")


if __name__ == "__main__":
    main()

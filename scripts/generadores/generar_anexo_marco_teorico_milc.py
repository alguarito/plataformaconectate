from __future__ import annotations

import json
import os
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "00_PLAN_AREA" / "anexo_marco_teorico_milc"
XELATEX = os.environ.get("XELATEX", "xelatex")
AUTHOR = "Dr. Álvaro Cárdenas Orozco"

PALETTE = {
    "magenta": "E6007E",
    "vino": "5A0038",
    "verde": "58A923",
    "turquesa": "0093A5",
    "mostaza": "E5B400",
    "gris": "F4F2F4",
    "negro": "241A1D",
}


@dataclass
class Chapter:
    title: str
    color: str
    paragraphs: list[str]
    bullets: list[str]
    synthesis: str


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in str(text))


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for name in ("pdf", "tex", "docx"):
        (OUT / name).mkdir(exist_ok=True)


def chapters() -> list[Chapter]:
    return [
        Chapter(
            title="1. Presentación y estatuto académico del anexo",
            color=PALETTE["vino"],
            paragraphs=[
                "Este anexo desarrolla el marco teórico que soporta el Plan de Área de Tecnología e Informática 2025 y, especialmente, la formalización inicial del Modelo MILC como propuesta pedagógico-investigativa original. Su finalidad no es decorar el documento institucional con autores, sino construir una base conceptual que explique por qué el área requiere una reorganización epistemológica, metodológica, evaluativa y cultural.",
                "MILC se presenta aquí como un modelo en proceso de consolidación teórica y validación pedagógica. Esto significa que no se formula como teoría cerrada ni como receta universal; se propone como arquitectura de trabajo que integra investigación escolar, liberación, tecnología, territorio, pensamiento computacional, socioemocionalidad, ciudadanía digital y evaluación transformadora.",
                "El anexo adopta una escritura académica de carácter argumentativo: identifica un problema educativo, delimita categorías, articula referentes, propone relaciones conceptuales y deriva consecuencias curriculares concretas para guías, proyectos, exámenes y evaluación.",
            ],
            bullets=[
                "Objeto del anexo: fundamentar teóricamente el Modelo MILC y su pertinencia curricular.",
                "Objeto pedagógico: transformar la clase de Tecnología e Informática en experiencia de investigación, creación y responsabilidad.",
                "Objeto institucional: ofrecer soporte académico para el plan de estudios y para futuras fases de validación.",
            ],
            synthesis="MILC debe leerse como un modelo original en formalización: situado, crítico, investigativo y orientado a la praxis tecnológica liberadora.",
        ),
        Chapter(
            title="2. Problema epistemológico: tecnología sin conciencia crítica",
            color=PALETTE["magenta"],
            paragraphs=[
                "La educación tecnológica escolar suele oscilar entre dos reduccionismos. El primero es instrumental: enseñar tecnología como dominio de herramientas, botones, programas o instrucciones. El segundo es tecnocrático: asumir que la innovación consiste en incorporar dispositivos o plataformas sin transformar las relaciones pedagógicas. Ambos reduccionismos producen estudiantes que ejecutan tareas, pero no necesariamente comprenden sistemas, formulan problemas, reconocen impactos o toman decisiones éticas.",
                "Este problema se agrava en la cultura digital contemporánea. Los estudiantes acceden a información abundante, interfaces intuitivas y sistemas automatizados, pero no siempre desarrollan criterio para validar fuentes, proteger datos, comprender sesgos algorítmicos, regular su atención, crear productos propios o reconocer la dimensión social de una tecnología. La alfabetización tecnológica, por tanto, no puede limitarse a operar recursos digitales; debe formar juicio.",
                "Desde esta perspectiva, el problema epistemológico del área consiste en definir qué tipo de conocimiento tecnológico se quiere construir. MILC responde que el conocimiento tecnológico escolar debe ser situado, problematizador, documentado, productivo, colaborativo y éticamente evaluable. No se aprende tecnología únicamente para usar artefactos; se aprende para leer, cuidar y transformar realidades mediadas por sistemas técnicos.",
            ],
            bullets=[
                "Reduccionismo instrumental: saber usar sin saber explicar, valorar o transformar.",
                "Reduccionismo tecnocrático: introducir tecnología sin cambiar la pedagogía.",
                "Reto curricular: pasar de usuarios competentes a sujetos críticos, creadores y responsables.",
            ],
            synthesis="La pregunta de fondo no es qué herramienta se enseña, sino qué sujeto se forma cuando se enseña tecnología.",
        ),
        Chapter(
            title="3. Crítica a la educación bancaria y fundamento freireano",
            color=PALETTE["verde"],
            paragraphs=[
                "Paulo Freire ofrece una clave imprescindible para comprender la necesidad de MILC. Su crítica a la educación bancaria denuncia una relación pedagógica en la que el docente deposita contenidos y el estudiante aparece como recipiente. En esa lógica, el aprendizaje se mide por repetición, obediencia y acumulación; la realidad del estudiante queda subordinada al programa y el conocimiento pierde su potencia transformadora.",
                "En Tecnología e Informática, la educación bancaria aparece cuando el estudiante copia procedimientos, memoriza definiciones, reproduce formatos o entrega productos sin comprender el problema que resuelve. También aparece cuando la evaluación premia la apariencia final del archivo, pero no la pregunta, el proceso, la toma de decisiones, la ética de la información ni la reflexión sobre impactos.",
                "MILC se alinea con una pedagogía problematizadora: parte de una situación real, exige escucha del contexto, promueve sistematización de evidencias, conduce a una praxis tecnológica y cierra con evaluación liberadora. El estudiante deja de ser receptor de instrucciones para convertirse en investigador escolar de su propia realidad tecnológica.",
                "La categoría freireana de praxis es central: reflexión y acción sobre el mundo para transformarlo. En MILC, la praxis se concreta en productos, prototipos, campañas, documentos, datos, programas, diagnósticos, diseños o soluciones que no son fines en sí mismos, sino mediaciones para pensar y actuar.",
            ],
            bullets=[
                "De depósito de contenidos a formulación de problemas.",
                "De copia de procedimientos a producción argumentada.",
                "De nota como cierre a evaluación como conciencia del proceso.",
                "De herramienta neutral a tecnología situada en relaciones sociales.",
            ],
            synthesis="MILC traduce la pedagogía problematizadora al aula de Tecnología e Informática.",
        ),
        Chapter(
            title="4. Desarraigo cultural, territorio y reconocimiento",
            color=PALETTE["turquesa"],
            paragraphs=[
                "Un currículo tecnológico puede fracasar cuando enseña herramientas globales sin reconocer realidades locales. El desarraigo curricular ocurre cuando los temas escolares no dialogan con la cultura, el territorio, los lenguajes, los oficios, las memorias, los riesgos y las condiciones materiales de los estudiantes. En ese escenario, la tecnología se presenta como mundo ajeno, externo y neutral.",
                "MILC propone una educación tecnológica con raíces. Esto implica reconocer que la escuela está situada en una comunidad concreta y que sus problemas no son abstractos: conectividad irregular, brechas de acceso, uso familiar del celular, comunicación institucional, exposición a fraudes, circulación de desinformación, relación con dispositivos, residuos electrónicos, economía local, identidad digital y proyectos de vida.",
                "La dimensión territorial no significa encerrar el currículo en lo local, sino partir de lo local para dialogar críticamente con lo global. El estudiante puede aprender inteligencia artificial, programación, diseño web o análisis de datos, pero debe hacerlo preguntándose por su comunidad, sus necesidades, sus saberes y sus formas de cuidado.",
                "Desde una perspectiva de reconocimiento cultural, MILC evita que el estudiante se sienta consumidor periférico de tecnologías diseñadas por otros. Lo invita a reconocerse como sujeto capaz de investigar, crear, narrar, diseñar, mejorar y proponer soluciones desde su lugar.",
            ],
            bullets=[
                "Territorio: contexto vivo de problemas, saberes y posibilidades.",
                "Cultura digital: prácticas reales de comunicación, consumo, creación y riesgo.",
                "Reconocimiento: el estudiante como sujeto con historia, lenguaje y agencia.",
                "Pertinencia: relación entre currículo, comunidad y proyecto de vida.",
            ],
            synthesis="Una educación tecnológica liberadora no desarraiga: conecta la técnica con la cultura y el territorio.",
        ),
        Chapter(
            title="5. Filosofía de la liberación y tecnología: aporte de Dussel",
            color=PALETTE["vino"],
            paragraphs=[
                "Enrique Dussel aporta a MILC una exigencia ética fundamental: partir del otro. La tecnología no puede evaluarse únicamente por eficiencia, velocidad, novedad o rendimiento; debe preguntarse por las vidas que afecta, por las exclusiones que reproduce y por las posibilidades de dignidad que abre o cierra.",
                "La filosofía de la liberación permite problematizar la neutralidad tecnológica. Todo sistema técnico organiza relaciones: quién accede, quién decide, quién queda vigilado, quién produce datos, quién se beneficia, quién asume costos ambientales y quién queda por fuera. Esta pregunta es especialmente importante en contextos escolares con desigualdades de conectividad, capital cultural y acompañamiento familiar.",
                "MILC incorpora esta perspectiva mediante la Escuta, entendida como escucha ética del contexto. Antes de diseñar una solución, el estudiante debe reconocer a los usuarios, las condiciones, los límites y las consecuencias. La tecnología se convierte así en mediación de responsabilidad y no solo en demostración de habilidad.",
                "La dimensión liberadora no significa rechazar la tecnología, sino someterla a juicio ético. Una solución tecnológica escolar debe ser útil, comprensible, segura, situada, inclusiva y respetuosa de la dignidad de las personas.",
            ],
            bullets=[
                "El otro como punto de partida de la innovación.",
                "Crítica a la neutralidad tecnológica.",
                "Tecnología como mediación de justicia, cuidado y responsabilidad.",
                "Escuta como principio ético antes de la producción técnica.",
            ],
            synthesis="Desde Dussel, MILC entiende que toda creación tecnológica debe responder a una responsabilidad con el otro.",
        ),
        Chapter(
            title="6. Infoesfera, vida onlife y ética de la información",
            color=PALETTE["magenta"],
            paragraphs=[
                "Luciano Floridi permite comprender que los estudiantes habitan una infoesfera: un entorno donde datos, identidades, relaciones, dispositivos y plataformas configuran la experiencia cotidiana. La vida digital no es un espacio separado de la vida real; es parte constitutiva de la vida social, escolar, afectiva y cultural.",
                "Esta comprensión modifica el sentido de Tecnología e Informática. No basta enseñar documentos, hojas de cálculo, programación o redes como contenidos técnicos. Es necesario formar conciencia sobre identidad digital, huella informacional, privacidad, calidad de datos, reputación, desinformación, inteligencia artificial y responsabilidad comunicativa.",
                "MILC incorpora la infoesfera como campo de investigación escolar. Los estudiantes pueden analizar cómo circula una noticia, cómo se valida una fuente, cómo se protege una contraseña, cómo se representa un dato, cómo un algoritmo recomienda contenido o cómo una decisión técnica afecta la convivencia.",
                "La ética de la información se convierte en un eje transversal: producir información exige honestidad; compartir información exige cuidado; usar datos exige consentimiento y pertinencia; automatizar decisiones exige criterio humano.",
            ],
            bullets=[
                "Infoesfera: entorno informacional donde se configuran identidad, relación y conocimiento.",
                "Vida onlife: continuidad entre experiencia física y digital.",
                "Soberanía informacional: capacidad de buscar, validar, proteger y producir información con criterio.",
                "Ética digital: cuidado de datos, fuentes, reputación, autoría e impacto.",
            ],
            synthesis="MILC forma estudiantes capaces de habitar la infoesfera con criterio, cuidado y agencia.",
        ),
        Chapter(
            title="7. Pensamiento computacional como estructura cognitiva",
            color=PALETTE["verde"],
            paragraphs=[
                "El pensamiento computacional, popularizado por Jeannette Wing, ofrece a MILC una estructura cognitiva para formular problemas y diseñar soluciones representables. Su importancia no radica únicamente en aprender programación, sino en desarrollar formas de razonamiento transferibles: descomponer, abstraer, reconocer patrones, modelar, automatizar, depurar y evaluar.",
                "En el plan de área, el pensamiento computacional aparece desde sexto hasta undécimo de manera progresiva. En los grados iniciales se vincula con secuencias, instrucciones, organización de información y solución de problemas cotidianos. En grados intermedios se relaciona con programación visual, datos, redes y prototipos. En la media se proyecta hacia diseño web, inteligencia artificial, proyectos tecnológicos y gestión de soluciones.",
                "MILC amplía el pensamiento computacional al integrarlo con contexto y ética. Formular un problema no es solo representarlo de manera computable; también implica preguntar si el problema es pertinente, si la solución es justa, si los datos son confiables, si el usuario fue escuchado y si el producto mejora una situación real.",
                "De este modo, el pensamiento computacional deja de ser una competencia abstracta y se convierte en herramienta de praxis: una manera de organizar la acción tecnológica para intervenir responsablemente el mundo.",
            ],
            bullets=[
                "Descomposición: dividir problemas complejos en partes abordables.",
                "Abstracción: identificar lo esencial y representar lo relevante.",
                "Patrones: reconocer regularidades y relaciones.",
                "Algoritmos: diseñar secuencias lógicas de acción.",
                "Depuración: revisar errores como oportunidad de aprendizaje.",
                "Evaluación: valorar eficiencia, claridad, impacto y pertinencia.",
            ],
            synthesis="MILC convierte el pensamiento computacional en pensamiento situado: técnico, crítico y ético.",
        ),
        Chapter(
            title="8. Educación socioemocional, cuidado y convivencia tecnológica",
            color=PALETTE["turquesa"],
            paragraphs=[
                "La tecnología no solo involucra habilidades técnicas; compromete emociones, vínculos, atención, identidad y convivencia. Por eso, una propuesta curricular robusta debe integrar educación socioemocional. Aprender con tecnología exige tolerancia a la frustración, regulación ante el error, escucha, empatía, colaboración, comunicación clara y toma responsable de decisiones.",
                "El marco CASEL identifica competencias socioemocionales como autoconciencia, autorregulación, conciencia social, habilidades relacionales y toma responsable de decisiones. Estas competencias dialogan directamente con MILC. La Escuta requiere conciencia social; la Sistematización exige autorregulación cognitiva; la Praxis demanda colaboración; la Evaluación liberadora exige autoconciencia, honestidad y apertura a la mejora.",
                "En el aula de Tecnología e Informática, la educación socioemocional se expresa en prácticas concretas: cuidado de equipos compartidos, respeto por roles de trabajo, escucha de usuarios, manejo de desacuerdos, reconocimiento del error, prevención del ciberacoso, comunicación responsable y protección de datos personales.",
                "MILC asume que no hay ciudadanía digital sin formación del carácter. El estudiante debe aprender a crear, pero también a cuidar; a argumentar, pero también a escuchar; a resolver, pero también a reconocer límites.",
            ],
            bullets=[
                "Autoconciencia: reconocer emociones, fortalezas y límites ante retos tecnológicos.",
                "Autorregulación: sostener esfuerzo, manejar error y actuar con responsabilidad.",
                "Conciencia social: escuchar usuarios, compañeros y comunidad.",
                "Habilidades relacionales: colaborar, comunicar y resolver conflictos.",
                "Decisión responsable: valorar consecuencias éticas, sociales y digitales.",
            ],
            synthesis="MILC integra tecnología y socioemocionalidad porque crear sin cuidado puede reproducir daño.",
        ),
        Chapter(
            title="9. Inteligencia artificial, ciudadanía digital y soberanía informacional",
            color=PALETTE["mostaza"],
            paragraphs=[
                "La inteligencia artificial introduce un nuevo umbral curricular. Los estudiantes ya no solo consumen información: interactúan con sistemas que clasifican, recomiendan, predicen, generan contenido y median decisiones. Por ello, el área debe preparar usuarios críticos y co-creadores responsables, no consumidores pasivos de automatización.",
                "El marco de competencias en IA para estudiantes de UNESCO plantea dimensiones como mentalidad centrada en el ser humano, ética de la IA, técnicas y aplicaciones de IA, y diseño de sistemas de IA. Estas dimensiones se alinean con MILC porque exigen comprender, aplicar y crear con responsabilidad.",
                "La soberanía informacional se vuelve una categoría decisiva. Implica que el estudiante pueda buscar información, valorar fuentes, detectar sesgos, proteger datos, reconocer límites de la automatización, citar autoría, evitar plagio, decidir cuándo usar IA y explicar por qué una salida generada requiere verificación humana.",
                "En lugar de prohibir o idealizar la IA, MILC propone investigarla: preguntar cómo funciona, qué datos usa, qué riesgos trae, qué posibilidades abre, cómo afecta la creatividad y qué responsabilidades exige a estudiantes y docentes.",
            ],
            bullets=[
                "Mentalidad humano-centrada: la IA debe ampliar agencia humana, no sustituir juicio.",
                "Ética de IA: sesgo, privacidad, transparencia, autoría y responsabilidad.",
                "Aplicaciones: uso crítico de herramientas para investigar, crear y mejorar.",
                "Diseño: prototipar soluciones inclusivas, sostenibles y verificables.",
            ],
            synthesis="MILC orienta la IA hacia ciudadanía, criterio y creación responsable.",
        ),
        Chapter(
            title="10. Formalización del Modelo MILC",
            color=PALETTE["vino"],
            paragraphs=[
                "MILC se define como un modelo pedagógico-investigativo original, en proceso de formalización, orientado a integrar investigación escolar, pensamiento tecnológico, conciencia crítica, reconocimiento cultural, educación socioemocional y evaluación liberadora. Su estructura se organiza en cuatro fases: Escuta, Sistematización, Praxis y Evaluación liberadora.",
                "Escuta es la fase de apertura ética y contextual. Consiste en escuchar la realidad antes de intervenirla: observar, dialogar, identificar necesidades, reconocer usuarios, comprender condiciones y formular preguntas significativas. Se diferencia de una simple motivación inicial porque constituye el fundamento epistemológico del problema.",
                "Sistematización es la fase de ordenamiento cognitivo. El estudiante organiza información, construye categorías, valida fuentes, reconoce patrones, selecciona criterios y documenta su proceso. Esta fase evita que la acción tecnológica sea improvisada.",
                "Praxis es la fase de acción transformadora. Implica diseñar, programar, construir, prototipar, comunicar, experimentar o intervenir. No es activismo sin reflexión; es acción guiada por el análisis previo.",
                "Evaluación liberadora es la fase de juicio crítico y mejora. Valora evidencias, impactos, aprendizajes, errores, decisiones éticas y compromisos. Libera porque permite al estudiante comprender su proceso, reconocer avances y proyectar nuevas acciones.",
            ],
            bullets=[
                "Principio de contexto: todo aprendizaje parte de una realidad escuchada.",
                "Principio de evidencia: todo proceso debe documentarse.",
                "Principio de praxis: todo saber debe poder convertirse en acción responsable.",
                "Principio de liberación: toda evaluación debe producir conciencia y posibilidad de mejora.",
                "Principio de cultura: toda tecnología debe dialogar con territorio, memoria y comunidad.",
            ],
            synthesis="MILC organiza la clase como ciclo de investigación, producción y conciencia crítica.",
        ),
        Chapter(
            title="11. Categorías estructurales del Modelo MILC",
            color=PALETTE["verde"],
            paragraphs=[
                "La formalización de un modelo pedagógico requiere declarar sus categorías estructurales. En MILC, estas categorías permiten pasar de una intuición metodológica a una arquitectura conceptual defendible. Las categorías no son términos ornamentales; son unidades de análisis que orientan la planeación, la observación, la evaluación y la futura investigación del modelo.",
                "La primera categoría es contexto. MILC parte de que ningún aprendizaje tecnológico ocurre en vacío: toda práctica se inscribe en una institución, una cultura, una historia, una infraestructura, una economía de acceso y unas relaciones de poder. El contexto define la pertinencia de la pregunta y evita que la tecnología sea enseñada como abstracción desarraigada.",
                "La segunda categoría es evidencia. En MILC, aprender no es afirmar que se sabe, sino producir rastros verificables del proceso: preguntas, bocetos, bitácoras, pruebas, errores, versiones, capturas, datos, reflexiones, productos y sustentaciones. La evidencia permite evaluar de forma justa y construir memoria pedagógica.",
                "La tercera categoría es mediación tecnológica. La tecnología no se entiende como fin, sino como mediación cultural, cognitiva y social. Una hoja de cálculo, un algoritmo, una presentación, una red o una herramienta de IA son medios para representar, comunicar, resolver, colaborar o decidir.",
                "La cuarta categoría es praxis. Esta categoría articula reflexión y acción. Una clase MILC debe conducir a una intervención: crear un producto, validar una fuente, construir un protocolo, programar una solución, diagnosticar un sistema, diseñar una campaña o sustentar una decisión.",
                "La quinta categoría es liberación. No se trata de un concepto retórico, sino de un criterio pedagógico: el estudiante debe ganar autonomía, criterio, voz, responsabilidad y capacidad de actuar sobre su realidad tecnológica. La liberación se expresa cuando el estudiante deja de depender de instrucciones y comienza a justificar decisiones.",
            ],
            bullets=[
                "Contexto: realidad situada que da sentido al problema.",
                "Evidencia: registro verificable del aprendizaje y de la toma de decisiones.",
                "Mediación tecnológica: herramienta como puente entre pensamiento, cultura y acción.",
                "Praxis: acción reflexiva que transforma una situación.",
                "Liberación: construcción de autonomía, criterio y responsabilidad.",
                "Cuidado: principio transversal de protección de personas, datos, equipos, vínculos y ambiente.",
            ],
            synthesis="Las categorías MILC permiten convertir el modelo en objeto de planeación, evaluación e investigación.",
        ),
        Chapter(
            title="12. Rol docente, rol estudiante y mediaciones didácticas",
            color=PALETTE["turquesa"],
            paragraphs=[
                "MILC redefine los roles del aula. El docente no desaparece ni se reduce a facilitador genérico; asume una función más compleja: diseña ambientes de investigación, selecciona problemas, regula el nivel de desafío, ofrece andamiajes, cuida la dimensión ética y produce retroalimentación formativa. Su autoridad no se funda en entregar respuestas, sino en sostener condiciones para pensar mejor.",
                "El estudiante, por su parte, deja de ser ejecutor de instrucciones para convertirse en sujeto investigador. Esto implica preguntar, registrar, comparar, ensayar, equivocarse, corregir, explicar y sustentar. El estudiante MILC no solo entrega productos; reconstruye el proceso por el cual llegó a ellos.",
                "Las mediaciones didácticas del modelo incluyen guías, bitácoras, rúbricas, mapas de problema, protocolos de cuidado, plantillas de diseño, laboratorios técnicos, ejercicios PRIMM, retos de datos, portafolios y proyectos integradores. Estas mediaciones cumplen una función de andamiaje: permiten que la autonomía se forme progresivamente y no sea exigida de manera improvisada.",
                "El aula MILC combina momentos de instrucción explícita con momentos de exploración, producción, diálogo y evaluación. No rechaza la explicación docente; la ubica dentro de una secuencia más amplia donde explicar sirve para investigar y producir, no para cerrar la pregunta.",
            ],
            bullets=[
                "Docente como diseñador curricular, mediador epistemológico y orientador ético.",
                "Estudiante como investigador escolar, productor de evidencias y sujeto de reflexión.",
                "Guía como unidad didáctica mínima del modelo.",
                "Proyecto como dispositivo de integración y transferencia.",
                "Portafolio como memoria del aprendizaje y base de evaluación liberadora.",
            ],
            synthesis="MILC transforma roles: enseñar es mediar investigación; aprender es producir evidencia con sentido.",
        ),
        Chapter(
            title="13. Sistema de evaluación del Modelo MILC",
            color=PALETTE["magenta"],
            paragraphs=[
                "La evaluación en MILC no se limita a comprobar resultados. Evalúa el proceso completo por el cual el estudiante escucha, organiza, actúa y reflexiona. Por ello, combina evaluación diagnóstica, formativa, sumativa, autoevaluativa y coevaluativa.",
                "En Escuta se evalúa la pertinencia de la pregunta, la escucha del usuario, la identificación del contexto y la sensibilidad ética. En Sistematización se evalúa la calidad de fuentes, la organización de información, la claridad conceptual y la documentación. En Praxis se evalúa la calidad del producto, el uso técnico, la colaboración, la creatividad y la funcionalidad. En Evaluación liberadora se evalúa la reflexión, la mejora, la sustentación y el compromiso.",
                "La ponderación orientativa por periodo puede organizarse así: guías y portafolio 30%, proyecto o producto integrador 30%, prueba aplicada o examen final 20%, participación, cuidado, colaboración y reflexión MILC 20%. Esta distribución debe ajustarse al SIEE institucional, pero ofrece coherencia entre proceso, producto, transferencia y actitud.",
                "Los instrumentos sugeridos son rúbricas, listas de cotejo, bitácoras, portafolios, guías, proyectos, pruebas aplicadas, sustentaciones, autoevaluación y coevaluación. La retroalimentación debe ser descriptiva y orientada a la mejora, no únicamente numérica.",
            ],
            bullets=[
                "Criterio cognitivo: comprensión, explicación, relación conceptual y juicio informacional.",
                "Criterio procedimental: producción, uso técnico, documentación, programación, análisis y prototipado.",
                "Criterio socioemocional: colaboración, autorregulación, comunicación, empatía y cuidado.",
                "Criterio ético-cultural: responsabilidad digital, reconocimiento del contexto, autoría y pertinencia social.",
                "Criterio de mejora: capacidad de revisar, corregir, sustentar y transferir aprendizajes.",
            ],
            synthesis="Evaluar en MILC es hacer visible el aprendizaje para transformarlo en autonomía y responsabilidad.",
        ),
        Chapter(
            title="14. Matriz de coherencia curricular del modelo",
            color=PALETTE["mostaza"],
            paragraphs=[
                "Una propuesta curricular doctoralmente defendible debe mostrar coherencia entre fundamentos, metodología, contenidos, recursos y evaluación. La matriz de coherencia de MILC permite revisar si cada decisión curricular responde a una arquitectura común o si aparece como actividad aislada.",
                "En la fase de Escuta, la coherencia se observa cuando la guía parte de un problema real o verosímil, reconoce usuarios, formula una pregunta orientadora y activa saberes previos. En Sistematización, se observa cuando el estudiante organiza información, construye criterios y produce una explicación. En Praxis, se evidencia cuando el estudiante diseña o produce algo que responde al problema. En Evaluación liberadora, se verifica cuando el estudiante revisa impactos, mejora el producto y explicita aprendizajes.",
                "Esta matriz también permite revisar proyectos integradores. Un proyecto sin Escuta corre el riesgo de ser producto decorativo; sin Sistematización, se vuelve improvisación; sin Praxis, queda en discurso; sin Evaluación liberadora, termina en entrega sin conciencia.",
                "La coherencia curricular exige que los instrumentos de evaluación no contradigan la metodología. Si MILC promueve investigación y creación, la evaluación debe valorar proceso, criterio, evidencia, producción, reflexión y mejora. Evaluar solo respuestas cerradas rompería la lógica del modelo.",
            ],
            bullets=[
                "Pregunta orientadora: conecta contexto, contenido y propósito.",
                "Conceptos clave: dan lenguaje para pensar el problema.",
                "Actividad: convierte comprensión en acción guiada.",
                "Producto: materializa la praxis tecnológica.",
                "Reflexión: permite conciencia crítica y transferencia.",
                "Rúbrica: traduce el modelo en criterios observables.",
            ],
            synthesis="La coherencia MILC se verifica cuando pregunta, contenido, actividad, producto y evaluación responden al mismo sentido formativo.",
        ),
        Chapter(
            title="15. Alcances, límites y criterios de validación futura",
            color=PALETTE["turquesa"],
            paragraphs=[
                "Como modelo original en construcción, MILC requiere una agenda de validación. Su valor inicial se encuentra en la coherencia teórica y en la capacidad de organizar recursos curriculares, pero su consolidación exige observar implementación, recoger evidencias, analizar resultados y ajustar categorías.",
                "Los alcances del modelo incluyen su capacidad para articular tecnología, investigación, territorio, pensamiento computacional, educación socioemocional y evaluación. Sus límites se relacionan con condiciones institucionales: tiempo, conectividad, formación docente, disponibilidad de equipos, tamaño de grupos y cultura evaluativa.",
                "La validación futura puede realizarse mediante investigación-acción, análisis de portafolios, entrevistas a estudiantes, diarios docentes, rúbricas comparadas, observación de clases y evaluación de productos. También puede construirse una matriz de indicadores para medir autonomía, calidad de evidencias, transferencia, criterio ético y participación.",
                "La formalización doctoral del modelo exigiría precisar sus categorías, construir instrumentos, delimitar población, diseñar ciclos de intervención, triangular datos y analizar transformaciones en prácticas de aula.",
            ],
            bullets=[
                "Coherencia interna: relación entre principios, fases, instrumentos y evidencias.",
                "Pertinencia contextual: capacidad de responder a problemas reales de la comunidad escolar.",
                "Transferencia: posibilidad de aplicar aprendizajes a situaciones nuevas.",
                "Aceptabilidad docente y estudiantil: claridad, usabilidad y sentido del modelo.",
                "Impacto formativo: autonomía, pensamiento crítico, producción tecnológica y responsabilidad digital.",
            ],
            synthesis="MILC debe crecer como modelo investigable: abierto a evidencia, crítica, ajuste y validación.",
        ),
    ]


def references() -> list[str]:
    return [
        "CASEL. (2024). What Is the CASEL Framework? Collaborative for Academic, Social, and Emotional Learning.",
        "Dussel, E. (1996). Filosofía de la liberación. Nueva América.",
        "Floridi, L. (2014). The Fourth Revolution: How the Infosphere is Reshaping Human Reality. Oxford University Press.",
        "Freire, P. (1970). Pedagogía del oprimido. Siglo XXI.",
        "Ministerio de Educación Nacional de Colombia. (2008). Guía No. 30: Ser competente en tecnología: una necesidad para el desarrollo.",
        "UNESCO. (2024). AI competency framework for students. UNESCO.",
        "Wing, J. M. (2006). Computational Thinking. Communications of the ACM, 49(3), 33-35.",
        "Wing, J. M. (2008). Computational thinking and thinking about computing. Philosophical Transactions of the Royal Society A, 366, 3717-3725.",
    ]


def color_name(hex_color: str) -> str:
    return {
        PALETTE["vino"]: "vino",
        PALETTE["magenta"]: "magentaMILC",
        PALETTE["verde"]: "verde",
        PALETTE["turquesa"]: "turquesa",
        PALETTE["mostaza"]: "mostaza",
    }.get(hex_color, "vino")


def tex_chapter(chapter: Chapter) -> str:
    c = color_name(chapter.color)
    pieces = [
        rf"\chapter*{{{tex_escape(chapter.title)}}}",
        rf"\addcontentsline{{toc}}{{chapter}}{{{tex_escape(chapter.title)}}}",
        rf"\begin{{tcolorbox}}[enhanced,colback={c}!6,colframe={c},arc=2mm,boxrule=.7pt,left=5mm,right=5mm,top=4mm,bottom=4mm]",
    ]
    for paragraph in chapter.paragraphs:
        pieces.append(tex_escape(paragraph) + "\n")
    if chapter.bullets:
        pieces.append(r"\begin{itemize}")
        for item in chapter.bullets:
            pieces.append(rf"\item {tex_escape(item)}")
        pieces.append(r"\end{itemize}")
    pieces.extend(
        [
            rf"\begin{{tcolorbox}}[colback=white,colframe={c},title=\textbf{{Síntesis}}]",
            tex_escape(chapter.synthesis),
            r"\end{tcolorbox}",
            r"\end{tcolorbox}",
        ]
    )
    return "\n".join(pieces)


def build_tex() -> Path:
    body = "\n\n".join(tex_chapter(chapter) for chapter in chapters())
    refs = "\n".join(rf"\item {tex_escape(ref)}" for ref in references())
    tex = rf"""
\documentclass[10pt,letterpaper]{{report}}
\usepackage[margin=1.65cm]{{geometry}}
\usepackage{{fontspec}}
\usepackage{{xcolor}}
\usepackage{{tikz}}
\usepackage[most]{{tcolorbox}}
\usepackage{{hyperref}}
\usepackage{{enumitem}}
\setmainfont{{Helvetica Neue}}
\setsansfont{{Helvetica Neue}}
\definecolor{{magentaMILC}}{{HTML}}{{{PALETTE['magenta']}}}
\definecolor{{vino}}{{HTML}}{{{PALETTE['vino']}}}
\definecolor{{verde}}{{HTML}}{{{PALETTE['verde']}}}
\definecolor{{turquesa}}{{HTML}}{{{PALETTE['turquesa']}}}
\definecolor{{mostaza}}{{HTML}}{{{PALETTE['mostaza']}}}
\definecolor{{gris}}{{HTML}}{{{PALETTE['gris']}}}
\definecolor{{negro}}{{HTML}}{{{PALETTE['negro']}}}
\hypersetup{{colorlinks=true,linkcolor=vino,urlcolor=turquesa}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{6pt}}
\setlist[itemize]{{leftmargin=*,itemsep=2pt,topsep=2pt}}
\color{{negro}}

\begin{{document}}
\begin{{titlepage}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[magentaMILC] (current page.south west) rectangle (current page.north east);
  \foreach \x in {{-1,1,...,23}}{{
    \foreach \y in {{-1,1,...,31}}{{
      \draw[vino, opacity=.12, line width=.6pt] (\x,\y) .. controls +(1,.6) and +(-1,-.6) .. +(2,1.4);
      \draw[vino, opacity=.10, line width=.6pt] (\x+.6,\y+.4) arc (210:510:.45);
    }}
  }}
\end{{tikzpicture}}
\vspace*{{.85cm}}
\begin{{center}}
\begin{{tcolorbox}}[
  enhanced,colback=vino,colframe=vino,arc=12mm,width=.86\textwidth,
  boxrule=0pt,left=12mm,right=12mm,top=11mm,bottom=11mm,center,
  overlay={{\draw[magentaMILC,line width=2pt,dash pattern=on 9pt off 8pt,rounded corners=10mm]
  ([xshift=7mm,yshift=-7mm]frame.north west) rectangle ([xshift=-7mm,yshift=7mm]frame.south east);}}]
  \centering
  {{\sffamily\bfseries\color{{white}}\fontsize{{14}}{{16}}\selectfont ANEXO ACADÉMICO · MODELO MILC}}\\[9mm]
  {{\sffamily\bfseries\color{{white}}\fontsize{{29}}{{32}}\selectfont Marco Teórico del Modelo MILC}}\\[4mm]
  {{\sffamily\bfseries\color{{mostaza}}\fontsize{{18}}{{22}}\selectfont Investigación Liberadora y Científica}}\\[8mm]
  {{\sffamily\color{{white}}\Large Plan de Área de Tecnología e Informática 2025}}\\[9mm]
  {{\sffamily\bfseries\color{{white}}\large Institución Educativa Sor María Juliana}}\\[2mm]
  {{\sffamily\color{{white}}Autor: {tex_escape(AUTHOR)}}}\\[4mm]
  {{\sffamily\itshape\color{{white}}Versión híbrida sobria MILC para soporte curricular y académico}}
\end{{tcolorbox}}
\vfill
\begin{{tcolorbox}}[enhanced,colback=vino!72!black,colframe=vino!72!black,arc=5mm,boxrule=0pt,width=.86\textwidth,left=6mm,right=6mm,top=4mm,bottom=4mm,center]
{{\color{{white}}\sffamily\small Escuta · Sistematización · Praxis · Evaluación liberadora}}
\end{{tcolorbox}}
\end{{center}}
\end{{titlepage}}
\tableofcontents
\newpage
{body}

\chapter*{{Referencias Bibliográficas}}
\addcontentsline{{toc}}{{chapter}}{{Referencias Bibliográficas}}
\begin{{tcolorbox}}[colback=gris,colframe=vino]
\begin{{itemize}}
{refs}
\end{{itemize}}
\end{{tcolorbox}}
\end{{document}}
"""
    path = OUT / "tex" / "anexo_marco_teorico_milc.tex"
    path.write_text(tex, encoding="utf-8")
    return path


def shade_paragraph(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def add_heading(doc: Document, text: str, color: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(p, color)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(9)


def add_bullet(doc: Document, text: str, color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    b = p.add_run("• ")
    b.bold = True
    b.font.color.rgb = RGBColor.from_string(color)
    b.font.size = Pt(9)
    r = p.add_run(text)
    r.font.size = Pt(8)


def build_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANEXO ACADÉMICO\nMARCO TEÓRICO DEL MODELO MILC")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(PALETTE["vino"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Investigación Liberadora y Científica · Plan de Área de Tecnología e Informática 2025")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(PALETTE["magenta"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Autor: {AUTHOR}").bold = True
    doc.add_page_break()

    for chapter in chapters():
        add_heading(doc, chapter.title, chapter.color)
        for paragraph in chapter.paragraphs:
            add_paragraph(doc, paragraph)
        for item in chapter.bullets:
            add_bullet(doc, item, chapter.color)
        add_heading(doc, "Síntesis", chapter.color)
        add_paragraph(doc, chapter.synthesis)
        doc.add_paragraph()

    add_heading(doc, "Referencias Bibliográficas", PALETTE["vino"])
    for ref in references():
        add_bullet(doc, ref, PALETTE["vino"])

    path = OUT / "docx" / "anexo_marco_teorico_milc.docx"
    doc.save(path)
    return path


def compile_pdf(tex_path: Path) -> Path:
    for _ in range(2):
        subprocess.run(
            [XELATEX, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=tex_path.parent,
            check=True,
            capture_output=True,
            text=True,
        )
    pdf_src = tex_path.with_suffix(".pdf")
    pdf_dst = OUT / "pdf" / pdf_src.name
    shutil.copy2(pdf_src, pdf_dst)
    return pdf_dst


def main() -> None:
    ensure_dirs()
    tex_path = build_tex()
    pdf_path = compile_pdf(tex_path)
    docx_path = build_docx()
    result = {
        "capitulos": len(chapters()),
        "tex": str(tex_path),
        "pdf": str(pdf_path),
        "docx": str(docx_path),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

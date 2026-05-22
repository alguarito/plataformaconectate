from __future__ import annotations

import csv
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_ROOT = Path(__file__).resolve().parents[1]
ROOT = Path(os.environ.get("MILC_RESOURCE_ROOT", SCRIPT_ROOT))
PLAN = Path(os.environ.get("MILC_PLAN", ROOT / "00_PLAN_AREA" / "plan_area_tecnologia_2025.docx"))
OLD_GUIDES = Path(os.environ.get("MILC_OLD_GUIDES", ROOT / "07_ANTERIORES_RESPALDO" / "Copia de 6"))
OUT = Path(os.environ.get("MILC_OUT", ROOT / "01_GUIAS_APRENDIZAJE" / "06_SEXTO"))
AUTHOR = "Dr. Álvaro Cárdenas Orozco"
GRADE = 6
XELATEX = os.environ.get("XELATEX", "xelatex")
TECTONIC = Path(os.environ.get("TECTONIC", "/Users/alvarocardenasorozco/.codex/plugins/cache/openai-bundled/latex-tectonic/0.1.0/bin/tectonic"))
GRADE_NAMES = {
    6: "sexto",
    7: "septimo",
    8: "octavo",
    9: "noveno",
    10: "decimo",
    11: "undecimo",
}

COLORS = {
    "magenta": "E6007E",
    "vino": "5A0038",
    "turquesa": "0093A5",
    "verde": "58A923",
    "mostaza": "E5B400",
    "ocre": "B86600",
    "negro": "241816",
    "gris": "F7F7F4",
}

GRADE_COVER_COLORS = {
    6: {
        "primary": "2B65F6",
        "dark": "064CCF",
        "soft": "EAF2FF",
        "accent": "FFE47A",
        "text": "FFFFFF",
        "docx": "2B65F6",
    },
    7: {
        "primary": "B6FA51",
        "dark": "4F8E00",
        "soft": "F2FFD9",
        "accent": "66B8FF",
        "text": "111807",
        "docx": "4F8E00",
    },
    8: {
        "primary": "ED7446",
        "dark": "B63E16",
        "soft": "FFF0E8",
        "accent": "CFE7FF",
        "text": "FFFFFF",
        "docx": "ED7446",
    },
    9: {
        "primary": "733DE3",
        "dark": "4B22A8",
        "soft": "F1E9FF",
        "accent": "9DEDF5",
        "text": "FFFFFF",
        "docx": "733DE3",
    },
    10: {
        "primary": "EB498A",
        "dark": "A90F58",
        "soft": "FFEAF3",
        "accent": "FFD24F",
        "text": "FFFFFF",
        "docx": "EB498A",
    },
    11: {
        "primary": "60D1F9",
        "dark": "007FA4",
        "soft": "E5F9FF",
        "accent": "FF6B35",
        "text": "070F12",
        "docx": "007FA4",
    },
}


@dataclass
class PeriodData:
    number: int
    name: str
    dba: str
    referentes: str
    estandares: str
    competencias: str
    aprendizajes: str
    cognitivos: str
    procedimentales: str
    actitudinales: str


@dataclass
class GuideData:
    number: int
    grade: int
    period: PeriodData
    title: str
    short_title: str
    focus: str
    question: str
    product: str


def clean(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split()).strip()


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+", clean(text))
    return [p.strip() for p in parts if p.strip()]


def grade_name(grade: int) -> str:
    return GRADE_NAMES.get(grade, str(grade))


def grade_display(grade: int) -> str:
    return f"grado {grade}"


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in str(text))


def period_name(number: int) -> str:
    return {1: "Primero", 2: "Segundo", 3: "Tercero"}[number]


def period_for_session(session: int) -> int:
    if session <= 10:
        return 1
    if session <= 20:
        return 2
    return 3


def extract_period_data() -> dict[int, PeriodData]:
    doc = Document(str(PLAN))
    result: dict[int, PeriodData] = {}
    for idx, period in enumerate([1, 2, 3]):
        table = doc.tables[idx]
        row1 = [clean(c.text) for c in table.rows[1].cells]
        row2 = [clean(c.text) for c in table.rows[2].cells]
        row5 = [clean(c.text) for c in table.rows[5].cells]
        result[period] = PeriodData(
            number=period,
            name=period_name(period),
            dba=row1[0].replace("DBA:", "").strip(),
            referentes=row2[0].replace("REFERENTES:", "").strip(),
            estandares=row5[0],
            competencias=row5[1],
            aprendizajes=row5[2],
            cognitivos=row5[3],
            procedimentales=row5[4],
            actitudinales=row5[5],
        )
    return result


def extract_title_from_docx(path: Path) -> str:
    doc = Document(str(path))
    candidates = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    for text in candidates:
        if "—" in text and len(text) > 20 and "GUÍA" not in text.upper():
            return text
    for text in candidates:
        if len(text) > 25 and not text.lower().startswith(("tecnología", "grado", "institución", "docente")):
            return text
    return path.stem


def extract_topics() -> dict[int, str]:
    current_temario = OUT / "fuente" / f"temario_{grade_name(GRADE)}.json"
    if current_temario.exists():
        data = json.loads(current_temario.read_text(encoding="utf-8"))
        topics = {int(item["guia"]): item for item in data if 1 <= int(item["guia"]) <= 30}
        missing = [i for i in range(1, 31) if i not in topics]
        if not missing:
            return topics

    topics: dict[int, str] = {}
    for path in OLD_GUIDES.glob("PERIODO */guia_aprendizaje_grado_sexto_sesion_*.docx"):
        match = re.search(r"sesion_(\d+)\.docx$", path.name)
        if not match:
            continue
        session = int(match.group(1))
        if 1 <= session <= 30:
            topics[session] = extract_title_from_docx(path)
    missing = [i for i in range(1, 31) if i not in topics]
    if missing:
        raise RuntimeError(f"Faltan títulos para sesiones: {missing}")
    return topics


def make_short_title(title: str) -> str:
    title = clean(title)
    def trim_words(text: str, max_chars: int = 92) -> str:
        text = clean(text)
        if len(text) <= max_chars:
            return text
        cut = text[:max_chars].rsplit(" ", 1)[0].rstrip(" ,;:-")
        return cut if len(cut) >= 40 else text[:max_chars].rstrip(" ,;:-")

    if "—" in title:
        left, right = [p.strip() for p in title.split("—", 1)]
        if len(left) <= 92:
            return left
        return trim_words(right, 82)
    return trim_words(title, 92)


def cover_title_size(title: str) -> tuple[int, int]:
    length = len(clean(title))
    if length > 62:
        return 22, 25
    if length > 48:
        return 25, 28
    if length > 34:
        return 28, 31
    return 31, 34


def docx_title_size(title: str) -> int:
    length = len(clean(title))
    if length > 62:
        return 15
    if length > 48:
        return 17
    return 20


def docx_cover_color(grade: int) -> str:
    return grade_cover_palette(grade).get("docx", COLORS["vino"])


def grade_cover_palette(grade: int) -> dict[str, str]:
    return GRADE_COVER_COLORS.get(grade, GRADE_COVER_COLORS[6])


def tex_grade_color_definitions(grade: int) -> str:
    palette = grade_cover_palette(grade)
    return "\n".join(
        [
            rf"\definecolor{{milcGradePrimary}}{{HTML}}{{{palette['primary']}}}",
            rf"\definecolor{{milcGradeDark}}{{HTML}}{{{palette['dark']}}}",
            rf"\definecolor{{milcGradeSoft}}{{HTML}}{{{palette['soft']}}}",
            rf"\definecolor{{milcGradeAccent}}{{HTML}}{{{palette['accent']}}}",
            rf"\definecolor{{milcGradeText}}{{HTML}}{{{palette['text']}}}",
        ]
    )


def web_cover_title_size(title: str) -> tuple[int, int]:
    length = len(clean(title))
    if length > 62:
        return 15, 18
    if length > 48:
        return 17, 20
    if length > 34:
        return 19, 22
    return 22, 25


def cover_title_linebreaks(title: str, max_chars: int = 32) -> str:
    words = clean(title).split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join(current + [word])
        if current and len(candidate) > max_chars:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return r"\\ ".join(tex_escape(line) for line in lines)


def make_focus(title: str) -> str:
    return clean(title.split("—", 1)[0])


def guide_family_from_text(text: str) -> str:
    focus = clean(text).lower()
    if any(term in focus for term in ["identidad digital", "huella digital", "oikeiôsis estoica", "oikeiosis estoica"]):
        return "digital_identity"
    if any(term in focus for term in ["correo institucional", "microsoft 365", "firma profesional", "bandeja de entrada", "netiqueta", "correo formal", "correos formales", "comunicación digital profesional", "comunicacion digital profesional"]):
        return "digital_communication"
    if any(term in focus for term in ["normas icontec", "normas apa", "presentación de documentos", "presentacion de documentos", "citación y referenciación", "citacion y referenciacion", "informe técnico con normas", "informe tecnico con normas", "normativa técnica", "normativa tecnica"]):
        return "document_standards"
    if any(term in focus for term in ["propiedad intelectual", "patentes", "licencias de software", "creative commons", "gpl", "software libre", "software propietario"]):
        return "intellectual_property"
    if any(term in focus for term in ["arquitectura del computador", "hardware interno", "periféricos", "perifericos", "software: definición", "software: definicion", "sistema operativo", "gestión de archivos", "gestion de archivos", "ensamble", "mantenimiento preventivo", "sistema computacional"]):
        return "computer_systems"
    if any(term in focus for term in ["procesador de texto", "word/google docs", "formato de fuente", "formato de párrafo", "formato de parrafo", "interlineado", "sangría", "sangria", "inserción de imágenes", "insercion de imagenes", "tablas y formas en documentos", "documento autobiográfico", "documento autobiografico", "documento final con normas básicas", "documento final con normas basicas"]):
        return "document_processing"
    if any(term in focus for term in ["motores de búsqueda", "motores de busqueda", "información confiable", "informacion confiable", "fuentes web", "fake news", "sesgo informacional", "investigación guiada", "investigacion guiada", "fuentes validadas"]):
        return "information_literacy"
    if any(term in focus for term in ["sala de sistemas", "diagnóstico tecnológico", "diagnostico tecnológico", "diagnostico tecnologico"]):
        return "systems_room"
    if any(term in focus for term in ["historia de la tecnología", "historia de la tecnologia", "revolución industrial", "revolucion industrial", "telégrafo", "telegrafo", "electricidad", "era electrónica", "era electronica", "revolución digital", "revolucion digital"]):
        return "tech_history"
    if any(term in focus for term in [
        "definición del problema o necesidad empresarial", "definicion del problema o necesidad empresarial",
        "investigación del usuario", "investigacion del usuario", "cliente o comunidad beneficiaria",
        "diseño de la solución", "diseno de la solucion", "producto, servicio, proceso o sistema",
        "integración del sitio web", "integracion del sitio web", "integración de datos",
        "integracion de datos", "dashboard o reporte final", "automatización o mejora de proceso",
        "automatizacion o mejora de proceso", "documentación técnica y comercial",
        "documentacion tecnica y comercial", "preparación del pitch", "preparacion del pitch",
        "sustentación final", "sustentacion final", "feria empresarial", "tribunal milc",
    ]):
        return "integrated_business_project"
    if any(term in focus for term in [
        "presencia digital empresarial", "identidad de marca", "arquitectura web empresarial",
        "html/css", "cms", "landing page", "portafolio comercial", "redacción comercial",
        "redaccion comercial", "diseño visual de marca", "diseno visual de marca",
        "seo básico", "seo basico", "contenido para redes", "ética digital empresarial",
        "etica digital empresarial", "sustentación del sitio", "sustentacion del sitio",
    ]):
        return "business_web_presence"
    if any(term in focus for term in [
        "procesos empresariales", "mapeo de procesos", "diagnosticar procesos",
        "formularios digitales", "organización de datos", "organizacion de datos",
        "automatización básica", "automatizacion basica", "atención al cliente",
        "atencion al cliente", "calidad y mejora continua", "pdca",
        "protección de datos", "proteccion de datos", "proceso automatizado",
    ]):
        return "business_process_automation"
    if any(term in focus for term in ["del libro al informe", "latex", "informe profesional", "informe técnico", "informe tecnico", "informe comercial", "estudio de mercado", "código documental", "codigo documental", "tema comercial o contable", "encuesta", "recolección de datos", "recoleccion de datos", "tablas en latex", "gráficos e imágenes en el informe", "graficos e imagenes en el informe", "ia para planear informes", "planear informes", "análisis comercial con apoyo de ia", "analisis comercial con apoyo de ia", "edición, citación y transparencia", "edicion, citacion y transparencia", "sustentación del informe comercial", "sustentacion del informe comercial"]):
        return "latex_business_report"
    if any(term in focus for term in ["libro", "escritura inversa", "asistente de escritura", "prompts básicos", "prompts basicos", "planeación editorial", "planeacion editorial", "autoría y voz", "autoria y voz", "primer capítulo", "primer capitulo", "edición con ia", "edicion con ia", "derechos de autor", "diseño editorial digital", "diseno editorial digital", "tabla de contenido", "feria editorial"]):
        return "ai_book_project"
    if any(term in focus for term in ["datos para tomar decisiones empresariales", "base de datos comercial", "fórmulas empresariales", "formulas empresariales", "limpieza y organización de datos", "limpieza y organizacion de datos", "gráficos comerciales", "graficos comerciales", "indicadores clave", "producto más vendido", "producto mas vendido", "cliente frecuente", "punto de equilibrio", "ia para interpretar datos", "ia para cargue", "cargue, limpieza", "visualización de datos", "visualizacion de datos", "dashboard", "reporte empresarial", "decisiones empresariales"]):
        return "business_dashboard_ai"
    if (
        re.search(r"\bia\b", focus)
        or "inteligencia artificial" in focus
        or "machine learning" in focus
        or "aprenden las máquinas" in focus
        or "aprenden las maquinas" in focus
        or "aprendizaje automático" in focus
        or "aprendizaje automatico" in focus
        or "sesgos algorítmicos" in focus
        or "sesgos algoritmicos" in focus
        or "algoritmos de recomendación" in focus
        or "algoritmos de recomendacion" in focus
        or "prompt" in focus
        or "generativa" in focus
        or "alucinaciones" in focus
        or "sesgos" in focus
        or "recomendadores" in focus
    ):
        return "ai"
    if any(term in focus for term in ["office 365", "google workspace", "documento compartido", "documentos compartidos", "trabajo en la nube", "almacenamiento en la nube", "almacenamiento remoto", "sincronización", "sincronizacion", "copias de seguridad", "permisos", "competencias colaborativas", "foro", "foros", "blog", "blogs", "lms", "liderazgo servicial", "cartografía social digital", "cartografia social digital", "portafolio digital del periodo"]):
        return "cloud_collaboration"
    if any(term in focus for term in ["raee", "e-waste", "ewaste", "residuos electrónicos", "residuos electronicos", "reciclaje de componentes", "auditoría de e-waste", "auditoria de e-waste", "impacto ambiental tecnológico", "impacto ambiental tecnologico", "tribunal analéctico", "tribunal analectico"]):
        return "ewaste"
    if any(term in focus for term in ["análisis de datos", "analisis de datos", "phronesis", "base de datos", "tabla de datos", "tablas de datos", "limpieza básica", "limpieza basica", "gráfico", "grafico", "gráficos", "graficos", "estudio de datos", "lectura crítica de datos", "lectura critica de datos", "conclusión prudente", "conclusion prudente", "recolección de datos", "recoleccion de datos"]):
        return "data_analysis"
    if any(term in focus for term in ["excel", "hoja de cálculo", "hoja de calculo", "celda", "celdas", "fórmula", "formula", "suma", "promedio", "referencia", "formato condicional", "validación de datos"]):
        return "spreadsheet"
    if any(term in focus for term in ["computación física", "computacion fisica", "sensor", "sensores", "makecode", "micro:bit", "microbit", "actuador", "actuadores", "led", "buzzer", "temperatura", "luz", "umbral", "umbrales", "monitoreo ambiental", "control físico", "control fisico", "sistema de alerta"]):
        return "physical_computing"
    if any(term in focus for term in ["scratch", "primm", "algoritmia", "algoritmo", "variables", "eventos", "bucles", "repeticiones", "condicionales simples", "depuración", "depuracion", "debugging", "narrativa digital interactiva", "narrativa digital", "competencias computacionales", "historia o juego"]):
        return "scratch_programming"
    programming_terms = ["lógica", "condicional", "operador", "algoritmo", "programación por bloques", "variable", "ciclo", "pseudocódigo", "pseudocodigo", "diagrama de flujo", "tabla de verdad", "tablas de verdad"]
    if any(term in focus for term in programming_terms) or re.search(r"\b(if|else|elif|else-if)\b", focus):
        return "programming"
    if any(term in focus for term in ["imagen", "sonido", "audio", "recurso visual"]):
        return "media"
    if any(term in focus for term in ["presentación", "presentaciones", "powerpoint", "canva", "diapositiva", "diapositivas", "hipervínculo"]):
        return "presentation"
    if any(term in focus for term in ["imagen", "sonido", "audio", "multimedia", "recurso visual"]):
        return "media"
    if any(term in focus for term in ["estética de la liberación", "estetica de la liberacion", "cultura popular", "belleza de la cultura", "representación cultural", "representacion cultural"]):
        return "cultural_media"
    if any(term in focus for term in ["seguridad", "ciberbullying", "sexting", "grooming", "phishing", "spam", "contraseñas", "privacidad en la red"]):
        return "safety"
    if "infoesfera" in focus and any(term in focus for term in ["proyecto", "socialización", "socializacion", "portafolio"]):
        return "info_project"
    if any(term in focus for term in ["proyecto", "sustentación", "portafolio", "bitácora"]):
        return "project"
    if any(term in focus for term in ["diseño editorial", "autoedición", "autoedicion", "retícula", "reticula", "boletín", "boletin", "folleto", "publicación", "publicacion", "jerarquía visual", "jerarquia visual", "maquetación", "maquetacion"]):
        return "editorial"
    if any(term in focus for term in ["redes", "wi-fi", "wifi", "bluetooth", "protocolo", "filius", "dirección ip", "direccion ip", "configuración de ip", "configuracion de ip", "topología", "topologia", "direccionamiento", "cliente-servidor", "conectividad", "emisor", "receptor", "canal", "infraestructura digital"]):
        return "network"
    return "general"


def guide_family(guide: GuideData) -> str:
    if guide.grade == 6:
        if guide.number == 2:
            return "digital_identity"
        if guide.number in {3, 4, 7}:
            return "digital_communication"
        if guide.number == 6:
            return "tech_history"
        if guide.number in {8, 27}:
            return "safety"
        if 11 <= guide.number <= 20:
            return "computer_systems"
        if guide.number in {21, 22, 23, 24, 29}:
            return "document_processing"
        if guide.number in {25, 26, 28}:
            return "information_literacy"
        if guide.number == 30:
            return "project"
    if guide.grade == 9:
        if guide.number == 3:
            return "tech_history"
        if guide.number in {4, 5, 6, 10}:
            return "document_standards"
        if guide.number in {7, 8}:
            return "intellectual_property"
        if guide.number in {11, 12, 13}:
            return "editorial"
    if guide.grade == 9 and guide.number == 11:
        return "editorial"
    if is_undecimo_process_inventory(guide):
        return "business_process_automation"
    return guide_family_from_text(guide.title)


def is_noveno_editorial_i(guide: GuideData) -> bool:
    return guide.grade == 9 and guide.number == 11


def is_undecimo_process_inventory(guide: GuideData) -> bool:
    return guide.grade == 11 and guide.number == 11


FAMILY_CONCEPT_ROWS: dict[str, list[tuple[str, str, str]]] = {
    "digital_identity": [
        ("Identidad digital", "Conjunto de datos, acciones y huellas que representan a una persona en entornos digitales.", "Nombre de usuario, foto, publicaciones, comentarios y archivos compartidos."),
        ("Huella digital", "Rastro que queda al navegar, publicar, buscar, comentar o iniciar sesión.", "Un comentario antiguo o una foto etiquetada."),
        ("Dato personal", "Información que permite identificar directa o indirectamente a una persona.", "Nombre, curso, ubicación, correo, rostro o teléfono."),
        ("Privacidad", "Capacidad de controlar qué información se comparte, con quién y para qué.", "No publicar horarios, dirección o contraseñas."),
        ("Reputación digital", "Imagen que otras personas construyen a partir de lo que encuentran sobre alguien.", "Un correo respetuoso fortalece confianza; una burla pública la daña."),
        ("Consentimiento", "Autorización clara antes de usar imagen, voz o información de otra persona.", "Pedir permiso antes de subir una foto grupal."),
        ("Oikeiôsis digital", "Cuidado de sí y de los otros al habitar espacios virtuales.", "Pensar si una publicación protege mi dignidad y la del grupo."),
        ("Autocuidado", "Decisión responsable para prevenir exposición, engaños o daño emocional.", "Revisar privacidad antes de aceptar contactos."),
    ],
    "digital_communication": [
        ("Correo institucional", "Canal formal de comunicación escolar que exige claridad, respeto y propósito.", "Mensaje al docente para pedir retroalimentación."),
        ("Asunto", "Frase breve que anticipa el motivo del mensaje.", "Solicitud de revisión de guía 3."),
        ("Destinatario", "Persona o grupo que debe recibir el mensaje.", "Enviar al docente correcto y no copiar a quien no corresponde."),
        ("Firma profesional", "Bloque final con nombre, curso y datos pertinentes.", "Álvaro Cárdenas, grado sexto, Tecnología."),
        ("Bandeja de entrada", "Espacio donde se reciben y organizan mensajes.", "Marcar como importante, archivar o responder a tiempo."),
        ("Netiqueta", "Normas de respeto y claridad en la comunicación digital.", "Saludar, explicar el motivo y despedirse."),
        ("Adjunto", "Archivo que acompaña un correo y debe nombrarse de forma clara.", "guia3_apellido_nombre.pdf."),
        ("Protección de datos", "Cuidado de información personal al enviar, reenviar o responder mensajes.", "No compartir correos de compañeros sin permiso."),
    ],
    "computer_systems": [
        ("Sistema computacional", "Conjunto de hardware, software, datos, usuario y procedimientos que trabajan juntos.", "Computador de sala con sistema operativo, programas y archivos."),
        ("Hardware", "Componentes físicos que se pueden tocar o conectar.", "CPU, RAM, disco, monitor, teclado y mouse."),
        ("CPU", "Procesador que ejecuta instrucciones y coordina operaciones.", "Abrir un programa o calcular una operación."),
        ("Memoria RAM", "Memoria temporal usada mientras el equipo trabaja.", "Mantener abiertas varias pestañas o aplicaciones."),
        ("Almacenamiento", "Medio donde se guardan archivos de forma permanente.", "HDD, SSD, USB o nube."),
        ("Periférico", "Dispositivo que permite entrada, salida o ambas.", "Teclado, impresora, pantalla táctil o audífonos."),
        ("Software", "Programas e instrucciones que permiten usar el equipo.", "Sistema operativo, navegador y procesador de texto."),
        ("Mantenimiento preventivo", "Acciones para cuidar el equipo antes de que falle.", "Limpieza externa, orden de cables y reporte oportuno."),
    ],
    "document_processing": [
        ("Procesador de texto", "Herramienta para crear, editar, revisar y presentar documentos.", "Word o Google Docs."),
        ("Formato de fuente", "Ajustes visuales de las letras.", "Tipo, tamaño, color, negrita o cursiva."),
        ("Formato de párrafo", "Ajustes que organizan bloques de texto.", "Alineación, interlineado, sangría y espaciado."),
        ("Estructura documental", "Orden interno que ayuda a leer con claridad.", "Título, subtítulos, párrafos, imágenes, tablas y fuentes."),
        ("Imagen con propósito", "Recurso visual que explica o complementa el texto.", "Foto citada que muestra el problema investigado."),
        ("Tabla", "Organización en filas y columnas para comparar información.", "Tabla de causas, evidencias y posibles soluciones."),
        ("Revisión", "Lectura final para corregir ortografía, formato, fuentes y coherencia.", "Usar lista de cotejo antes de entregar."),
        ("Presentación básica", "Normas mínimas para que el documento sea legible y profesional.", "Márgenes, títulos claros, alineación consistente y créditos."),
    ],
    "information_literacy": [
        ("Pregunta de búsqueda", "Interrogante claro que orienta qué información se necesita.", "¿Qué causas tiene el desperdicio de agua en mi colegio?"),
        ("Palabra clave", "Término preciso usado para encontrar mejores resultados.", "agua colegio desperdicio solución."),
        ("Fuente confiable", "Publicación con autor, fecha, propósito y evidencia verificable.", "Sitio institucional, biblioteca digital o medio reconocido."),
        ("Fake news", "Información falsa o manipulada que se presenta como verdadera.", "Imagen sacada de contexto para causar alarma."),
        ("Sesgo informacional", "Inclinación que resalta una postura y oculta otras evidencias.", "Página que solo muestra beneficios y omite riesgos."),
        ("Contraste", "Comparación entre varias fuentes antes de aceptar una idea.", "Revisar si dos instituciones coinciden en el dato."),
        ("Cita básica", "Registro mínimo de origen de una idea, imagen o dato.", "Autor, título, sitio, fecha y enlace."),
        ("Conclusión prudente", "Respuesta que usa evidencia sin exagerar lo que los datos permiten decir.", "La muestra es pequeña, por eso la conclusión es inicial."),
    ],
    "document_standards": [
        ("Norma documental", "Conjunto de reglas para presentar trabajos con orden y rigor.", "ICONTEC o APA."),
        ("Estructura", "Partes que organizan un documento académico o técnico.", "Portada, introducción, desarrollo, conclusiones y referencias."),
        ("Formato", "Configuración visible del documento.", "Márgenes, fuente, interlineado, numeración y títulos."),
        ("Citación", "Forma de reconocer una idea tomada de otra fuente.", "Cita breve en el texto con autor y año."),
        ("Referenciación", "Lista ordenada de fuentes consultadas.", "Autor, fecha, título, editorial o enlace."),
        ("Informe técnico", "Documento que presenta problema, proceso, evidencia y conclusión.", "Informe sobre una tecnología regional."),
        ("Rigor", "Cuidado para que forma, fuentes y contenido sean verificables.", "No inventar bibliografía ni mezclar estilos."),
        ("Conciencia histórica", "Lectura de la tecnología como proceso social con beneficios, exclusiones y responsabilidades.", "Explicar quién ganó y quién quedó por fuera de un invento."),
    ],
    "intellectual_property": [
        ("Propiedad intelectual", "Protección y reconocimiento de creaciones, ideas expresadas e innovaciones.", "Texto, imagen, software, diseño o invención."),
        ("Patente", "Derecho sobre una invención técnica novedosa y aplicable.", "Dispositivo o procedimiento registrado."),
        ("Derecho de autor", "Reconocimiento de quien crea una obra literaria, artística, científica o digital.", "Créditos de una imagen o texto citado."),
        ("Licencia", "Permiso que define cómo puede usarse una obra o software.", "Creative Commons, GPL o licencia propietaria."),
        ("Software libre", "Programa que permite usar, estudiar, modificar y compartir bajo ciertas condiciones.", "Aplicación distribuida con código abierto y licencia libre."),
        ("Software propietario", "Programa cuyo uso, copia o modificación está restringido por el titular.", "Aplicación comercial cerrada."),
        ("Plagio", "Presentar trabajo ajeno como propio.", "Copiar texto, imagen o código sin citar."),
        ("Justicia tecnológica", "Pregunta por quién puede acceder, crear, beneficiarse y ser reconocido.", "Valorar saberes locales y reconocer exclusiones históricas."),
    ],
}


FAMILY_CASE_TEXT: dict[str, str] = {
    "digital_identity": "Un estudiante encuentra su nombre en una publicación antigua con comentarios burlones. También nota que su perfil muestra datos personales que no necesita compartir. Debe decidir qué cambiar, qué reportar y cómo cuidar su identidad sin afectar a otros.",
    "digital_communication": "Un correo al docente llega sin asunto, sin saludo, con abreviaturas y un archivo llamado documento-final-final.pdf. El mensaje genera confusión porque no explica qué se solicita ni protege los datos de quienes aparecen copiados.",
    "computer_systems": "Un equipo de la sala presenta fallas. Algunos quieren desconectar cables, otros culpan al computador sin revisar componentes. El grupo debe identificar partes del sistema, registrar síntomas y proponer acciones seguras de diagnóstico.",
    "document_processing": "Un documento contiene buena información, pero mezcla fuentes, imágenes sin crédito, párrafos extensos y tablas desordenadas. El lector no sabe qué mirar primero y el equipo debe mejorar estructura, formato y revisión.",
    "information_literacy": "Un grupo investiga un problema comunitario y encuentra una noticia llamativa, pero no hay autor, fecha ni evidencia. Antes de usarla, debe contrastar fuentes, reconocer sesgos y escribir una conclusión prudente.",
    "document_standards": "Un informe técnico tiene ideas valiosas, pero no sigue normas, no cita fuentes y mezcla ICONTEC con APA sin criterio. El equipo debe ordenar estructura, formato, citas y referencias para que el trabajo sea verificable.",
    "intellectual_property": "Un equipo usa imágenes de Internet, copia fragmentos de una página y descarga software sin revisar licencia. El producto parece completo, pero no respeta autoría, permisos ni justicia tecnológica.",
}


FAMILY_CONCEPTUAL_PARAGRAPHS: dict[str, list[str]] = {
    "digital_identity": [
        "La identidad digital no se construye solo cuando publicamos; también aparece en cuentas, búsquedas, fotos, archivos compartidos y comentarios de otras personas.",
        "La huella digital puede abrir oportunidades o generar riesgos. Por eso conviene revisar qué datos compartimos, qué permisos damos y qué publicaciones nos representan.",
        "La autonomía digital implica detenerse antes de publicar, pedir consentimiento, cuidar la reputación propia y respetar la dignidad de otras personas.",
        "La Oikeiôsis aplicada al mundo virtual invita a cuidar lo cercano: mi cuerpo, mi nombre, mi grupo, mi comunidad y la forma como habitamos los espacios digitales.",
    ],
    "digital_communication": [
        "Un correo institucional es una comunicación formal: debe tener propósito claro, tono respetuoso, asunto informativo, mensaje ordenado, firma y adjuntos nombrados.",
        "La netiqueta no es una lista de modales vacíos. Ayuda a evitar malentendidos, proteger datos y construir confianza en comunidades escolares.",
        "Escribir bien un correo exige pensar en el destinatario: qué necesita saber, qué acción se solicita y qué evidencia se adjunta.",
        "La autonomía se demuestra cuando el estudiante revisa su mensaje antes de enviarlo y corrige asunto, saludo, claridad, ortografía, archivos y privacidad.",
    ],
    "computer_systems": [
        "Un computador es un sistema: sus partes físicas, programas, archivos y usuarios se relacionan. Una falla puede venir de hardware, software, energía, conexión o uso inadecuado.",
        "Identificar componentes permite diagnosticar con seguridad. No se trata de abrir equipos sin permiso, sino de observar, clasificar, nombrar funciones y reportar síntomas.",
        "El mantenimiento preventivo evita daños: limpieza externa, postura adecuada, ventilación, orden de cables, apagado correcto y cuidado de archivos.",
        "La autonomía técnica empieza con una pregunta responsable: qué puedo revisar sin riesgo, qué debo documentar y cuándo necesito ayuda del docente o soporte técnico.",
    ],
    "document_processing": [
        "Un procesador de texto no solo sirve para escribir. Permite organizar ideas, mejorar lectura, integrar imágenes y tablas, revisar ortografía y preparar documentos formales.",
        "El formato debe ayudar al contenido. Fuente, tamaño, alineación, interlineado y sangría hacen que el lector encuentre el mensaje sin esfuerzo.",
        "Las imágenes y tablas deben tener propósito: explicar, comparar o evidenciar. Si solo decoran o no tienen fuente, debilitan el documento.",
        "La autonomía editorial consiste en planear, escribir, revisar con criterios y corregir antes de entregar.",
    ],
    "information_literacy": [
        "Buscar información no es escribir la primera palabra que se ocurre y copiar el primer resultado. Es formular una pregunta, elegir palabras clave y revisar la calidad de las fuentes.",
        "Una fuente confiable muestra autor, fecha, propósito, evidencia y relación con el tema. Si falta esa información, se debe contrastar antes de usarla.",
        "Las fake news y los sesgos informacionales pueden parecer convincentes. Por eso se comparan fuentes, se revisan imágenes y se evita compartir sin verificar.",
        "Investigar con autonomía exige registrar de dónde salió cada dato y escribir conclusiones prudentes, sin exagerar ni ocultar límites.",
    ],
    "document_standards": [
        "Las normas ICONTEC y APA ayudan a presentar documentos con orden, trazabilidad y respeto por las fuentes. No son adornos: hacen que el trabajo pueda revisarse.",
        "Un informe técnico debe mostrar problema, propósito, desarrollo, evidencia, conclusiones y referencias. Si una parte falta, el lector no puede seguir el proceso.",
        "Citar y referenciar evita plagio, reconoce autoría y permite que otra persona encuentre la fuente original.",
        "La conciencia histórica exige que el documento no solo enumere tecnologías, sino que explique necesidades, impactos, exclusiones y responsabilidades.",
    ],
    "intellectual_property": [
        "La propiedad intelectual protege creaciones y reconoce a quienes producen conocimiento, arte, software o soluciones técnicas.",
        "No todo lo que está en Internet es libre de usar. Una licencia indica si se puede copiar, modificar, compartir o usar comercialmente.",
        "El software libre y el propietario no se diferencian por ser gratis o pago, sino por los permisos para estudiar, modificar y distribuir.",
        "La justicia tecnológica pregunta quién crea, quién se beneficia, quién queda excluido y cómo reconocer autorías sin frenar el aprendizaje compartido.",
    ],
}


FAMILY_GUIDED_PRACTICE: dict[str, list[str]] = {
    "digital_identity": [
        "Escribe tres datos que sí puedes compartir en contexto escolar y tres que debes proteger.",
        "Dibuja un mapa de huella digital: cuenta, publicación, comentario, imagen, archivo y permiso.",
        "Analiza una situación: qué riesgo aparece, a quién afecta y qué acción responsable conviene.",
        "Redacta una regla personal para publicar o comentar con cuidado.",
        "Revisa si una publicación respeta consentimiento, privacidad y dignidad.",
        "Cierra con una mejora concreta para cuidar tu identidad digital esta semana.",
    ],
    "digital_communication": [
        "Lee el caso y define propósito, destinatario y acción esperada del correo.",
        "Redacta un asunto breve que permita entender el mensaje antes de abrirlo.",
        "Escribe saludo, cuerpo y despedida con tono respetuoso y claro.",
        "Nombra el adjunto con curso, guía, apellido y versión.",
        "Revisa netiqueta, ortografía, datos personales y destinatarios antes de enviar.",
        "Compara un correo débil y uno mejorado; explica qué cambió.",
    ],
    "computer_systems": [
        "Identifica si cada elemento pertenece a hardware, software, archivo, usuario o procedimiento.",
        "Relaciona componente y función: CPU, RAM, almacenamiento, periférico y sistema operativo.",
        "Observa un síntoma de falla y registra qué información se necesita antes de actuar.",
        "Clasifica acciones: seguras para estudiante, requieren docente o deben evitarse.",
        "Diseña una ficha de diagnóstico con estado, síntoma, posible causa y reporte.",
        "Escribe una norma de mantenimiento preventivo para la sala.",
    ],
    "document_processing": [
        "Define propósito, audiencia y estructura mínima del documento.",
        "Organiza título, subtítulos y párrafos antes de aplicar formato.",
        "Aplica una regla consistente de fuente, tamaño, alineación e interlineado.",
        "Inserta una imagen o tabla solo si aporta evidencia o claridad.",
        "Agrega fuente o crédito de los recursos usados.",
        "Revisa ortografía, márgenes, legibilidad y coherencia antes de exportar.",
    ],
    "information_literacy": [
        "Convierte el tema en una pregunta de investigación concreta.",
        "Elige cinco palabras clave y prueba combinaciones más precisas.",
        "Evalúa dos fuentes con autor, fecha, propósito, evidencia y confiabilidad.",
        "Contrasta un dato en al menos dos fuentes diferentes.",
        "Identifica un posible sesgo, exageración o información faltante.",
        "Escribe una conclusión prudente con fuente y límite.",
    ],
    "document_standards": [
        "Identifica qué norma o criterio usará el documento: ICONTEC, APA o instrucción del docente.",
        "Ordena el informe en partes: portada, introducción, desarrollo, conclusiones y referencias.",
        "Transforma una idea tomada de fuente externa en cita o paráfrasis reconocida.",
        "Construye una referencia básica con autor, fecha, título y origen.",
        "Revisa coherencia entre citas en el texto y lista de referencias.",
        "Explica cómo la norma mejora rigor, justicia con las fuentes y lectura del informe.",
    ],
    "intellectual_property": [
        "Diferencia obra, idea, invención, software y recurso digital.",
        "Clasifica ejemplos según patente, derecho de autor, licencia o dominio público cuando aplique.",
        "Lee una licencia y subraya qué permite y qué restringe.",
        "Revisa un producto propio: qué imágenes, textos, datos o códigos necesitan crédito.",
        "Propón una forma ética de usar un recurso ajeno sin plagio.",
        "Relaciona propiedad intelectual con justicia: acceso, reconocimiento y exclusión.",
    ],
}


FAMILY_PRODUCT_CRITERIA: dict[str, list[str]] = {
    "digital_identity": ["Mapa de huella digital con mínimo seis elementos", "Tres datos protegidos y tres datos compartibles según contexto", "Análisis de un riesgo con acción responsable", "Regla personal de autocuidado digital", "Reflexión sobre dignidad, consentimiento y comunidad"],
    "digital_communication": ["Correo modelo con asunto, saludo, cuerpo, despedida y firma", "Adjunto nombrado correctamente o simulación del envío", "Revisión de netiqueta, ortografía y privacidad", "Comparación entre versión inicial y versión mejorada", "Explicación del propósito y destinatario"],
    "computer_systems": ["Ficha de componentes con función y ejemplo", "Clasificación de hardware, software, periféricos y almacenamiento", "Registro de un síntoma o caso de diagnóstico", "Acciones seguras de mantenimiento preventivo", "Reporte claro sin manipulación riesgosa"],
    "document_processing": ["Documento con título, subtítulos y párrafos ordenados", "Formato de fuente y párrafo consistente", "Imagen o tabla con propósito y fuente", "Lista de revisión aplicada antes de entregar", "Exportación o presentación final legible"],
    "information_literacy": ["Pregunta de investigación clara", "Lista de palabras clave usadas", "Dos fuentes evaluadas con criterios de confiabilidad", "Dato contrastado y fuente registrada", "Conclusión prudente con límite reconocido"],
    "document_standards": ["Documento o esquema con estructura normativa", "Formato básico coherente con ICONTEC/APA o criterio indicado", "Citas y referencias revisables", "Informe técnico con problema, evidencia y conclusión", "Reflexión sobre rigor, autoría y conciencia histórica"],
    "intellectual_property": ["Cuadro comparativo entre patente, derecho de autor y licencia", "Análisis de una licencia de software o recurso digital", "Créditos de recursos usados en un producto", "Caso de plagio corregido con citación o permiso", "Reflexión sobre justicia tecnológica y reconocimiento"],
}


FAMILY_EVALUATION_ROWS: dict[str, list[tuple[str, str, str, str]]] = {
    "digital_identity": [
        ("Comprensión", "Explico identidad, huella, privacidad y consentimiento con ejemplos propios.", "Reconozco los conceptos, pero algunos ejemplos son generales.", "Confundo identidad digital con uso de redes solamente."),
        ("Autocuidado", "Propongo acciones claras para proteger datos y reputación.", "Propongo cuidados básicos, pero sin justificar riesgos.", "No identifico datos sensibles ni acciones preventivas."),
        ("Análisis ético", "Relaciono mis decisiones con dignidad propia y cuidado de otros.", "Menciono respeto, pero falta profundidad o evidencia.", "No considero impacto sobre otras personas."),
        ("Producto", "El mapa y la reflexión son completos, claros y revisados.", "El producto existe, pero le faltan detalles o revisión.", "El producto queda incompleto o desordenado."),
    ],
    "digital_communication": [
        ("Claridad comunicativa", "El mensaje tiene propósito, destinatario y solicitud comprensibles.", "El mensaje se entiende, pero falta precisión.", "El correo resulta confuso o incompleto."),
        ("Netiqueta", "Uso tono respetuoso, asunto, saludo, firma y despedida adecuados.", "Cumplo algunos elementos formales.", "Escribo sin estructura formal ni cuidado del tono."),
        ("Privacidad", "Reviso destinatarios, copias, adjuntos y datos antes de enviar.", "Reconozco cuidados, pero los aplico parcialmente.", "Comparto o expongo información sin criterio."),
        ("Revisión", "Comparo versiones y justifico mejoras de escritura.", "Corrijo algunos errores, pero explico poco.", "Entrego sin revisar."),
    ],
    "computer_systems": [
        ("Conceptos técnicos", "Relaciono componente, función y ejemplo con precisión.", "Identifico componentes, pero confundo algunas funciones.", "Enumero partes sin explicar su utilidad."),
        ("Diagnóstico", "Registro síntomas y posibles causas sin manipular riesgosamente.", "Observo algunas fallas, pero con poca evidencia.", "Propongo acciones sin observar ni reportar."),
        ("Mantenimiento", "Planteo cuidados preventivos viables para la sala.", "Menciono cuidados generales.", "Ignoro normas de seguridad o cuidado."),
        ("Producto", "La ficha es clara, ordenada y útil para otra persona.", "La ficha se entiende parcialmente.", "La ficha no permite comprender el sistema."),
    ],
    "document_processing": [
        ("Estructura", "Organizo documento con título, subtítulos y párrafos coherentes.", "La estructura existe, pero puede leerse mejor.", "El documento queda como texto suelto."),
        ("Formato", "Aplico fuente, párrafo, tablas e imágenes con consistencia.", "Uso formato, pero con inconsistencias.", "El formato dificulta la lectura."),
        ("Autoría", "Registro fuentes de imágenes, datos o textos usados.", "Cito algunos recursos.", "Uso recursos sin crédito."),
        ("Revisión", "Corrijo ortografía, legibilidad y presentación antes de entregar.", "Reviso parcialmente.", "Entrego sin revisión visible."),
    ],
    "information_literacy": [
        ("Búsqueda", "Formulo pregunta y palabras clave precisas.", "Busco información, pero la pregunta es amplia.", "Copio resultados sin orientar la búsqueda."),
        ("Evaluación de fuentes", "Reviso autor, fecha, propósito, evidencia y confiabilidad.", "Reviso algunos criterios.", "Acepto fuentes sin verificar."),
        ("Contraste", "Comparo datos y reconozco sesgos o límites.", "Contrasto parcialmente.", "Uso una sola fuente como verdad final."),
        ("Conclusión", "Concluyo con evidencia, fuente y prudencia.", "Concluyo con datos, pero sin límites claros.", "Afirmo sin sustento suficiente."),
    ],
    "document_standards": [
        ("Norma y formato", "Aplico estructura, formato y partes del documento con coherencia.", "Cumplo partes básicas, pero con inconsistencias.", "No sigo criterios normativos."),
        ("Citación", "Cito y referencio fuentes de manera revisable.", "Incluyo fuentes, pero con errores o faltantes.", "Uso información sin reconocer autoría."),
        ("Informe técnico", "Relaciono problema, evidencia, análisis y conclusión.", "El informe tiene ideas, pero falta conexión.", "Presento datos sin estructura ni conclusión."),
        ("Conciencia histórica", "Reconozco impactos, necesidades, exclusiones y responsabilidades.", "Menciono impactos generales.", "Enumero hechos sin análisis social."),
    ],
    "intellectual_property": [
        ("Comprensión legal y ética", "Diferencio patente, autoría, licencia y plagio con ejemplos.", "Reconozco conceptos, pero con imprecisiones.", "Confundo permiso, gratuidad y autoría."),
        ("Uso de recursos", "Cito, acredita o elijo recursos según licencia.", "Cito parcialmente o falta revisar permisos.", "Uso recursos ajenos sin reconocimiento."),
        ("Software", "Comparo software libre y propietario por permisos y responsabilidades.", "Reconozco diferencias generales.", "Pienso que libre solo significa gratis."),
        ("Justicia tecnológica", "Analizo acceso, reconocimiento y exclusión con postura propia.", "Menciono justicia, pero sin evidencia.", "No relaciono el tema con impacto social."),
    ],
}


FAMILY_LEARNING_GOAL: dict[str, str] = {
    "digital_identity": "Al terminar la sesión podrás reconocer tu identidad y huella digital, proteger datos personales y tomar decisiones de autocuidado con respeto por la comunidad.",
    "digital_communication": "Al terminar la sesión podrás redactar y revisar comunicaciones digitales formales con asunto, propósito, netiqueta, firma, adjuntos y protección de datos.",
    "computer_systems": "Al terminar la sesión podrás identificar componentes de un sistema computacional, explicar funciones básicas, registrar síntomas y proponer acciones seguras de cuidado.",
    "document_processing": "Al terminar la sesión podrás construir o mejorar un documento con estructura, formato, imágenes, tablas, fuentes y revisión autónoma.",
    "information_literacy": "Al terminar la sesión podrás buscar, evaluar y contrastar información confiable para investigar un problema sin copiar ni compartir datos dudosos.",
    "document_standards": "Al terminar la sesión podrás aplicar criterios básicos de normas documentales, citación, referenciación e informe técnico con rigor y conciencia histórica.",
    "intellectual_property": "Al terminar la sesión podrás reconocer derechos de autor, patentes y licencias para usar recursos digitales con ética, créditos y justicia tecnológica.",
}


FAMILY_HOW_TO_USE: dict[str, str] = {
    "digital_identity": "Trabaja como ciudadano digital en formación: observa tu propia huella, protege datos, analiza riesgos y escribe decisiones que puedas aplicar sin esperar instrucciones externas.",
    "digital_communication": "Trabaja como comunicador responsable: antes de escribir define propósito, destinatario, tono, adjunto y dato sensible. Revisa la versión final con la lista de netiqueta.",
    "computer_systems": "Trabaja como técnico aprendiz: observa, clasifica, pregunta por la función de cada componente y reporta con evidencia sin manipular partes riesgosas.",
    "document_processing": "Trabaja como editor de documentos: primero ordena ideas, luego aplica formato y finalmente revisa ortografía, fuentes, legibilidad y coherencia.",
    "information_literacy": "Trabaja como investigador autónomo: formula una pregunta, prueba palabras clave, evalúa fuentes y contrasta antes de concluir.",
    "document_standards": "Trabaja como autor riguroso: organiza el informe, registra fuentes, cita con honestidad y usa la norma para que otra persona pueda revisar tu trabajo.",
    "intellectual_property": "Trabaja como creador responsable: pregunta qué recurso es propio, cuál es ajeno, qué licencia tiene y cómo reconocer a quien lo creó.",
}


FAMILY_ESCUTA_CHALLENGE: dict[str, str] = {
    "digital_identity": "Revisa una situación cotidiana de identidad digital: una foto, nombre de usuario, comentario, perfil o archivo compartido. Pregúntate qué dice de ti y qué datos deberías proteger.",
    "digital_communication": "Recuerda un mensaje escolar que haya sido claro o confuso. Observa asunto, tono, destinatario, adjuntos y datos compartidos.",
    "computer_systems": "Observa tu estación de trabajo como sistema. Identifica qué partes reciben datos, cuáles procesan, cuáles muestran resultados y qué acciones ayudan a conservarlas.",
    "document_processing": "Mira un documento escolar real o imaginado. Identifica qué facilita la lectura y qué la dificulta: títulos, párrafos, imágenes, tablas, ortografía o fuentes.",
    "information_literacy": "Piensa en una información que viste en Internet y que pudo ser dudosa. Antes de creerla, decide qué señales revisarías.",
    "document_standards": "Observa un trabajo escrito. Pregúntate si otra persona podría saber de dónde salió la información, qué norma sigue y qué conclusión defiende.",
    "intellectual_property": "Elige una imagen, texto, canción, software o diseño que uses en clase. Pregúntate quién lo creó y bajo qué permiso podrías usarlo.",
}


FAMILY_ESCUTA_CHECK_ITEMS: dict[str, list[str]] = {
    "digital_identity": ["Reconozco datos personales que debo proteger.", "Identifico una acción que deja huella digital.", "Formulo una pregunta sobre privacidad o consentimiento."],
    "digital_communication": ["Identifico propósito y destinatario del mensaje.", "Reconozco una regla de netiqueta necesaria.", "Detecto un dato o adjunto que debo revisar."],
    "computer_systems": ["Identifico componentes físicos y lógicos.", "Relaciono una parte con su función.", "Reconozco una acción segura y una que debo evitar."],
    "document_processing": ["Identifico propósito y audiencia del documento.", "Reconozco un problema de formato o estructura.", "Ubico una fuente, imagen o tabla que requiere revisión."],
    "information_literacy": ["Formulo una pregunta de búsqueda.", "Identifico señales de confiabilidad o duda.", "Reconozco la necesidad de contrastar fuentes."],
    "document_standards": ["Reconozco la estructura del documento.", "Identifico una fuente que debe citarse.", "Relaciono norma documental con rigor y lectura."],
    "intellectual_property": ["Distingo recurso propio y recurso ajeno.", "Identifico una licencia o permiso posible.", "Reconozco un riesgo de plagio o uso indebido."],
}


FAMILY_ESCUTA_RESPONSE_PROMPTS: dict[str, tuple[str, str]] = {
    "digital_identity": ("Un dato o huella digital que debo cuidar:", "Una decisión responsable que puedo tomar:"),
    "digital_communication": ("El propósito del mensaje que debo escribir o mejorar:", "Un cuidado de netiqueta o privacidad que debo aplicar:"),
    "computer_systems": ("El componente o función que observo:", "Una falla, riesgo o cuidado que debo reportar:"),
    "document_processing": ("El documento o parte que debo organizar:", "Un criterio de formato, fuente o revisión que debo aplicar:"),
    "information_literacy": ("La pregunta que orienta mi búsqueda:", "Una señal de confiabilidad o duda que debo verificar:"),
    "document_standards": ("La parte del informe o norma que debo cuidar:", "Una fuente o idea que necesita cita:"),
    "intellectual_property": ("El recurso o creación que quiero usar:", "El permiso, crédito o riesgo ético que debo revisar:"),
}


FAMILY_PERIOD_IDEA: dict[str, str] = {
    "digital_identity": "El periodo inicia con ciudadanía digital: identidad, huella, comunicación formal, seguridad, memoria de los medios y socialización de hallazgos.",
    "digital_communication": "El periodo fortalece comunicación digital formal con correo institucional, netiqueta, redacción, adjuntos, privacidad y propósito comunicativo.",
    "computer_systems": "El periodo desarrolla comprensión del computador como sistema: arquitectura, hardware, software, archivos, ensamble virtual, diagnóstico y cuidado preventivo.",
    "document_processing": "El periodo trabaja producción documental: procesador de texto, formato, imágenes, tablas, normas básicas, investigación y producto final revisado.",
    "information_literacy": "El periodo fortalece investigación autónoma: búsqueda, evaluación de fuentes, fake news, seguridad digital y conclusiones con evidencia.",
    "document_standards": "El periodo articula historia de la tecnología con normas documentales, citación, informes técnicos, propiedad intelectual y conciencia histórica.",
    "intellectual_property": "El periodo aborda creación y justicia tecnológica mediante patentes, derechos de autor, licencias, software libre y uso responsable de recursos.",
}


FAMILY_PRAXIS_TASK: dict[str, str] = {
    "digital_identity": "Elabora un mapa de tu huella digital escolar. Incluye datos, cuentas, publicaciones, archivos y permisos. Luego marca qué protegerías, qué ajustarías y qué compromiso aplicarás.",
    "digital_communication": "Redacta o corrige un correo institucional completo. Debe tener asunto, saludo, propósito, solicitud, adjunto nombrado, despedida, firma y revisión de privacidad.",
    "computer_systems": "Construye una ficha de diagnóstico del sistema computacional: componentes, función, síntoma observado, acción segura, acción que debe evitarse y reporte al docente.",
    "document_processing": "Elabora o mejora una página de documento con título, subtítulo, dos párrafos, una imagen o tabla con fuente, formato consistente y lista de revisión aplicada.",
    "information_literacy": "Investiga una pregunta sencilla sobre un problema cercano. Evalúa dos fuentes, contrasta un dato, registra créditos y escribe una conclusión prudente.",
    "document_standards": "Organiza un esquema de informe técnico con estructura normativa, una cita, una referencia y una explicación de cómo la norma mejora la confiabilidad del trabajo.",
    "intellectual_property": "Analiza tres recursos digitales que podrías usar en un proyecto. Indica autor, tipo de licencia o permiso, uso permitido, riesgo y forma correcta de crédito.",
}


def tech_history_variant(guide: GuideData) -> str:
    focus = clean(guide.title).lower()
    if any(term in focus for term in ["era electrónica", "era electronica", "revolución digital", "revolucion digital", "internet", "computador"]):
        return "digital"
    return "industrial"


def make_question(title: str, period: PeriodData) -> str:
    focus = make_focus(title).lower()
    family = guide_family_from_text(title)
    if family == "digital_identity":
        return "¿Cómo puedo cuidar mi identidad y huella digital sin afectar mi dignidad ni la de otras personas?"
    if family == "digital_communication":
        return "¿Cómo escribir comunicaciones digitales formales, claras y respetuosas que protejan datos y eviten malentendidos?"
    if family == "computer_systems":
        return "¿Cómo reconocer las partes de un sistema computacional, explicar su función y cuidarlas con autonomía?"
    if family == "document_processing":
        return "¿Cómo crear documentos claros, legibles y bien revisados usando estructura, formato, imágenes, tablas y fuentes?"
    if family == "information_literacy":
        return "¿Cómo buscar, evaluar y contrastar información confiable antes de usarla en una investigación escolar?"
    if family == "document_standards":
        return "¿Cómo aplicar normas documentales, citas y referencias para que un informe sea riguroso y verificable?"
    if family == "intellectual_property":
        return "¿Cómo usar creaciones, software e imágenes respetando autoría, licencias y justicia tecnológica?"
    if family == "systems_room":
        return "¿Cómo puedo reconocer de manera autónoma la sala de sistemas, identificar riesgos y proponer acuerdos para cuidarla?"
    if family == "tech_history":
        if any(term in focus for term in ["era electrónica", "era electronica", "revolución digital", "revolucion digital"]):
            return "¿Cómo transformaron la electrónica, los computadores e Internet la comunicación, el trabajo y la vida social?"
        return "¿Cómo transformaron la Revolución Industrial, el telégrafo y la electricidad la producción, la comunicación y la vida cotidiana?"
    if family == "cloud_collaboration":
        return f"¿Cómo usar {focus} para colaborar en la nube con roles, permisos, evidencia y cuidado de datos?"
    if family == "ewaste":
        return f"¿Cómo analizar {focus} para reducir impacto ambiental y proponer acciones responsables?"
    if family == "business_web_presence":
        return f"¿Cómo usar {focus} para construir una presencia digital empresarial clara, ética y verificable?"
    if family == "business_process_automation":
        return f"¿Cómo usar {focus} para mejorar procesos empresariales con datos, automatización e IA responsable?"
    if family == "integrated_business_project":
        return f"¿Cómo integrar {focus} en una solución empresarial sustentada con IA, datos y evidencias?"
    if family == "ai_book_project":
        return f"¿Cómo usar {focus} para escribir, editar y publicar un libro con IA sin perder autoría ni ética?"
    if family == "latex_business_report":
        return f"¿Cómo usar {focus} para producir informes comerciales o estudios de mercado claros, citados y compilables?"
    if family == "business_dashboard_ai":
        return f"¿Cómo usar {focus} para transformar datos empresariales en visualizaciones, indicadores y decisiones sustentadas?"
    if family == "ai":
        return f"¿Cómo comprender y usar {focus} de manera crítica, ética y verificable?"
    if family == "data_analysis":
        return f"¿Cómo usar {focus} para leer datos con prudencia, evidencia y conclusiones responsables?"
    if family == "spreadsheet":
        return f"¿Cómo usar {focus} para organizar datos, comprobar cálculos y tomar decisiones informadas?"
    if family == "physical_computing":
        return f"¿Cómo usar {focus} para leer señales del entorno, programar respuestas y comprobar resultados?"
    if family == "scratch_programming":
        return f"¿Cómo usar {focus} para construir una narrativa digital interactiva con lógica, pruebas y depuración?"
    if family == "programming":
        return f"¿Cómo usar {focus} para tomar decisiones claras en un programa y explicar su impacto?"
    if family == "presentation":
        return f"¿Cómo comunicar una idea tecnológica con {focus} de forma clara, ética y visualmente ordenada?"
    if family == "editorial":
        return f"¿Cómo usar {focus} para ordenar información, cuidar autoría y mejorar la lectura?"
    if family == "network":
        return f"¿Cómo comprender {focus} para explicar conexiones, probar comunicación y cuidar la información?"
    if family == "info_project":
        return f"¿Cómo integrar {focus} para comunicar una necesidad escolar con evidencias, pruebas y responsabilidad digital?"
    if family == "media":
        return f"¿Cómo mejorar recursos digitales de {focus} para comunicar mejor sin perder sentido ni respeto?"
    if family == "cultural_media":
        return f"¿Cómo crear {focus} con respeto por la cultura popular, la autoría y la dignidad de las personas?"
    if family == "safety":
        return f"¿Cómo reconocer riesgos de {focus} y actuar con cuidado, evidencia y rutas de apoyo?"
    if "sala de sistemas" in focus:
        return "¿Cómo podemos reconocer y cuidar la sala de sistemas para aprender mejor con otros?"
    if "identidad digital" in focus or "huella" in focus:
        return "¿Qué dice nuestra huella digital sobre quiénes somos y cómo podemos cuidarla?"
    if "correo" in focus:
        return "¿Cómo usar el correo institucional para comunicarnos con claridad, respeto y propósito?"
    if "hardware" in focus or "computador" in focus or "periféricos" in focus or "perifericos" in focus or "ensamble" in focus:
        return "¿Cómo reconocer los componentes del computador nos ayuda a cuidarlo y usarlo mejor?"
    if "software" in focus:
        return "¿Cómo distinguir tipos de software para usar el computador con más autonomía y criterio?"
    if "sistema operativo" in focus:
        return "¿Cómo reconocer las funciones del sistema operativo para manejar recursos y archivos con autonomía?"
    if "gestión de archivos" in focus or "gestion de archivos" in focus:
        return "¿Cómo organizar archivos y carpetas para encontrar, proteger y compartir información de forma responsable?"
    if "internet" in focus or "fuentes" in focus or "fake news" in focus:
        return "¿Cómo decidir si una información en Internet es confiable y útil para aprender?"
    if "procesador de texto" in focus or "documento" in focus:
        return "¿Cómo producir documentos claros, ordenados y útiles para comunicar una idea?"
    return f"¿Cómo se relaciona {focus} con el uso responsable y creativo de la tecnología?"


def make_product(title: str, session: int) -> str:
    focus = make_focus(title).lower()
    if session in {10, 20, 30}:
        return "Portafolio de evidencias, autoevaluación y socialización de aprendizajes."
    family = guide_family_from_text(title)
    if family == "digital_identity":
        return "Mapa de huella digital con datos protegidos, riesgo identificado, acción responsable y compromiso de autocuidado."
    if family == "digital_communication":
        return "Correo institucional o simulación revisada con asunto, netiqueta, firma, adjunto claro y protección de datos."
    if family == "computer_systems":
        return "Ficha de sistema computacional con componentes, funciones, diagnóstico seguro y cuidado preventivo."
    if family == "document_processing":
        return "Documento editado con estructura, formato consistente, recurso visual o tabla citada y lista de revisión."
    if family == "information_literacy":
        return "Ficha de investigación con pregunta, palabras clave, fuentes evaluadas, contraste y conclusión prudente."
    if family == "document_standards":
        return "Esquema o informe técnico con estructura normativa, cita, referencia y reflexión sobre rigor documental."
    if family == "intellectual_property":
        return "Análisis de recursos y licencias con créditos, permisos, riesgos de plagio y reflexión de justicia tecnológica."
    if family == "systems_room":
        return "Ficha autónoma de diagnóstico de la sala: mapa de riesgos, normas de uso y compromiso verificable."
    if family == "tech_history":
        if any(term in focus for term in ["era electrónica", "era electronica", "revolución digital", "revolucion digital"]):
            return "Línea de tiempo de la era electrónica y digital con impactos sociales, beneficios, riesgos y brechas."
        return "Línea de tiempo de la Revolución Industrial al telégrafo/electricidad con cambios técnicos y sociales."
    if family == "cloud_collaboration":
        return "Evidencia colaborativa en la nube con roles, permisos, historial/comentarios y reflexión sobre datos."
    if family == "ewaste":
        return "Informe o auditoría RAEE con inventario, clasificación, riesgos, propuesta comunitaria y sustentación."
    if family == "business_web_presence":
        return "Evidencia de presencia digital empresarial: marca, público, arquitectura, textos, SEO, pieza de red o sustentación ética."
    if family == "business_process_automation":
        return "Evidencia de proceso mejorado: mapa, formulario, base, flujo, protocolo de datos, FAQ o presentación de automatización."
    if family == "integrated_business_project":
        return "Proyecto empresarial integrado con problema, usuario, solución, presencia digital, datos, dashboard, automatización, documentación y pitch."
    if family == "ai_book_project":
        return "Evidencia editorial del libro: prompts, borrador, reescritura, créditos, declaración de IA y reflexión autoral."
    if family == "latex_business_report":
        return "Sección de informe profesional en LaTeX con estructura, tabla o gráfico, fuente, análisis y declaración de IA."
    if family == "business_dashboard_ai":
        return "Dashboard o reporte empresarial con datos limpios, indicadores, visualizaciones, hallazgos, recomendación y sustentación."
    if family == "ai":
        return "Evidencia crítica sobre IA: explicación, ejemplo aplicado, verificación y reflexión ética."
    if family == "data_analysis":
        return "Mini análisis de datos con tabla, cálculo o gráfico, conclusión prudente y evidencia de revisión."
    if family == "spreadsheet":
        return "Hoja de cálculo con datos organizados, fórmula comprobada, explicación y conclusión breve."
    if family == "physical_computing":
        return "Simulación o algoritmo de computación física con sensor, condición, respuesta y prueba documentada."
    if family == "scratch_programming":
        return "Proyecto Scratch o análisis PRIMM con algoritmo, bloques, pruebas, depuración y reflexión."
    if family == "programming":
        return "Algoritmo comentado o diagrama de flujo con prueba de casos y reflexión sobre decisiones."
    if family == "presentation":
        return "Presentación de 5 a 7 diapositivas con mensaje claro, diseño consistente y evidencia de revisión."
    if family == "editorial":
        return "Publicación digital breve con retícula, jerarquía visual, fuentes citadas y revisión de legibilidad."
    if family == "network":
        return "Esquema o simulación de red con componentes, dirección o protocolo, prueba y explicación de resultados."
    if family == "info_project":
        return "Portafolio integrador con publicación digital, esquema de red, pruebas, reflexión ética y mejora sustentada."
    if family == "media":
        return "Recurso visual o sonoro editado con versión inicial, versión final y justificación de cambios."
    if family == "cultural_media":
        return "Pieza visual o sonora con intención estética, créditos, consentimiento y reflexión crítica sobre representación."
    if family == "safety":
        return "Protocolo o infografía de prevención con señales de alerta, rutas de ayuda y compromiso personal."
    if family == "project":
        return "Producto integrador con problema, propuesta, evidencia técnica, reflexión ética y socialización."
    if "línea de tiempo" in focus or "historia" in focus:
        return "Línea de tiempo visual con explicación breve y reflexión crítica."
    if "correo" in focus or "netiqueta" in focus:
        return "Mensaje o simulación de comunicación digital profesional."
    if "hardware" in focus or "componentes" in focus or "ensamble" in focus or "periféricos" in focus or "perifericos" in focus:
        return "Mapa o ficha de identificación de componentes y cuidados básicos."
    if "software" in focus:
        return "Cuadro de clasificación de software con ejemplos de sistema operativo, aplicación y programación."
    if "sistema operativo" in focus:
        return "Ficha de exploración del sistema operativo: interfaz, recursos, ventanas, archivos y acciones básicas."
    if "gestión de archivos" in focus or "gestion de archivos" in focus:
        return "Árbol de carpetas organizado con nombres claros, ubicación de archivos y explicación del criterio usado."
    if "documento" in focus or "word" in focus or "texto" in focus:
        return "Documento editado con formato claro, imágenes o tablas según la consigna."
    if "seguridad" in focus or "contraseñas" in focus or "phishing" in focus:
        return "Protocolo visual de prevención y respuesta ante riesgos digitales."
    return "Evidencia de aprendizaje MILC: análisis, producción breve y compromiso personal."


def make_question_for_guide(guide: GuideData) -> str:
    family = guide_family(guide)
    if family == "systems_room":
        return "¿Cómo reconocer la sala de sistemas, cuidar sus equipos y trabajar con autonomía y convivencia?"
    if family == "digital_identity":
        return "¿Cómo puedo cuidar mi identidad y huella digital sin afectar mi dignidad ni la de otras personas?"
    if family == "digital_communication":
        return "¿Cómo escribir comunicaciones digitales formales, claras y respetuosas que protejan datos y eviten malentendidos?"
    if family == "computer_systems":
        return "¿Cómo reconocer las partes de un sistema computacional, explicar su función y cuidarlas con autonomía?"
    if family == "document_processing":
        return "¿Cómo crear documentos claros, legibles y bien revisados usando estructura, formato, imágenes, tablas y fuentes?"
    if family == "information_literacy":
        return "¿Cómo buscar, evaluar y contrastar información confiable antes de usarla en una investigación escolar?"
    if family == "document_standards":
        return "¿Cómo aplicar normas documentales, citas y referencias para que un informe sea riguroso y verificable?"
    if family == "intellectual_property":
        return "¿Cómo usar creaciones, software e imágenes respetando autoría, licencias y justicia tecnológica?"
    if family == "editorial":
        return "¿Cómo aplicar diseño editorial digital para comunicar con claridad, autoría, accesibilidad y revisión responsable?"
    if family == "cloud_collaboration":
        return "¿Cómo colaborar en la nube con roles, permisos, evidencias y cuidado de datos para construir un producto común?"
    if family == "scratch_programming":
        return "¿Cómo comprender, probar y mejorar un algoritmo o proyecto interactivo con PRIMM, lógica y depuración?"
    if family == "ai":
        return "¿Cómo usar IA con prompts claros, verificación, privacidad y criterio humano en una necesidad escolar?"
    if family == "ewaste":
        return "¿Cómo diagnosticar residuos tecnológicos y proponer acciones responsables para reducir impacto ambiental?"
    if family == "data_analysis":
        return "¿Cómo organizar datos reales, representarlos y concluir con prudencia según la evidencia?"
    if family == "spreadsheet":
        return "¿Cómo usar hojas de cálculo para organizar datos, aplicar fórmulas y comprobar conclusiones?"
    if family == "programming":
        return "¿Cómo expresar decisiones lógicas con condiciones, operadores y pruebas de caso?"
    if family == "physical_computing":
        return "¿Cómo leer señales del entorno, programar respuestas y comprobar un sistema físico o simulado?"
    if family == "presentation":
        return "¿Cómo comunicar una idea tecnológica mediante una presentación clara, visual y sustentada?"
    if family == "media":
        return "¿Cómo editar recursos visuales o sonoros para comunicar mejor respetando autoría y sentido?"
    if family == "cultural_media":
        return "¿Cómo crear contenidos digitales que representen cultura popular con respeto, autoría y reflexión crítica?"
    if family == "safety":
        return "¿Cómo reconocer riesgos digitales, proteger evidencias y activar rutas de ayuda sin amplificar el daño?"
    if family == "network":
        return "¿Cómo explicar, simular o diagnosticar una red reconociendo nodos, protocolos, dirección y seguridad?"
    if family == "info_project":
        return "¿Cómo integrar publicación digital, red escolar y ética de la información en un proyecto verificable?"
    if family == "ai_book_project":
        return "¿Cómo crear y editar un libro con IA conservando voz propia, ética, derechos de autor y evidencia del proceso?"
    if family == "latex_business_report":
        return "¿Cómo convertir la escritura asistida en un informe comercial en LaTeX con datos, fuentes y análisis?"
    if family == "business_dashboard_ai":
        return "¿Cómo transformar datos empresariales en indicadores, visualizaciones y decisiones sustentadas con IA verificada?"
    if family == "business_web_presence":
        return "¿Cómo construir una presencia digital empresarial clara, ética, accesible y orientada al cliente?"
    if family == "business_process_automation":
        return "¿Cómo diagnosticar y mejorar procesos empresariales con datos, automatización básica, IA verificada y protección de información?"
    if family == "integrated_business_project":
        return "¿Cómo integrar problema, usuario, solución digital, datos, automatización y pitch en un proyecto empresarial sustentable?"
    if family == "project":
        return "¿Cómo organizar y sustentar un portafolio de evidencias con criterios de calidad, reflexión ética y mejoras verificables?"
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return "¿Cómo transformaron los medios electrónicos y digitales la comunicación, el acceso a información y las desigualdades?"
        return "¿Cómo han respondido los inventos a necesidades sociales y qué personas o comunidades pudieron quedar excluidas?"
    return make_question(guide.title, guide.period)


def make_product_for_guide(guide: GuideData) -> str:
    family = guide_family(guide)
    if guide.number in {10, 20, 30}:
        return "Portafolio de evidencias, autoevaluación y socialización de aprendizajes."
    if family == "systems_room":
        return "Ficha autónoma de diagnóstico de la sala: mapa de riesgos, normas de uso y compromiso verificable."
    if family == "digital_identity":
        return "Mapa de huella digital con datos protegidos, riesgo identificado, acción responsable y compromiso de autocuidado."
    if family == "digital_communication":
        return "Correo institucional o simulación revisada con asunto, netiqueta, firma, adjunto claro y protección de datos."
    if family == "computer_systems":
        return "Ficha de sistema computacional con componentes, funciones, diagnóstico seguro y cuidado preventivo."
    if family == "document_processing":
        return "Documento editado con estructura, formato consistente, recurso visual o tabla citada y lista de revisión."
    if family == "information_literacy":
        return "Ficha de investigación con pregunta, palabras clave, fuentes evaluadas, contraste y conclusión prudente."
    if family == "document_standards":
        return "Esquema o informe técnico con estructura normativa, cita, referencia y reflexión sobre rigor documental."
    if family == "intellectual_property":
        return "Análisis de recursos y licencias con créditos, permisos, riesgos de plagio y reflexión de justicia tecnológica."
    if family == "editorial":
        return "Publicación digital breve con retícula, jerarquía visual, fuentes citadas y revisión de legibilidad."
    if family == "cloud_collaboration":
        return "Evidencia colaborativa en la nube con roles, permisos, historial o comentarios y reflexión sobre datos."
    if family == "scratch_programming":
        return "Proyecto Scratch o análisis PRIMM con algoritmo, bloques, pruebas, depuración y reflexión."
    if family == "ai":
        return "Evidencia crítica sobre IA: explicación, ejemplo aplicado, verificación y reflexión ética."
    if family == "ewaste":
        return "Informe o auditoría RAEE con inventario, clasificación, riesgos, propuesta comunitaria y sustentación."
    if family == "data_analysis":
        return "Mini análisis de datos con tabla, cálculo o gráfico, conclusión prudente y evidencia de revisión."
    if family == "spreadsheet":
        return "Hoja de cálculo con datos organizados, fórmula comprobada, explicación y conclusión breve."
    if family == "programming":
        return "Algoritmo comentado o diagrama de flujo con prueba de casos y reflexión sobre decisiones."
    if family == "physical_computing":
        return "Simulación o algoritmo de computación física con sensor, condición, respuesta y prueba documentada."
    if family == "presentation":
        return "Presentación de 5 a 7 diapositivas con mensaje claro, diseño consistente y evidencia de revisión."
    if family == "media":
        return "Recurso visual o sonoro editado con versión inicial, versión final y justificación de cambios."
    if family == "cultural_media":
        return "Pieza visual o sonora con intención estética, créditos, consentimiento y reflexión crítica sobre representación."
    if family == "safety":
        return "Protocolo o infografía de prevención con señales de alerta, rutas de ayuda y compromiso personal."
    if family == "network":
        return "Esquema o simulación de red con componentes, dirección o protocolo, prueba y explicación de resultados."
    if family == "info_project":
        return "Portafolio integrador con publicación digital, esquema de red, pruebas, reflexión ética y mejora sustentada."
    if family == "ai_book_project":
        return "Evidencia editorial del libro: prompts, borrador, reescritura, créditos, declaración de IA y reflexión autoral."
    if family == "latex_business_report":
        return "Sección de informe profesional en LaTeX con estructura, tabla o gráfico, fuente, análisis y declaración de IA."
    if family == "business_dashboard_ai":
        return "Dashboard o reporte empresarial con datos limpios, indicadores, visualizaciones, hallazgos, recomendación y sustentación."
    if family == "business_web_presence":
        return "Evidencia de presencia digital empresarial: marca, público, arquitectura, textos, SEO, pieza de red o sustentación ética."
    if family == "business_process_automation":
        return "Evidencia de proceso mejorado: mapa, formulario, base, flujo, protocolo de datos, FAQ o presentación de automatización."
    if family == "integrated_business_project":
        return "Proyecto empresarial integrado con problema, usuario, solución, presencia digital, datos, dashboard, automatización, documentación y pitch."
    if family == "project":
        return "Producto integrador con problema, propuesta, evidencia técnica, reflexión ética y socialización."
    if family == "tech_history":
        if "necesidades sociales" in guide.focus.lower() or "excluidos" in guide.focus.lower() or "dusseliano" in guide.focus.lower():
            return "Matriz crítica de invento, necesidad social, beneficiados, excluidos e impacto tecnológico."
        return make_product(guide.title, guide.number)
    return make_product(guide.title, guide.number)


def build_guides(periods: dict[int, PeriodData], topics: dict[int, object]) -> list[GuideData]:
    guides = []
    for number in range(1, 31):
        period = periods[period_for_session(number)]
        entry = topics[number]
        if isinstance(entry, dict):
            title = clean(entry.get("titulo", ""))
            short_title = clean(entry.get("titulo_corto", "")) or make_short_title(title)
        else:
            title = clean(str(entry))
            short_title = make_short_title(title)
        probe = GuideData(
            number=number,
            grade=GRADE,
            period=period,
            title=title,
            short_title=short_title,
            focus=make_focus(title),
            question="",
            product="",
        )
        question = make_question_for_guide(probe)
        product = make_product_for_guide(probe)
        guides.append(
            GuideData(
                number=number,
                grade=GRADE,
                period=period,
                title=title,
                short_title=short_title,
                focus=make_focus(title),
                question=question,
                product=product,
            )
        )
    return guides


def concept_rows(guide: GuideData) -> list[tuple[str, str, str]]:
    focus = guide.focus.lower()
    family = guide_family(guide)
    if family in FAMILY_CONCEPT_ROWS:
        return FAMILY_CONCEPT_ROWS[family]
    if family == "systems_room":
        return [
            ("Sala de sistemas", "Espacio común donde se aprende con equipos, normas y colaboración.", "Entrar con cuidado, escuchar indicaciones y ubicar el puesto asignado."),
            ("Estación de trabajo", "Conjunto de elementos que usa un estudiante en clase.", "Computador, monitor, teclado, mouse, silla y mesa."),
            ("Diagnóstico tecnológico", "Observación ordenada para reconocer estado, riesgos y necesidades antes de actuar.", "Revisar si el mouse funciona y reportar si no responde."),
            ("Riesgo físico", "Situación del espacio que puede causar accidente o daño al equipo.", "Cable atravesado, líquido cerca del teclado o silla mal ubicada."),
            ("Riesgo digital", "Situación que expone cuentas, datos o información personal.", "Encontrar una sesión abierta o una contraseña visible."),
            ("Norma de uso", "Acuerdo claro que orienta una acción segura y respetuosa.", "No consumir alimentos cerca de los equipos."),
            ("Reporte responsable", "Aviso breve y verificable sobre una falla, sin manipular elementos peligrosos.", "Informar al docente: puesto 12, teclado desconectado, antes de mover cables."),
            ("Autonomía", "Capacidad de leer, observar, decidir y pedir ayuda cuando corresponde.", "Marco la lista de cotejo antes de preguntar qué debo hacer."),
        ]
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                ("Electrónica", "Uso de componentes que controlan señales eléctricas para procesar información.", "Radio, televisión, circuitos, transistores y tarjetas."),
                ("Transistor", "Componente que permitió equipos más pequeños, rápidos y confiables.", "Base de computadores, radios portátiles y dispositivos digitales."),
                ("Computador", "Máquina programable que procesa datos y automatiza tareas.", "Calcular, escribir, simular, almacenar y comunicar información."),
                ("Internet", "Red global que conecta dispositivos, personas y servicios.", "Correo, páginas web, plataformas educativas y comercio digital."),
                ("Revolución digital", "Cambio social producido por datos, software, redes y dispositivos conectados.", "Trabajo remoto, banca digital, redes sociales y aprendizaje en línea."),
                ("Automatización", "Uso de sistemas para realizar procesos con menor intervención humana.", "Cajeros, inventarios, producción y atención digital."),
                ("Brecha digital", "Desigualdad en acceso, habilidades o uso significativo de la tecnología.", "No tener conexión, equipo o formación para participar plenamente."),
                ("Impacto social", "Consecuencia de una tecnología en relaciones, trabajo, cultura y poder.", "Nuevos oficios, vigilancia, dependencia o acceso a información."),
            ]
        return [
            ("Revolución Industrial", "Periodo de cambios técnicos y sociales ligados a máquinas, fábricas y nuevas fuentes de energía.", "Producción textil mecanizada y fábricas urbanas."),
            ("Mecanización", "Uso de máquinas para reemplazar o ampliar trabajo manual.", "Telar mecánico o máquina de vapor."),
            ("Energía", "Capacidad que permite mover máquinas, iluminar espacios o transmitir señales.", "Vapor, carbón y electricidad."),
            ("Telégrafo", "Sistema de comunicación a distancia mediante señales codificadas.", "Enviar mensajes por código Morse entre ciudades."),
            ("Electricidad", "Fuente que transformó iluminación, comunicación, industria y vida cotidiana.", "Alumbrado público, motores eléctricos y redes de energía."),
            ("Infraestructura", "Conjunto de redes y obras que permiten funcionar a una tecnología.", "Vías férreas, postes, cables, centrales y estaciones."),
            ("Urbanización", "Crecimiento de ciudades asociado a fábricas, servicios y migración.", "Barrios obreros cerca de zonas industriales."),
            ("Cambio social", "Transformación en trabajo, tiempo, familia, educación o desigualdad.", "Nuevas jornadas laborales y organización obrera."),
        ]
    if family == "cloud_collaboration":
        return [
            ("Nube", "Servicio que permite crear, guardar y compartir archivos por Internet.", "Drive, OneDrive, Docs o Word en línea."),
            ("Documento compartido", "Archivo que varias personas pueden editar o comentar según permisos.", "Un informe grupal con historial de cambios."),
            ("Permiso", "Nivel de acceso que define qué puede hacer cada persona.", "Ver, comentar, editar o administrar."),
            ("Rol colaborativo", "Responsabilidad concreta dentro de un equipo.", "Coordinador, relator, verificador o editor."),
            ("Historial", "Registro de cambios que permite revisar aportes y recuperar versiones.", "Ver quién editó un párrafo y cuándo."),
            ("Comentario", "Intervención que sugiere, pregunta o retroalimenta sin borrar el texto.", "Pedir fuente para un dato dudoso."),
            ("Sincronización", "Actualización de archivos entre dispositivos y cuentas.", "Un cambio aparece en el computador y en el celular."),
            ("Soberanía de datos", "Capacidad de decidir dónde, cómo y con quién se comparte información.", "No subir datos personales sin permiso."),
        ]
    if family == "ewaste":
        return [
            ("RAEE", "Residuos de aparatos eléctricos y electrónicos que requieren manejo especial.", "Celulares, cables, monitores, pilas o teclados dañados."),
            ("E-waste", "Nombre en inglés para residuos tecnológicos con partes aprovechables y riesgos ambientales.", "Tarjetas, baterías o componentes obsoletos."),
            ("Obsolescencia", "Pérdida de utilidad por daño, incompatibilidad, moda o falta de soporte.", "Un equipo funciona, pero ya no recibe actualizaciones."),
            ("Riesgo ambiental", "Daño posible por sustancias, metales o mala disposición.", "Baterías arrojadas a la basura común."),
            ("Rethink", "Repensar si realmente se necesita comprar o desechar un dispositivo.", "Reparar antes de reemplazar."),
            ("Refuse", "Rechazar consumo tecnológico innecesario o de corta vida.", "No cambiar audífonos solo por moda."),
            ("Reduce", "Disminuir residuos mediante cuidado, reparación y uso prolongado.", "Proteger cargadores y reutilizar componentes."),
            ("Auditoría", "Revisión ordenada para identificar dispositivos, estado, riesgos y acciones.", "Inventario de equipos obsoletos del aula."),
        ]
    if family == "ai_book_project":
        return [
            ("Escritura inversa", "Método que imagina primero el libro final para planear hacia atrás el proceso.", "Definir lector, propósito, índice y tono antes de redactar."),
            ("Prompt editorial", "Instrucción que orienta a la IA con rol, contexto, tarea, formato y criterios.", "Pedir una escaleta con capítulos, público lector y estilo."),
            ("Voz propia", "Manera personal de pensar, narrar, argumentar y elegir palabras.", "Reescribir una sugerencia de IA con experiencia y ejemplos propios."),
            ("Borrador asistido", "Texto inicial generado o mejorado con ayuda de IA, aún no listo para entregar.", "Primer capítulo que requiere lectura, edición y verificación."),
            ("Edición humana", "Revisión consciente de coherencia, ritmo, ortografía, estructura y sentido.", "Reducir repeticiones y fortalecer una idea débil."),
            ("Derechos de autor", "Reconocimiento y respeto por obras, imágenes, ideas y fuentes ajenas.", "Citar una imagen usada en la portada."),
            ("Declaración de IA", "Nota honesta que explica cómo se usó IA en el proceso.", "La IA ayudó a proponer títulos; el autor revisó y reescribió."),
            ("Producto editorial", "Libro digital con portada, índice, capítulos, créditos y formato consistente.", "PDF final presentado en la feria editorial MILC."),
        ]
    if family == "business_web_presence":
        return [
            ("Presencia digital", "Conjunto de espacios, mensajes y recursos con los que una marca se presenta en línea.", "Landing page, portafolio comercial, perfil o página de servicios."),
            ("Propuesta de valor", "Razón clara por la que un usuario debería interesarse por el producto, servicio o marca.", "Entrega rápida, precio justo, asesoría personalizada o calidad local."),
            ("Público objetivo", "Grupo de personas para quienes se diseña el mensaje y la experiencia digital.", "Estudiantes, familias, comerciantes del barrio o clientes frecuentes."),
            ("Identidad de marca", "Decisiones de nombre, tono, colores, imágenes y promesa que hacen reconocible la propuesta.", "Una marca cercana, confiable y orientada al servicio."),
            ("Arquitectura web", "Organización de secciones para guiar al usuario desde la información hasta la acción.", "Inicio, beneficios, servicios, evidencias, contacto y preguntas frecuentes."),
            ("Llamado a la acción", "Indicación concreta para que el usuario dé un siguiente paso.", "Solicitar información, hacer pedido, descargar catálogo o escribir por WhatsApp."),
            ("SEO básico", "Ajustes de títulos, palabras clave y descripciones para que el contenido sea fácil de encontrar.", "Título claro con producto, lugar y beneficio principal."),
            ("Ética digital empresarial", "Cuidado de imagen, datos, publicidad, autoría y transparencia cuando se publica una oferta.", "No prometer resultados falsos ni usar imágenes sin permiso."),
        ]
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                ("Proceso empresarial", "Secuencia de actividades que transforma una entrada en una salida útil para un cliente, usuario o equipo.", "Recibir pedido, registrar datos, preparar entrega y confirmar al cliente."),
                ("Inventario de procesos", "Lista organizada de procesos existentes, con responsables, entradas, salidas y problemas observados.", "Ventas, pedidos, inventario, atención y seguimiento."),
                ("Entrada", "Solicitud, dato, recurso o situación que inicia el proceso.", "Mensaje de WhatsApp, formulario, producto faltante o pedido nuevo."),
                ("Actividad", "Acción concreta que ocurre dentro del proceso.", "Revisar disponibilidad, registrar cantidad, responder una pregunta o actualizar una tabla."),
                ("Responsable", "Persona o rol que ejecuta, revisa o aprueba una parte del proceso.", "Ventas, atención, inventario, contabilidad o coordinación."),
                ("Salida", "Resultado verificable que deja el proceso.", "Pedido confirmado, cliente informado, inventario actualizado o reporte de encuesta."),
                ("Punto crítico", "Lugar del proceso donde se repiten errores, retrasos, pérdidas de datos o confusiones.", "Pedidos duplicados, respuestas tardías o datos incompletos."),
                ("Oportunidad digital", "Posible mejora con formulario, hoja de cálculo, plantilla, IA o automatización futura.", "Pasar pedidos del chat a un formulario con campos obligatorios."),
                ("Dato sensible", "Información personal o comercial que requiere cuidado, consentimiento y acceso limitado.", "Teléfono, dirección, nombre del cliente, pago o historial de compra."),
            ]
        return [
            ("Proceso empresarial", "Secuencia de actividades que transforma una entrada en una salida útil para un negocio o servicio.", "Recibir pedido, confirmar datos, preparar entrega y registrar pago."),
            ("Entrada", "Información, solicitud o recurso que inicia el proceso.", "Formulario de pedido, mensaje de cliente o registro de inventario."),
            ("Actividad", "Acción concreta realizada por una persona, herramienta o sistema.", "Clasificar respuestas, enviar confirmación o actualizar una tabla."),
            ("Responsable", "Persona o rol encargado de ejecutar, revisar o aprobar una parte del proceso.", "Atención al cliente, ventas, inventario o contabilidad."),
            ("Salida", "Resultado verificable que deja el proceso.", "Pedido confirmado, factura, reporte, respuesta o base actualizada."),
            ("Automatización", "Uso de herramientas digitales para repetir tareas, reducir errores y ahorrar tiempo.", "Respuesta automática, filtro, plantilla o flujo de notificación."),
            ("Formulario digital", "Instrumento para capturar datos con campos, validaciones y consentimiento.", "Encuesta de satisfacción o formulario de pedidos."),
            ("PDCA", "Ciclo de mejora continua: planear, hacer, verificar y actuar.", "Probar un flujo, medir errores y ajustar el procedimiento."),
            ("Protección de datos", "Cuidado legal y ético de información personal recolectada en procesos digitales.", "Pedir autorización y no publicar datos de clientes."),
        ]
    if family == "integrated_business_project":
        return [
            ("Problema empresarial", "Necesidad comercial, contable, administrativa o de servicio que se puede mejorar con una solución.", "Baja visibilidad, pedidos desordenados o falta de control de inventario."),
            ("Usuario o cliente", "Persona beneficiaria cuyas necesidades orientan el diseño.", "Comprador, emprendedor escolar, familia, curso o comunidad."),
            ("Solución", "Producto, servicio, proceso o sistema que responde al problema con criterios verificables.", "Página de ventas con formulario, dashboard y protocolo de atención."),
            ("Presencia digital", "Canal donde se comunica la solución, su valor y la forma de contacto.", "Landing page, portafolio o ficha comercial."),
            ("Dato e indicador", "Evidencia organizada para medir viabilidad, utilidad, costo, preferencia o impacto.", "Ventas simuladas, encuesta, margen, demanda o satisfacción."),
            ("Dashboard", "Tablero que sintetiza indicadores y visualizaciones para tomar decisiones.", "Gráficos de demanda, costos, ingresos y recomendaciones."),
            ("Automatización", "Mejora parcial que agiliza una tarea del proyecto.", "Formulario con confirmación, plantilla de respuesta o flujo de seguimiento."),
            ("Documentación", "Registro técnico y comercial para que otra persona entienda, use y evalúe la solución.", "Manual breve, costos, fuentes, datos, permisos y declaración de IA."),
            ("Pitch", "Presentación breve y convincente que defiende problema, solución, evidencia, impacto y sostenibilidad.", "Exposición ante feria empresarial o tribunal MILC."),
        ]
    if family == "latex_business_report":
        return [
            ("Informe profesional", "Documento que comunica un problema, método, datos, análisis y conclusión.", "Informe comercial sobre preferencias de clientes."),
            ("LaTeX", "Lenguaje de marcado para codificar documentos y compilar PDF con estructura estable.", "\\section, \\subsection, listas, tablas y figuras."),
            ("Código documental", "Texto fuente que define contenido y formato del documento.", "documentclass, begin document, title y author."),
            ("Caso comercial", "Situación de producto, servicio, emprendimiento o proceso contable que se analiza.", "Ventas de una tienda escolar por categoría."),
            ("Fuente de datos", "Origen verificable de información usada en el informe.", "Encuesta, registro de ventas, tabla de costos o formulario."),
            ("Tabla", "Organización de datos en filas y columnas para comparar resultados.", "Frecuencias de preferencias por producto."),
            ("Gráfico", "Visualización que resume patrones o comparaciones importantes.", "Barras de ventas por mes o circular de preferencias."),
            ("Citación", "Registro de fuentes, imágenes, datos y apoyos usados.", "Fuente: encuesta aplicada al curso, abril de 2026."),
            ("Declaración de IA", "Explicación transparente del apoyo recibido y de la revisión humana.", "IA usada para sugerir estructura; datos verificados por el equipo."),
        ]
    if family == "business_dashboard_ai":
        return [
            ("Dato empresarial", "Registro útil para comprender ventas, costos, clientes, tiempos o preferencias.", "Fecha, producto, cantidad, precio y cliente."),
            ("Indicador", "Medida que resume desempeño para tomar decisiones.", "Ingreso total, utilidad, margen, producto más vendido."),
            ("Limpieza de datos", "Corrección de duplicados, vacíos, formatos y errores antes de analizar.", "Unificar fechas y separar números de texto."),
            ("Visualización", "Gráfico o tabla que permite ver patrones con rapidez.", "Barras para productos, líneas para tendencia mensual."),
            ("Dashboard", "Tablero que reúne indicadores, gráficos y filtros para leer una situación.", "Panel de ventas con total, margen y tendencia."),
            ("Reporte asistido", "Narrativa que explica hallazgos con ayuda de IA y verificación humana.", "Resumen de oportunidades basado en los gráficos."),
            ("Decisión empresarial", "Elección comercial, contable o administrativa sustentada en evidencia.", "Aumentar inventario de un producto rentable."),
            ("Riesgo", "Límite o consecuencia que debe considerarse antes de decidir.", "Muestra pequeña, datos incompletos o sesgo de encuesta."),
            ("Sustentación", "Defensa oral o escrita de datos, visualizaciones, hallazgos y recomendación.", "Explicar por qué una decisión conviene y qué falta verificar."),
        ]
    if family == "ai":
        return [
            ("Inteligencia artificial", "Sistema que realiza tareas que suelen requerir aprendizaje, reconocimiento de patrones o generación de respuestas.", "Un asistente que resume un texto o responde preguntas."),
            ("Datos", "Información usada para entrenar, consultar o evaluar un sistema de IA.", "Textos, imágenes, clics, respuestas o ejemplos."),
            ("Modelo", "Sistema entrenado con datos para producir predicciones, clasificaciones o contenidos.", "Un modelo que recomienda videos o genera texto."),
            ("Prompt", "Instrucción que damos a una IA para orientar su respuesta.", "Rol, contexto, tarea, formato y criterios de calidad."),
            ("IA generativa", "IA que crea textos, imágenes, ideas, códigos o borradores a partir de instrucciones.", "Generar un borrador que luego reviso y corrijo."),
            ("Alucinación", "Respuesta falsa o inventada que parece segura o convincente.", "Citar una fuente que no existe."),
            ("Sesgo", "Resultado parcial o injusto por datos, diseño o uso inadecuado.", "Asociar profesiones solo con un género."),
            ("Verificación", "Proceso de contrastar una respuesta con fuentes, evidencias y criterio humano.", "Comparar la respuesta con una página institucional o libro confiable."),
            ("Privacidad", "Cuidado de datos personales propios y de otras personas.", "No escribir nombres completos, direcciones o información sensible en un prompt."),
        ]
    if family == "data_analysis":
        return [
            ("Dato", "Registro de una característica observada o medida.", "Cantidad de residuos, edad, tiempo o respuesta de encuesta."),
            ("Variable", "Característica que puede cambiar entre casos.", "Tipo de transporte, nota, temperatura o asistencia."),
            ("Registro", "Fila que reúne datos de un caso o persona sin exponer información sensible.", "Un día de medición o una respuesta anónima."),
            ("Campo", "Columna que define qué se está observando.", "Fecha, categoría, cantidad o comentario."),
            ("Limpieza", "Revisión para corregir datos repetidos, vacíos o mal escritos.", "Unificar 'si', 'sí' y 'SI' antes de contar."),
            ("Cálculo", "Operación que resume o compara datos.", "Total, promedio, porcentaje, máximo o mínimo."),
            ("Gráfico", "Representación visual que ayuda a comparar patrones.", "Barras para categorías o líneas para cambios en el tiempo."),
            ("Phronesis", "Prudencia para decidir con datos sin ignorar contexto, límites y consecuencias.", "No concluir con una encuesta incompleta."),
        ]
    if family == "spreadsheet":
        return [
            ("Dato", "Valor que se registra para analizar una situación.", "Cantidad de residuos por día."),
            ("Celda", "Espacio de la hoja donde se escribe un dato o fórmula.", "B2 contiene el precio unitario."),
            ("Fórmula", "Expresión que calcula un resultado a partir de datos.", "=B2*C2 calcula un subtotal."),
            ("Función", "Fórmula predefinida para resolver una tarea frecuente.", "=PROMEDIO(B2:B8)."),
            ("Referencia", "Dirección de una celda usada en un cálculo.", "B2 cambia; $B$2 queda fija."),
            ("Validación", "Regla que controla qué datos se pueden ingresar.", "Solo números entre 1 y 5."),
        ]
    if family == "physical_computing":
        return [
            ("Computación física", "Uso de código para leer señales del entorno y producir respuestas.", "Un sensor detecta luz y activa un LED."),
            ("Sensor", "Dispositivo o entrada que mide una condición.", "Sensor de luz, temperatura, botón o micrófono."),
            ("Actuador", "Salida que realiza una acción visible, sonora o mecánica.", "LED, buzzer, pantalla o motor."),
            ("Variable", "Dato que cambia durante la ejecución del programa.", "temperatura, nivel_luz o alerta."),
            ("Umbral", "Valor límite usado para decidir una acción.", "Si temperatura > 30, mostrar alerta."),
            ("Evento", "Situación que inicia una acción del programa.", "Al presionar A, al cambiar luz o al iniciar."),
            ("Condición", "Pregunta lógica que decide qué respuesta activar.", "Si luz < 40, encender LED."),
            ("Prueba", "Verificación con valores esperados, contrarios y límite.", "Probar 39, 40 y 41 cuando el umbral es 40."),
            ("Depuración", "Proceso de encontrar y corregir errores de lógica, conexión o lectura.", "Cambiar el umbral porque la alerta se activa siempre."),
        ]
    if family == "scratch_programming":
        return [
            ("Algoritmo", "Secuencia ordenada de pasos para resolver una tarea.", "Mover, preguntar, decidir y mostrar respuesta."),
            ("Sprite", "Objeto o personaje programable dentro de Scratch.", "Un personaje que salta o habla."),
            ("Evento", "Acción que inicia un bloque de código.", "Al presionar bandera verde o tecla espacio."),
            ("Variable", "Dato que cambia durante el programa.", "puntaje, vidas, turno o respuesta."),
            ("Bucle", "Repetición de instrucciones mientras se cumple una condición o durante cierto número de veces.", "Repetir 10 veces mover 5 pasos."),
            ("Condicional", "Bloque que decide qué camino tomar.", "Si toca borde, rebotar."),
            ("PRIMM", "Ruta para aprender programación: predecir, ejecutar, investigar, modificar y crear.", "Antes de cambiar un juego, se predice qué hará."),
            ("Depuración", "Búsqueda y corrección de errores en el programa.", "Detectar por qué el sprite no responde al evento."),
        ]
    if family == "programming":
        return [
            ("Condición", "Pregunta lógica que decide qué camino sigue un programa.", "Si la nota es mayor o igual a 3.0, aprueba."),
            ("Rama", "Bloque de instrucciones que se ejecuta cuando se cumple un caso.", "if, else if y else organizan alternativas."),
            ("Anidación", "Una decisión ubicada dentro de otra decisión.", "Si hay permiso, revisar si también hay cupo."),
            ("Operador lógico", "Conector que combina condiciones para tomar mejores decisiones.", "AND exige dos verdades; OR acepta una."),
            ("Orden de evaluación", "Secuencia en que el programa revisa las condiciones.", "Primero revisa if; luego else if; al final else."),
            ("Prueba de casos", "Revisión con valores distintos para encontrar errores.", "Probar frontera, caso esperado y caso contrario."),
        ]
    if family == "presentation":
        return [
            ("Mensaje central", "Idea principal que debe recordar la audiencia.", "Una frase que resume el proyecto."),
            ("Jerarquía visual", "Orden de importancia entre títulos, imágenes y datos.", "Título grande, apoyo breve, evidencia clara."),
            ("Diapositiva", "Unidad visual que desarrolla una sola idea.", "Problema, propuesta, evidencia o cierre."),
            ("Transición", "Cambio entre diapositivas que ayuda al recorrido.", "Usarla solo cuando aporta claridad."),
            ("Hipervínculo", "Enlace que conecta con una fuente o recurso.", "Botón hacia el video o formulario."),
            ("Coherencia", "Relación entre colores, fuentes, imágenes y tono.", "La misma paleta en toda la presentación."),
        ]
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                ("Audiencia", "Grupo de personas para quienes se diseña la publicación.", "Estudiantes de sexto que necesitan entender una norma de convivencia."),
                ("Propósito", "Intención comunicativa que orienta las decisiones de diseño.", "Informar, invitar, prevenir, orientar o convencer con evidencia."),
                ("Mensaje central", "Idea principal que la audiencia debe recordar después de leer.", "Cuidar los datos personales también protege al grupo."),
                ("Retícula", "Estructura de márgenes, columnas y zonas que ordena la página.", "Dos columnas: una para texto breve y otra para imagen o dato clave."),
                ("Jerarquía visual", "Orden de importancia visible entre título, subtítulo, cuerpo e imagen.", "Título grande, subtítulo orientador, párrafos cortos y fuente al final."),
                ("Legibilidad", "Facilidad para leer sin esfuerzo ni confusión.", "Buen contraste, letra suficiente y poco texto por bloque."),
                ("Imagen con licencia", "Recurso visual usado con permiso, crédito o fuente verificable.", "Imagen Creative Commons citada debajo o al cierre del boletín."),
                ("Revisión editorial", "Lectura final con criterios antes de publicar o entregar.", "Verificar ortografía, autoría, contraste, orden y propósito."),
            ]
        return [
            ("Publicación digital", "Producto que organiza información para una audiencia concreta.", "Folleto, boletín o ficha informativa escolar."),
            ("Retícula", "Estructura invisible que ordena columnas, márgenes y bloques.", "Dos columnas con margen constante."),
            ("Jerarquía visual", "Diferencia de tamaño, peso o posición para guiar la lectura.", "Título principal, subtítulo y cuerpo de texto."),
            ("Legibilidad", "Facilidad con que se lee un contenido.", "Contraste suficiente y párrafos breves."),
            ("Imagen con propósito", "Recurso visual que aporta información y no solo adorno.", "Fotografía citada que explica un problema."),
            ("Autoría y licencia", "Reconocimiento de fuentes y permisos de uso.", "Imagen Creative Commons citada al final."),
            ("Accesibilidad", "Condición para que más personas puedan comprender el contenido.", "Texto alternativo, buen contraste y lenguaje claro."),
            ("Revisión editorial", "Lectura final para corregir forma, sentido y derechos.", "Lista de cotejo antes de publicar."),
        ]
    if family == "network":
        return [
            ("Red", "Conjunto de dispositivos conectados para compartir datos o servicios.", "Computadores conectados al router de la sala."),
            ("Nodo", "Dispositivo o punto que participa en la red.", "Computador, celular, impresora, servidor o router."),
            ("Emisor y receptor", "Quien envía y quien recibe un mensaje o paquete de datos.", "Un celular envía una foto y otro la recibe."),
            ("Canal", "Medio por el que viaja la información.", "Cable, Wi-Fi, Bluetooth o radio."),
            ("Protocolo", "Regla que permite que los dispositivos se entiendan.", "IP organiza direcciones; Bluetooth empareja equipos cercanos."),
            ("Dirección IP", "Identificación lógica de un dispositivo dentro de una red.", "192.168.1.10 en una red local."),
            ("Topología", "Forma en que se conectan los dispositivos.", "Estrella, bus, malla o árbol."),
            ("Prueba de conectividad", "Verificación de que dos nodos pueden comunicarse.", "Enviar ping o revisar si un servicio responde."),
        ]
    if family == "info_project":
        return [
            ("Infoesfera", "Entorno de información donde personas, medios y redes conviven e influyen entre sí.", "Boletín escolar que circula por canales digitales."),
            ("Necesidad", "Situación que requiere informar, orientar o prevenir con claridad.", "Explicar normas de uso seguro de la sala."),
            ("Publicación", "Producto editorial que organiza el mensaje para una audiencia.", "Boletín, folleto, infografía o ficha."),
            ("Red escolar", "Conjunto de conexiones y servicios por donde circula información.", "Equipos, router, servidor o canal de publicación."),
            ("Evidencia", "Prueba que demuestra cómo se diseñó, configuró o verificó algo.", "Captura de simulación, diagrama o prueba de conectividad."),
            ("Criterio ético", "Regla para proteger datos, autoría, inclusión y convivencia.", "No exponer información personal ni usar imágenes sin permiso."),
            ("Mejora", "Cambio justificado después de revisar evidencia o recibir retroalimentación.", "Aumentar contraste o corregir un rango IP."),
            ("Sustentación", "Explicación oral o escrita de decisiones, límites y aprendizajes.", "Defender por qué el diseño y la red responden a la necesidad."),
        ]
    if family == "media":
        return [
            ("Recurso original", "Archivo de partida antes de editar.", "Foto, captura o audio grabado."),
            ("Propósito", "Razón comunicativa de la edición.", "Informar, narrar, alertar o explicar."),
            ("Recorte", "Selección de la parte más importante.", "Quitar espacio que distrae."),
            ("Mejora", "Ajuste técnico que aumenta claridad.", "Brillo, contraste, volumen o limpieza."),
            ("Exportación", "Guardado final en formato adecuado.", "PNG para imagen; MP3 o WAV para audio."),
            ("Derechos", "Respeto por autoría, consentimiento y uso.", "Citar fuente y evitar exponer personas."),
        ]
    if family == "cultural_media":
        return [
            ("Cultura popular", "Saberes, símbolos, prácticas y expresiones de una comunidad.", "Música, oficios, fiestas, dichos, comidas o relatos locales."),
            ("Estética", "Forma sensible de expresar belleza, sentido e identidad.", "Color, ritmo, composición, voz o imagen."),
            ("Representación", "Manera en que se muestra a una persona, grupo o territorio.", "Retratar una tradición con respeto y contexto."),
            ("Estereotipo", "Imagen simplificada o injusta que reduce una cultura.", "Mostrar una comunidad como si todas las personas fueran iguales."),
            ("Consentimiento", "Autorización informada para usar imagen, voz o historia de alguien.", "Pedir permiso antes de grabar o publicar."),
            ("Autoría", "Reconocimiento de quien creó, narró o aportó un recurso.", "Créditos de música, foto, relato o entrevista."),
            ("Voz comunitaria", "Participación de quienes viven o conocen la experiencia representada.", "Incluir una cita o revisión de la persona entrevistada."),
            ("Edición responsable", "Mejora técnica que no borra sentido ni manipula la realidad.", "Ajustar audio sin cambiar el testimonio."),
        ]
    if family == "safety":
        return [
            ("Riesgo digital", "Situación en línea que puede causar daño.", "Acoso, engaño o exposición de datos."),
            ("Señal de alerta", "Indicio que pide detenerse y buscar apoyo.", "Amenaza, insistencia o solicitud íntima."),
            ("Privacidad", "Control sobre datos, imágenes e identidad.", "No compartir ubicación o fotos privadas."),
            ("Evidencia", "Registro que ayuda a pedir apoyo responsable.", "Captura, fecha, usuario y enlace."),
            ("Ruta de ayuda", "Persona o institución a quien acudir.", "Familia, docente, orientación o autoridad."),
            ("Cuidado colectivo", "Acción para protegerse y proteger a otros.", "No reenviar, acompañar y reportar."),
        ]
    if family == "project":
        return [
            ("Problema", "Necesidad concreta que se quiere comprender.", "Una situación escolar o social verificable."),
            ("Evidencia", "Dato o registro que sustenta la propuesta.", "Encuesta, observación, captura o prueba."),
            ("Propuesta", "Respuesta posible al problema identificado.", "Presentación, prototipo o campaña."),
            ("Criterio", "Condición para decidir si el producto funciona.", "Claridad, utilidad, ética y viabilidad."),
            ("Iteración", "Mejora después de recibir retroalimentación.", "Cambiar una diapositiva confusa."),
            ("Socialización", "Comunicación pública del proceso y resultado.", "Exponer hallazgos y decisiones."),
        ]
    if "correo" in focus or "netiqueta" in focus:
        return [
            ("Correo institucional", "Canal formal para comunicarse en la escuela.", "Enviar una consulta clara al docente."),
            ("Asunto", "Frase breve que anticipa el propósito del mensaje.", "Solicitud de revisión de tarea."),
            ("Destinatario", "Persona o grupo que debe recibir el mensaje.", "Escribir al docente correcto y evitar copias innecesarias."),
            ("Netiqueta", "Normas de respeto en comunicación digital.", "Saludar, escribir claro y despedirse."),
            ("Privacidad", "Cuidado de datos, archivos y conversaciones.", "No reenviar correos con información personal."),
            ("Evidencia", "Producto que muestra lo aprendido.", "Captura, borrador o correo revisado."),
        ]
    if "hardware" in focus or "computador" in focus or "periféricos" in focus:
        return [
            ("Hardware", "Parte física del sistema computacional.", "CPU, RAM, teclado o pantalla."),
            ("Entrada", "Dispositivo o acción que permite ingresar datos.", "Teclado, mouse, micrófono o cámara."),
            ("Procesamiento", "Trabajo interno que transforma datos en resultados.", "La CPU ejecuta instrucciones del programa."),
            ("Función", "Tarea que cumple un componente.", "La RAM ayuda a trabajar con programas abiertos."),
            ("Cuidado", "Acción que protege el equipo.", "No mover cables ni golpear periféricos."),
            ("Reporte", "Aviso responsable de una falla.", "Informar pantalla sin imagen antes de manipular."),
        ]
    if "internet" in focus or "fuentes" in focus or "fake news" in focus:
        return [
            ("Fuente", "Lugar de donde proviene la información.", "Página institucional o artículo confiable."),
            ("Búsqueda", "Uso de palabras clave para encontrar información útil.", "Combinar tema, lugar y fecha."),
            ("Sesgo", "Inclinación que puede afectar la información.", "Texto que solo muestra una postura."),
            ("Validación", "Revisión de confiabilidad y pertinencia.", "Comparar autor, fecha y evidencia."),
            ("Contraste", "Comparación entre varias fuentes antes de concluir.", "Revisar si dos sitios serios coinciden."),
            ("Soberanía informacional", "Capacidad de decidir con criterio.", "No compartir información dudosa."),
        ]
    if "texto" in focus or "documento" in focus or "word" in focus:
        return [
            ("Formato", "Decisiones visuales que ordenan el documento.", "Fuente, tamaño, alineación y espaciado."),
            ("Estructura", "Organización de partes del texto.", "Título, subtítulos, párrafos e imágenes."),
            ("Párrafo", "Unidad de ideas que desarrolla un punto concreto.", "Una idea principal con apoyo breve."),
            ("Claridad", "Facilidad para leer y comprender.", "Ideas cortas con buena presentación."),
            ("Revisión", "Lectura final para corregir forma y contenido.", "Verificar ortografía, imágenes y márgenes."),
            ("Propósito", "Razón por la que se escribe.", "Informar, narrar, explicar o argumentar."),
        ]
    return [
        ("Tecnología", "Uso creativo y responsable de conocimientos y recursos.", "Resolver una necesidad de la escuela."),
        ("Necesidad", "Situación que pide una respuesta o mejora.", "Organizar información del curso."),
        ("Proceso", "Secuencia de acciones para lograr un resultado.", "Planear, crear, revisar y compartir."),
        ("Riesgo", "Situación que puede causar daño o pérdida.", "Cuenta abierta, cable suelto o dato expuesto."),
        ("Criterio", "Razón para tomar una decisión responsable.", "Elegir una acción segura y respetuosa."),
        ("Compromiso", "Acción concreta para mejorar el aprendizaje común.", "Cuidar equipos y respetar acuerdos."),
    ]


def case_text(guide: GuideData) -> str:
    focus = guide.focus.lower()
    family = guide_family(guide)
    if family in FAMILY_CASE_TEXT:
        return FAMILY_CASE_TEXT[family]
    if family == "systems_room":
        return "Llegas por primera vez a una estación de trabajo. El monitor está apagado, hay un cable cerca del paso, un compañero quiere mover conexiones para avanzar rápido y en otro equipo aparece una cuenta abierta. Antes de actuar, debes observar, decidir qué puedes resolver con seguridad y qué debes reportar."
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return "Un grupo afirma que la revolución digital solo trajo progreso porque ahora todo es más rápido. Al revisar el caso, aparecen beneficios como comunicación inmediata y acceso a información, pero también brecha digital, vigilancia, dependencia de plataformas y cambios en el trabajo."
        return "Un equipo construye una línea de tiempo sobre la Revolución Industrial, pero solo anota inventos. Al revisar, descubre que cada tecnología también cambió jornadas laborales, ciudades, comunicación, energía, desigualdades y formas de organización social."
    if family == "cloud_collaboration":
        return "Un equipo trabaja en un documento compartido. Una persona borra aportes sin avisar, otra cambia permisos a público y nadie registra roles ni comentarios. El grupo debe reorganizar el trabajo, proteger datos y dejar evidencia de colaboración."
    if family == "ewaste":
        return "En la institución aparecen cables, teclados y equipos dañados guardados sin clasificación. Algunas personas quieren botarlos a la basura común, pero el grupo debe identificar riesgos, posibilidades de reutilización y una ruta responsable de gestión RAEE."
    if family == "ai_book_project":
        return "Una estudiante quiere publicar un libro con tema libre usando IA. Tiene ideas, un borrador y una portada, pero algunas partes suenan impersonales, no hay créditos de imágenes y no queda claro qué escribió ella y qué sugirió la IA. Debe reorganizar el proceso para conservar autoría, ética y calidad editorial."
    if family == "business_web_presence":
        return "Un emprendimiento escolar publica una página con textos genéricos, colores incoherentes, imágenes sin permiso y sin llamado a la acción. Aunque la idea es valiosa, el usuario no entiende qué ofrece, para quién es ni cómo contactarlo. El equipo debe rediseñar la presencia digital con estrategia, ética y evidencia."
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return "Un emprendimiento escolar vende productos durante los descansos. Los pedidos llegan por chat, algunos se anotan en cuadernos, otros quedan en la memoria del equipo, se mezclan pagos, cantidades y datos de clientes. Antes de proponer una herramienta, el grupo debe inventariar los procesos, reconocer puntos críticos y priorizar una oportunidad de mejora."
        return "Un negocio escolar recibe pedidos por chat. Se pierden datos, se duplican solicitudes, algunas respuestas llegan tarde y nadie sabe si el cliente autorizó el uso de su información. El equipo debe mapear el proceso, detectar tareas repetitivas y proponer una mejora digital con protección de datos."
    if family == "integrated_business_project":
        return "Un equipo llega a la feria empresarial con sitio web, datos, dashboard, automatización y pitch, pero cada pieza parece hecha por separado. La solución no demuestra bien el problema, el usuario, los indicadores ni el impacto. El reto es integrar todo en una propuesta coherente, verificable y defendible."
    if family == "latex_business_report":
        return "Un equipo debe entregar un informe comercial en PDF. Tiene una encuesta, una tabla de costos y varias ideas, pero el documento está desordenado, no cita fuentes, el gráfico no tiene interpretación y el archivo LaTeX no compila por errores de estructura."
    if family == "business_dashboard_ai":
        return "Un emprendimiento escolar registra ventas, costos y clientes en una hoja de cálculo. El equipo quiere tomar una decisión rápida, pero hay datos duplicados, una visualización confusa y una recomendación de IA que no coincide del todo con los números."
    if family == "ai":
        return "Un equipo usa una herramienta de IA para avanzar rápido en una tarea. La respuesta parece completa, pero incluye datos sin fuente, una afirmación dudosa y un ejemplo que podría exponer información personal. El grupo debe decidir qué conservar, qué verificar, qué corregir y cómo declarar el uso de IA."
    if family == "data_analysis":
        return "Un equipo reúne datos sobre una situación escolar, pero mezcla categorías, deja celdas vacías, elige un gráfico que exagera el resultado y quiere sacar una conclusión general sin revisar la calidad de la información."
    if family == "spreadsheet":
        return "Un grupo registra datos en una hoja de cálculo, pero mezcla textos con números, copia fórmulas sin revisar referencias y obtiene conclusiones que no coinciden con los datos."
    if family == "physical_computing":
        return "Un grupo programa una alerta ambiental con sensor simulado. La alerta se activa cuando no corresponde, nadie registró el umbral usado y el equipo no sabe si el error está en la condición, la lectura del sensor o la respuesta del actuador."
    if family == "scratch_programming":
        return "Un equipo modifica un proyecto de Scratch sin predecir qué hará el código. El personaje no responde al evento, el bucle nunca termina y nadie documenta qué bloque causó el error."
    if family == "programming":
        return "Un equipo crea un programa que clasifica resultados, pero ubica mal un else-if. Algunos casos funcionan y otros quedan en una categoría incorrecta porque no probaron valores límite."
    if family == "presentation":
        return "Un equipo tiene buena información, pero su presentación usa demasiado texto, animaciones innecesarias y no deja claro cuál es la propuesta principal."
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return "El consejo estudiantil quiere publicar un boletín para explicar una campaña de cuidado escolar. Tiene datos útiles y varias imágenes, pero la primera versión no define audiencia, mezcla tamaños de letra, usa una foto sin crédito y no deja claro qué debe leer primero la comunidad."
        return "Un grupo diseña un boletín con información útil, pero mezcla muchas fuentes, usa imágenes sin crédito, deja poco contraste y no define qué debe leer primero la audiencia."
    if family == "network":
        return "Un equipo arma o simula una red, pero no registra direcciones, confunde canal con protocolo y dice que la red funciona sin hacer una prueba de conectividad."
    if family == "info_project":
        return "Un equipo quiere publicar una campaña escolar, pero el boletín no tiene jerarquía visual, el esquema de red no muestra pruebas y algunas imágenes o datos personales podrían compartirse sin permiso."
    if family == "media":
        return "Una estudiante edita una imagen o audio para su exposición, pero elimina información importante, usa baja calidad y no explica por qué hizo los cambios."
    if family == "cultural_media":
        return "Un equipo quiere crear una pieza sobre cultura popular, pero usa imágenes sin permiso, repite estereotipos y no reconoce a las personas que aportaron relatos, música o fotografías."
    if family == "safety":
        return "En un chat aparece una situación incómoda. Algunas personas quieren reenviar capturas, otras no saben a quién acudir y la persona afectada necesita apoyo sin ser expuesta."
    if family == "project":
        return "Un equipo prepara un proyecto final, pero su evidencia no demuestra el problema, la propuesta no responde a la necesidad y la exposición queda desconectada de la reflexión ética."
    if "correo" in focus or "netiqueta" in focus:
        return "Un estudiante escribe al docente sin asunto, sin saludo y con palabras incompletas. Otro compañero propone reenviar el mensaje con copia a todo el curso sin revisar la información."
    if "hardware" in focus or "computador" in focus or "periféricos" in focus:
        return "Un equipo no enciende. Un estudiante mueve varios cables sin avisar, otro golpea el mouse porque no responde y el grupo pierde tiempo sin registrar lo ocurrido."
    if "internet" in focus or "fuentes" in focus or "fake news" in focus:
        return "Un grupo encuentra una página con información llamativa, pero no identifica autor, fecha ni fuente. Aun así quieren usarla como base de su trabajo."
    if "texto" in focus or "documento" in focus or "word" in focus:
        return "Un documento tiene muchos tipos de letra, imágenes sin relación y párrafos sin orden. El contenido existe, pero el lector no entiende con facilidad el mensaje."
    return "Durante la actividad, un equipo quiere avanzar rápido sin leer instrucciones, otro estudiante queda sin participar y aparece una decisión que puede afectar el cuidado del espacio o la información."


def conceptual_paragraphs(guide: GuideData) -> list[str]:
    family = guide_family(guide)
    focus = guide.focus
    if family in FAMILY_CONCEPTUAL_PARAGRAPHS:
        return [paragraph.format(focus=focus) for paragraph in FAMILY_CONCEPTUAL_PARAGRAPHS[family]]
    if family == "systems_room":
        return [
            "En esta primera guía no se espera que domines todos los equipos. El propósito es aprender a observar de manera autónoma: leer la consigna, mirar el puesto, reconocer señales de cuidado y registrar lo que encuentras antes de actuar.",
            "Un diagnóstico tecnológico es una revisión ordenada. Sirve para saber cómo está un espacio, qué elementos funcionan, qué riesgos existen y qué acciones son seguras para un estudiante.",
            "La sala de sistemas es un espacio compartido. Aunque cada estudiante usa un puesto, sus decisiones afectan a todo el grupo: un cable mal movido, una cuenta abierta o una norma ignorada pueden causar accidentes, pérdida de información o problemas de convivencia.",
            "Para trabajar con autonomía usa este ciclo: leo la tarea, observo sin manipular, comparo con la lista de cotejo, decido si puedo actuar o debo reportar, registro evidencia y escribo un compromiso verificable.",
        ]
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                f"En {focus}, la atención se centra en la era electrónica y digital: circuitos, transistores, computadores, redes e Internet cambiaron la forma de producir, estudiar, comprar, trabajar y relacionarnos.",
                "La revolución digital no es solo tener dispositivos. Es la transformación de actividades sociales mediante datos, software, conectividad y plataformas.",
                "El análisis debe reconocer beneficios y tensiones: acceso a información, automatización y colaboración, pero también brecha digital, dependencia tecnológica, vigilancia y nuevos conflictos laborales.",
                "Una línea de tiempo crítica no enumera aparatos: relaciona cada avance con un cambio social, una oportunidad, un riesgo y una pregunta sobre justicia tecnológica.",
            ]
        return [
            f"En {focus}, la atención se centra en la Revolución Industrial y en dos tecnologías decisivas: el telégrafo y la electricidad.",
            "La máquina de vapor, la mecanización y las fábricas modificaron la producción, el tiempo de trabajo, la organización de las ciudades y las desigualdades sociales.",
            "El telégrafo acortó distancias comunicativas: por primera vez un mensaje podía viajar más rápido que una persona, un caballo o un barco.",
                "La electricidad amplió la jornada, iluminó espacios, impulsó motores y creó nuevas infraestructuras. Por eso debe analizarse como cambio técnico y también como cambio social.",
            ]
    if family == "cloud_collaboration":
        return [
            f"En {focus}, colaborar no significa que todos escriban al mismo tiempo sin acuerdos. Significa coordinar roles, permisos, versiones y evidencias para construir un producto común.",
            "La nube facilita el trabajo porque permite compartir documentos, comentar, revisar historial y recuperar versiones, pero también exige cuidado con datos, enlaces y permisos.",
            "Una colaboración responsable deja huella: quién aportó, qué se modificó, qué fuentes se usaron, qué comentarios se atendieron y qué decisiones tomó el equipo.",
            "El liderazgo servicial ayuda a que la tecnología no sea competencia individual sino organización del bien común: escuchar, repartir tareas, cuidar ritmos y responder por el grupo.",
        ]
    if family == "ewaste":
        return [
            f"En {focus}, los residuos tecnológicos no son basura común. Muchos contienen partes reutilizables y materiales que pueden afectar ambiente y salud si se manejan mal.",
            "Una auditoría RAEE empieza por observar: tipo de dispositivo, estado, posible reparación, riesgo, responsable y ruta de disposición.",
            "Las estrategias Rethink, Refuse y Reduce ayudan a actuar antes del reciclaje: repensar consumo, rechazar compras innecesarias y reducir residuos con cuidado y reparación.",
            "La responsabilidad comunitaria consiste en convertir el diagnóstico en propuesta: inventario, clasificación, campaña, punto de acopio o recomendación institucional.",
        ]
    if family == "ai_book_project":
        return [
            f"En {focus}, la IA puede acompañar la escritura, pero no reemplaza la mirada del autor. El estudiante decide tema, lector, propósito, tono, estructura y versión final.",
            "La escritura inversa ayuda a trabajar con autonomía: se imagina el libro terminado, se define qué debe contener y luego se planean los pasos para llegar a ese producto.",
            "Un prompt editorial debe ser específico: rol, contexto, tarea, formato, extensión, tono, criterios y límites éticos. Un prompt vago produce textos genéricos difíciles de defender.",
            "La calidad aparece en la reescritura humana: verificar, corregir, ampliar, reducir, ordenar capítulos y agregar memoria, ejemplos o postura propia.",
            "Todo libro asistido por IA debe reconocer fuentes, derechos de autor, imágenes, licencias y declaración de uso de IA para evitar plagio y fortalecer transparencia.",
        ]
    if family == "business_web_presence":
        return [
            f"En {focus}, el propósito es transformar una idea empresarial en una presencia digital comprensible: quién ofrece, a quién sirve, qué problema resuelve y qué acción espera del usuario.",
            "La IA puede ayudar a explorar nombres, tonos, textos comerciales, palabras clave y calendario de publicaciones, pero el estudiante debe verificar que la marca sea auténtica, honesta y coherente con el contexto.",
            "Una página empresarial funciona cuando su arquitectura guía la lectura: propuesta de valor, beneficios, servicios, evidencia, contacto, preguntas frecuentes y llamado a la acción.",
            "El diseño visual no es decoración. Color, tipografía, imagen, contraste y accesibilidad deben apoyar la confianza y facilitar que cualquier usuario comprenda la oferta.",
            "Publicar también exige responsabilidad: derechos de imagen, datos personales, publicidad responsable, créditos de recursos y declaración transparente del apoyo de IA.",
        ]
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                f"En {focus}, el primer paso no es automatizar: es mirar cómo trabaja realmente un emprendimiento, negocio escolar o dependencia institucional. Un proceso se entiende cuando se puede explicar qué lo inicia, qué acciones ocurren, quién responde y qué resultado deja.",
                "Un inventario de procesos permite comparar varias rutinas: ventas, atención, inventario, encuestas, pedidos, pagos, seguimiento o entrega. Así el equipo evita escoger una herramienta solo porque parece novedosa.",
                "La oportunidad de mejora aparece cuando se identifican puntos críticos: datos repetidos, información incompleta, retrasos, tareas manuales innecesarias, pérdida de evidencias o falta de responsable.",
                "La IA puede ayudar a ordenar observaciones, sugerir preguntas o clasificar problemas, pero sus recomendaciones deben validarse con evidencia del proceso y con revisión humana.",
                "Como el periodo trabaja procesos empresariales, todo inventario debe cuidar información personal y comercial: solo se registran datos necesarios, se evita exponer clientes reales y se define quién puede acceder a la evidencia.",
            ]
        return [
            f"En {focus}, primero se comprende el proceso y después se automatiza. Si no se sabe cuál es la entrada, la actividad, el responsable y la salida, cualquier herramienta puede ordenar el error en lugar de resolverlo.",
            "La IA puede ayudar a diagnosticar tareas repetitivas, cuellos de botella, errores frecuentes y oportunidades de mejora, pero sus sugerencias deben comprobarse con datos, observación y criterio humano.",
            "Formularios, hojas de cálculo, filtros, plantillas y respuestas automáticas permiten capturar información, organizar registros y reducir trabajo repetitivo cuando se diseñan con campos claros.",
            "Todo proceso digital debe proteger datos: consentimiento, finalidad, acceso limitado, no publicación de información sensible y revisión antes de compartir resultados.",
            "La mejora continua se trabaja con PDCA: planear el cambio, hacerlo en pequeño, verificar evidencias y actuar ajustando el proceso con responsabilidad.",
        ]
    if family == "integrated_business_project":
        return [
            f"En {focus}, el grado undécimo integra lo aprendido: presencia digital, datos, IA, automatización, documentación y sustentación empresarial.",
            "Un proyecto final sólido nace de un problema real y de un usuario concreto. Antes de diseñar, se investiga, se formulan criterios de éxito y se reconocen límites del alcance.",
            "La presencia digital comunica la solución, los datos permiten evaluarla, el dashboard ayuda a decidir y la automatización mejora una parte del proceso.",
            "La IA puede apoyar ideación, análisis, redacción, visualización y preparación del pitch, pero cada uso debe quedar declarado, revisado y contrastado con evidencia.",
            "La sustentación final debe mostrar una línea clara: problema, usuario, solución, datos, decisión, riesgos, ética, sostenibilidad y mejora posible.",
        ]
    if family == "latex_business_report":
        return [
            f"En {focus}, el estudiante pasa de escribir textos creativos a producir documentos profesionales con estructura, datos y decisiones verificables.",
            "LaTeX permite codificar documentos: el contenido se organiza con comandos, se compila a PDF y conserva una presentación estable para informes largos.",
            "Un informe comercial o estudio de mercado no se limita a describir: formula un problema, explica método, presenta datos, analiza hallazgos y propone una decisión prudente.",
            "La IA puede ayudar a planear estructura, revisar claridad o sugerir análisis, pero las fuentes, datos y conclusiones deben verificarse con evidencia.",
            "La modalidad empresarial, comercial y contable exige cuidar precisión: unidades, costos, precios, muestra, fuentes, anexos y declaración honesta del proceso.",
        ]
    if family == "business_dashboard_ai":
        return [
            f"En {focus}, los datos empresariales se convierten en decisiones cuando están limpios, organizados y representados con indicadores comprensibles.",
            "Un dashboard no es una colección de gráficos: es una lectura ejecutiva que responde preguntas concretas sobre ventas, costos, clientes, márgenes o tendencias.",
            "La IA puede orientar cargue, limpieza, visualización y redacción de hallazgos, pero cada sugerencia debe contrastarse con la tabla y con el contexto del negocio.",
            "Un buen reporte combina evidencia y prudencia: muestra qué dicen los datos, qué no permiten afirmar, qué riesgos existen y qué decisión se recomienda.",
            "La sustentación empresarial exige argumentar con indicadores, gráficos, límites y consecuencias éticas, no solo leer una conclusión generada automáticamente.",
        ]
    if family == "ai":
        return [
            f"En {focus}, la meta no es creerle todo a una herramienta ni rechazarla por miedo. La meta es comprender qué puede hacer, qué no puede hacer y cómo usarla con criterio humano.",
            "Una IA trabaja con datos y patrones. Por eso sus respuestas pueden ser útiles, pero también incompletas, sesgadas o falsas. Toda producción asistida necesita lectura crítica, verificación y mejora humana.",
            "Para entender un recomendador sencillo se puede dibujar un diagrama de decisión: si una condición es verdadera, un IF activa una recomendación; si es falsa, ELSE muestra otra. Los operadores lógicos ayudan a leer reglas simples, aunque no explican por completo los modelos actuales.",
            "Un buen uso escolar de IA deja huella del proceso: propósito, prompt, respuesta obtenida, revisión, fuentes consultadas, correcciones y decisión final del estudiante.",
            "La autonomía consiste en formular mejores preguntas, proteger datos personales, contrastar información y explicar con honestidad cuándo y cómo se usó IA.",
        ]
    if family == "data_analysis":
        return [
            f"En {focus}, trabajar con datos no significa llenar una tabla: significa formular una pregunta, registrar información comparable y revisar si la evidencia permite responderla.",
            "Una tabla útil diferencia campos, registros, tipos de datos y unidades. Si los datos están incompletos o mezclados, el cálculo puede ser correcto pero la conclusión será débil.",
            "Los gráficos ayudan a ver patrones, pero también pueden confundir si se elige una escala, categoría o tipo de gráfico inadecuado.",
            "La phronesis exige prudencia: antes de decidir con datos, se revisa de dónde vienen, qué límites tienen, a quién afectan y qué no se puede afirmar todavía.",
        ]
    if family == "spreadsheet":
        return [
            f"En {focus}, la hoja de cálculo no se usa solo para escribir números: sirve para organizar datos, calcular con reglas claras y revisar si una conclusión tiene soporte.",
            "Antes de aplicar una fórmula, conviene reconocer qué representa cada columna, qué tipo de dato acepta y qué resultado se espera obtener.",
            "Una buena hoja permite rastrear el cálculo: si cambia un dato, el resultado debe actualizarse sin perder sentido."
        ]
    if family == "physical_computing":
        return [
            f"En {focus}, la programación se conecta con el mundo físico o simulado: una entrada se lee, una condición se evalúa y una salida responde.",
            "El centro no es solo hacer que algo se encienda. El centro es explicar qué sensor se lee, qué umbral se eligió, qué condición decide y qué actuador comunica la respuesta.",
            "Todo sistema de control necesita pruebas: valores normales, valores de alerta y valores límite. Sin esas pruebas no sabemos si el programa funciona o solo parece funcionar.",
            "La responsabilidad aparece al interpretar datos del entorno: una alerta debe ser comprensible, no generar pánico y proteger información o condiciones de la comunidad.",
        ]
    if family == "scratch_programming":
        return [
            f"En {focus}, programar es construir una secuencia de acciones que otra persona pueda probar, entender y mejorar.",
            "Scratch permite aprender lógica con bloques: eventos, movimiento, apariencia, sonido, variables, bucles y condicionales se combinan para crear una interacción.",
            "Un diagrama de flujo ayuda a leer la decisión antes de mover bloques: si la condición es verdadera, el bloque IF ejecuta una acción; si es falsa, ELSE permite otra respuesta. El operador compara datos como puntaje, vidas o tiempo.",
            "La ruta PRIMM evita copiar sin comprender: primero se predice, luego se ejecuta, se investiga cómo funciona, se modifica con control y finalmente se crea algo propio.",
            "Un buen proyecto interactivo incluye pruebas y depuración: se revisa qué debería pasar, qué pasó realmente, dónde está el error y qué ajuste mejora el resultado.",
        ]
    if family == "programming":
        return [
            f"En {focus}, programar significa ordenar decisiones para que el computador sepa qué hacer en cada caso.",
            "Una condición es una frase lógica que se puede responder como verdadero o falso. Para escribirla se usan variables, símbolos de comparación y operadores lógicos.",
            "Una estructura if-else-if revisa condiciones en orden: cuando una condición se cumple, ejecuta esa rama y deja de revisar las siguientes.",
            "Por eso el orden importa. Las condiciones más específicas suelen ir antes que las generales, y siempre conviene probar casos límite."
        ]
    if family == "presentation":
        return [
            f"En {focus}, una presentación no es un documento pegado en diapositivas: es una ruta visual para que otra persona comprenda una idea.",
            "Cada diapositiva debe desarrollar una sola intención: presentar un problema, mostrar evidencia, explicar una propuesta o cerrar con una conclusión.",
            "El diseño ayuda cuando guía la lectura; distrae cuando usa efectos, textos o colores sin propósito."
        ]
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                f"En {focus}, el estudiante asume el rol de editor aprendiz: observa una necesidad de comunicación, define a quién va dirigido el mensaje y decide cómo ordenar texto, imagen y evidencia para que otra persona pueda comprender sin explicación adicional.",
                "Diseñar un folleto o boletín no es decorar una hoja. Es tomar decisiones: qué información va primero, qué se resume, qué imagen aporta evidencia, qué fuente se cita y qué espacio se deja para que la lectura respire.",
                "La retícula funciona como una estructura invisible. Ayuda a ubicar márgenes, columnas, títulos, imágenes y créditos con orden. Cuando la retícula falla, el lector se pierde aunque la información sea correcta.",
                "La autoría también hace parte del diseño. Toda imagen, dato o texto tomado de otra fuente debe revisarse, citarse y usarse con permiso o licencia adecuada. Publicar sin créditos debilita la confianza y puede vulnerar derechos de autor.",
                "La autonomía editorial se demuestra cuando el estudiante no espera que le digan cada paso: compara ejemplos, boceta, revisa criterios, corrige una versión y justifica por qué cada cambio mejora la lectura.",
            ]
        return [
            f"En {focus}, diseñar no es decorar: es ordenar información para que una persona pueda leer, comparar y actuar con claridad.",
            "Una publicación digital necesita retícula, jerarquía visual y legibilidad. Esas decisiones hacen visible qué es importante, qué es apoyo y qué evidencia sostiene el mensaje.",
            "También requiere responsabilidad: citar fuentes, respetar licencias, cuidar imágenes de personas y revisar accesibilidad antes de compartir.",
            "La autonomía editorial consiste en planear, bocetar, revisar con criterios y justificar cada cambio con relación al propósito comunicativo.",
        ]
    if family == "network":
        return [
            f"En {focus}, una red no se entiende solo nombrando dispositivos: hay que explicar cómo viaja la información y qué reglas permiten la comunicación.",
            "Para analizar una red se identifican nodos, medio de transmisión, protocolo, dirección o servicio y prueba de funcionamiento.",
            "Una simulación o diagrama vale más cuando muestra evidencia: qué se configuró, qué se probó, qué falló y cómo se corrigió.",
            "La responsabilidad digital aparece cuando se cuidan datos, accesos, nombres de equipos y decisiones que afectan a otros usuarios de la red.",
        ]
    if family == "info_project":
        return [
            f"En {focus}, la infoesfera se trabaja como convivencia entre mensajes, personas, tecnologías y redes. No basta comunicar: hay que cuidar cómo circula la información.",
            "El proyecto integra dos dimensiones: una publicación clara y una explicación técnica de la red o canal por donde esa información puede circular.",
            "Toda sustentación necesita evidencias: bocetos, fuentes, licencias, diagrama de red, configuración, prueba de conectividad, errores corregidos y decisiones de mejora.",
            "La responsabilidad aparece en tres preguntas: qué datos protegemos, qué autorías reconocemos y qué impacto puede tener el mensaje en la comunidad escolar.",
        ]
    if family == "media":
        return [
            f"En {focus}, editar no significa adornar por adornar. Cada cambio debe mejorar comprensión, calidad o intención comunicativa.",
            "Antes de modificar un recurso conviene decidir qué debe entender la audiencia y qué partes del archivo aportan a ese propósito.",
            "También se cuida la ética: autoría, consentimiento, privacidad y representación respetuosa de las personas."
        ]
    if family == "cultural_media":
        return [
            f"En {focus}, crear contenido digital no consiste en usar la cultura como decoración. Consiste en reconocer dignidad, memoria, belleza y contexto.",
            "Una pieza responsable evita estereotipos, pide consentimiento cuando usa imagen o voz, reconoce autorías y explica por qué eligió ciertos símbolos o sonidos.",
            "La estética de la liberación valora lo popular sin ridiculizarlo ni convertirlo en mercancía vacía: pregunta quién habla, quién aparece y quién se beneficia del mensaje.",
            "La autonomía se muestra cuando el estudiante revisa su obra con criterios de respeto, claridad, créditos y participación de la comunidad representada.",
        ]
    if family == "safety":
        return [
            f"En {focus}, aprender tecnología implica reconocer riesgos, detener acciones dañinas y buscar ayuda a tiempo.",
            "No toda situación digital se resuelve respondiendo de inmediato. A veces la acción responsable es guardar evidencia, no reenviar y pedir acompañamiento.",
            "La meta no es culpar a la víctima, sino proteger la dignidad, la privacidad y la seguridad de la comunidad."
        ]
    if family == "project":
        return [
            f"En {focus}, el producto final debe mostrar un proceso: problema, evidencia, decisiones, resultado y reflexión.",
            "Un proyecto es más sólido cuando se puede explicar por qué se eligió una solución y cómo se revisó antes de presentarla.",
            "La socialización no es solo exponer: es escuchar preguntas, justificar decisiones y reconocer mejoras posibles."
        ]
    return [
        f"En {focus}, aprender tecnología significa comprender un concepto, aplicarlo en una situación real y revisar sus consecuencias.",
        "La actividad debe conectar saber, hacer y cuidar: no basta con terminar rápido si no se puede explicar el proceso.",
        "Un buen trabajo tecnológico deja evidencia clara de decisiones, pruebas y aprendizajes."
    ]


def worked_examples(guide: GuideData) -> list[tuple[str, str]]:
    family = guide_family(guide)
    if family == "systems_room":
        return [
            ("Ejemplo 1: cable en zona de paso", "Observación: hay un cable atravesado cerca de una silla. Acción autónoma: no halarlo ni pisarlo; avisar al docente y señalar el riesgo para que nadie tropiece."),
            ("Ejemplo 2: cuenta abierta", "Observación: aparece el correo o archivo de otra persona. Acción autónoma: no leer ni modificar nada; levantar la mano y reportar que hay una sesión abierta."),
            ("Ejemplo 3: equipo que no responde", "Observación: el mouse no funciona. Acción autónoma: revisar si está bien ubicado y reportar puesto, elemento y falla; no golpear ni desconectar cables."),
        ]
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                ("Ejemplo 1: transistor", "El transistor permitió dispositivos más pequeños y confiables. Impacto: computadores personales, electrónica de consumo y nuevas industrias."),
                ("Ejemplo 2: Internet", "Internet conectó información y servicios. Impacto: educación en línea y comunicación inmediata, pero también desinformación y brecha de acceso."),
                ("Ejemplo 3: automatización", "Un sistema digital puede controlar inventarios o pagos. Impacto: eficiencia, nuevos empleos técnicos y pérdida de algunas tareas repetitivas."),
            ]
        return [
            ("Ejemplo 1: máquina de vapor", "Aumentó la producción y el transporte, pero también concentró trabajo en fábricas y cambió la vida urbana."),
            ("Ejemplo 2: telégrafo", "Permitió enviar mensajes codificados a distancia y aceleró comercio, prensa, coordinación política y transporte."),
                ("Ejemplo 3: electricidad", "Transformó iluminación, motores y comunicaciones; exigió redes, centrales, cables y nuevas formas de mantenimiento."),
            ]
    if family == "cloud_collaboration":
        return [
            ("Ejemplo 1: permisos", "Para revisar un texto se puede dar permiso de comentario; para coeditar se da permiso de edición solo al equipo."),
            ("Ejemplo 2: historial", "Si se borra un párrafo por error, el historial permite recuperar la versión anterior y dialogar sobre el cambio."),
            ("Ejemplo 3: rol POGIL", "El relator escribe, el verificador revisa fuentes, el coordinador cuida tiempos y el portavoz socializa."),
        ]
    if family == "ewaste":
        return [
            ("Ejemplo 1: clasificación", "Un teclado dañado puede registrarse como periférico; una batería inflada requiere ruta especial por riesgo."),
            ("Ejemplo 2: reducción", "Antes de desechar un mouse, se revisa limpieza, cable, puerto y posibilidad de reparación."),
            ("Ejemplo 3: propuesta", "La auditoría puede terminar en campaña de acopio, inventario institucional o guía de cuidado de equipos."),
        ]
    if family == "ai_book_project":
        return [
            ("Ejemplo 1: prompt editorial", "Prompt débil: 'hazme un libro'. Prompt mejorado: 'Actúa como editor juvenil. Ayúdame a planear un libro breve sobre fútbol barrial para lectores de grado décimo, con 5 capítulos, tono cercano y criterios de originalidad'."),
            ("Ejemplo 2: reescritura con voz propia", "Si la IA propone un párrafo genérico, el estudiante agrega una experiencia, cambia el tono, elimina frases repetidas y explica qué decisión autoral tomó."),
            ("Ejemplo 3: créditos y transparencia", "Una portada con imagen externa debe incluir fuente o licencia. La declaración de IA indica qué partes fueron sugeridas, revisadas y reescritas."),
        ]
    if family == "business_web_presence":
        return [
            ("Ejemplo 1: prompt de marca", "Prompt débil: 'hazme una marca'. Prompt mejorado: 'Actúa como asesor de marca escolar. Propón 5 nombres para un servicio de meriendas saludables, público estudiantes de bachillerato, tono cercano, sin prometer beneficios falsos'."),
            ("Ejemplo 2: arquitectura de landing", "Una página mínima puede incluir inicio, problema, propuesta de valor, servicios, evidencias, preguntas frecuentes, contacto y botón de acción. Cada sección responde una duda del usuario."),
            ("Ejemplo 3: revisión ética", "Si la página usa fotos de personas, se pide consentimiento o se reemplaza por recursos con licencia. Si se usó IA para textos, se declara y se revisa que no haya publicidad engañosa."),
        ]
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                ("Ejemplo 1: venta escolar", "Entrada: pedido de un estudiante. Actividades: registrar producto, revisar disponibilidad, confirmar precio y entregar. Salida: pedido atendido y registro de venta."),
                ("Ejemplo 2: punto crítico", "Si tres personas reciben pedidos por chat y nadie unifica la información, aparecen duplicados, olvidos y respuestas tardías. Esa situación se marca como oportunidad de mejora."),
                ("Ejemplo 3: dato sensible", "Un número de teléfono o dirección no debe ponerse en una cartelera ni compartirse con todo el grupo. Se registra solo si es necesario, con finalidad y acceso limitado."),
                ("Ejemplo 4: uso prudente de IA", "La IA puede proponer categorías para clasificar errores, pero el equipo debe comprobar si esas categorías corresponden al proceso observado."),
            ]
        return [
            ("Ejemplo 1: mapa de pedido", "Entrada: formulario de pedido. Actividad: revisar datos. Responsable: ventas. Herramienta: hoja de cálculo. Salida: pedido confirmado y registro actualizado."),
            ("Ejemplo 2: formulario responsable", "Un formulario de clientes solicita solo datos necesarios, explica finalidad, evita información sensible y usa validación para correo, producto y cantidad."),
            ("Ejemplo 3: IA con revisión humana", "La IA propone respuestas para preguntas frecuentes, pero el equipo verifica precios, tiempos, condiciones y casos que deben pasar a una persona responsable."),
        ]
    if family == "integrated_business_project":
        return [
            ("Ejemplo 1: problema a solución", "Problema: pedidos de un emprendimiento se pierden. Solución: landing con catálogo, formulario, base de datos y confirmación. Evidencia: registros antes/después y satisfacción del usuario."),
            ("Ejemplo 2: dashboard a decisión", "Si el dashboard muestra alta demanda y bajo margen en un producto, el equipo puede recomendar ajustar precio, proveedor o promoción explicando riesgos y límites."),
            ("Ejemplo 3: pitch sustentado", "Un pitch sólido conecta usuario, costo, impacto, datos, automatización, ética y sostenibilidad en lugar de mostrar piezas sueltas sin argumento."),
        ]
    if family == "latex_business_report":
        return [
            ("Ejemplo 1: estructura LaTeX", "Un informe mínimo incluye título, autor, resumen, secciones, tabla de resultados, análisis y fuentes. Si falta cerrar un entorno, el PDF no compila."),
            ("Ejemplo 2: tabla comercial", "Una tabla de precios compara producto, costo, precio de venta y margen. El análisis explica qué producto conviene revisar y por qué."),
            ("Ejemplo 3: gráfico con lectura", "No basta insertar una imagen del gráfico. Debe tener título, fuente y un comentario que conecte el patrón con una decisión comercial."),
        ]
    if family == "business_dashboard_ai":
        return [
            ("Ejemplo 1: limpieza antes de decidir", "Si un producto aparece como 'cuaderno', 'Cuaderno' y 'CUADERNO', primero se unifica la categoría antes de calcular ventas."),
            ("Ejemplo 2: indicador interpretable", "El margen se calcula como utilidad sobre ingreso. Un margen alto puede orientar promoción, pero se revisa volumen y costo antes de decidir."),
            ("Ejemplo 3: IA verificada", "La IA sugiere que suban precios, pero el equipo revisa el gráfico, identifica baja demanda y propone primero ajustar inventario y promoción."),
        ]
    if family == "ai":
        return [
            ("Ejemplo 1: prompt mejorado", "Prompt débil: 'Explícame la IA'. Prompt mejorado: 'Actúa como docente de grado noveno. Explica qué es la IA con tres ejemplos cotidianos, un riesgo y una pregunta de verificación'."),
            ("Ejemplo 2: respuesta verificada", "Si la IA afirma un dato histórico o técnico, el estudiante busca una fuente confiable, compara la información y marca qué parte confirma, corrige o descarta."),
            ("Ejemplo 3: privacidad", "En lugar de escribir nombres reales de compañeros, se usan datos ficticios o descripciones generales para evitar exponer información personal."),
        ]
    if family == "data_analysis":
        return [
            ("Ejemplo 1: dato limpio", "Si una columna registra 'sí', 'SI' y 'si', primero se unifica la categoría antes de contar respuestas."),
            ("Ejemplo 2: gráfico adecuado", "Para comparar categorías se usan barras; para ver cambios por fechas se puede usar línea. Elegir mal el gráfico cambia la lectura."),
            ("Ejemplo 3: conclusión prudente", "Si solo respondieron 8 estudiantes, la conclusión puede decir 'en esta muestra' y no 'todos los estudiantes del colegio'."),
        ]
    if family == "spreadsheet":
        return [
            ("Ejemplo 1: fórmula revisable", "Si B2 contiene la cantidad y C2 el precio, la fórmula =B2*C2 calcula el subtotal. Para comprobarla, cambia B2 y verifica que el subtotal también cambie."),
            ("Ejemplo 2: conclusión con datos", "Si el promedio de asistencia baja en una semana, no basta decir 'fue mala'. Hay que revisar datos, comparar días y explicar una causa posible.")
        ]
    if family == "physical_computing":
        return [
            ("Ejemplo 1: umbral de luz", "Si nivel_luz < 40, encender LED. Se prueba con 39, 40 y 41 para revisar si el límite está bien definido."),
            ("Ejemplo 2: alerta de temperatura", "Si temperatura > 30, mostrar 'alerta'; si no, mostrar 'normal'. Luego se prueba con 29, 30 y 31."),
            ("Ejemplo 3: depuración", "Si el buzzer suena siempre, se revisa si la condición está invertida, si el umbral es demasiado bajo o si el evento se repite sin control."),
        ]
    if family == "scratch_programming":
        return [
            ("Ejemplo 1: predecir", "Antes de ejecutar, el estudiante lee: 'al presionar bandera, repetir 10, mover 5'. Predice que el sprite avanzará 50 pasos."),
            ("Ejemplo 2: modificar", "Si se cambia repetir 10 por repetir 20, se registra qué cambia y si afecta el objetivo del programa."),
            ("Ejemplo 3: depurar", "Si el personaje no salta, se revisa si el evento correcto está conectado y si el bloque de movimiento está dentro de la secuencia."),
        ]
    if family == "programming":
        return [
            ("Ejemplo 1: condición simple", "La frase 'la nota es 3.0 o más' se escribe nota >= 3.0. Si nota vale 2.9, la condición es falsa; si vale 3.0, es verdadera."),
            ("Ejemplo 2: condiciones combinadas", "La frase 'la edad está entre 12 y 17' se escribe edad >= 12 AND edad <= 17. Las dos partes deben ser verdaderas."),
            ("Ejemplo 3: orden de condiciones", "Para clasificar temperatura: si temperatura < 18, 'frío'; else if temperatura <= 28, 'templado'; else, 'calor'. Si se pone primero temperatura <= 28, el caso 'frío' queda mal clasificado.")
        ]
    if family == "presentation":
        return [
            ("Ejemplo 1: una idea por diapositiva", "Una diapositiva dice 'Problema: desperdicio de agua' y muestra un dato. La explicación oral amplía; la diapositiva no se llena de párrafos."),
            ("Ejemplo 2: hipervínculo útil", "Un botón 'Ver evidencia' lleva a una encuesta o video. El enlace se prueba antes de presentar para evitar interrupciones.")
        ]
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                ("Ejemplo 1: mensaje central", "Tema amplio: cuidado de la sala. Mensaje central: 'Entrar con alimentos puede dañar equipos y afectar a todos'. Esa frase ayuda a elegir imagen, título y llamado a la acción."),
                ("Ejemplo 2: jerarquía visual", "Si título, subtítulo y párrafo tienen el mismo tamaño, el lector no sabe por dónde empezar. Se corrige diferenciando niveles: título principal, subtítulo, cuerpo breve y fuente."),
                ("Ejemplo 3: imagen responsable", "Una foto de Internet no se pega sin revisar. Se busca autor o licencia, se cita la fuente y se verifica que la imagen aporte información al mensaje."),
                ("Ejemplo 4: revisión", "Antes de entregar, el editor revisa cinco puntos: propósito, orden de lectura, contraste, ortografía y créditos. Luego anota qué cambió y por qué."),
            ]
        return [
            ("Ejemplo 1: jerarquía visual", "Si todo el texto tiene el mismo tamaño, la audiencia no sabe por dónde empezar. Se corrige usando título, subtítulo, cuerpo y pie de fuente."),
            ("Ejemplo 2: imagen citada", "Una imagen tomada de Internet no se pega sin más. Se revisa licencia, se cita la fuente y se explica por qué ayuda al mensaje."),
            ("Ejemplo 3: retícula simple", "Un folleto de dos columnas se lee mejor cuando mantiene márgenes, alineación y espacios constantes entre bloques."),
        ]
    if family == "network":
        return [
            ("Ejemplo 1: emisor, canal y receptor", "Cuando un celular envía un archivo por Bluetooth, el celular origen es emisor, el destino es receptor y el canal es la conexión inalámbrica cercana."),
            ("Ejemplo 2: IP coherente", "Si dos equipos están en la misma red local, sus direcciones deben pertenecer al mismo rango, por ejemplo 192.168.1.10 y 192.168.1.11."),
            ("Ejemplo 3: prueba antes de concluir", "No basta dibujar una red. Se hace ping, se revisa respuesta o se explica qué evidencia confirma la comunicación."),
        ]
    if family == "info_project":
        return [
            ("Ejemplo 1: publicación con red", "Un boletín sobre cuidado de datos incluye diseño editorial claro y un diagrama que muestra por qué canal se comparte y quién puede acceder."),
            ("Ejemplo 2: prueba técnica", "Si el proyecto usa una red simulada, el equipo registra IP, servicio y prueba de conectividad antes de afirmar que la comunicación funciona."),
            ("Ejemplo 3: mejora ética", "Si una imagen muestra estudiantes identificables, el equipo pide permiso, anonimiza o reemplaza el recurso y explica la decisión."),
        ]
    if family == "media":
        return [
            ("Ejemplo 1: imagen con propósito", "Recortar una foto para destacar el objeto principal mejora la lectura; saturarla sin razón puede ocultar información."),
            ("Ejemplo 2: audio comprensible", "Reducir ruido de fondo y ajustar volumen ayuda a que la voz se entienda, pero cortar frases cambia el sentido del mensaje.")
        ]
    if family == "cultural_media":
        return [
            ("Ejemplo 1: crédito visible", "Si se usa una foto de una fiesta local, se cita autor, lugar y permiso de uso en la pieza o en el portafolio."),
            ("Ejemplo 2: evitar estereotipo", "En lugar de mostrar una comunidad con una sola imagen repetida, se explican varias voces, oficios o prácticas."),
            ("Ejemplo 3: edición cuidadosa", "Si se limpia un audio de entrevista, no se cortan frases que cambien el sentido de lo dicho."),
        ]
    if family == "safety":
        return [
            ("Ejemplo 1: no reenviar daño", "Si llega una captura ofensiva, compartirla aumenta el daño. La acción correcta es guardar evidencia y buscar apoyo adulto."),
            ("Ejemplo 2: señal de alerta", "Si alguien insiste en pedir fotos privadas o secretos, se detiene la conversación, se bloquea/reporta y se avisa a una persona de confianza.")
        ]
    if family == "project":
        return [
            ("Ejemplo 1: evidencia antes de propuesta", "Antes de diseñar una campaña, el equipo muestra datos: encuesta, observación o registro del problema."),
            ("Ejemplo 2: mejora por retroalimentación", "Si la audiencia no entiende la solución, el equipo ajusta una diapositiva, agrega evidencia y ensaya una explicación más clara.")
        ]
    return [
        ("Ejemplo 1: decisión responsable", "Si una herramienta permite hacer algo rápido, el estudiante revisa si la información es correcta, segura y respetuosa."),
        ("Ejemplo 2: evidencia clara", "Un trabajo mejora cuando incluye proceso, resultado y explicación de las decisiones tomadas.")
    ]


def common_error(guide: GuideData) -> str:
    family = guide_family(guide)
    if family == "systems_room":
        return "Error común: mover cables, cambiar configuraciones o usar una cuenta abierta antes de observar. Corrección: detenerse, revisar la lista de cotejo, registrar el problema y reportar sin manipular elementos de riesgo."
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return "Error común: presentar la revolución digital como progreso automático. Corrección: analizar beneficios, riesgos, brecha digital e impacto social con ejemplos."
        return "Error común: listar inventos sin explicar cambios sociales. Corrección: relacionar cada avance con producción, comunicación, trabajo, ciudad e infraestructura."
    if family == "cloud_collaboration":
        return "Error común: compartir enlaces abiertos y editar sin roles. Corrección: definir permisos, roles, comentarios, historial y acuerdos de datos antes de trabajar."
    if family == "ewaste":
        return "Error común: tratar residuos tecnológicos como basura común. Corrección: clasificar, identificar riesgos, revisar reutilización y proponer ruta RAEE responsable."
    if family == "ai_book_project":
        return "Error común: copiar capítulos generados por IA y presentarlos como propios. Corrección: conservar bitácora de prompts, verificar, reescribir con voz propia, citar recursos y declarar el uso de IA."
    if family == "business_web_presence":
        return "Error común: publicar una marca con textos genéricos, imágenes sin permiso y promesas poco verificables. Corrección: definir público, propuesta de valor, arquitectura, créditos, accesibilidad y declaración de IA antes de publicar."
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return "Error común: elegir una herramienta digital antes de entender el proceso. Corrección: observar, inventariar, identificar entrada, responsable, salida, punto crítico, datos sensibles y oportunidad de mejora antes de pensar en IA o automatización."
        return "Error común: automatizar una tarea sin mapear el proceso ni proteger datos. Corrección: identificar entrada, actividad, responsable, salida, consentimiento, prueba y revisión humana."
    if family == "integrated_business_project":
        return "Error común: presentar sitio, datos, dashboard y pitch como piezas aisladas. Corrección: conectar problema, usuario, solución, indicadores, automatización, ética y decisión empresarial en una sola narrativa."
    if family == "latex_business_report":
        return "Error común: concentrarse en que el PDF se vea bonito sin revisar datos, fuentes ni estructura. Corrección: compilar, leer, citar, interpretar tablas/gráficos y verificar conclusiones."
    if family == "business_dashboard_ai":
        return "Error común: aceptar una recomendación de IA sin revisar la base de datos. Corrección: limpiar registros, comprobar fórmulas, elegir visualización adecuada y sustentar la decisión con indicadores."
    if family == "ai":
        return "Error común: copiar la primera respuesta de la IA como si fuera propia y verdadera. Corrección: declarar el uso de IA, verificar fuentes, corregir errores y reescribir con voz propia."
    if family == "data_analysis":
        return "Error común: sacar conclusiones fuertes con datos desordenados o insuficientes. Corrección: limpiar la tabla, revisar muestra, elegir gráfico adecuado y escribir límites de la conclusión."
    if family == "spreadsheet":
        return "Error común: copiar fórmulas sin revisar referencias. Corrección: verificar qué celdas cambian, cuáles deben quedar fijas y probar con datos simples."
    if family == "physical_computing":
        return "Error común: programar la alerta sin probar valores límite. Corrección: registrar sensor, umbral, condición, actuador y tres pruebas antes de presentar el sistema."
    if family == "scratch_programming":
        return "Error común: modificar bloques al azar hasta que algo funcione. Corrección: predecir, ejecutar, investigar, cambiar una cosa por vez y documentar la depuración."
    if family == "programming":
        return "Error común: poner primero una condición general y después una específica. Corrección: ordenar las condiciones, probar valores límite y explicar qué rama se ejecuta."
    if family == "presentation":
        return "Error común: llenar diapositivas con texto y efectos. Corrección: dejar una idea principal, evidencia breve y diseño consistente."
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return "Error común: empezar a diseñar agregando colores, imágenes y texto sin saber para quién se comunica. Corrección: definir audiencia, propósito y mensaje central; luego bocetar una retícula, ubicar jerarquía visual, citar recursos y revisar legibilidad antes de entregar."
        return "Error común: llenar la página con texto e imágenes sin orden. Corrección: usar retícula, jerarquía visual, contraste y fuentes citadas antes de publicar."
    if family == "network":
        return "Error común: afirmar que una red funciona solo porque el dibujo se ve correcto. Corrección: registrar configuración, prueba de conectividad y explicación del resultado."
    if family == "info_project":
        return "Error común: presentar un producto bonito sin evidencias técnicas ni revisión ética. Corrección: anexar boceto, fuentes, diagrama, prueba, mejora y reflexión sobre datos/autoría."
    if family == "media":
        return "Error común: editar por estética sin cuidar el mensaje. Corrección: comparar antes/después y justificar cada cambio por claridad, calidad o ética."
    if family == "cultural_media":
        return "Error común: usar símbolos culturales como adorno sin contexto ni permiso. Corrección: investigar, pedir consentimiento, citar autorías y revisar posibles estereotipos."
    if family == "safety":
        return "Error común: reenviar capturas o responder con agresión. Corrección: no amplificar el daño, guardar evidencia y activar una ruta de ayuda."
    if family == "project":
        return "Error común: presentar una solución sin evidencia. Corrección: mostrar el problema, explicar datos y relacionar cada decisión con una necesidad real."
    return "Error común: hacer la actividad sin explicar el proceso. Corrección: registrar decisiones, evidencia y aprendizaje obtenido."


def guided_practice_items(guide: GuideData) -> list[str]:
    family = guide_family(guide)
    if family in FAMILY_GUIDED_PRACTICE:
        return FAMILY_GUIDED_PRACTICE[family]
    if family == "systems_room":
        return [
            "Lee la pregunta guía y subraya dos palabras clave: autonomía y cuidado.",
            "Observa tu estación de trabajo sin mover cables ni cambiar configuraciones.",
            "Marca en tu cuaderno o guía el estado de monitor, teclado, mouse, silla, cables y espacio de paso.",
            "Identifica dos riesgos físicos y un riesgo digital posible en la sala.",
            "Clasifica cada acción: puedo hacerla solo, debo pedir ayuda o debo evitarla.",
            "Escribe tres normas positivas para que el grupo aprenda mejor en la sala.",
        ]
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                "Ubica cinco hitos entre electrónica, computación, Internet y revolución digital.",
                "Para cada hito escribe qué problema resolvió o qué posibilidad abrió.",
                "Agrega un impacto social: trabajo, comunicación, educación, comercio o cultura.",
                "Identifica un riesgo o tensión: brecha digital, vigilancia, dependencia o desinformación.",
                "Ordena los hitos en una línea de tiempo con fecha aproximada y explicación breve.",
                "Cierra con una conclusión sobre cómo decidir responsablemente en la era digital.",
            ]
        return [
            "Ubica cinco hitos entre Revolución Industrial, telégrafo y electricidad.",
            "Para cada hito escribe qué necesidad técnica o social atendió.",
            "Agrega un cambio en producción, comunicación, transporte, ciudad o trabajo.",
            "Identifica quiénes se beneficiaron y quiénes pudieron quedar excluidos o afectados.",
            "Ordena los hitos en una línea de tiempo con fecha aproximada y explicación breve.",
                "Cierra con una reflexión sobre tecnología, progreso y responsabilidad social.",
            ]
    if family == "cloud_collaboration":
        return [
            "Define el producto colaborativo y la necesidad que atenderá el equipo.",
            "Asigna roles claros: coordinación, redacción, verificación, edición y socialización.",
            "Configura permisos adecuados para ver, comentar o editar.",
            "Usa comentarios para pedir cambios sin borrar aportes de otros.",
            "Revisa historial o versiones para reconocer aportes y corregir errores.",
            "Cierra con un acuerdo sobre cuidado de datos, enlaces y publicación.",
        ]
    if family == "ewaste":
        return [
            "Identifica el dispositivo o residuo tecnológico observado.",
            "Clasifica su estado: funciona, falla, está obsoleto, requiere reparación o debe gestionarse como RAEE.",
            "Registra riesgo ambiental, dato faltante o cuidado necesario.",
            "Aplica Rethink, Refuse o Reduce para proponer una acción antes de desechar.",
            "Organiza evidencia: foto, ficha, conteo, ubicación o responsable.",
            "Formula una propuesta comunitaria verificable y respetuosa.",
        ]
    if family == "business_web_presence":
        return [
            "Define producto, servicio o marca personal que vas a comunicar.",
            "Describe público objetivo, necesidad principal y propuesta de valor en una frase clara.",
            "Diseña la arquitectura: secciones, orden de lectura y llamado a la acción.",
            "Crea o revisa textos comerciales con IA y reescríbelos con voz propia.",
            "Elige criterios visuales: color, tipografía, imagen, contraste y accesibilidad.",
            "Registra fuentes, permisos, derechos de imagen y declaración de uso de IA.",
        ]
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                "Selecciona un emprendimiento, negocio escolar o caso simulado que puedas observar sin exponer datos personales.",
                "Identifica al menos tres procesos: ventas, pedidos, inventario, atención, encuesta, pago, entrega o seguimiento.",
                "Para cada proceso escribe entrada, actividades principales, responsable, herramienta actual y salida verificable.",
                "Marca un punto crítico: demora, error, duplicado, dato faltante, pérdida de evidencia o falta de responsable.",
                "Reconoce qué datos son sensibles y cómo deberían protegerse antes de usar herramientas digitales o IA.",
                "Elige el proceso con mayor necesidad de mejora usando tres criterios: impacto, facilidad y riesgo.",
                "Propón una oportunidad digital inicial sin desarrollarla todavía: formulario, hoja, plantilla, IA, tablero o flujo futuro.",
                "Cierra con una pregunta de investigación que oriente la siguiente guía de mapeo de procesos.",
            ]
        return [
            "Selecciona un proceso real o simulado: ventas, pedidos, inventario, encuesta, atención o seguimiento.",
            "Mapea entrada, actividad, responsable, herramienta y salida.",
            "Detecta tareas repetitivas, errores, tiempos muertos o datos faltantes.",
            "Diseña una mejora con formulario, hoja de cálculo, plantilla, filtro, respuesta o flujo básico.",
            "Incluye consentimiento, finalidad de datos, acceso responsable y revisión humana.",
            "Aplica PDCA: planear, probar, verificar evidencia y ajustar el proceso.",
        ]
    if family == "integrated_business_project":
        return [
            "Formula el problema empresarial, comercial o contable en una oración verificable.",
            "Identifica usuario, cliente o comunidad beneficiaria y una evidencia de su necesidad.",
            "Diseña la solución con componentes: presencia digital, datos, dashboard y mejora de proceso.",
            "Define indicadores, fuente de datos y decisión que el proyecto debe sustentar.",
            "Integra automatización, documentación, declaración de IA y protección de datos.",
            "Prepara el pitch conectando problema, solución, evidencia, impacto, costo, riesgo y sostenibilidad.",
        ]
    if family == "ai_book_project":
        return [
            "Define tema libre, lector, propósito, género y tono del libro.",
            "Formula o mejora un prompt con rol, contexto, tarea, formato, extensión y criterios.",
            "Compara una respuesta de IA con tu intención autoral y marca qué sirve, qué sobra y qué falta.",
            "Reescribe un fragmento con voz propia, ejemplos personales y coherencia con el índice.",
            "Registra fuentes, imágenes, licencias y declaración de uso de IA.",
            "Prepara una evidencia editorial: capítulo, escaleta, créditos, portada o bitácora según la sesión.",
        ]
    if family == "latex_business_report":
        return [
            "Define el problema comercial o contable y la pregunta que orienta el informe.",
            "Escribe la estructura del documento: título, autor, secciones, tabla, gráfico, análisis y fuentes.",
            "Codifica o revisa el fragmento LaTeX necesario para la sesión.",
            "Organiza datos en tabla o visualización con título, unidad y fuente.",
            "Redacta una interpretación prudente con hallazgo, evidencia, límite y posible decisión.",
            "Compila o revisa el PDF y documenta un error corregido o una mejora editorial.",
        ]
    if family == "business_dashboard_ai":
        return [
            "Define la pregunta empresarial que debe responder el tablero o reporte.",
            "Revisa campos, registros, duplicados, vacíos, formatos y datos sensibles.",
            "Calcula o identifica indicadores: ventas, costos, utilidad, margen, frecuencia o tendencia.",
            "Elige una visualización coherente y escribe qué patrón permite leer.",
            "Usa IA para proponer limpieza, interpretación o narrativa, y verifica cada sugerencia con los datos.",
            "Formula una recomendación con evidencia, riesgo, límite y decisión sustentable.",
        ]
    if family == "ai":
        return [
            "Lee el propósito de la guía y escribe qué tarea concreta quieres resolver con IA.",
            "Formula un prompt con rol, contexto, tarea, formato y criterio de calidad.",
            "Revisa la respuesta: subraya una idea útil, una idea dudosa y una mejora necesaria.",
            "Contrasta al menos un dato con una fuente confiable o con conocimiento de clase.",
            "Reescribe el resultado con tus palabras y declara cómo usaste la IA.",
            "Cierra con una decisión ética: qué dato no debes compartir, qué sesgo debes vigilar o qué límite debes reconocer.",
        ]
    if family == "data_analysis":
        return [
            "Formula una pregunta que pueda responderse con datos observables.",
            "Define campos, registros, unidades y tipo de dato antes de llenar la tabla.",
            "Revisa datos repetidos, vacíos o escritos de varias formas.",
            "Aplica un cálculo o gráfico que responda la pregunta inicial.",
            "Escribe una conclusión prudente con evidencia y límite de alcance.",
            "Explica qué decisión responsable se puede tomar y qué información falta.",
        ]
    if family == "spreadsheet":
        return [
            "Identifica tres columnas necesarias para registrar datos del caso.",
            "Escribe una fórmula o función que pueda revisarse con datos de prueba.",
            "Explica qué dato cambiarías para comprobar si el resultado es correcto."
        ]
    if family == "physical_computing":
        return [
            "Identifica la entrada del sistema: sensor, botón o lectura simulada.",
            "Define el umbral o condición que activa una respuesta.",
            "Escribe el pseudocódigo o bloque lógico con if, else-if o else.",
            "Indica el actuador o salida: LED, sonido, pantalla, mensaje o motor.",
            "Prueba al menos tres valores: esperado, contrario y límite.",
            "Registra un error encontrado y la mejora aplicada.",
        ]
    if family == "scratch_programming":
        return [
            "Lee el programa o idea y predice qué debería ocurrir.",
            "Ejecuta o simula la secuencia y compara con tu predicción.",
            "Identifica eventos, variables, bucles, condicionales o mensajes usados.",
            "Modifica un parámetro a la vez y registra qué cambió.",
            "Prueba un caso esperado, un caso contrario y un caso límite cuando aplique.",
            "Crea o ajusta una escena interactiva y documenta un error corregido.",
        ]
    if family == "programming":
        return [
            "Elige una variable del caso, por ejemplo nota, edad, temperatura, luz o permiso.",
            "Traduce una frase del caso usando un símbolo lógico: >, <, >=, <=, == o !=.",
            "Combina dos condiciones con AND u OR cuando necesites verificar más de una cosa.",
            "Prueba un caso esperado, un caso contrario y un caso límite antes de escribir el algoritmo final."
        ]
    if family == "presentation":
        return [
            "Define el mensaje central en una sola frase.",
            "Planea tres diapositivas: problema, evidencia y propuesta.",
            "Decide qué imagen, dato o enlace ayuda realmente a comprender."
        ]
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                "Elige un tema libre pero útil para la comunidad escolar: convivencia, ambiente, tecnología, salud, cultura, deporte o emprendimiento.",
                "Define en una frase la audiencia, el propósito y el mensaje central del folleto o boletín.",
                "Observa dos publicaciones reales y anota una decisión de diseño que ayuda y una que confunde.",
                "Dibuja una retícula inicial con márgenes, dos o tres columnas y zonas para título, texto, imagen, dato y créditos.",
                "Redacta un título, un subtítulo y tres bloques breves de texto; cada bloque debe tener una sola idea.",
                "Selecciona una imagen con propósito, registra fuente, autor o licencia y explica por qué aporta al mensaje.",
                "Revisa contraste, tamaño de letra, alineación, ortografía y cantidad de texto por bloque.",
                "Compara versión 1 y versión 2; escribe qué cambiaste y cómo ese cambio mejora la lectura.",
            ]
        return [
            "Define audiencia, propósito y mensaje central de la publicación.",
            "Dibuja una retícula simple con márgenes, columnas y zonas de imagen.",
            "Ubica título, subtítulo, cuerpo de texto, imagen y fuente de información.",
            "Revisa contraste, tamaño de letra y cantidad de texto por bloque.",
            "Cita las fuentes usadas y verifica permisos o licencias.",
            "Ajusta una versión final explicando qué cambio mejoró la lectura.",
        ]
    if family == "network":
        return [
            "Identifica los nodos de la red y la función de cada uno.",
            "Dibuja o simula el canal de comunicación entre emisor y receptor.",
            "Registra dirección, protocolo o servicio usado según el caso.",
            "Realiza una prueba de conectividad o explica cómo se verificaría.",
            "Anota un error posible y una corrección responsable.",
            "Relaciona la red con cuidado de datos, accesos o convivencia digital.",
        ]
    if family == "info_project":
        return [
            "Define la necesidad escolar o comunitaria que el proyecto quiere comunicar.",
            "Elige una audiencia y redacta el mensaje central de la publicación.",
            "Diseña un boceto con retícula, jerarquía visual, fuentes e imágenes citadas.",
            "Construye un esquema o simulación de red que explique circulación, acceso o servicio.",
            "Registra una prueba técnica y una mejora aplicada después de revisar errores.",
            "Prepara una sustentación con decisiones de diseño, decisiones técnicas, límites y reflexión ética.",
        ]
    if family == "media":
        return [
            "Describe qué comunica el recurso original.",
            "Elige dos mejoras técnicas y justifica su propósito.",
            "Compara antes/después y verifica que el sentido no cambie."
        ]
    if family == "cultural_media":
        return [
            "Elige una expresión de cultura popular que conozcas o puedas investigar con respeto.",
            "Define qué belleza, memoria o problema quieres comunicar.",
            "Identifica fuentes, personas o recursos y registra permisos o créditos necesarios.",
            "Crea un boceto visual o sonoro que evite estereotipos y cuide la dignidad.",
            "Revisa la pieza con criterios de claridad, autoría, consentimiento y representación.",
            "Escribe una reflexión sobre qué aprendiste de la cultura representada y qué límites debes reconocer.",
        ]
    if family == "safety":
        return [
            "Reconoce la señal de alerta del caso.",
            "Define qué evidencia se debe conservar sin exponer a la persona.",
            "Propón una ruta de ayuda concreta y respetuosa."
        ]
    if family == "project":
        return [
            "Formula el problema en una oración verificable.",
            "Elige una evidencia que demuestre la necesidad.",
            "Propón una mejora y explica cómo la evaluarías."
        ]
    return [
        "Reconoce el concepto central de la guía.",
        "Describe una situación real donde se aplique.",
        "Explica una decisión responsable relacionada con el tema."
    ]


def product_criteria(guide: GuideData) -> list[str]:
    family = guide_family(guide)
    if family in FAMILY_PRODUCT_CRITERIA:
        return FAMILY_PRODUCT_CRITERIA[family]
    if family == "systems_room":
        return [
            "Ficha con fecha, puesto observado y estado de los elementos principales",
            "Mapa o esquema de la sala con al menos cuatro objetos o zonas señaladas",
            "Tres riesgos identificados con acción segura: prevenir, reportar o evitar",
            "Cinco normas de uso escritas en forma positiva y comprensible",
            "Compromiso personal verificable para la próxima clase",
        ]
    if family == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                "Cinco hitos de la era electrónica y digital ordenados cronológicamente",
                "Explicación de impacto social para cada hito",
                "Identificación de beneficios, riesgos y brecha digital",
                "Conclusión crítica sobre decisiones responsables en la cultura digital",
            ]
        return [
            "Cinco hitos desde la Revolución Industrial hasta telégrafo/electricidad",
            "Relación entre invento, necesidad, infraestructura y cambio social",
            "Identificación de beneficiados, afectados o excluidos",
            "Reflexión crítica sobre progreso tecnológico y responsabilidad social",
        ]
    if family == "cloud_collaboration":
        return [
            "Producto colaborativo con propósito y audiencia definidos",
            "Roles del equipo registrados y cumplidos",
            "Permisos, comentarios o historial usados responsablemente",
            "Evidencia de aportes, revisión y acuerdos de datos",
            "Reflexión sobre liderazgo servicial y cuidado de la información",
        ]
    if family == "ewaste":
        return [
            "Inventario o ficha RAEE con dispositivo, estado y ubicación",
            "Clasificación de riesgos y posibilidades de reutilización o reparación",
            "Aplicación de Rethink, Refuse o Reduce",
            "Propuesta comunitaria con responsables, acciones y evidencia",
            "Reflexión sobre impacto ambiental y consumo tecnológico",
        ]
    if family == "ai_book_project":
        return [
            "Tema libre, lector, propósito, género y tono definidos",
            "Prompts registrados con rol, contexto, tarea, formato y criterios",
            "Borrador o fragmento revisado y reescrito con voz propia",
            "Decisiones de edición, estructura o diseño justificadas",
            "Fuentes, créditos, derechos de autor y declaración de uso de IA visibles",
        ]
    if family == "business_web_presence":
        return [
            "Propósito, público objetivo y propuesta de valor definidos",
            "Arquitectura de página o presencia digital con secciones y llamado a la acción",
            "Textos comerciales revisados con IA, voz propia y verificación humana",
            "Decisiones visuales con coherencia de marca, contraste y accesibilidad",
            "Créditos, permisos, derechos de imagen y declaración de IA documentados",
        ]
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                "Inventario de mínimo tres procesos del caso o emprendimiento observado",
                "Cada proceso incluye entrada, actividades, responsable, herramienta actual y salida",
                "Identificación de punto crítico: error, demora, duplicado, dato faltante o riesgo",
                "Reconocimiento de datos sensibles, finalidad y cuidado de acceso",
                "Priorización de un proceso con criterios de impacto, facilidad y riesgo",
                "Oportunidad digital o de IA formulada como mejora futura verificable",
            ]
        return [
            "Proceso seleccionado con entrada, actividades, responsables, herramientas y salidas",
            "Diagnóstico de errores, tiempos muertos, tareas repetitivas o datos faltantes",
            "Formulario, hoja, flujo, plantilla o respuesta automatizada con evidencia de prueba",
            "Protocolo de consentimiento, privacidad y revisión humana",
            "Mejora explicada con PDCA, indicadores o comparación antes/después",
        ]
    if family == "integrated_business_project":
        return [
            "Problema, usuario, alcance y criterios de éxito claramente formulados",
            "Solución integrada con presencia digital, datos, dashboard y mejora de proceso",
            "Indicadores, visualizaciones y recomendación sustentados con evidencia",
            "Automatización, documentación técnica-comercial y declaración de IA incluidas",
            "Pitch final con impacto, costos, riesgos, ética, sostenibilidad y mejora posible",
        ]
    if family == "latex_business_report":
        return [
            "Problema comercial o contable y pregunta de informe definidos",
            "Estructura LaTeX codificada con secciones, tabla, gráfico o anexos según la sesión",
            "Datos organizados con fuente, unidad y explicación",
            "Análisis con hallazgo, evidencia, límite y recomendación",
            "PDF compilado o revisado con citación y declaración de uso de IA",
        ]
    if family == "business_dashboard_ai":
        return [
            "Base de datos empresarial con campos, registros y formatos consistentes",
            "Limpieza documentada de vacíos, duplicados, categorías o errores",
            "Indicadores calculados e interpretados",
            "Dashboard o reporte con visualizaciones pertinentes y narrativa clara",
            "Decisión sustentada con evidencia, riesgos, límites y verificación de IA",
        ]
    if family == "ai":
        return [
            "Propósito claro del uso de IA",
            "Prompt registrado con contexto, tarea, formato y criterios",
            "Respuesta revisada con al menos una corrección o mejora humana",
            "Verificación de información con fuente o evidencia",
            "Reflexión ética sobre privacidad, sesgos, autoría o límites de la IA",
        ]
    if family == "data_analysis":
        return [
            "Pregunta de análisis, campos y registros definidos",
            "Datos revisados, limpios y organizados en tabla",
            "Cálculo o gráfico coherente con la pregunta",
            "Conclusión prudente con evidencia y límites",
            "Decisión o recomendación responsable basada en datos",
        ]
    if family == "spreadsheet":
        return ["Datos organizados con encabezados claros", "Fórmula o función visible y comprobada", "Conclusión basada en los resultados", "Explicación breve del procedimiento"]
    if family == "physical_computing":
        return ["Entrada o sensor identificado", "Condición, umbral y respuesta programada", "Actuador o salida explicada", "Tres pruebas documentadas", "Mejora y reflexión sobre impacto"]
    if family == "scratch_programming":
        return ["Predicción y ejecución registradas", "Eventos, variables, bucles o condicionales identificados", "Modificación controlada y prueba documentada", "Error depurado con explicación", "Proyecto interactivo o análisis PRIMM socializable"]
    if family == "programming":
        return ["Algoritmo o diagrama con condiciones ordenadas", "Uso correcto de if, else-if, else u operadores", "Tres casos de prueba con resultado esperado", "Reflexión sobre la decisión automatizada"]
    if family == "presentation":
        return ["Mensaje central visible", "5 a 7 diapositivas con jerarquía visual", "Evidencia o fuente revisable", "Transiciones o enlaces probados"]
    if family == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                "Tema libre con audiencia, propósito y mensaje central claramente escritos",
                "Boceto con retícula visible: márgenes, columnas, título, texto, imagen, dato y créditos",
                "Jerarquía visual aplicada: título, subtítulo, cuerpo breve y llamado a la acción",
                "Imagen o recurso visual con fuente, autor o licencia registrada",
                "Revisión de legibilidad: contraste, tamaño de letra, alineación y ortografía",
                "Justificación final de dos decisiones de diseño y una mejora realizada",
            ]
        return ["Audiencia y propósito definidos", "Retícula y jerarquía visual aplicadas", "Imágenes y fuentes citadas", "Revisión de legibilidad, contraste y accesibilidad"]
    if family == "network":
        return ["Nodos, canal y protocolo identificados", "Direcciones o configuración registradas", "Prueba de conectividad documentada", "Reflexión sobre seguridad y cuidado de datos"]
    if family == "info_project":
        return ["Necesidad, audiencia y mensaje central definidos", "Publicación con retícula, jerarquía y fuentes citadas", "Esquema o simulación de red con prueba documentada", "Sustentación de decisiones, límites, ética y mejora"]
    if family == "media":
        return ["Archivo original y versión final", "Dos mejoras técnicas justificadas", "Respeto por autoría y privacidad", "Exportación en formato adecuado"]
    if family == "cultural_media":
        return ["Expresión cultural investigada con respeto", "Intención estética y mensaje definidos", "Créditos, permisos o fuentes registrados", "Revisión de estereotipos, dignidad y representación", "Reflexión crítica sobre impacto cultural"]
    if family == "safety":
        return ["Definición clara del riesgo", "Señales de alerta", "Ruta de ayuda concreta", "Compromiso de no amplificar el daño"]
    if family == "project":
        return ["Problema y evidencia", "Propuesta o producto", "Criterios de evaluación", "Reflexión ética y mejoras"]
    return ["Concepto explicado", "Evidencia de proceso", "Producto revisable", "Conclusión personal"]


def evaluation_rows(guide: GuideData) -> list[tuple[str, str, str, str]]:
    family = guide_family(guide)
    if family in FAMILY_EVALUATION_ROWS:
        return FAMILY_EVALUATION_ROWS[family]
    if guide_family(guide) == "systems_room":
        return [
            ("Lectura autónoma", "Sigo la guía, marco avances y pregunto solo después de observar.", "Sigo algunas instrucciones, pero necesito apoyo frecuente.", "Espero instrucciones para iniciar y no uso la guía como apoyo."),
            ("Diagnóstico de la sala", "Identifico puesto, elementos, riesgos físicos y riesgos digitales con ejemplos concretos.", "Reconozco algunos elementos, pero mis riesgos son incompletos o poco claros.", "No logro diferenciar elementos, riesgos o acciones seguras."),
            ("Cuidado y convivencia", "Propongo normas positivas, respeto turnos y explico cómo mis acciones afectan al grupo.", "Reconozco normas, pero me falta aplicarlas o justificarlas.", "No relaciono el uso de la sala con cuidado del espacio y de los demás."),
            ("Reporte y compromiso", "Reporto fallas sin manipular riesgos y formulo un compromiso verificable.", "Reporto de forma general o mi compromiso es poco observable.", "Manipulo antes de reportar o escribo compromisos que no se pueden comprobar."),
        ]
    if guide_family(guide) == "tech_history":
        if tech_history_variant(guide) == "digital":
            return [
                ("Comprensión histórica", "Explico hitos de electrónica, computación e Internet con relación clara.", "Reconozco hitos, pero falta conexión entre ellos.", "Enumero tecnologías sin explicar su evolución."),
                ("Impacto social", "Relaciono cada hito con beneficios, riesgos y cambios sociales.", "Menciono impactos generales, pero poco sustentados.", "No conecto tecnología con vida social."),
                ("Lectura crítica", "Reconozco brecha digital, dependencia, vigilancia o desinformación.", "Identifico algunos riesgos sin profundizar.", "Presento la tecnología como progreso automático."),
                ("Producto histórico", "La línea de tiempo es ordenada, explicada y tiene conclusión crítica.", "La línea de tiempo existe, pero le faltan explicación o cierre.", "La línea de tiempo es incompleta o solo decorativa."),
            ]
        return [
            ("Comprensión histórica", "Explico Revolución Industrial, telégrafo y electricidad con relaciones claras.", "Reconozco algunos hitos, pero falta conexión entre ellos.", "Enumero inventos sin explicar su importancia."),
            ("Cambio social", "Relaciono tecnología con trabajo, ciudad, comunicación e infraestructura.", "Menciono cambios generales, pero poco sustentados.", "No conecto inventos con transformaciones sociales."),
            ("Lectura crítica", "Identifico beneficiados, afectados o excluidos por los cambios técnicos.", "Reconozco desigualdades de forma general.", "Presento el progreso sin tensiones ni preguntas."),
            ("Producto histórico", "La línea de tiempo es ordenada, explicada y tiene reflexión crítica.", "La línea de tiempo existe, pero le faltan explicación o cierre.", "La línea de tiempo es incompleta o solo decorativa."),
        ]
    if guide_family(guide) == "cloud_collaboration":
        return [
            ("Colaboración", "Trabajo con roles claros, escucha y aportes visibles en el producto.", "Participo, pero mi rol o aporte no queda del todo claro.", "Trabajo de forma aislada o sin acuerdos de equipo."),
            ("Uso de nube", "Configuro permisos, comentarios e historial con cuidado.", "Uso la nube, pero con permisos o evidencias incompletas.", "Comparto sin controlar accesos o pierdo versiones."),
            ("Producto digital", "El documento, blog, mapa o portafolio responde a la necesidad definida.", "El producto existe, pero falta claridad, revisión o audiencia.", "El producto queda incompleto o sin propósito."),
            ("Datos y liderazgo", "Protejo información y practico liderazgo servicial.", "Reconozco cuidados, pero los aplico parcialmente.", "Expongo datos o no contribuyo al cuidado del grupo."),
        ]
    if guide_family(guide) == "ewaste":
        return [
            ("Diagnóstico RAEE", "Identifico dispositivos, estado, riesgos y evidencias con claridad.", "Registro algunos elementos, pero faltan datos o clasificación.", "No diferencio residuos, riesgos o estado de los equipos."),
            ("Criterio ambiental", "Aplico Rethink, Refuse o Reduce con argumentos.", "Menciono estrategias, pero sin justificar bien.", "Propongo desechar sin análisis ambiental."),
            ("Propuesta comunitaria", "Planteo acciones viables con responsables y evidencia.", "La propuesta es útil, pero general o incompleta.", "No conecto el diagnóstico con una acción concreta."),
            ("Reflexión crítica", "Relaciono consumo tecnológico, ambiente y responsabilidad social.", "Mi reflexión reconoce impactos, pero falta profundidad.", "No explico el impacto ambiental o social."),
        ]
    if guide_family(guide) == "ai_book_project":
        return [
            ("Diseño del libro", "Defino lector, propósito, género, tono e índice con coherencia.", "Tengo una idea general, pero falta precisar lector, tono o estructura.", "Escribo sin plan editorial ni propósito claro."),
            ("Uso de prompts", "Registro prompts claros y explico cómo orientaron planeación, escritura o edición.", "Uso IA, pero mis instrucciones o registros son incompletos.", "Dependo de respuestas vagas sin controlar el proceso."),
            ("Autoría y edición", "Reescribo con voz propia, verifico y justifico cambios editoriales.", "Edito parcialmente, pero quedan partes genéricas o poco revisadas.", "Copio texto asistido sin intervención autoral visible."),
            ("Ética editorial", "Cito fuentes, respeto derechos y declaro el uso de IA con transparencia.", "Reconozco riesgos, pero faltan créditos o declaración precisa.", "Uso recursos o IA sin citar, verificar ni declarar."),
        ]
    if guide_family(guide) == "business_web_presence":
        return [
            ("Estrategia digital", "Defino público, necesidad, propuesta de valor y llamado a la acción con coherencia.", "La estrategia existe, pero falta precisión en público, valor o acción esperada.", "Publico contenido sin saber para quién, para qué o qué acción busca."),
            ("Contenido y marca", "Uso IA para explorar y mejorar textos, pero verifico, reescribo y mantengo voz propia.", "Uso IA con revisión parcial o textos todavía genéricos.", "Copio textos automáticos sin criterio ni autenticidad."),
            ("Diseño y accesibilidad", "Aplico arquitectura, jerarquía visual, contraste, imágenes pertinentes y navegación clara.", "El diseño comunica, pero tiene fallas de orden, contraste o accesibilidad.", "El sitio o pieza es confuso, saturado o difícil de usar."),
            ("Ética y transparencia", "Cuido derechos de imagen, datos, publicidad responsable, créditos y declaración de IA.", "Reconozco algunos cuidados, pero faltan permisos, créditos o declaración completa.", "Publico recursos, datos o promesas sin autorización ni revisión ética."),
        ]
    if guide_family(guide) == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return [
                ("Mirada empresarial", "Observo el caso como sistema de procesos y no como tareas sueltas.", "Reconozco algunos procesos, pero me cuesta relacionarlos.", "Solo describo actividades aisladas sin estructura empresarial."),
                ("Inventario de procesos", "Registro mínimo tres procesos con entrada, actividades, responsable, herramienta y salida.", "Registro procesos, pero faltan elementos clave o claridad.", "El inventario no permite entender cómo funciona el caso."),
                ("Diagnóstico inicial", "Identifico puntos críticos, datos sensibles y oportunidades de mejora con evidencia.", "Identifico problemas generales, pero falta evidencia o cuidado de datos.", "Propongo mejoras sin observar errores, datos ni riesgos."),
                ("Priorización responsable", "Elijo un proceso a mejorar usando impacto, facilidad, riesgo y revisión humana/IA.", "Elijo una mejora, pero justifico poco los criterios.", "Escojo una herramienta sin explicar por qué ni para qué sirve."),
            ]
        return [
            ("Mapeo del proceso", "Represento entrada, actividades, responsables, herramientas, salidas y puntos de mejora.", "El proceso se entiende, pero faltan responsables, herramientas o salidas claras.", "No logro explicar cómo funciona el proceso ni dónde mejora."),
            ("Mejora digital", "Diseño formulario, base, flujo, plantilla o respuesta que reduce errores o tiempo.", "La mejora ayuda parcialmente, pero falta prueba o ajuste.", "La herramienta no responde al problema o aumenta la confusión."),
            ("Datos y privacidad", "Recojo solo datos necesarios, explico finalidad, consentimiento y acceso responsable.", "Cuido algunos datos, pero el protocolo queda incompleto.", "Solicito o comparto información sin cuidado ni autorización."),
            ("Verificación y PDCA", "Pruebo el proceso, documento evidencia, identifico límites y propongo ajustes.", "Verifico de manera parcial o sin indicadores suficientes.", "Afirmo que funciona sin prueba ni mejora documentada."),
        ]
    if guide_family(guide) == "integrated_business_project":
        return [
            ("Problema y usuario", "Formulo necesidad, usuario, contexto, alcance y criterios de éxito con evidencia.", "La necesidad está planteada, pero falta evidencia, usuario o alcance.", "Presento una idea sin problema claro ni usuario definido."),
            ("Integración de solución", "Conecto presencia digital, datos, dashboard, automatización y documentación en una solución coherente.", "Integro algunas piezas, pero la relación entre ellas es débil.", "Presento productos aislados que no responden al problema."),
            ("Decisión con datos e IA", "Uso indicadores, visualizaciones e IA verificada para sustentar una recomendación.", "Uso datos o IA, pero falta verificar hallazgos o reconocer límites.", "Propongo decisiones sin evidencia suficiente o copiadas de IA."),
            ("Pitch y ética", "Sustento impacto, costos, riesgos, protección de datos, autoría, transparencia y sostenibilidad.", "La sustentación comunica la idea, pero faltan riesgos, costos o ética.", "Expongo sin defender decisiones ni reconocer responsabilidades."),
        ]
    if guide_family(guide) == "latex_business_report":
        return [
            ("Estructura profesional", "Organizo el informe con problema, método, datos, análisis y cierre.", "El informe tiene partes útiles, pero falta orden o propósito.", "Presento contenido suelto sin estructura profesional."),
            ("Codificación LaTeX", "Uso comandos, secciones, tablas o figuras y reviso compilación.", "Codifico parcialmente, pero quedan errores o poca claridad.", "No logro relacionar código fuente con PDF final."),
            ("Análisis comercial", "Interpreto datos con fuente, evidencia, límites y recomendación.", "Presento datos, pero el análisis es general o incompleto.", "Muestro tablas/gráficos sin explicar qué significan."),
            ("Citación y transparencia", "Registro fuentes, anexos, derechos y uso de IA con honestidad.", "Incluyo algunas fuentes, pero faltan detalles o revisión.", "No cito, no declaro IA o afirmo sin evidencia."),
        ]
    if guide_family(guide) == "business_dashboard_ai":
        return [
            ("Calidad de datos", "Organizo, limpio y explico campos, registros, errores y límites.", "La base funciona, pero faltan validaciones o limpieza completa.", "Analizo datos desordenados o poco confiables."),
            ("Indicadores y visualización", "Selecciono indicadores y gráficos adecuados para la pregunta empresarial.", "Los indicadores o gráficos son útiles, pero no explican bien el caso.", "Elijo visualizaciones sin relación clara con la decisión."),
            ("IA verificada", "Uso IA para orientar análisis y verifico cada hallazgo con la tabla.", "Uso IA con revisión parcial o sin registrar bien el proceso.", "Acepto recomendaciones automáticas sin comprobarlas."),
            ("Sustentación empresarial", "Defiendo una decisión con evidencia, riesgos, límites y consecuencia ética.", "Propongo una decisión, pero falta sustento o reconocimiento de límites.", "Presento una recomendación sin datos ni argumentos suficientes."),
        ]
    if guide_family(guide) == "ai":
        return [
            ("Comprensión de IA", "Explico el concepto trabajado con ejemplos y límites claros.", "Reconozco ideas generales, pero me falta precisión o ejemplos.", "Confundo IA con cualquier herramienta digital."),
            ("Uso de prompts", "Formulo instrucciones claras con rol, contexto, tarea, formato y criterios.", "Uso prompts básicos, pero no controlo bien el resultado.", "Pido respuestas vagas y dependo totalmente de la herramienta."),
            ("Verificación crítica", "Contrasto información, detecto errores y corrijo con evidencia.", "Reviso parcialmente, pero dejo afirmaciones sin comprobar.", "Acepto o copio respuestas sin verificar."),
            ("Ética y autoría", "Protejo datos, declaro el uso de IA y reescribo con voz propia.", "Reconozco riesgos éticos, pero mi declaración o reescritura es incompleta.", "Expongo datos, oculto el uso de IA o entrego producción sin intervención propia."),
        ]
    if guide_family(guide) == "data_analysis":
        return [
            ("Pregunta y datos", "Formulo una pregunta clara y organizo campos, registros y unidades.", "La pregunta o la tabla son comprensibles, pero incompletas.", "Recojo datos sin propósito o sin estructura."),
            ("Limpieza y cálculo", "Reviso datos, aplico cálculo/gráfico adecuado y explico el procedimiento.", "Hago cálculos, pero falta revisar calidad o coherencia.", "Los datos o cálculos no permiten responder la pregunta."),
            ("Conclusión prudente", "Concluyo con evidencia, límites y cuidado por el contexto.", "Mi conclusión usa datos, pero exagera o deja límites sin explicar.", "Afirmo sin evidencia suficiente."),
            ("Decisión responsable", "Propongo una acción posible sin ocultar riesgos, sesgos o información faltante.", "Propongo una acción general con poca conexión a los datos.", "No relaciono los datos con decisiones responsables."),
        ]
    if guide_family(guide) == "physical_computing":
        return [
            ("Lectura del entorno", "Identifico sensor/entrada, variable y umbral con claridad.", "Reconozco la entrada, pero el umbral o variable queda poco claro.", "No diferencio sensor, variable o condición."),
            ("Lógica de control", "Programa o explica condición y salida con if/else o bloques equivalentes.", "La lógica funciona parcialmente o tiene ramas incompletas.", "La respuesta no depende correctamente de la condición."),
            ("Prueba y depuración", "Documento valores esperado, contrario y límite, y corrijo errores.", "Pruebo algunos valores, pero no explico bien la mejora.", "Afirmo que funciona sin probar."),
            ("Impacto y cuidado", "Relaciono la alerta o sistema con utilidad, límites y cuidado de la comunidad.", "Mi reflexión es general o no reconoce límites.", "No conecto el sistema con responsabilidad o contexto."),
        ]
    if guide_family(guide) == "scratch_programming":
        return [
            ("Comprensión del código", "Predigo, ejecuto e identifico eventos, variables, bucles o condiciones.", "Reconozco algunos bloques, pero me falta explicar relaciones.", "Ejecuto sin comprender qué hace el programa."),
            ("Modificación PRIMM", "Cambio un elemento con propósito y documento el efecto.", "Modifico el programa, pero explico poco el cambio.", "Cambio bloques al azar sin control."),
            ("Prueba y depuración", "Registro pruebas, error, causa probable y corrección.", "Pruebo parcialmente, pero falta evidencia o mejora.", "Entrego sin probar ni depurar."),
            ("Creación interactiva", "La narrativa o juego funciona, tiene propósito y se puede socializar.", "La creación funciona en parte, pero falta cierre o claridad.", "El proyecto no responde a la consigna o no es interactivo."),
        ]
    if guide_family(guide) == "cultural_media":
        return [
            ("Comprensión cultural", "Explico la expresión cultural con contexto, respeto y voces reconocidas.", "Reconozco la expresión, pero falta contexto o profundidad.", "Uso la cultura como adorno sin comprenderla."),
            ("Creación estética", "La pieza comunica belleza, memoria o crítica con decisiones visuales/sonoras claras.", "La pieza comunica una idea, pero sus decisiones no están bien justificadas.", "El producto no tiene intención estética clara."),
            ("Autoría y consentimiento", "Registro fuentes, permisos, créditos y cuidado de imagen o voz.", "Cito parcialmente o me falta comprobar permisos.", "Uso recursos o personas sin reconocimiento ni consentimiento."),
            ("Representación responsable", "Evito estereotipos, reconozco límites y reflexiono sobre impacto.", "Identifico algunos riesgos, pero mi revisión es incompleta.", "Reproduzco estereotipos o no reviso el impacto."),
        ]
    if guide_family(guide) == "editorial":
        if is_noveno_editorial_i(guide):
            return [
                ("Autonomía editorial", "Sigo la ruta, investigo ejemplos, tomo decisiones y reviso antes de pedir aprobación final.", "Avanzo con la guía, pero necesito apoyo frecuente para decidir o revisar.", "Espero instrucciones paso a paso y no uso criterios para mejorar."),
                ("Propósito y audiencia", "Defino tema libre, audiencia, propósito y mensaje central con claridad.", "El tema existe, pero audiencia o propósito quedan generales.", "Diseño sin saber para quién ni para qué comunica."),
                ("Composición visual", "Uso retícula, jerarquía, contraste y espacios para orientar la lectura.", "Aplico algunos recursos, pero hay saturación o desorden.", "La página no permite identificar qué leer primero."),
                ("Autoría y revisión", "Registro fuentes/licencias, corrijo ortografía y justifico dos mejoras del boceto.", "Cito o reviso de forma parcial y explico poco los cambios.", "Uso recursos sin crédito o entrego sin revisión editorial."),
            ]
        return [
            ("Propósito editorial", "Defino audiencia, mensaje y formato con claridad.", "Reconozco el tema, pero mi audiencia o propósito queda general.", "Diseño sin saber para quién ni para qué."),
            ("Composición", "Uso retícula, jerarquía, contraste y espacios para orientar la lectura.", "Aplico algunos recursos, pero hay saturación o desorden.", "La página no permite identificar qué leer primero."),
            ("Autoría y accesibilidad", "Cito fuentes, cuido licencias y reviso legibilidad/accesibilidad.", "Cito parcialmente o reviso solo algunos aspectos.", "Uso recursos sin crédito o con baja legibilidad."),
            ("Revisión y mejora", "Comparo versiones y justifico cambios por claridad y responsabilidad.", "Hago cambios, pero explico poco su propósito.", "Entrego sin revisar ni justificar decisiones."),
        ]
    if guide_family(guide) == "network":
        return [
            ("Comprensión de red", "Identifico nodos, canal, protocolo y función de cada elemento.", "Reconozco partes, pero confundo algunas funciones.", "No diferencio componentes básicos de la red."),
            ("Configuración o diagrama", "Registro direcciones, conexiones o servicios de forma coherente.", "El esquema funciona parcialmente o tiene datos incompletos.", "El diagrama no muestra cómo se comunican los equipos."),
            ("Prueba y diagnóstico", "Documento prueba, resultado, error y corrección con evidencia.", "Hago una prueba, pero explico poco el resultado.", "Afirmo que funciona sin comprobar."),
            ("Seguridad y convivencia", "Relaciono la red con cuidado de datos, accesos y uso responsable.", "Menciono cuidado digital de forma general.", "No reconozco riesgos ni responsabilidades de conexión."),
        ]
    if guide_family(guide) == "info_project":
        return [
            ("Integración del proyecto", "Relaciono necesidad, publicación, red y audiencia con coherencia.", "Integro algunos elementos, pero falta conexión entre diseño y red.", "Presento piezas sueltas sin propósito común."),
            ("Evidencias y pruebas", "Muestro bocetos, fuentes, diagrama, configuración o prueba técnica.", "Incluyo evidencias parciales o poco explicadas.", "No demuestro el proceso ni verifico funcionamiento."),
            ("Ética y autoría", "Protejo datos, cito fuentes, respeto licencias y considero impacto comunitario.", "Reconozco algunos riesgos, pero mi revisión queda incompleta.", "Uso datos o recursos sin cuidado ni reconocimiento."),
            ("Sustentación y mejora", "Justifico decisiones, reconozco límites y explico mejoras aplicadas.", "Explico el producto, pero justifico poco las decisiones.", "Expongo resultados sin analizar errores ni mejoras."),
        ]
    return [
        ("Comprensión del tema", "Explico conceptos y ejemplos con mis palabras.", "Reconozco ideas, pero me falta precisión.", "Necesito repasar conceptos base."),
        ("Aplicación práctica", "Uso el concepto en un producto revisable.", "Aplico una parte, pero faltan ajustes.", "Necesito guía para iniciar o corregir."),
        ("Revisión y mejora", "Pruebo, comparo o corrijo mi trabajo con evidencia.", "Reviso de forma parcial.", "Entrego sin comprobar resultados."),
        ("Reflexión ética", "Relaciono mi producto con cuidado y responsabilidad.", "Mi reflexión es general.", "No conecto el trabajo con su impacto."),
    ]


def logic_symbol_rows(guide: GuideData) -> list[tuple[str, str, str]]:
    family = guide_family(guide)
    if family not in {"programming", "physical_computing", "scratch_programming", "ai"}:
        return []
    if family == "ai":
        return [
            (">", "mayor que", "interes > 70"),
            ("<", "menor que", "riesgo < 30"),
            (">=", "mayor o igual que", "confianza >= 80"),
            ("<=", "menor o igual que", "edad <= 13"),
            ("==", "igual a", "categoria == 'educativo'"),
            ("!=", "diferente de", "fuente != 'desconocida'"),
            ("AND", "se cumplen dos condiciones", "interes > 70 AND riesgo < 30"),
            ("OR", "se cumple al menos una condición", "tema == 'arte' OR tema == 'ciencia'"),
            ("NOT", "niega una condición", "NOT dato_personal"),
        ]
    if family == "scratch_programming":
        return [
            (">", "mayor que", "puntaje > 10"),
            ("<", "menor que", "vidas < 1"),
            (">=", "mayor o igual que", "nivel >= 3"),
            ("<=", "menor o igual que", "tiempo <= 0"),
            ("==", "igual a", "respuesta == 'sí'"),
            ("!=", "diferente de", "personaje != 'oculto'"),
            ("AND", "se cumplen dos condiciones", "puntaje > 10 AND vidas > 0"),
            ("OR", "se cumple al menos una condición", "tecla == 'A' OR tecla == 'B'"),
            ("NOT", "niega una condición", "NOT tocando_borde"),
        ]
    if family == "physical_computing":
        return [
            (">", "mayor que", "temperatura > 30"),
            ("<", "menor que", "luz < 40"),
            (">=", "mayor o igual que", "humedad >= 70"),
            ("<=", "menor o igual que", "sonido <= 20"),
            ("==", "igual a", "boton == presionado"),
            ("!=", "diferente de", "estado != 'normal'"),
            ("AND", "se cumplen dos condiciones", "temperatura > 30 AND luz < 40"),
            ("OR", "se cumple al menos una condición", "temperatura > 30 OR humedad > 70"),
            ("NOT", "niega una condición", "NOT alerta_activa"),
        ]
    return [
        (">", "mayor que", "edad > 12"),
        ("<", "menor que", "temperatura < 18"),
        (">=", "mayor o igual que", "nota >= 3.0"),
        ("<=", "menor o igual que", "temperatura <= 28"),
        ("==", "igual a", "opcion == 1"),
        ("!=", "diferente de", "respuesta != ''"),
        ("AND", "se cumplen dos condiciones", "edad >= 12 AND permiso == true"),
        ("OR", "se cumple al menos una condición", "dia == 'sábado' OR dia == 'domingo'"),
        ("NOT", "niega una condición", "NOT tiene_permiso"),
    ]


def translation_rows(guide: GuideData) -> list[tuple[str, str, str]]:
    family = guide_family(guide)
    if family not in {"programming", "physical_computing", "scratch_programming", "ai"}:
        return []
    if family == "ai":
        return [
            ("El contenido tiene alto interés.", "interes, >", "interes > 70"),
            ("La fuente no es desconocida.", "fuente, !=", "fuente != 'desconocida'"),
            ("Recomendar solo si interesa y no hay alto riesgo.", "interes, riesgo, AND", "interes > 70 AND riesgo < 30"),
            ("No debe incluir datos personales.", "dato_personal, NOT", "NOT dato_personal"),
        ]
    if family == "scratch_programming":
        return [
            ("El puntaje supera 10.", "puntaje, >", "puntaje > 10"),
            ("Ya no quedan vidas.", "vidas, <=", "vidas <= 0"),
            ("Tiene puntos y aún tiene vidas.", "puntaje, vidas, AND", "puntaje > 10 AND vidas > 0"),
            ("No está tocando el borde.", "tocando_borde, NOT", "NOT tocando_borde"),
        ]
    if family == "physical_computing":
        return [
            ("La temperatura supera el límite.", "temperatura, >", "temperatura > 30"),
            ("La luz está baja.", "luz, <", "luz < 40"),
            ("Hay calor y poca luz.", "temperatura, luz, AND", "temperatura > 30 AND luz < 40"),
            ("No hay alerta activa.", "alerta_activa, NOT", "NOT alerta_activa"),
        ]
    return [
        ("La nota es 3.0 o más.", "nota, >=", "nota >= 3.0"),
        ("La edad está entre 12 y 17.", "edad, >=, <=, AND", "edad >= 12 AND edad <= 17"),
        ("No tiene permiso.", "tiene_permiso, NOT", "NOT tiene_permiso"),
        ("La luz es baja o la temperatura es alta.", "luz, temperatura, OR", "luz < 30 OR temperatura > 28"),
    ]


def test_case_rows(guide: GuideData) -> list[tuple[str, str, str]]:
    family = guide_family(guide)
    if family not in {"programming", "physical_computing", "scratch_programming", "ai"}:
        return []
    if family == "ai":
        return [
            ("Caso esperado", "interes = 85 y riesgo = 20", "La regla es verdadera: se podría recomendar y luego verificar."),
            ("Caso contrario", "interes = 40 y riesgo = 20", "La regla es falso: no alcanza el interés definido."),
            ("Caso límite", "interes = 70 y riesgo = 30", "Con > y < es falso; con >= y <= cambiaría la decisión."),
        ]
    if family == "scratch_programming":
        return [
            ("Caso esperado", "puntaje = 12", "puntaje > 10 es verdadero: pasa de nivel"),
            ("Caso contrario", "puntaje = 8", "puntaje > 10 es falso: sigue jugando"),
            ("Caso límite", "puntaje = 10", "puntaje > 10 es falso; con >= sería verdadero"),
        ]
    if family == "physical_computing":
        return [
            ("Caso esperado", "temperatura = 31", "temperatura > 30 es verdadero: activar alerta"),
            ("Caso contrario", "temperatura = 28", "temperatura > 30 es falso: estado normal"),
            ("Caso límite", "temperatura = 30", "temperatura > 30 es falso; con >= sería verdadero"),
        ]
    return [
        ("Caso esperado", "nota = 4.0", "nota >= 3.0 es verdadero"),
        ("Caso contrario", "nota = 2.5", "nota >= 3.0 es falso"),
        ("Caso límite", "nota = 3.0", "nota >= 3.0 es verdadero"),
    ]


def autonomous_checklist_rows(guide: GuideData) -> list[tuple[str, str, str]]:
    if guide_family(guide) != "systems_room":
        return []
    return [
        ("Monitor y pantalla", "Enciende, no presenta golpes visibles y está ubicado de forma segura.", "Observo y reporto si no prende o se ve dañado."),
        ("Teclado y mouse", "Están completos, limpios y responden al uso normal.", "Pruebo con cuidado y no golpeo si fallan."),
        ("Cables y tomas", "No atraviesan zonas de paso ni están sueltos o tensos.", "No desconecto; aviso si veo riesgo."),
        ("Silla y mesa", "Permiten una postura cómoda y dejan espacio para moverse.", "Ajusto mi puesto sin empujar ni invadir otro espacio."),
        ("Cuenta o sesión", "No hay información personal abierta de otra persona.", "No leo ni cambio archivos; reporto sesión abierta."),
        ("Alimentos y líquidos", "No están cerca de equipos eléctricos.", "Los retiro del puesto siguiendo la norma de la sala."),
        ("Ruido y turnos", "El volumen y la participación permiten que todos aprendan.", "Uso voz moderada, respeto turnos y comparto el equipo."),
    ]


def learning_goal_text(guide: GuideData) -> str:
    family = guide_family(guide)
    if family in FAMILY_LEARNING_GOAL:
        return FAMILY_LEARNING_GOAL[family]
    if is_noveno_editorial_i(guide):
        return "Al terminar la sesión podrás aprender de manera autónoma a analizar una publicación real, definir audiencia y propósito, bocetar un folleto o boletín con retícula, citar recursos y justificar mejoras editoriales."
    if is_undecimo_process_inventory(guide):
        return "Al terminar la sesión podrás inventariar procesos empresariales de un caso o emprendimiento, identificar entradas, responsables, salidas, puntos críticos, datos sensibles y una oportunidad de mejora digital o con IA."
    if family == "systems_room":
        return "Al terminar la sesión podrás avanzar con autonomía: leer instrucciones, observar tu puesto, identificar riesgos, reportar fallas y proponer normas para cuidar la sala de sistemas."
    if family == "business_web_presence":
        return "Al terminar la sesión podrás fortalecer una presencia digital empresarial con marca, público, textos, diseño, SEO básico y criterios éticos de publicación con IA."
    if family == "business_process_automation":
        return "Al terminar la sesión podrás mapear, mejorar o automatizar parcialmente un proceso empresarial con datos, herramientas digitales, IA verificada y protección de información."
    if family == "integrated_business_project":
        return "Al terminar la sesión podrás integrar evidencias de un proyecto empresarial final: problema, usuario, solución, datos, dashboard, automatización, documentación, IA ética y pitch."
    if family == "ai_book_project":
        return "Al terminar la sesión podrás avanzar en tu libro con apoyo de IA, conservando voz propia, evidencia del proceso, edición humana y criterios éticos de autoría."
    if family == "latex_business_report":
        return "Al terminar la sesión podrás construir una parte de un informe comercial o contable en LaTeX, con datos, fuentes, análisis y revisión profesional."
    if family == "business_dashboard_ai":
        return "Al terminar la sesión podrás organizar datos empresariales, usar IA con verificación y sustentar indicadores, visualizaciones o decisiones de negocio."
    return "Al terminar la sesión podrás explicar ideas centrales del tema, desarrollar una evidencia práctica y asumir un compromiso de uso responsable de la tecnología."


def how_to_use_text(guide: GuideData) -> str:
    family = guide_family(guide)
    if family in FAMILY_HOW_TO_USE:
        return FAMILY_HOW_TO_USE[family]
    if is_noveno_editorial_i(guide):
        return "Trabaja como editor aprendiz: lee la meta, analiza ejemplos, toma notas, boceta antes de diseñar, revisa con la lista de calidad y deja evidencia de tus decisiones. El docente acompaña, pero la guía debe permitirte avanzar por ti mismo."
    if is_undecimo_process_inventory(guide):
        return "Trabaja como consultor junior: observa un caso, no propongas herramienta de inmediato, registra evidencias, identifica datos sensibles, compara procesos y prioriza una oportunidad de mejora que puedas defender."
    if guide_family(guide) == "systems_room":
        return "Trabaja como investigador de la sala: lee, observa sin manipular, marca la lista de cotejo, decide qué puedes hacer con seguridad, reporta lo que sea necesario y revisa tu producto antes de entregarlo."
    if guide_family(guide) in {"business_web_presence", "business_process_automation", "integrated_business_project"}:
        return "Trabaja como equipo consultor: lee el caso, define el usuario, registra decisiones, verifica cualquier apoyo de IA y deja evidencias que puedan sustentarse ante una audiencia empresarial."
    return "Lee cada fase, completa las marcas de verificación y escribe respuestas breves pero completas. Si aparece un espacio amplio, úsalo para dibujar, esquematizar o representar tu comprensión."


def escuta_challenge_text(guide: GuideData) -> str:
    family = guide_family(guide)
    if family in FAMILY_ESCUTA_CHALLENGE:
        return FAMILY_ESCUTA_CHALLENGE[family]
    if is_noveno_editorial_i(guide):
        return "Busca o recuerda un folleto, afiche, boletín o publicación escolar que hayas visto. Obsérvalo como editor: identifica qué se entiende rápido, qué distrae, qué falta citar y qué mejorarías para que otra persona aprenda sin ayuda."
    if is_undecimo_process_inventory(guide):
        return "Observa un emprendimiento escolar, tienda, servicio institucional o caso simulado. Antes de pensar en IA o automatización, identifica qué procesos existen, dónde se repite trabajo y qué información debe cuidarse."
    if guide_family(guide) == "systems_room":
        return "Antes de pedir instrucciones, observa tu puesto y la sala como si fueras responsable de dejarla lista para que otra persona aprenda allí. No manipules cables ni cuentas: primero mira, piensa y registra."
    if guide_family(guide) == "business_web_presence":
        return "Observa una marca o emprendimiento cercano. Pregúntate si entiendes qué ofrece, para quién es, por qué confiar y qué acción deberías realizar después de leer su presencia digital."
    if guide_family(guide) == "business_process_automation":
        return "Piensa en un proceso cotidiano del colegio o de un negocio: pedido, encuesta, inventario, respuesta o seguimiento. Ubica dónde se repite trabajo, se pierden datos o aparecen errores."
    if guide_family(guide) == "integrated_business_project":
        return "Mira tu proyecto como una solución completa: problema, usuario, evidencia, canal digital, datos, mejora del proceso y argumento final. Identifica qué pieza todavía no conversa con las demás."
    return "Conecta el tema con tu experiencia. Observa tu entorno, recuerda situaciones cercanas y formula preguntas iniciales antes de buscar respuestas."


def escuta_check_items(guide: GuideData) -> list[str]:
    family = guide_family(guide)
    if family in FAMILY_ESCUTA_CHECK_ITEMS:
        return FAMILY_ESCUTA_CHECK_ITEMS[family]
    if is_noveno_editorial_i(guide):
        return [
            "Identifico audiencia, propósito y mensaje central de una publicación.",
            "Reconozco una decisión visual que ayuda a leer y una que confunde.",
            "Formulo una pregunta de investigación sobre diseño, autoría o legibilidad.",
        ]
    if is_undecimo_process_inventory(guide):
        return [
            "Identifico al menos tres procesos de un caso empresarial o escolar.",
            "Reconozco entradas, responsables, salidas y herramientas usadas.",
            "Ubico un punto crítico y un dato sensible que requiere cuidado.",
        ]
    if guide_family(guide) == "systems_room":
        return [
            "Leo la pregunta guía antes de iniciar.",
            "Observo mi puesto sin mover cables ni cambiar configuraciones.",
            "Distingo entre lo que puedo hacer, lo que debo reportar y lo que debo evitar.",
        ]
    if guide_family(guide) == "business_web_presence":
        return [
            "Identifico público, necesidad y propuesta de valor.",
            "Reviso si los textos, imágenes y diseño comunican confianza.",
            "Reconozco un cuidado ético antes de publicar.",
        ]
    if guide_family(guide) == "business_process_automation":
        return [
            "Identifico entrada, actividad, responsable y salida del proceso.",
            "Ubico una tarea repetitiva o un error que puede mejorarse.",
            "Reconozco qué datos requieren consentimiento o cuidado.",
        ]
    if guide_family(guide) == "integrated_business_project":
        return [
            "Conecto el problema con un usuario real o simulado.",
            "Reconozco qué evidencia sustentará la solución.",
            "Identifico cómo IA, datos y automatización deben ser verificados.",
        ]
    return [
        "Identifico lo que ya sé sobre el tema.",
        "Reconozco una situación de mi vida escolar relacionada con el tema.",
        "Escribo una pregunta que me ayude a investigar mejor.",
    ]


def escuta_response_prompts(guide: GuideData) -> tuple[str, str]:
    family = guide_family(guide)
    if family in FAMILY_ESCUTA_RESPONSE_PROMPTS:
        return FAMILY_ESCUTA_RESPONSE_PROMPTS[family]
    if is_noveno_editorial_i(guide):
        return (
            "Lo que observo en una publicación real:",
            "La pregunta que necesito investigar antes de diseñar:",
        )
    if is_undecimo_process_inventory(guide):
        return (
            "Procesos que observo en el caso o emprendimiento:",
            "Un punto crítico o dato sensible que debo investigar:",
        )
    if family == "systems_room":
        return (
            "Lo que observo en mi puesto:",
            "Una duda o riesgo que debo aclarar:",
        )
    if family == "business_web_presence":
        return ("La propuesta de valor que debo comunicar:", "Un cuidado ético o de autoría que debo revisar:")
    if family == "business_process_automation":
        return ("El proceso que puedo mejorar:", "Un dato, permiso o riesgo que debo cuidar:")
    if family == "integrated_business_project":
        return ("El problema o usuario que orienta mi proyecto:", "La evidencia que necesito para sustentar la solución:")
    if family == "ai_book_project":
        return ("Lo que quiero comunicar en mi libro:", "Una decisión de autoría o ética que debo cuidar:")
    if family == "latex_business_report":
        return ("El problema comercial o contable que puedo analizar:", "Un dato, fuente o estructura que debo verificar:")
    if family == "business_dashboard_ai":
        return ("La decisión empresarial que podría tomar con datos:", "Un riesgo o dato dudoso que debo revisar:")
    return ("Lo que ya sé:", "Una situación cercana:")


def period_idea_text(guide: GuideData) -> str:
    family = guide_family(guide)
    if family in FAMILY_PERIOD_IDEA:
        return FAMILY_PERIOD_IDEA[family]
    if family == "systems_room":
        return "El periodo inicia con autonomía, cuidado de la sala, convivencia tecnológica y reconocimiento de normas antes de avanzar hacia identidad digital, correo institucional y netiqueta."
    if family == "cloud_collaboration":
        return "El periodo fortalece colaboración en la nube mediante roles, permisos, comentarios, historial, liderazgo servicial y evidencias de trabajo compartido."
    if family == "scratch_programming":
        return "El periodo desarrolla pensamiento algorítmico con Scratch, PRIMM, eventos, variables, bucles, condicionales, pruebas y depuración documentada."
    if family == "ai":
        return "El periodo introduce IA con ejemplos cotidianos, prompts responsables, verificación, privacidad, sesgos, autoría y criterio humano."
    if family == "ewaste":
        return "El periodo relaciona tecnología y ambiente mediante gestión RAEE, auditoría de e-waste, reducción de residuos y propuestas responsables para la institución."
    if family == "business_web_presence":
        return "El periodo desarrolla presencia digital empresarial con IA: identidad de marca, arquitectura web, textos comerciales, SEO básico, contenido para redes, accesibilidad y ética."
    if family == "business_process_automation":
        return "El periodo mejora procesos empresariales con IA: mapeo, formularios, datos, automatización básica, atención al cliente, PDCA, privacidad y presentación de resultados."
    if family == "integrated_business_project":
        return "El periodo integra proyecto final empresarial asistido con IA: problema, usuario, solución, presencia digital, datos, dashboard, automatización, documentación, pitch y sustentación."
    if family == "ai_book_project":
        return "El periodo desarrolla Escritura Inversa: libro de tema libre con IA, prompts, voz propia, edición, derechos de autor y declaración transparente."
    if family == "latex_business_report":
        return "El periodo proyecta la escritura asistida hacia informes comerciales en LaTeX, datos, tablas, gráficos, citación y análisis de mercado."
    if family == "business_dashboard_ai":
        return "El periodo integra datos empresariales, hojas de cálculo, IA, visualizaciones, dashboard, reportes y sustentación de decisiones comerciales o contables."
    ideas = split_sentences(guide.period.aprendizajes)
    return ideas[0] if ideas else guide.period.aprendizajes


def praxis_task_text(guide: GuideData) -> str:
    family = guide_family(guide)
    if family in FAMILY_PRAXIS_TASK:
        return FAMILY_PRAXIS_TASK[family]
    if family == "systems_room":
        return "Elabora un mapa sencillo de tu estación o zona de sala. Señala al menos cuatro elementos, marca riesgos con un símbolo y escribe una acción segura para cada riesgo."
    if family == "business_web_presence":
        return "Construye o mejora una evidencia de presencia digital: ficha de marca, arquitectura, textos, diseño visual, SEO, pieza de red o sustentación. Explica qué aportó la IA, qué verificaste y qué decisión ética aplicaste."
    if family == "business_process_automation":
        if is_undecimo_process_inventory(guide):
            return "Elabora un inventario de procesos del caso elegido. Incluye mínimo tres procesos con entrada, actividades, responsable, herramienta actual y salida. Luego selecciona uno para mejorar, explica el punto crítico, qué dato sensible debe cuidarse y qué oportunidad digital o de IA podría explorarse en las siguientes guías."
        return "Desarrolla una evidencia del proceso: mapa, diagnóstico, formulario, base, flujo, FAQ, protocolo de datos o presentación. Señala una tarea repetitiva, una mejora, una prueba y un límite ético."
    if family == "integrated_business_project":
        return "Integra una pieza del proyecto final: problema, usuario, solución, sitio, datos, dashboard, automatización, documentación o pitch. Conecta la evidencia con una decisión empresarial y registra el uso responsable de IA."
    if family == "ai_book_project":
        return "Produce una evidencia de tu libro según la sesión: plan editorial, prompt comentado, fragmento reescrito, créditos, portada o bitácora. Explica qué decidió la IA, qué decidiste tú y qué criterio ético aplicaste."
    if family == "latex_business_report":
        return "Construye una parte verificable del informe: estructura LaTeX, tabla, gráfico, análisis, fuente o anexo. Señala un error corregido y una decisión profesional tomada."
    if family == "business_dashboard_ai":
        return "Elabora una ficha de análisis empresarial: dato revisado, indicador, visualización, hallazgo, recomendación, riesgo y verificación de cualquier apoyo de IA."
    if is_noveno_editorial_i(guide):
        return "Elabora el boceto de un folleto o boletín de tema libre. Debe incluir título, subtítulo, tres bloques breves de texto, una zona de imagen, una fuente citada, un llamado a la acción y una mejora justificada después de revisar la primera versión."
    return f"Representa una escena relacionada con {guide.focus}. Incluye una acción correcta y una acción que debe evitarse."


def tex_itemize(items: list[str]) -> str:
    return "\\begin{itemize}\n" + "\n".join(f"\\item {tex_escape(item)}" for item in items) + "\n\\end{itemize}"


def tex_optional_table(title: str, headers: tuple[str, ...], rows: list[tuple[str, ...]], color: str = "milcTurquesa") -> str:
    if not rows:
        return ""
    columns = ">" + r"{\bfseries\RaggedRight\arraybackslash}p{3.0cm}" + "Y" * (len(headers) - 1)
    header_line = " & ".join(rf"\color{{white}}\bfseries {tex_escape(header)}" for header in headers) + r"\\"
    body = "\n".join(" & ".join(tex_escape(cell) for cell in row) + r"\\" for row in rows)
    return rf"""
\begin{{drawbox}}{{{color}}}{{{tex_escape(title)}}}
\small
\begin{{tabularx}}{{\textwidth}}{{{columns}}}
\rowcolor{{{color}!90}}{header_line}
{body}
\end{{tabularx}}
\normalsize
\end{{drawbox}}
"""


def tex_extra_after_how_to_use(guide: GuideData) -> str:
    if is_noveno_editorial_i(guide):
        return r"""
\begin{drawbox}{milcVerde}{Aprendizaje autónomo}
\small
\begin{tabularx}{\textwidth}{>{\bfseries}p{3.2cm}Y}
Meta personal & Diseñar un boceto editorial que otra persona pueda entender sin que yo tenga que explicarlo oralmente.\\
Recursos & Una publicación real para observar, la tabla de conceptos, una imagen con fuente verificable y esta lista de criterios.\\
Evidencia mínima & Boceto con retícula, mensaje central, imagen citada, revisión de legibilidad y justificación de dos decisiones.\\
Control de avance & Marco cada paso, comparo mi boceto con los criterios y corrijo antes de entregar.\\
\end{tabularx}
\normalsize
\end{drawbox}
"""
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{drawbox}{milcVerde}{Aprendizaje autónomo}
\small
\begin{tabularx}{\textwidth}{>{\bfseries\RaggedRight\arraybackslash}p{3.2cm}Y}
Meta personal & Construir un inventario de procesos que permita decidir qué parte de un emprendimiento o caso escolar merece mejora.\\
Recursos & Observación, entrevista breve, tabla de conceptos, matriz de procesos y criterios de privacidad.\\
Evidencia mínima & Tres procesos descritos, un proceso priorizado, un punto crítico, un dato sensible y una oportunidad digital o con IA.\\
Control de avance & No propongo herramienta hasta completar entrada, actividad, responsable, salida, riesgo y evidencia.\\
\end{tabularx}
\normalsize
\end{drawbox}
"""
    return ""


def tex_extra_after_escuta(guide: GuideData) -> str:
    if is_noveno_editorial_i(guide):
        return r"""
\begin{drawbox}{milcVerde}{Mini investigación}
Elige una publicación real: folleto, afiche, boletín, infografía o publicación institucional. Puede ser impresa o digital. Obsérvala durante tres minutos y responde con evidencia.
\begin{enumerate}
\item ¿Quién parece ser la audiencia?
\item ¿Qué mensaje central se entiende primero?
\item ¿Qué recurso visual ayuda y cuál distrae?
\item ¿Aparecen fuente, autor, fecha o licencia?
\end{enumerate}
\textbf{Dos hallazgos de mi investigación:}
\respuesta{2}
\textbf{Una decisión que puedo aplicar en mi propio boceto:}
\respuesta{1}
\end{drawbox}
"""
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{drawbox}{milcVerde}{Mini investigación empresarial}
Observa un proceso de un emprendimiento escolar, tienda familiar, servicio institucional o caso simulado. También puedes entrevistar brevemente a una persona que conozca el proceso.
\begin{enumerate}
\item ¿Qué acción o dato inicia el proceso?
\item ¿Quién realiza o revisa cada parte?
\item ¿Dónde se repiten errores, demoras o confusiones?
\item ¿Qué información personal o comercial debe protegerse?
\end{enumerate}
\textbf{Dos evidencias de mi observación o entrevista:}
\respuesta{2}
\textbf{Una pregunta para investigar en la siguiente guía:}
\respuesta{1}
\end{drawbox}
"""
    return ""


def tex_extra_after_practice(guide: GuideData) -> str:
    if is_noveno_editorial_i(guide):
        return r"""
\begin{softbox}{milcMostaza}{milcGris}{Auto chequeo antes de la praxis}
\checkbox Mi tema libre es útil para una audiencia concreta.\\
\checkbox Mi mensaje central cabe en una frase clara.\\
\checkbox Mi boceto deja ver qué se lee primero, segundo y tercero.\\
\checkbox Mis imágenes o datos tendrán fuente, autor o licencia.\\
\checkbox Ya sé qué cambio haré si el diseño queda saturado o confuso.\\
\end{softbox}
"""
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{drawbox}{milcMostaza}{Matriz de inventario inicial}
\small
\begin{tabularx}{\textwidth}{>{\bfseries\RaggedRight\arraybackslash}p{2.2cm}YYYY}
\rowcolor{milcMostaza!80}\color{white}\bfseries Proceso & \color{white}\bfseries Entrada & \color{white}\bfseries Responsable & \color{white}\bfseries Salida & \color{white}\bfseries Punto crítico\\
Proceso 1 & & & &\\[6mm]
\rowcolor{milcGris}Proceso 2 & & & &\\[6mm]
Proceso 3 & & & &\\[6mm]
\end{tabularx}
\normalsize
\end{drawbox}

\begin{softbox}{milcMostaza}{milcGris}{Auto chequeo antes de priorizar}
\checkbox Identifiqué mínimo tres procesos reales o simulados.\\
\checkbox Cada proceso tiene entrada, responsable y salida verificable.\\
\checkbox Reconocí un punto crítico con evidencia.\\
\checkbox Señalé datos sensibles y forma de protegerlos.\\
\checkbox Aún no elegí herramienta sin justificar la necesidad.\\
\end{softbox}
"""
    return ""


def tex_extra_after_phase2_bridge(guide: GuideData) -> str:
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{drawbox}{milcTurquesa}{Bitácora de decisión}
Antes de pasar a la praxis, convierte tu observación en una decisión empresarial. No se trata de escoger la herramienta más llamativa, sino de justificar qué proceso merece atención primero.
\begin{enumerate}
\item Hallazgo más importante observado o escuchado:
\item Evidencia que lo demuestra:
\item Dato sensible que debe protegerse:
\item Proceso que parece más urgente de revisar:
\item Pregunta que le haría a una IA para ordenar la información, sin entregar datos reales:
\end{enumerate}
\textbf{Mi criterio para elegir el proceso a trabajar:}
\respuesta{2}
\end{drawbox}

\begin{softbox}{milcMostaza}{milcGris}{Semáforo de preparación}
\checkbox Puedo explicar el proceso sin depender de una herramienta.\\
\checkbox Tengo al menos una evidencia: observación, entrevista, registro o ejemplo simulado.\\
\checkbox Reconozco quién participa y quién recibe el resultado.\\
\checkbox Sé qué información no debo publicar ni copiar en una IA.\\
\checkbox La mejora que imagino responde a un problema real del caso.\\
\end{softbox}
"""
    return ""


def tex_extra_after_praxis(guide: GuideData) -> str:
    if is_noveno_editorial_i(guide):
        return r"""
\begin{drawbox}{milcMagenta}{Bitácora de edición}
\small
\begin{tabularx}{\textwidth}{>{\bfseries}p{3.0cm}Y}
Versión 1 & ¿Qué problema tiene mi primer boceto?\\
\rowcolor{milcGris}Mejora aplicada & ¿Qué cambié en jerarquía, retícula, texto, imagen o crédito?\\
Criterio usado & ¿Por qué ese cambio mejora la lectura o la responsabilidad editorial?\\
\rowcolor{milcGris}Evidencia & ¿Qué parte del boceto muestra la mejora?\\
\end{tabularx}
\normalsize
\respuesta{2}
\end{drawbox}
"""
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{drawbox}{milcMagenta}{Priorización de mejora}
\small
\begin{tabularx}{\textwidth}{>{\bfseries\RaggedRight\arraybackslash}p{3.0cm}Y}
Proceso elegido & ¿Cuál proceso conviene mejorar primero?\\
\rowcolor{milcGris}Impacto & ¿Qué problema reduce para cliente, equipo, ventas, inventario o seguimiento?\\
Facilidad & ¿Qué tan posible es probar una mejora en clase con recursos disponibles?\\
\rowcolor{milcGris}Riesgo y datos & ¿Qué dato sensible o riesgo debe cuidarse antes de digitalizar?\\
Oportunidad futura & ¿Formulario, hoja, plantilla, IA, dashboard o flujo? ¿Por qué?\\
\end{tabularx}
\normalsize
\respuesta{2}
\end{drawbox}
"""
    return ""


def tex_extra_after_product(guide: GuideData) -> str:
    if is_noveno_editorial_i(guide):
        return r"""
\begin{softbox}{milcVino}{milcGris}{Declaración de autoría y responsabilidad}
Mi boceto es de tema libre, pero mis decisiones deben ser honestas: reconozco fuentes, evito copiar diseños completos, no uso imágenes de personas sin permiso y explico qué aprendí al revisar mi propia publicación.
\end{softbox}
"""
    if is_undecimo_process_inventory(guide):
        return r"""
\begin{softbox}{milcVino}{milcGris}{Protocolo mínimo de datos e IA}
Antes de usar formularios, hojas de cálculo o IA, el equipo debe definir finalidad, recolectar solo datos necesarios, evitar información sensible innecesaria, limitar accesos y verificar cualquier recomendación generada por IA con evidencia del proceso.
\end{softbox}
"""
    return ""


def tex_classic_cover(guide: GuideData, title_font: int, title_leading: int) -> str:
    period = guide.period
    return rf"""
\begin{{document}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[white] (current page.south west) rectangle (current page.north east);
  \fill[milcMagenta] (current page.north west) rectangle ([yshift=-1.35cm]current page.north east);
  \begin{{scope}}
    \clip (current page.north west) rectangle ([yshift=-1.35cm]current page.north east);
    \foreach \x in {{0,1.4,...,23}}{{
      \draw[milcVino,opacity=.16,line width=.5pt]
        ([xshift=\x cm,yshift=-.20cm]current page.north west)
        .. controls +(0.6cm,-.30cm) and +(-0.6cm,.30cm) ..
        ([xshift=\x cm,yshift=-1.10cm]current page.north west);
      \draw[white,opacity=.12,line width=.4pt]
        ([xshift=\x cm,yshift=-.05cm]current page.north west) -- ++(1.4cm,-1.30cm);
    }}
  \end{{scope}}
  \draw[milcVino,opacity=.10,line width=.6pt]
    ([yshift=-1.35cm]current page.north west) -- ([yshift=-1.35cm]current page.north east);
\end{{tikzpicture}}

\vspace*{{1.25cm}}
\begin{{tcolorbox}}[
  enhanced,colback=milcVino,colframe=milcVino,arc=12mm,width=.84\textwidth,
  boxrule=0pt,left=10mm,right=10mm,top=11mm,bottom=11mm,center,
  overlay={{\draw[milcMagenta,line width=2pt,dash pattern=on 9pt off 8pt,rounded corners=10mm]
  ([xshift=7mm,yshift=-7mm]frame.north west) rectangle ([xshift=-7mm,yshift=7mm]frame.south east);}}]
  \centering
  {{\sffamily\bfseries\color{{white}}\fontsize{{14}}{{16}}\selectfont SERIE GUÍAS · TECNOLOGÍA E INFORMÁTICA}}\\[8mm]
  \begin{{minipage}}{{.94\linewidth}}\centering\sloppy
  {{\sffamily\bfseries\color{{white}}\fontsize{{{title_font}}}{{{title_leading}}}\selectfont {tex_escape(guide.short_title)}}}
  \end{{minipage}}\\[4mm]
  {{\sffamily\bfseries\color{{milcMostaza}}\fontsize{{24}}{{27}}\selectfont Guía {guide.number}-{guide.grade}-TIC}}\\[8mm]
  {{\sffamily\color{{white}}\Large Grado {guide.grade} · Periodo {tex_escape(period.name)}}}\\[1.0cm]
  {{\sffamily\bfseries\color{{white}}\large Institución Educativa Sor María Juliana}}\\[2mm]
  {{\sffamily\color{{white}}Autor: {tex_escape(AUTHOR)}}}\\[4mm]
  {{\sffamily\itshape\color{{white}}Lo que necesitamos saber, cuidar y saber hacer}}
\end{{tcolorbox}}
\vfill
\begin{{tcolorbox}}[enhanced,colback=milcVino!72!black,colframe=milcVino!72!black,arc=5mm,boxrule=0pt,width=.86\textwidth,left=6mm,right=6mm,top=4mm,bottom=4mm,center]
\centering{{\color{{white}}\sffamily\small Metodología MILC · Escucha · Sistematización · Praxis · Evaluación liberadora\\Producto: {tex_escape(guide.product)}}}
\end{{tcolorbox}}

\newpage
"""


def tex_cover_icon(grade: int) -> str:
    if grade == 7:
        return r"""
  \begin{scope}[shift={([xshift=-3.05cm,yshift=-2.55cm]current page.north east)}]
    \draw[milcGradeText,line width=1.2pt,rounded corners=2mm] (-.58,-.38) rectangle (.58,.38);
    \fill[milcGradeText] (-.25,.05) circle (.07);
    \fill[milcGradeText] (.25,.05) circle (.07);
    \draw[milcGradeText,line width=.9pt] (-.28,-.16) -- (.28,-.16);
    \draw[milcGradeText,line width=.9pt] (0,.38) -- (0,.62);
    \fill[milcGradeAccent] (0,.72) circle (.10);
    \draw[milcGradeText,line width=.9pt] (-.76,.14) -- (-.58,.14) (.58,.14) -- (.76,.14);
  \end{scope}
"""
    if grade == 8:
        return r"""
  \begin{scope}[shift={([xshift=-3.05cm,yshift=-2.55cm]current page.north east)}]
    \fill[white,opacity=.80] (-.62,-.52) rectangle (.62,.52);
    \draw[milcGradeText,line width=1pt] (-.62,-.52) rectangle (.62,.52);
    \fill[milcGradeDark] (-.42,-.38) rectangle (-.22,.06);
    \fill[milcGradeAccent] (-.10,-.38) rectangle (.10,.24);
    \fill[milcGradePrimary!60!black] (.22,-.38) rectangle (.42,.38);
  \end{scope}
"""
    if grade == 9:
        return r"""
  \begin{scope}[shift={([xshift=-3.05cm,yshift=-2.55cm]current page.north east)}]
    \draw[milcGradeAccent,line width=1.2pt] (0,0) circle (.55);
    \draw[milcGradeAccent,line width=.8pt] (-.55,0) -- (.55,0);
    \draw[milcGradeAccent,line width=.8pt] (0,-.55) -- (0,.55);
    \draw[milcGradeAccent,line width=.8pt] (-.38,-.38) .. controls (0,-.18) .. (.38,-.38);
    \draw[milcGradeAccent,line width=.8pt] (-.38,.38) .. controls (0,.18) .. (.38,.38);
  \end{scope}
"""
    if grade == 10:
        return r"""
  \begin{scope}[shift={([xshift=-3.05cm,yshift=-2.55cm]current page.north east)}]
    \draw[milcGradeAccent,line width=1.3pt,rounded corners=1.5mm] (-.46,-.42) rectangle (.46,.24);
    \draw[milcGradeAccent,line width=1.3pt] (-.28,.24) .. controls (-.28,.72) and (.28,.72) .. (.28,.24);
    \fill[milcGradeAccent] (0,-.10) circle (.06);
    \draw[milcGradeAccent,line width=.9pt] (0,-.10) -- (0,-.26);
  \end{scope}
"""
    if grade == 11:
        return r"""
  \begin{scope}[shift={([xshift=-3.05cm,yshift=-2.55cm]current page.north east)},rotate=35]
    \fill[milcGradeText] (-.18,-.62) -- (.18,-.62) -- (.28,.32) -- (0,.72) -- (-.28,.32) -- cycle;
    \fill[milcGradeAccent] (0,-.86) -- (.17,-.58) -- (-.17,-.58) -- cycle;
    \fill[milcGradeSoft] (0,.28) circle (.09);
  \end{scope}
"""
    return r"""
  \shade[inner color=white,outer color=milcGradeAccent]
    ([xshift=-3.05cm,yshift=-2.35cm]current page.north east) circle (.62cm);
  \draw[white,line width=1.3pt,opacity=.72]
    ([xshift=-3.05cm,yshift=-1.74cm]current page.north east) .. controls +(0,-.55cm) and +(0,.28cm) ..
    ([xshift=-3.05cm,yshift=-2.85cm]current page.north east);
  \fill[milcGradeDark]
    ([xshift=-3.34cm,yshift=-3.02cm]current page.north east) rectangle
    ([xshift=-2.76cm,yshift=-3.34cm]current page.north east);
  \draw[white,opacity=.55,line width=.8pt]
    ([xshift=-3.34cm,yshift=-3.09cm]current page.north east) -- ([xshift=-2.76cm,yshift=-3.09cm]current page.north east)
    ([xshift=-3.34cm,yshift=-3.22cm]current page.north east) -- ([xshift=-2.76cm,yshift=-3.22cm]current page.north east);
  \fill[milcGradeDark!80!black]
    ([xshift=-3.25cm,yshift=-3.34cm]current page.north east) -- ([xshift=-2.85cm,yshift=-3.34cm]current page.north east) --
    ([xshift=-3.00cm,yshift=-3.52cm]current page.north east) -- cycle;
"""


def tex_grade_web_cover(guide: GuideData) -> str:
    period = guide.period
    title_font, title_leading = web_cover_title_size(guide.short_title)
    cover_title = cover_title_linebreaks(guide.short_title)
    grade_font = 124 if guide.grade < 10 else 104
    degree_x = 4.52 if guide.grade < 10 else 6.55
    icon = tex_cover_icon(guide.grade)
    return rf"""
\begin{{document}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[white] (current page.south west) rectangle (current page.north east);
  \path[fill=milcGradePrimary,rounded corners=16mm]
    ([xshift=.55cm,yshift=-.55cm]current page.north west) rectangle
    ([xshift=-.55cm,yshift=3.65cm]current page.south east);

  \node[anchor=north west,text=milcGradeText] at ([xshift=2.0cm,yshift=-1.85cm]current page.north west)
    {{\sffamily\bfseries\fontsize{{14}}{{16}}\selectfont GRADO}};
  \node[anchor=north west,text=milcGradeText,opacity=.82] at ([xshift=2.0cm,yshift=-2.75cm]current page.north west)
    {{\sffamily\bfseries\fontsize{{9}}{{11}}\selectfont SERIE GUÍAS · TECNOLOGÍA E INFORMÁTICA}};

{icon}

  \node[anchor=west,text=milcGradeText] at ([xshift=1.9cm,yshift=-12.85cm]current page.north west)
    {{\sffamily\bfseries\fontsize{{{grade_font}}}{{{grade_font}}}\selectfont {guide.grade}}};
  \draw[milcGradeText,line width=1.7pt]
    ([xshift={degree_x}cm,yshift=-11.68cm]current page.north west) circle (.13cm);

  \node[anchor=west,text=milcGradeText,text width=8.7cm,align=left] at ([xshift=10.35cm,yshift=-11.6cm]current page.north west)
    {{
     {{\sffamily\bfseries\fontsize{{{title_font}}}{{{title_leading}}}\selectfont {cover_title}}}\\[4mm]
     {{\sffamily\bfseries\fontsize{{23}}{{26}}\selectfont Guía {guide.number}-{guide.grade}-TIC}}\\[2mm]
     {{\sffamily\fontsize{{13}}{{16}}\selectfont Periodo {tex_escape(period.name)}}}\\[5mm]
     {{\sffamily\bfseries\fontsize{{11}}{{13}}\selectfont Institución Educativa Sor María Juliana}}\\[1mm]
     {{\sffamily\fontsize{{10}}{{12}}\selectfont Autor: {tex_escape(AUTHOR)}}}
    }};

  \path[fill=milcGradeSoft,rounded corners=7mm]
    ([xshift=1.35cm,yshift=.85cm]current page.south west) rectangle
    ([xshift=-1.35cm,yshift=3.15cm]current page.south east);
  \draw[milcGradeDark,line width=.65pt,rounded corners=7mm]
    ([xshift=1.35cm,yshift=.85cm]current page.south west) rectangle
    ([xshift=-1.35cm,yshift=3.15cm]current page.south east);
  \node[anchor=center,text=milcNegro,text width=17.0cm,align=center] at ([yshift=2.0cm]current page.south)
    {{\sffamily\fontsize{{10}}{{12}}\selectfont Metodología MILC · Escucha · Sistematización · Praxis · Evaluación liberadora\\
    \textcolor{{milcGradeDark}}{{\bfseries Producto:}} {tex_escape(guide.product)}}};
\end{{tikzpicture}}
\mbox{{}}
\newpage
"""


def tex_cover(guide: GuideData, title_font: int, title_leading: int) -> str:
    return tex_grade_web_cover(guide)


def tex_template(guide: GuideData) -> str:
    concepts = "\n".join(
        f"{tex_escape(a)} & {tex_escape(b)} & {tex_escape(c)}\\\\"
        for a, b, c in concept_rows(guide)
    )
    explanation = "\n\n".join(tex_escape(p) for p in conceptual_paragraphs(guide))
    logic_table = tex_optional_table("Lenguaje lógico mínimo", ("Símbolo", "Se lee", "Ejemplo"), logic_symbol_rows(guide))
    translation_table = tex_optional_table("De frase a condición", ("Frase natural", "Pistas", "Condición escrita"), translation_rows(guide))
    autonomous_table = tex_optional_table("Lista de cotejo autónoma", ("Reviso", "Señal esperada", "Acción autónoma"), autonomous_checklist_rows(guide), "milcVerde")
    examples = "\n\n".join(
        rf"\textbf{{{tex_escape(title)}}}: {tex_escape(body)}"
        for title, body in worked_examples(guide)
    )
    test_table = tex_optional_table("Mini tabla de pruebas", ("Tipo", "Valor", "Resultado esperado"), test_case_rows(guide), "milcMostaza")
    practice = tex_itemize(guided_practice_items(guide))
    criteria = tex_itemize(product_criteria(guide))
    escuta_checks = "\n".join(rf"\checkbox {tex_escape(item)}\\" for item in escuta_check_items(guide))
    prompt_a, prompt_b = escuta_response_prompts(guide)
    praxis_box_title = "Mapa de diagnóstico" if guide_family(guide) == "systems_room" else "Dibuja y explica"
    if is_noveno_editorial_i(guide):
        praxis_box_title = "Boceto editorial"
    praxis_title_label = "Título de mi dibujo:"
    praxis_explanation_label = "Mi explicación:"
    evaluation_sense = "Evaluar no es solo revisar una nota. Es reconocer qué aprendí, cómo aporto al grupo y qué compromiso asumo con el cuidado de un espacio común."
    if is_noveno_editorial_i(guide):
        praxis_title_label = "Título de mi publicación:"
        praxis_explanation_label = "Justificación de mi diseño:"
        evaluation_sense = "Evaluar no es solo revisar una nota. Es reconocer cómo leo, diseño, cito, reviso y mejoro una publicación para comunicar con claridad y responsabilidad."
    if is_undecimo_process_inventory(guide):
        praxis_box_title = "Inventario de procesos"
        praxis_title_label = "Proceso priorizado:"
        praxis_explanation_label = "Justificación empresarial:"
        evaluation_sense = "Evaluar no es solo revisar una nota. Es reconocer cómo observo procesos, protejo datos, priorizo mejoras y sustento oportunidades digitales o con IA."
    extra_how_to_use = tex_extra_after_how_to_use(guide)
    extra_escuta = tex_extra_after_escuta(guide)
    extra_practice = tex_extra_after_practice(guide)
    extra_phase2_bridge = tex_extra_after_phase2_bridge(guide)
    extra_praxis = tex_extra_after_praxis(guide)
    extra_product = tex_extra_after_product(guide)
    evaluation = "\n".join(
        f"{tex_escape(a)} & {tex_escape(b)} & {tex_escape(c)} & {tex_escape(d)}\\\\"
        for a, b, c, d in evaluation_rows(guide)
    )
    period = guide.period
    title_font, title_leading = cover_title_size(guide.short_title)
    cover = tex_cover(guide, title_font, title_leading)
    grade_colors = tex_grade_color_definitions(guide.grade)
    return rf"""
\documentclass[11pt,letterpaper]{{article}}
\usepackage[margin=1.75cm]{{geometry}}
\usepackage[table]{{xcolor}}
\usepackage{{fontspec}}
\usepackage{{tikz}}
\usepackage[most]{{tcolorbox}}
\usepackage{{tabularx}}
\usepackage{{array}}
\usepackage{{fancyhdr}}
\usepackage{{titlesec}}
\usepackage{{ragged2e}}
\usepackage{{enumitem}}

\setmainfont{{Helvetica Neue}}
\setsansfont{{Helvetica Neue}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{6pt}}
\setlength{{\emergencystretch}}{{3em}}
\hyphenpenalty=10000
\exhyphenpenalty=10000
\pretolerance=10000
\renewcommand{{\arraystretch}}{{1.25}}
\setlist{{leftmargin=5mm,itemsep=1mm,topsep=1mm}}
\newcolumntype{{Y}}{{>{{\RaggedRight\arraybackslash}}X}}

\definecolor{{milcMagenta}}{{HTML}}{{E6007E}}
\definecolor{{milcVino}}{{HTML}}{{5A0038}}
\definecolor{{milcTurquesa}}{{HTML}}{{0093A5}}
\definecolor{{milcVerde}}{{HTML}}{{58A923}}
\definecolor{{milcMostaza}}{{HTML}}{{E5B400}}
\definecolor{{milcOcre}}{{HTML}}{{B86600}}
\definecolor{{milcNegro}}{{HTML}}{{241816}}
\definecolor{{milcGris}}{{HTML}}{{F7F7F4}}
\definecolor{{milcLinea}}{{HTML}}{{E8E2DD}}
{grade_colors}
\color{{milcNegro}}

\pagestyle{{fancy}}
\fancyhf{{}}
\lhead{{\small Tecnología e Informática · Grado {guide.grade}}}
\rhead{{\small Guía {guide.number} · MILC}}
\cfoot{{\small\thepage}}
\renewcommand{{\headrulewidth}}{{0pt}}

\newtcolorbox{{titlebox}}[1]{{
  enhanced,colback=#1,colframe=#1,arc=7mm,boxrule=0pt,
  left=7mm,right=7mm,top=3mm,bottom=3mm,
  before skip=8pt,after skip=8pt,
  fontupper=\sffamily\bfseries\Large\color{{white}}
}}

\newtcolorbox{{softbox}}[3]{{
  enhanced,colback=#2,colframe=#1,borderline west={{4pt}}{{0pt}}{{#1}},
  arc=2mm,boxrule=.4pt,left=4mm,right=4mm,top=3mm,bottom=3mm,
  before skip=6pt,after skip=8pt,title={{#3}},coltitle=white,
  colbacktitle=#1,fonttitle=\sffamily\bfseries,fontupper=\RaggedRight,
  attach boxed title to top left={{xshift=3mm,yshift=-2mm}},
  boxed title style={{arc=2mm, boxrule=0pt}}
}}

\newtcolorbox{{drawbox}}[2]{{
  enhanced,colback=white,colframe=#1,arc=4mm,boxrule=1pt,
  left=5mm,right=5mm,top=4mm,bottom=4mm,before skip=8pt,after skip=8pt,
  title={{#2}},coltitle=white,colbacktitle=#1,fonttitle=\sffamily\bfseries,
  fontupper=\RaggedRight,
  attach boxed title to top left={{xshift=4mm,yshift=-2mm}},
  boxed title style={{arc=3mm, boxrule=0pt}}
}}

\newcommand{{\respuesta}}[1]{{\par\vspace{{2mm}}\foreach \n in {{1,...,#1}}{{\noindent\color{{milcLinea}}\rule{{\linewidth}}{{0.5pt}}\par\vspace{{5mm}}}}\color{{milcNegro}}}}
\newcommand{{\checkbox}}{{\textcolor{{milcGradeDark}}{{\fbox{{\phantom{{x}}}}}}\hspace{{5pt}}}}

\newcommand{{\phasecard}}[4]{{
\begin{{tcolorbox}}[enhanced,colback=#3,colframe=#2,arc=3mm,boxrule=.5pt,height=3.05cm,valign=center,halign=center,left=1.5mm,right=1.5mm,top=2mm,bottom=2mm]
{{\sffamily\bfseries\fontsize{{22}}{{24}}\selectfont\textcolor{{#2}}{{#1}}}}\par
{{\sffamily\bfseries\small #4}}
\end{{tcolorbox}}
}}

{cover}
\section*{{El contexto de esta cartilla}}
Esta guía hace parte del proceso de Tecnología e Informática de {grade_display(guide.grade)}. El tema de hoy, \textbf{{{tex_escape(guide.focus)}}}, nos ayuda a usar la tecnología con más criterio, cuidado y propósito.

El aprendizaje no se limita a repetir instrucciones. En MILC observamos el contexto, organizamos ideas, actuamos sobre situaciones reales y reflexionamos sobre el sentido humano de lo que hacemos.

\begin{{softbox}}{{milcVerde}}{{milcGris}}{{Saber y saber hacer}}
{tex_escape(learning_goal_text(guide))}
\end{{softbox}}

\begin{{softbox}}{{milcMostaza}}{{milcGris}}{{Pregunta guía}}
{tex_escape(guide.question)}
\end{{softbox}}

\small
\begin{{tabularx}}{{\textwidth}}{{>{{\bfseries}}p{{3.0cm}}Y}}
\rowcolor{{milcGradePrimary!12}}Autor & {tex_escape(AUTHOR)}\\
DBA & {tex_escape(period.dba)}\\
Referentes & {tex_escape(period.referentes)}\\
Producto final & {tex_escape(guide.product)}\\
\end{{tabularx}}
\normalsize

\begin{{titlebox}}{{milcGradeDark}}
Ruta MILC: así vamos a aprender
\end{{titlebox}}
\begin{{tabularx}}{{\textwidth}}{{>{{\centering\arraybackslash}}X>{{\centering\arraybackslash}}X>{{\centering\arraybackslash}}X>{{\centering\arraybackslash}}X}}
\phasecard{{1}}{{milcVerde}}{{milcGris}}{{Escucha\\[-1mm]\small Observo, escucho y pregunto.}} &
\phasecard{{2}}{{milcTurquesa}}{{milcGris}}{{Sistematización\\[-1mm]\small Organizo ideas y conceptos.}} &
\phasecard{{3}}{{milcMagenta}}{{milcGris}}{{Praxis\\[-1mm]\small Actúo, resuelvo y construyo.}} &
\phasecard{{4}}{{milcVino}}{{milcGris}}{{Evaluación\\[-1mm]\small Reflexiono y me comprometo.}}
\end{{tabularx}}

\begin{{softbox}}{{milcTurquesa}}{{milcGris}}{{Cómo usar esta guía}}
{tex_escape(how_to_use_text(guide))}
\end{{softbox}}

\newpage
{extra_how_to_use}
\begin{{titlebox}}{{milcVerde}}
Fase 1 · Escucha
\end{{titlebox}}
\begin{{softbox}}{{milcVerde}}{{milcGris}}{{Reto de observación}}
{tex_escape(escuta_challenge_text(guide))}
\end{{softbox}}
{escuta_checks}

\textbf{{{tex_escape(prompt_a)}}}
\respuesta{{2}}
\textbf{{{tex_escape(prompt_b)}}}
\respuesta{{2}}

\begin{{softbox}}{{milcVerde}}{{milcGris}}{{Escuchar también es observar}}
En MILC, la escuta no es solo oír. Es leer el contexto, reconocer necesidades y abrir espacio para que distintas voces ayuden a comprender mejor.
\end{{softbox}}
{extra_escuta}

\begin{{titlebox}}{{milcTurquesa}}
Fase 2 · Sistematización
\end{{titlebox}}
\begin{{softbox}}{{milcTurquesa}}{{milcGris}}{{Explicación clave}}
{explanation}
\end{{softbox}}
{logic_table}
{translation_table}
\small
\begin{{tabularx}}{{\textwidth}}{{>{{\bfseries\RaggedRight\arraybackslash}}p{{3.25cm}}YY}}
\rowcolor{{milcTurquesa!90}}\color{{white}}Concepto & \color{{white}}\bfseries Qué significa & \color{{white}}\bfseries Ejemplo\\
{concepts}
\end{{tabularx}}
\normalsize
{autonomous_table}

\begin{{drawbox}}{{milcTurquesa}}{{Ejemplos trabajados}}
{examples}
\end{{drawbox}}

\begin{{softbox}}{{milcMostaza}}{{milcGris}}{{Error común y corrección}}
{tex_escape(common_error(guide))}
\end{{softbox}}
{test_table}

\begin{{drawbox}}{{milcTurquesa}}{{Práctica guiada}}
{practice}
\textbf{{Respuesta breve o esquema:}}
\respuesta{{2}}
\end{{drawbox}}
{extra_practice}

\begin{{softbox}}{{milcTurquesa}}{{milcGris}}{{Ideas del periodo}}
{tex_escape(period_idea_text(guide))}
\end{{softbox}}
\textbf{{Mi idea más importante de esta fase:}}
\respuesta{{2}}
{extra_phase2_bridge}

\newpage
\begin{{titlebox}}{{milcMagenta}}
Fase 3 · Praxis
\end{{titlebox}}
\begin{{softbox}}{{milcMagenta}}{{milcGris}}{{Caso}}
{tex_escape(case_text(guide))}
\end{{softbox}}
\begin{{tabularx}}{{\textwidth}}{{YY}}
\textbf{{Problema identificado:}}
\respuesta{{1}}
\textbf{{Acción responsable:}}
\respuesta{{2}}
&
\textbf{{A quién afecta:}}
\respuesta{{1}}
\textbf{{Cómo prevenirlo:}}
\respuesta{{2}}
\end{{tabularx}}

\begin{{drawbox}}{{milcMagenta}}{{{tex_escape(praxis_box_title)}}}
{tex_escape(praxis_task_text(guide))}
\vspace{{4.7cm}}

{{\small\bfseries {tex_escape(praxis_title_label)}}} \textcolor{{milcLinea}}{{\rule{{.68\linewidth}}{{0.5pt}}}}

{{\small\bfseries {tex_escape(praxis_explanation_label)}}}
\respuesta{{2}}
\end{{drawbox}}
{extra_praxis}

\begin{{titlebox}}{{milcMostaza}}
Producto de la sesión
\end{{titlebox}}
{tex_escape(guide.product)}
\begin{{softbox}}{{milcMostaza}}{{milcGris}}{{Criterios de calidad}}
{criteria}
\end{{softbox}}
{extra_product}
\textbf{{Plan rápido de mi producto:}}
\respuesta{{2}}

\newpage
\begin{{titlebox}}{{milcVino}}
Fase 4 · Evaluación liberadora
\end{{titlebox}}
\begin{{softbox}}{{milcVino}}{{milcGris}}{{Sentido de la evaluación}}
{tex_escape(evaluation_sense)}
\end{{softbox}}
\small
\begin{{tabularx}}{{\textwidth}}{{>{{\RaggedRight\arraybackslash}}p{{3.0cm}}YYY}}
\rowcolor{{milcVino}}\color{{white}}\bfseries Criterio & \color{{white}}\bfseries Lo logré & \color{{white}}\bfseries En proceso & \color{{white}}\bfseries Necesito apoyo\\
{evaluation}
\end{{tabularx}}
\normalsize
\textbf{{Hoy entendí que:}}
\respuesta{{2}}
\textbf{{Mi compromiso para la próxima clase es:}}
\respuesta{{2}}
\begin{{drawbox}}{{milcVino}}{{Mi símbolo de compromiso}}
Dibuja un pequeño símbolo que represente cómo aplicarás este aprendizaje.
\vspace{{3.0cm}}

{{\small\bfseries Mi símbolo significa:}} \textcolor{{milcLinea}}{{\rule{{.65\linewidth}}{{0.5pt}}}}
\end{{drawbox}}
\end{{document}}
"""


def rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "241816", size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_box(doc: Document, title: str, body: str, color: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.cell(0, 0), table.cell(0, 1)
    shade_cell(left, color)
    shade_cell(right, "F7F7F4")
    set_cell_text(left, title, bold=True, color="FFFFFF", size=10)
    set_cell_text(right, body, size=10)
    doc.add_paragraph()


def add_docx_lines(doc: Document, count: int) -> None:
    for _ in range(count):
        doc.add_paragraph("_" * 72)


def write_docx(guide: GuideData, path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{guide.number}-{guide.grade}-TIC · {guide.short_title}")
    run.bold = True
    run.font.size = Pt(docx_title_size(guide.short_title))
    run.font.color.rgb = rgb(docx_cover_color(guide.grade))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Grado {guide.grade} · Periodo {guide.period.name} · Autor: {AUTHOR}").bold = True

    add_docx_box(doc, "Pregunta guía", guide.question, COLORS["mostaza"])
    add_docx_box(doc, "Producto", guide.product, COLORS["magenta"])
    add_docx_box(doc, "DBA", guide.period.dba, COLORS["vino"])
    add_docx_box(doc, "Saber y saber hacer", learning_goal_text(guide), COLORS["verde"])
    add_docx_box(doc, "Cómo usar esta guía", how_to_use_text(guide), COLORS["turquesa"])
    if is_noveno_editorial_i(guide):
        add_docx_box(
            doc,
            "Aprendizaje autónomo",
            "Meta: diseñar un boceto editorial entendible sin explicación oral. Evidencia: retícula, mensaje central, imagen citada, revisión de legibilidad y dos decisiones justificadas.",
            COLORS["verde"],
        )
    if is_undecimo_process_inventory(guide):
        add_docx_box(
            doc,
            "Aprendizaje autónomo",
            "Meta: inventariar procesos de un caso o emprendimiento. Evidencia: tres procesos, un proceso priorizado, punto crítico, dato sensible y oportunidad digital o con IA.",
            COLORS["verde"],
        )

    for label, color, body in [
        ("Fase 1 · Escucha", COLORS["verde"], escuta_challenge_text(guide)),
        ("Fase 2 · Sistematización", COLORS["turquesa"], "Organiza conceptos, ejemplos y relaciones importantes."),
        ("Fase 3 · Praxis", COLORS["magenta"], "Aplica lo aprendido en una situación concreta y produce una evidencia."),
        ("Fase 4 · Evaluación liberadora", COLORS["vino"], "Reflexiona sobre lo aprendido y formula un compromiso responsable."),
    ]:
        add_docx_box(doc, label, body, color)
        if "Escucha" in label:
            for item in escuta_check_items(guide):
                doc.add_paragraph(f"☐ {item}")
            prompt_a, prompt_b = escuta_response_prompts(guide)
            doc.add_paragraph(prompt_a)
            add_docx_lines(doc, 2)
            doc.add_paragraph(prompt_b)
            add_docx_lines(doc, 1)
            if is_noveno_editorial_i(guide):
                doc.add_paragraph("Mini investigación:").runs[0].bold = True
                for item in [
                    "Observa una publicación real e identifica audiencia, mensaje central, recurso visual y fuente.",
                    "Escribe dos hallazgos y una decisión que aplicarás en tu boceto.",
                ]:
                    doc.add_paragraph(item, style=None)
                add_docx_lines(doc, 3)
            if is_undecimo_process_inventory(guide):
                doc.add_paragraph("Mini investigación empresarial:").runs[0].bold = True
                for item in [
                    "Observa o entrevista sobre un proceso: qué lo inicia, quién responde, dónde falla y qué dato debe cuidarse.",
                    "Registra dos evidencias y una pregunta para la siguiente guía de mapeo.",
                ]:
                    doc.add_paragraph(item, style=None)
                add_docx_lines(doc, 3)
        elif "Sistematización" in label:
            doc.add_paragraph("Explicación clave:").runs[0].bold = True
            for paragraph in conceptual_paragraphs(guide):
                doc.add_paragraph(paragraph)
            if logic_symbol_rows(guide):
                doc.add_paragraph("Lenguaje lógico mínimo:").runs[0].bold = True
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                for i, h in enumerate(["Símbolo", "Se lee", "Ejemplo"]):
                    shade_cell(table.cell(0, i), COLORS["turquesa"])
                    set_cell_text(table.cell(0, i), h, True, "FFFFFF")
                for row in logic_symbol_rows(guide):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        set_cell_text(cells[i], value, bold=(i == 0))
            if translation_rows(guide):
                doc.add_paragraph("De frase a condición:").runs[0].bold = True
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                for i, h in enumerate(["Frase natural", "Pistas", "Condición escrita"]):
                    shade_cell(table.cell(0, i), COLORS["turquesa"])
                    set_cell_text(table.cell(0, i), h, True, "FFFFFF")
                for row in translation_rows(guide):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        set_cell_text(cells[i], value, bold=(i == 0))
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for i, h in enumerate(["Concepto", "Qué significa", "Ejemplo"]):
                shade_cell(table.cell(0, i), COLORS["turquesa"])
                set_cell_text(table.cell(0, i), h, True, "FFFFFF")
            for row in concept_rows(guide):
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    set_cell_text(cells[i], value, bold=(i == 0))
            if autonomous_checklist_rows(guide):
                doc.add_paragraph("Lista de cotejo autónoma:").runs[0].bold = True
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                for i, h in enumerate(["Reviso", "Señal esperada", "Acción autónoma"]):
                    shade_cell(table.cell(0, i), COLORS["verde"])
                    set_cell_text(table.cell(0, i), h, True, "FFFFFF")
                for row in autonomous_checklist_rows(guide):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        set_cell_text(cells[i], value, bold=(i == 0))
            doc.add_paragraph("Ejemplos trabajados:").runs[0].bold = True
            for title, body in worked_examples(guide):
                p = doc.add_paragraph()
                p.add_run(f"{title}: ").bold = True
                p.add_run(body)
            add_docx_box(doc, "Error común", common_error(guide), COLORS["mostaza"])
            if test_case_rows(guide):
                doc.add_paragraph("Mini tabla de pruebas:").runs[0].bold = True
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                for i, h in enumerate(["Tipo", "Valor", "Resultado esperado"]):
                    shade_cell(table.cell(0, i), COLORS["mostaza"])
                    set_cell_text(table.cell(0, i), h, True, "FFFFFF")
                for row in test_case_rows(guide):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        set_cell_text(cells[i], value, bold=(i == 0))
            doc.add_paragraph("Práctica guiada:").runs[0].bold = True
            for item in guided_practice_items(guide):
                doc.add_paragraph(item, style=None)
            if is_noveno_editorial_i(guide):
                doc.add_paragraph("Auto chequeo antes de la praxis:").runs[0].bold = True
                for item in [
                    "Tema útil para una audiencia concreta.",
                    "Mensaje central escrito en una frase clara.",
                    "Boceto con orden de lectura visible.",
                    "Imagen o dato con fuente, autor o licencia.",
                    "Mejora prevista si el diseño queda saturado o confuso.",
                ]:
                    doc.add_paragraph(f"☐ {item}")
            if is_undecimo_process_inventory(guide):
                doc.add_paragraph("Matriz de inventario inicial:").runs[0].bold = True
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                for i, h in enumerate(["Proceso", "Entrada", "Responsable", "Salida", "Punto crítico"]):
                    shade_cell(table.cell(0, i), COLORS["mostaza"])
                    set_cell_text(table.cell(0, i), h, True, "FFFFFF")
                for label in ["Proceso 1", "Proceso 2", "Proceso 3"]:
                    cells = table.add_row().cells
                    set_cell_text(cells[0], label, True)
                    for cell in cells[1:]:
                        set_cell_text(cell, "")
                doc.add_paragraph("Auto chequeo antes de priorizar:").runs[0].bold = True
                for item in [
                    "Mínimo tres procesos identificados.",
                    "Cada proceso tiene entrada, responsable y salida verificable.",
                    "Punto crítico respaldado con evidencia.",
                    "Datos sensibles y acceso responsable señalados.",
                    "Herramienta futura justificada por necesidad real.",
                ]:
                    doc.add_paragraph(f"☐ {item}")
            add_docx_box(doc, "Ideas del periodo", period_idea_text(guide), COLORS["turquesa"])
            doc.add_paragraph("Mi idea más importante:")
            add_docx_lines(doc, 2)
            if is_undecimo_process_inventory(guide):
                add_docx_box(
                    doc,
                    "Bitácora de decisión",
                    "Antes de pasar a la praxis, convierte tu observación en una decisión empresarial. "
                    "Justifica qué proceso merece atención primero antes de escoger una herramienta.",
                    COLORS["turquesa"],
                )
                for item in [
                    "Hallazgo más importante observado o escuchado.",
                    "Evidencia que lo demuestra.",
                    "Dato sensible que debe protegerse.",
                    "Proceso que parece más urgente de revisar.",
                    "Pregunta que le haría a una IA para ordenar la información, sin entregar datos reales.",
                ]:
                    doc.add_paragraph(item)
                    add_docx_lines(doc, 1)
                doc.add_paragraph("Semáforo de preparación:").runs[0].bold = True
                for item in [
                    "Puedo explicar el proceso sin depender de una herramienta.",
                    "Tengo al menos una evidencia: observación, entrevista, registro o ejemplo simulado.",
                    "Reconozco quién participa y quién recibe el resultado.",
                    "Sé qué información no debo publicar ni copiar en una IA.",
                    "La mejora que imagino responde a un problema real del caso.",
                ]:
                    doc.add_paragraph(f"☐ {item}")
        elif "Praxis" in label:
            add_docx_box(doc, "Caso", case_text(guide), COLORS["magenta"])
            doc.add_paragraph(praxis_task_text(guide))
            add_docx_lines(doc, 8)
            if is_noveno_editorial_i(guide):
                doc.add_paragraph("Bitácora de edición:").runs[0].bold = True
                for item in ["Problema de la versión 1", "Mejora aplicada", "Criterio usado", "Evidencia"]:
                    doc.add_paragraph(item)
                    add_docx_lines(doc, 1)
            if is_undecimo_process_inventory(guide):
                doc.add_paragraph("Priorización de mejora:").runs[0].bold = True
                for item in ["Proceso elegido", "Impacto", "Facilidad", "Riesgo y datos", "Oportunidad futura"]:
                    doc.add_paragraph(item)
                    add_docx_lines(doc, 1)
        else:
            doc.add_paragraph("Criterios específicos del producto:").runs[0].bold = True
            for item in product_criteria(guide):
                doc.add_paragraph(item, style=None)
            if is_undecimo_process_inventory(guide):
                add_docx_box(
                    doc,
                    "Protocolo mínimo de datos e IA",
                    "Antes de usar formularios, hojas de cálculo o IA, el equipo debe definir finalidad, "
                    "recolectar solo datos necesarios, evitar información sensible innecesaria, limitar accesos "
                    "y verificar cualquier recomendación generada por IA con evidencia del proceso.",
                    COLORS["vino"],
                )
            doc.add_paragraph("Autoevaluación:").runs[0].bold = True
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["Criterio", "Lo logré", "En proceso", "Necesito apoyo"]):
                shade_cell(table.cell(0, i), COLORS["vino"])
                set_cell_text(table.cell(0, i), h, True, "FFFFFF")
            for row in evaluation_rows(guide):
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    set_cell_text(cells[i], value, bold=(i == 0))
            doc.add_paragraph("Hoy entendí que:")
            add_docx_lines(doc, 2)
            doc.add_paragraph("Mi compromiso:")
            add_docx_lines(doc, 2)
            doc.add_paragraph("Mi símbolo de compromiso:")
            add_docx_lines(doc, 5)
    doc.save(path)


def setup_dirs() -> dict[str, Path]:
    scripts_dir = ROOT / "05_SCRIPTS_GENERADORES" if (ROOT / "05_SCRIPTS_GENERADORES").exists() else OUT.parent / "scripts"
    dirs = {
        "base": OUT,
        "fuente": OUT / "fuente",
        "tex": OUT / "tex",
        "pdf": OUT / "pdf",
        "docx": OUT / "docx",
        "doc": OUT / "doc",
        "build": OUT / "build",
        "scripts": scripts_dir,
    }
    for directory in dirs.values():
        directory.mkdir(parents=True, exist_ok=True)
    return dirs


def convert_docx_to_doc(docx_path: Path, doc_path: Path) -> tuple[bool, str]:
    cmd = ["textutil", "-convert", "doc", "-output", str(doc_path), str(docx_path)]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        return False, (proc.stderr or proc.stdout or "textutil error")[-1000:]
    return True, "ok"


def compile_tex(tex_path: Path, pdf_out: Path, build_dir: Path) -> tuple[bool, str]:
    job_dir = build_dir / tex_path.stem
    if job_dir.exists():
        shutil.rmtree(job_dir)
    job_dir.mkdir(parents=True)
    work_tex = job_dir / tex_path.name
    shutil.copy2(tex_path, work_tex)
    tectonic_message = ""
    if TECTONIC.exists():
        proc = subprocess.run([str(TECTONIC), "--outdir", str(job_dir), work_tex.name], cwd=job_dir, capture_output=True, text=True)
        combined_log = (proc.stdout + "\n" + proc.stderr)[-4000:]
        if proc.returncode == 0:
            built_pdf = job_dir / f"{tex_path.stem}.pdf"
            if built_pdf.exists():
                shutil.copy2(built_pdf, pdf_out)
                return True, "ok"
            tectonic_message = "Tectonic terminó sin generar PDF."
        else:
            tectonic_message = combined_log[-1200:]

    cmd = [XELATEX, "-interaction=nonstopmode", "-halt-on-error", work_tex.name]
    combined_log = ""
    for _ in range(2):
        proc = subprocess.run(cmd, cwd=job_dir, capture_output=True, text=True)
        combined_log = (combined_log + "\n" + proc.stdout + "\n" + proc.stderr)[-4000:]
        if proc.returncode != 0:
            message = combined_log[-2000:]
            if tectonic_message:
                message = f"Tectonic falló; XeLaTeX también falló.\n{tectonic_message}\n{message}"[-3000:]
            return False, message
    built_pdf = job_dir / f"{tex_path.stem}.pdf"
    shutil.copy2(built_pdf, pdf_out)
    return True, "ok"


def main() -> None:
    dirs = setup_dirs()
    periods = extract_period_data()
    topics = extract_topics()
    guides = build_guides(periods, topics)
    current_grade_name = grade_name(GRADE)

    temario = [
        {
            "guia": g.number,
            "grado": g.grade,
            "periodo": g.period.number,
            "periodo_nombre": g.period.name,
            "titulo": g.title,
            "titulo_corto": g.short_title,
            "pregunta": g.question,
            "producto": g.product,
        }
        for g in guides
    ]
    (dirs["fuente"] / f"temario_{current_grade_name}.json").write_text(json.dumps(temario, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest_rows = []
    for guide in guides:
        stem = f"{guide.number}-{guide.grade}-TIC"
        tex_path = dirs["tex"] / f"{stem}.tex"
        pdf_path = dirs["pdf"] / f"{stem}.pdf"
        docx_path = dirs["docx"] / f"{stem}.docx"
        doc_path = dirs["doc"] / f"{stem}.doc"
        tex_path.write_text(tex_template(guide), encoding="utf-8")
        write_docx(guide, docx_path)
        doc_ok, doc_message = convert_docx_to_doc(docx_path, doc_path)
        ok, message = compile_tex(tex_path, pdf_path, dirs["build"])
        manifest_rows.append(
            {
                "guia": stem,
                "periodo": guide.period.name,
                "titulo": guide.title,
                "tex": str(tex_path),
                "pdf": str(pdf_path) if ok else "",
                "docx": str(docx_path),
                "doc": str(doc_path) if doc_ok else "",
                "estado_pdf": "ok" if ok else "error",
                "estado_doc": "ok" if doc_ok else "error",
                "mensaje": message if ok else message,
                "mensaje_doc": doc_message,
            }
        )
        print(f"{stem}: {'PDF ok' if ok else 'PDF error'} · DOCX ok · {'DOC ok' if doc_ok else 'DOC error'}")

    manifest_path = OUT / f"manifest_{current_grade_name}.csv"
    with manifest_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(manifest_rows[0].keys()))
        writer.writeheader()
        writer.writerows(manifest_rows)

    script_copy = dirs["scripts"] / f"generar_guias_{current_grade_name}.py"
    if Path(__file__).resolve() != script_copy.resolve():
        shutil.copy2(Path(__file__), script_copy)
    print(f"\nSalida: {OUT}")
    print(f"Manifest: {manifest_path}")


if __name__ == "__main__":
    main()

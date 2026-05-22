from __future__ import annotations

import csv
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PLAN = Path(os.environ.get("MILC_PLAN", ROOT / "00_PLAN_AREA" / "plan_area_tecnologia_2025.docx"))
RESOURCES = ROOT
OUT = ROOT / "00_PLAN_AREA" / "plan_area_2025_editorial_milc"
XELATEX = os.environ.get("XELATEX", "xelatex")
AUTHOR = "Dr. Álvaro Cárdenas Orozco"
PLATFORM_URL = "https://alguarito.github.io/plataformaconectate/"

GRADE_NAMES = {
    6: "sexto",
    7: "septimo",
    8: "octavo",
    9: "noveno",
    10: "decimo",
    11: "undecimo",
}
RESOURCE_GRADE_FOLDERS = {
    6: "06_SEXTO",
    7: "07_SEPTIMO",
    8: "08_OCTAVO",
    9: "09_NOVENO",
    10: "10_DECIMO",
    11: "11_UNDECIMO",
}
GRADE_LABELS = {
    6: "Sexto",
    7: "Séptimo",
    8: "Octavo",
    9: "Noveno",
    10: "Décimo",
    11: "Undécimo",
}
TABLE_START = {6: 0, 7: 3, 8: 6, 9: 9, 10: 12, 11: 15}
PERIOD_NAMES = {1: "Primero", 2: "Segundo", 3: "Tercero"}

PALETTE = {
    "magenta": "E6007E",
    "vino": "5A0038",
    "verde": "58A923",
    "turquesa": "0093A5",
    "mostaza": "E5B400",
    "gris": "F4F2F4",
    "negro": "241A1D",
}

GRADE_BRAND_COLORS = {
    6: {
        "name": "Azul plataforma",
        "hex": "2B65F6",
        "text": "FFFFFF",
        "focus": "Alfabetización digital y comunicación",
    },
    7: {
        "name": "Verde plataforma",
        "hex": "B6FA51",
        "text": "111807",
        "focus": "Trabajo colaborativo, algoritmia e IA",
    },
    8: {
        "name": "Naranja plataforma",
        "hex": "ED7446",
        "text": "FFFFFF",
        "focus": "Excel, lógica de control y multimedia",
    },
    9: {
        "name": "Violeta plataforma",
        "hex": "733DE3",
        "text": "FFFFFF",
        "focus": "Historia de la tecnología, diseño y datos",
    },
    10: {
        "name": "Fucsia plataforma",
        "hex": "EB498A",
        "text": "FFFFFF",
        "focus": "Ciberseguridad, arquitectura y Python",
    },
    11: {
        "name": "Cian plataforma",
        "hex": "60D1F9",
        "text": "070F12",
        "focus": "Desarrollo web, calidad y proyecto final",
    },
}

COMPONENTS = [
    "Naturaleza y evolución de la tecnología",
    "Apropiación y uso de la tecnología",
    "Solución de problemas con tecnología",
    "Tecnología y sociedad",
]

BIBLIOGRAPHY = [
    "Cárdenas Orozco, Á. (2026). ConectaTE: Conectando Tecnología con Educación. Plataforma educativa institucional. https://alguarito.github.io/plataformaconectate/",
    "Collaborative for Academic, Social, and Emotional Learning. (2020). CASEL's SEL framework: What are the core competence areas and where are they promoted?",
    "Computer Science Teachers Association. (2017). CSTA K-12 Computer Science Standards.",
    "Dussel, E. (1998). Ética de la liberación en la edad de la globalización y de la exclusión. Trotta.",
    "Floridi, L. (2014). The Fourth Revolution: How the Infosphere is Reshaping Human Reality. Oxford University Press.",
    "Freire, P. (2005). Pedagogía del oprimido. Siglo XXI Editores. Obra original publicada en 1970.",
    "International Organization for Standardization. (2015). ISO 9001:2015 Quality management systems - Requirements.",
    "International Organization for Standardization. (2018). ISO 21001:2018 Educational organizations - Management systems for educational organizations.",
    "International Society for Technology in Education. (2016). ISTE Standards for Students.",
    "Ministerio de Educación Nacional de Colombia. (2008). Orientaciones generales para la educación en tecnología: Ser competente en tecnología. Guía No. 30.",
    "República de Colombia. (1994). Ley 115 de 1994, Ley General de Educación.",
    "República de Colombia. (2012). Ley 1581 de 2012, régimen general de protección de datos personales.",
    "UNESCO. (2023). Guidance for generative AI in education and research.",
    "Wing, J. M. (2006). Computational thinking. Communications of the ACM, 49(3), 33-35.",
    "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2.",
]


@dataclass
class PeriodData:
    grade: int
    period: int
    asignacion: str
    dba: str
    referentes: str
    estandares: str
    competencias: str
    aprendizajes: str
    cognitivos: str
    procedimentales: str
    actitudinales: str
    topics: list[str]


@dataclass
class ComponentRow:
    standard: str
    competencias: str
    aprendizajes: str
    cognitivos: str
    procedimentales: str
    actitudinales: str


@dataclass
class PrelimSection:
    title: str
    color: str
    paragraphs: list[str]
    bullets: list[str]
    callout_title: str = ""
    callout: str = ""


def clean(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split()).strip()


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in str(text))


def compact(text: str, max_len: int = 520) -> str:
    text = clean(text)
    if len(text) <= max_len:
        return text
    cut = text[:max_len].rsplit(" ", 1)[0]
    return cut.rstrip(".,;:") + "."


def sentences(text: str, limit: int = 2) -> str:
    parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", clean(text)) if p.strip()]
    return " ".join(parts[:limit]) if parts else clean(text)


def load_topics(grade: int, period: int) -> list[str]:
    folder = GRADE_NAMES[grade]
    resource_folder = RESOURCE_GRADE_FOLDERS[grade]
    path = RESOURCES / "01_GUIAS_APRENDIZAJE" / resource_folder / "fuente" / f"temario_{folder}.json"
    if not path.exists() and grade == 6:
        path = RESOURCES / "01_GUIAS_APRENDIZAJE" / folder / "fuente" / "temario_sexto.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    return [clean(item["titulo"]) for item in data if int(item["periodo"]) == period]


def extract_periods() -> list[PeriodData]:
    doc = Document(str(PLAN))
    periods: list[PeriodData] = []
    for grade in range(6, 12):
        start = TABLE_START[grade]
        for period in (1, 2, 3):
            table = doc.tables[start + period - 1]
            row0 = [clean(c.text) for c in table.rows[0].cells]
            row1 = [clean(c.text) for c in table.rows[1].cells]
            row2 = [clean(c.text) for c in table.rows[2].cells]
            row5 = [clean(c.text) for c in table.rows[5].cells]
            periods.append(
                PeriodData(
                    grade=grade,
                    period=period,
                    asignacion=row0[0],
                    dba=row1[0].replace("DBA:", "").strip(),
                    referentes=row2[0].replace("REFERENTES:", "").strip(),
                    estandares=row5[0],
                    competencias=row5[1],
                    aprendizajes=row5[2],
                    cognitivos=row5[3],
                    procedimentales=row5[4],
                    actitudinales=row5[5],
                    topics=load_topics(grade, period),
                )
            )
    return periods


def infer_emphasis(pdata: PeriodData) -> str:
    special = {
        (9, 3): "alfabetización básica en IA; uso crítico de prompts, verificación y privacidad; lectura ética de sesgos, límites e impacto escolar.",
        (10, 1): "producción editorial digital con IA; construcción de autoría, voz propia y derechos de autor; edición, diseño y publicación responsable.",
        (10, 2): "documentación técnica y comercial con LaTeX; escritura profesional asistida por IA; recolección, análisis y presentación de datos de mercado.",
        (10, 3): "analítica empresarial con hojas de cálculo e IA; visualización mediante dashboards; toma de decisiones comerciales y contables basada en evidencia.",
        (11, 1): "presencia digital empresarial con IA; identidad de marca, arquitectura web y comunicación comercial; SEO, accesibilidad y ética digital.",
        (11, 2): "automatización de procesos empresariales con IA; formularios, bases de datos y flujos de trabajo; calidad, atención al cliente y protección de datos.",
        (11, 3): "proyecto empresarial integrado asistido con IA; presencia digital, datos, dashboard, automatización y pitch; impacto, costos y sostenibilidad.",
    }
    if (pdata.grade, pdata.period) in special:
        return special[(pdata.grade, pdata.period)]

    text = " ".join(pdata.topics + [pdata.aprendizajes, pdata.competencias]).lower()
    emphasis: list[str] = []
    if any(w in text for w in ["hardware", "computador", "sistema", "redes", "internet", "mantenimiento"]):
        emphasis.append("comprensión de sistemas tecnológicos y sus relaciones")
    if any(w in text for w in ["word", "excel", "documento", "ofimática", "datos", "programación", "scratch", "hoja de cálculo"]):
        emphasis.append("uso apropiado de herramientas digitales para producir evidencias")
    if any(w in text for w in ["proyecto", "solución", "problema", "algorit", "prototipo", "diseño"]):
        emphasis.append("resolución de problemas mediante diseño, prototipado y mejora")
    if any(w in text for w in ["ética", "seguridad", "ciudadanía", "sociedad", "impacto", "responsable", "fuentes"]):
        emphasis.append("lectura crítica del impacto social, ético y ambiental de la tecnología")
    if not emphasis:
        emphasis = [
            "integración de los cuatro componentes MEN en situaciones de aprendizaje contextualizadas"
        ]
    return "; ".join(emphasis[:3]) + "."


def active_strategy(pdata: PeriodData) -> str:
    special = {
        (9, 3): "Indagación guiada sobre IA, laboratorio de prompts, verificación de fuentes y reflexión ética.",
        (10, 1): "Escritura Inversa, taller editorial, bitácora de prompts, revisión por pares y feria MILC.",
        (10, 2): "Taller de codificación documental, indagación comercial, recolección de datos y sustentación técnica.",
        (10, 3): "Analítica empresarial, aprendizaje basado en datos, visualización, dashboard y sustentación con evidencia.",
        (11, 1): "Design Thinking, prototipado web, revisión de marca, estrategia de contenidos y sustentación digital.",
        (11, 2): "Mapeo de procesos, ciclo PDCA, automatización guiada, pruebas de flujo y mejora iterativa.",
        (11, 3): "ABP empresarial, integración de evidencias, prototipado, pitch y tribunal MILC.",
    }
    if (pdata.grade, pdata.period) in special:
        return special[(pdata.grade, pdata.period)]

    text = " ".join(pdata.topics).lower()
    if any(w in text for w in ["programación", "scratch", "algorit"]):
        return "PRIMM, retos de depuración, prototipado por bloques y socialización de soluciones."
    if any(w in text for w in ["proyecto", "emprendimiento", "innovación", "prototipo"]):
        return "Design Thinking, aprendizaje basado en proyectos, validación con usuarios y mejora iterativa."
    if any(w in text for w in ["excel", "datos", "hoja"]):
        return "Aprendizaje basado en datos, experimento escolar, visualización y toma de decisiones."
    if any(w in text for w in ["hardware", "redes", "mantenimiento", "computador"]):
        return "Laboratorio técnico, diagnóstico guiado, bitácora de evidencias y protocolo de cuidado."
    return "Indagación guiada MILC, trabajo colaborativo, producción situada y reflexión ética."


def simplified_topics(pdata: PeriodData) -> list[str]:
    items: list[str] = []
    for topic in pdata.topics:
        base = topic.split("—", 1)[0].strip()
        if base and base not in items:
            items.append(base)
    return items


def period_focus(pdata: PeriodData) -> str:
    special = {
        (9, 3): "inteligencia artificial básica, prompts, verificación, sesgos, privacidad y uso responsable",
        (10, 1): "escritura de un libro breve con IA, autoría, edición, derechos de autor y diseño editorial",
        (10, 2): "informes comerciales en LaTeX, datos de mercado, citación y transparencia en el uso de IA",
        (10, 3): "analítica empresarial, visualización de datos, dashboards y decisiones comerciales o contables",
        (11, 1): "presencia digital empresarial, identidad de marca, arquitectura web, SEO y ética con IA",
        (11, 2): "automatización de procesos empresariales, formularios, bases de datos y protección de datos",
        (11, 3): "proyecto empresarial integrado con IA, presencia digital, datos, dashboard, automatización y pitch",
    }
    if (pdata.grade, pdata.period) in special:
        return special[(pdata.grade, pdata.period)]
    match = re.search(r"Tema Central:\s*([^.;]+)", pdata.aprendizajes)
    if match:
        return clean(match.group(1)).rstrip(".")
    topics = simplified_topics(pdata)
    if topics:
        return topics[0].lower()
    return clean(infer_emphasis(pdata)).rstrip(".")


def select_component_topics(pdata: PeriodData, component: str, limit: int = 3) -> list[str]:
    patterns = {
        COMPONENTS[0]: [
            "historia",
            "evolución",
            "concepto",
            "introducción",
            "fundamentos",
            "arquitectura",
            "hardware",
            "sistema",
            "red",
            "datos",
            "procesos",
            "inteligencia artificial",
            "latex",
        ],
        COMPONENTS[1]: [
            "taller",
            "creación",
            "configuración",
            "uso",
            "word",
            "excel",
            "office",
            "google",
            "scratch",
            "makecode",
            "latex",
            "prompts",
            "dashboard",
            "formulario",
            "presentación",
        ],
        COMPONENTS[2]: [
            "proyecto",
            "reto",
            "problema",
            "prototipo",
            "simulación",
            "sistema",
            "diseño",
            "depuración",
            "decisiones",
            "solución",
            "sustentación",
        ],
        COMPONENTS[3]: [
            "ética",
            "seguridad",
            "privacidad",
            "propiedad",
            "derechos",
            "sesgos",
            "impacto",
            "ciudadanía",
            "raee",
            "accesibilidad",
            "licencias",
            "fuentes",
            "protección",
            "transparencia",
        ],
    }
    topics = simplified_topics(pdata)
    selected = [
        topic
        for topic in topics
        if any(pattern in topic.lower() for pattern in patterns.get(component, []))
    ]
    for topic in topics:
        if len(selected) >= limit:
            break
        if topic not in selected:
            selected.append(topic)
    return selected[:limit]


def readable_list(items: list[str]) -> str:
    if not items:
        return "los contenidos del periodo"
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + " y " + items[-1]


def component_rows(pdata: PeriodData) -> list[ComponentRow]:
    focus = period_focus(pdata)
    nature = readable_list(select_component_topics(pdata, COMPONENTS[0]))
    use = readable_list(select_component_topics(pdata, COMPONENTS[1]))
    problem = readable_list(select_component_topics(pdata, COMPONENTS[2]))
    society = readable_list(select_component_topics(pdata, COMPONENTS[3]))

    return [
        ComponentRow(
            standard=COMPONENTS[0],
            competencias=f"Comprende principios, componentes y relaciones del eje del periodo: {focus}.",
            aprendizajes=f"Reconoce conceptos clave en torno a {nature} y los conecta con necesidades reales del contexto escolar o productivo.",
            cognitivos=f"Explica elementos, relaciones y antecedentes que permiten entender {focus}.",
            procedimentales="Elabora esquemas, diagnósticos o registros que ordenan la información técnica estudiada.",
            actitudinales="Demuestra curiosidad, precisión conceptual y disposición para corregir ideas previas con evidencia.",
        ),
        ComponentRow(
            standard=COMPONENTS[1],
            competencias=f"Usa de manera responsable herramientas, lenguajes o plataformas asociadas a {focus}.",
            aprendizajes=f"Aplica procedimientos de trabajo en actividades relacionadas con {use}.",
            cognitivos="Selecciona herramientas y criterios de uso según el propósito, la información y la calidad esperada.",
            procedimentales="Configura, produce, documenta, revisa y entrega recursos digitales con criterios de calidad.",
            actitudinales="Cuida equipos, archivos, datos y tiempos de trabajo; solicita apoyo cuando lo necesita y respeta acuerdos de uso.",
        ),
        ComponentRow(
            standard=COMPONENTS[2],
            competencias=f"Formula y desarrolla respuestas tecnológicas a problemas relacionados con {focus}.",
            aprendizajes=f"Diseña soluciones o productos vinculados con {problem}, integrando investigación, prueba, mejora y sustentación.",
            cognitivos="Analiza necesidades, restricciones, datos y criterios de éxito antes de decidir.",
            procedimentales="Planea, construye, prueba, corrige y comunica una evidencia o producto mediante la ruta MILC del periodo.",
            actitudinales="Asume el error como oportunidad de mejora, coopera con el equipo y persevera hasta obtener una solución defendible.",
        ),
        ComponentRow(
            standard=COMPONENTS[3],
            competencias=f"Valora implicaciones éticas, sociales, ambientales, culturales o empresariales de {focus}.",
            aprendizajes=f"Relaciona situaciones como {society} con privacidad, derechos, inclusión, sostenibilidad, seguridad o impacto comunitario.",
            cognitivos="Identifica riesgos, beneficios, actores afectados y responsabilidades asociadas al uso de la tecnología.",
            procedimentales="Sustenta decisiones, declara fuentes o uso de IA cuando aplique y propone acciones de cuidado o mejora social.",
            actitudinales="Actúa con honestidad intelectual, respeto por otros, sensibilidad frente al contexto y compromiso con un uso responsable.",
        ),
    ]


def resource_row(pdata: PeriodData) -> tuple[str, str, str]:
    start = (pdata.period - 1) * 10 + 1
    end = pdata.period * 10
    guides = ", ".join(f"{i}-{pdata.grade}-TIC" for i in range(start, end + 1))
    project = f"proyecto-{pdata.period}-{pdata.grade}-TIC"
    exam = f"examen-{pdata.period}-{pdata.grade}-TIC"
    return guides, project, exam


def resource_summary(pdata: PeriodData) -> tuple[str, str, str]:
    start = (pdata.period - 1) * 10 + 1
    end = pdata.period * 10
    guides = f"{start}-{pdata.grade}-TIC a {end}-{pdata.grade}-TIC"
    project = f"proyecto-{pdata.period}-{pdata.grade}-TIC"
    exam = f"examen-{pdata.period}-{pdata.grade}-TIC"
    return guides, project, exam


def academic_preliminaries() -> list[PrelimSection]:
    return [
        PrelimSection(
            title="Presentación institucional del área",
            color=PALETTE["vino"],
            paragraphs=[
                "El Plan de Área de Tecnología e Informática 2025 de la Institución Educativa Sor María Juliana se concibe como un documento rector para la formación tecnológica integral de estudiantes de básica secundaria y media. Su propósito no se reduce al uso operativo de programas o dispositivos: organiza una trayectoria de alfabetización tecnológica, pensamiento computacional, ciudadanía digital, producción ética de información y solución situada de problemas.",
                "La propuesta reconoce que los estudiantes habitan una infoesfera donde la vida escolar, familiar, social y productiva se encuentra atravesada por datos, plataformas, algoritmos, redes, automatización e inteligencia artificial. Por ello, aprender tecnología implica comprender sistemas, tomar decisiones informadas, cuidar la identidad digital, crear productos con sentido y evaluar las consecuencias humanas, ambientales y sociales de cada solución.",
            ],
            bullets=[
                "Área: Tecnología e Informática.",
                "Año lectivo: 2025.",
                "Cobertura: grados sexto a undécimo.",
                "Autor: Dr. Álvaro Cárdenas Orozco.",
                "Enfoque: formación tecnológica crítica, creativa, ética y liberadora.",
            ],
            callout_title="Idea fuerza",
            callout="La tecnología no se enseña como un inventario de herramientas, sino como una práctica cultural para saber, cuidar, producir y transformar responsablemente el entorno.",
        ),
        PrelimSection(
            title="Fundamentación curricular y normativa",
            color=PALETTE["turquesa"],
            paragraphs=[
                "El diseño curricular articula los lineamientos nacionales del Ministerio de Educación Nacional, especialmente las Orientaciones Generales para la Educación en Tecnología, conocidas como Guía No. 30, con referentes contemporáneos de alfabetización digital, pensamiento computacional, ciudadanía digital, competencias en inteligencia artificial, gestión de calidad y sostenibilidad tecnológica.",
                "La estructura del área asume los cuatro componentes de la Guía No. 30 como ejes permanentes de lectura curricular: naturaleza y evolución de la tecnología; apropiación y uso de la tecnología; solución de problemas con tecnología; y tecnología y sociedad. Estos componentes no aparecen como contenidos aislados, sino como lentes para leer cada periodo, cada guía, cada proyecto y cada evidencia de evaluación.",
            ],
            bullets=[
                "Constitución Política de Colombia: derecho a la educación, acceso a la ciencia, la cultura y el conocimiento.",
                "Ley 115 de 1994: Tecnología e Informática como área fundamental y obligatoria.",
                "Guía No. 30 MEN: formación en competencias tecnológicas para comprender, usar, transformar y valorar la tecnología.",
                "Ley 1581 de 2012: protección de datos personales y responsabilidad frente a la información.",
                "Referentes actuales: pensamiento computacional, inteligencia artificial, conectividad escolar, cultura maker, sostenibilidad y ciudadanía digital.",
            ],
            callout_title="Criterio curricular",
            callout="Cada unidad debe permitir que el estudiante explique cómo funciona una tecnología, la use con criterio, resuelva un problema y valore su impacto social.",
        ),
        PrelimSection(
            title="Horizonte pedagógico: Modelo MILC",
            color=PALETTE["magenta"],
            paragraphs=[
                "El Modelo MILC, entendido como Modelo de Investigación Liberadora y Científica, es la columna vertebral metodológica del plan. Su valor radica en que convierte la clase de Tecnología e Informática en un proceso de investigación escolar: el estudiante escucha el contexto, organiza lo que encuentra, actúa mediante productos o prototipos y evalúa críticamente el sentido de lo realizado.",
                "MILC evita que la tecnología se vuelva una práctica mecánica, dependiente de instrucciones o de la fascinación acrítica por la novedad. Su orientación liberadora parte de reconocer que toda tecnología configura relaciones de poder: quién accede, quién queda por fuera, quién produce información, quién la interpreta, quién se beneficia y quién asume los riesgos. Por eso, la alfabetización tecnológica debe producir autonomía, no dependencia.",
                "La dimensión científica de MILC exige observación, formulación de preguntas, registro de evidencias, análisis de información, experimentación, prototipado, validación y comunicación de resultados. La dimensión liberadora exige responsabilidad ética, lectura del contexto, cuidado de sí y de los otros, soberanía informacional y compromiso con la comunidad.",
            ],
            bullets=[
                "Escucha: escuchar, observar y problematizar la realidad tecnológica del estudiante y su comunidad.",
                "Sistematización: ordenar información, reconocer patrones, construir criterios, documentar procesos y formular explicaciones.",
                "Praxis: diseñar, experimentar, programar, prototipar, comunicar o intervenir una situación concreta.",
                "Evaluación liberadora: valorar evidencias, impactos, aprendizajes, errores, mejoras y compromisos éticos.",
            ],
            callout_title="MILC como principio de diseño",
            callout="Toda guía debe tener pregunta, contexto, producción, reflexión y evidencia; todo proyecto debe pasar por escucha, análisis, creación y evaluación social.",
        ),
        PrelimSection(
            title="Diagnóstico crítico: educación bancaria, desarraigo y cultura digital",
            color=PALETTE["vino"],
            paragraphs=[
                "El punto de partida de esta propuesta es una crítica a la persistencia de prácticas escolares bancarias: clases centradas en transmisión, copia, obediencia procedimental y evaluación memorística. En ese modelo, el estudiante aparece como receptor de contenidos y no como sujeto que interpreta su realidad, formula preguntas, construye conocimiento y transforma su contexto. En Tecnología e Informática, esta educación bancaria suele tomar una forma muy concreta: aprender a seguir pasos sin comprender sistemas, repetir comandos sin diseñar soluciones, consumir plataformas sin preguntarse por sus efectos y entregar productos sin reflexión ética.",
                "La educación bancaria también produce desarraigo cultural cuando separa el currículo de la historia, los lenguajes, los problemas y las prácticas de la comunidad. Una tecnología enseñada sin territorio puede convertirse en imitación acrítica de modelos externos: se enseña la herramienta, pero no se pregunta para quién sirve, qué problema local atiende, qué saberes comunitarios reconoce, qué exclusiones reproduce o qué formas de cuidado exige. Por ello, el plan propone una educación tecnológica con raíces: situada en la escuela, en Cartago, en el Valle del Cauca, en las familias, en los oficios, en los modos de comunicación, en los riesgos digitales reales y en las posibilidades creativas de los estudiantes.",
                "MILC responde a esta tensión al desplazar el centro de la clase: de la instrucción al problema, de la obediencia al criterio, de la copia a la producción, de la acumulación de tareas al portafolio con sentido, de la calificación aislada a la evaluación liberadora. La tecnología se convierte en mediación para leer el mundo y no únicamente en objeto de entrenamiento instrumental.",
            ],
            bullets=[
                "Problema pedagógico: aprendizaje pasivo, baja autonomía y dependencia de instrucciones.",
                "Problema cultural: currículo descontextualizado, débil reconocimiento del territorio y de los saberes de la comunidad.",
                "Problema tecnológico: consumo acrítico de plataformas, fragilidad en seguridad digital y baja soberanía informacional.",
                "Respuesta MILC: escucha del contexto, investigación escolar, creación situada y evaluación de impactos.",
            ],
            callout_title="Tesis de transformación",
            callout="El plan propone pasar de una tecnología enseñada como receta a una tecnología vivida como investigación, creación cultural y responsabilidad social.",
        ),
        PrelimSection(
            title="Necesidad de cambio educativo y pertinencia institucional",
            color=PALETTE["turquesa"],
            paragraphs=[
                "La escuela contemporánea enfrenta una paradoja: los estudiantes están rodeados de tecnología, pero no siempre poseen criterios para comprenderla, cuidarse en ella o producir con ella. La abundancia de dispositivos y plataformas no garantiza alfabetización tecnológica. Por el contrario, puede profundizar nuevas formas de dependencia: atención fragmentada, desinformación, copia automática, exposición de datos, baja tolerancia al error, consumo pasivo de contenido y debilitamiento de la conversación cara a cara.",
                "Este plan asume que el cambio educativo no consiste en añadir más herramientas digitales, sino en reorganizar la experiencia de aprendizaje. La innovación curricular debe modificar la pregunta de fondo: no basta preguntar qué aplicación se usará, sino qué realidad se va a comprender, qué problema se va a investigar, qué producto se va a crear, qué evidencia se va a valorar y qué responsabilidad ética se va a asumir.",
                "La pertinencia institucional se expresa en una secuencia completa de recursos: guías, proyectos, exámenes, portafolios y mallas. Esta secuencia permite que la metodología no dependa de esfuerzos aislados del docente, sino de una arquitectura pedagógica estable, replicable y evaluable.",
            ],
            bullets=[
                "Cambio de foco: de cobertura de temas a construcción de competencias.",
                "Cambio de rol docente: de transmisor de pasos a mediador de investigación, diseño y reflexión.",
                "Cambio de rol estudiantil: de usuario obediente a sujeto creador, crítico y responsable.",
                "Cambio evaluativo: de nota final a evidencia de proceso, mejora y transferencia.",
            ],
            callout_title="Pertinencia",
            callout="La institución necesita un plan que conecte tecnología, cultura, ética, proyecto de vida, pensamiento computacional y cuidado socioemocional.",
        ),
        PrelimSection(
            title="Marco teórico ampliado",
            color=PALETTE["verde"],
            paragraphs=[
                "El marco teórico de este plan se organiza como una constelación de referentes que sostienen la metodología MILC. No se trata de yuxtaponer autores, sino de construir una base coherente para comprender por qué Tecnología e Informática debe enseñarse como formación crítica, cultural, científica, socioemocional y ética.",
                "Paulo Freire ofrece el primer fundamento: la crítica a la educación bancaria y la defensa de una educación problematizadora. En la lógica bancaria, el estudiante recibe depósitos de información; en la lógica liberadora, interpreta su realidad y actúa sobre ella mediante diálogo, conciencia crítica y praxis. MILC traduce este principio al campo tecnológico: el estudiante no memoriza herramientas, sino que investiga problemas, produce evidencias y reflexiona sobre el sentido de sus decisiones.",
                "Enrique Dussel amplía esta perspectiva desde la filosofía de la liberación. La tecnología no puede pensarse únicamente desde la eficiencia, el mercado o la fascinación por la novedad; debe pensarse desde el otro, especialmente desde quien queda excluido, vigilado, precarizado o silenciado por sistemas técnicos. Por eso, cada proyecto tecnológico debe preguntar por el usuario real, la comunidad, el acceso, la justicia, el cuidado y la dignidad.",
                "Luciano Floridi permite comprender la condición informacional de la vida contemporánea. La noción de infoesfera muestra que los estudiantes no están simplemente frente a dispositivos: habitan entornos donde identidad, reputación, datos, comunicación, consumo, aprendizaje y participación social se entrelazan. Esta mirada justifica la inclusión de soberanía informacional, seguridad digital, ética de datos, inteligencia artificial y ciudadanía onlife como dimensiones curriculares.",
                "El pensamiento computacional, popularizado por Jeannette Wing, aporta una gramática cognitiva para resolver problemas: descomponer, reconocer patrones, abstraer, diseñar algoritmos, depurar, modelar y evaluar soluciones. En este plan, el pensamiento computacional no se reduce a programar; se comprende como forma de razonar aplicable a documentos, datos, redes, proyectos, automatización, inteligencia artificial y toma de decisiones.",
                "La educación socioemocional aporta una dimensión indispensable. Aprender tecnología exige manejar frustración, error, incertidumbre, colaboración, comunicación, empatía y toma responsable de decisiones. Las competencias socioemocionales permiten que el trabajo con tecnología no deteriore la convivencia ni se convierta en competencia individualista, sino que fortalezca autocuidado, autorregulación, conciencia social, habilidades relacionales y responsabilidad.",
                "La perspectiva cultural y territorial recuerda que todo currículo porta una idea de ser humano y de comunidad. Un plan de Tecnología e Informática situado debe reconocer lenguajes, experiencias, problemas locales, memorias familiares, prácticas productivas, condiciones de conectividad, desigualdades de acceso y formas de creatividad propias del contexto. La tecnología se aprende mejor cuando dialoga con la vida cotidiana y no cuando se presenta como un mundo ajeno.",
                "Finalmente, la ética práctica y la gestión de calidad aportan hábitos de mejora: planear, hacer, verificar, actuar, documentar, corregir y volver a intentar. Estos hábitos permiten que el estudiante comprenda el error como dato para mejorar, no como fracaso; y que el producto tecnológico sea evaluado por su claridad, utilidad, pertinencia, seguridad, accesibilidad e impacto.",
            ],
            bullets=[
                "Freire: crítica a la educación bancaria, diálogo, problematización, conciencia crítica y praxis transformadora.",
                "Dussel: liberación, exterioridad, reconocimiento del otro, crítica a la dominación y tecnología con justicia.",
                "Floridi: infoesfera, identidad digital, ética de la información y ciudadanía onlife.",
                "Wing: pensamiento computacional como habilidad fundamental para formular problemas y representar soluciones.",
                "CASEL: educación socioemocional como base para autorregulación, empatía, colaboración y decisiones responsables.",
                "UNESCO: ciudadanía responsable y creativa en escenarios de inteligencia artificial.",
                "Guía No. 30 MEN: alfabetización tecnológica, cuatro componentes y relación tecnología-sociedad.",
                "MILC: síntesis metodológica que integra investigación, liberación, creación, evidencia y evaluación crítica.",
            ],
            callout_title="Síntesis formativa",
            callout="El estudiante aprende a pensar técnicamente, sentir responsablemente, reconocer su cultura y actuar éticamente sobre problemas reales.",
        ),
        PrelimSection(
            title="Objetivos formativos del área",
            color=PALETTE["mostaza"],
            paragraphs=[
                "Los objetivos del área se organizan desde una progresión que va del reconocimiento y cuidado de sistemas tecnológicos en sexto y séptimo, hacia la producción de soluciones, análisis de datos, programación, diseño web, innovación y gestión de proyectos en la educación media. La progresión busca que el estudiante avance desde el uso guiado hacia la creación autónoma y argumentada.",
                "Al finalizar la básica secundaria, se espera que el estudiante comprenda sistemas tecnológicos, produzca documentos y recursos digitales de calidad, valide información, aplique principios de seguridad digital, use datos para tomar decisiones y participe en proyectos colaborativos. Al finalizar la media, deberá diseñar soluciones tecnológicas con mayor autonomía, integrar programación, gestión de proyectos, comunicación digital, ética de datos e impacto social.",
            ],
            bullets=[
                "Desarrollar competencias tecnológicas integrales: saber, saber hacer, saber cuidar y saber convivir en entornos digitales.",
                "Promover pensamiento computacional y cultura de resolución de problemas desde situaciones reales.",
                "Fortalecer soberanía informacional: búsqueda, validación, organización, protección y producción de información.",
                "Articular tecnología, sociedad, ambiente y ética en proyectos con impacto escolar o comunitario.",
                "Consolidar portafolios de evidencias, productos y reflexiones que hagan visible la progresión del aprendizaje.",
            ],
            callout_title="Meta de egreso",
            callout="Un estudiante capaz de comprender la tecnología, usarla críticamente, crear con ella y evaluar sus efectos en la vida común.",
        ),
        PrelimSection(
            title="Diseño metodológico y estrategias activas",
            color=PALETTE["turquesa"],
            paragraphs=[
                "La metodología combina MILC con estrategias activas de aprendizaje para que las clases sean experiencias de producción y pensamiento, no sesiones de repetición técnica. El docente actúa como diseñador de ambientes de aprendizaje, mediador epistemológico, editor pedagógico de evidencias y orientador ético de las decisiones tecnológicas.",
                "Las guías de aprendizaje funcionan como cartillas de trabajo. Los proyectos integradores permiten cerrar cada periodo con un producto de mayor complejidad. Los exámenes finales se conciben como instrumentos de síntesis, aplicación y argumentación, no como ejercicios de memoria aislada.",
            ],
            bullets=[
                "Design Thinking: empatizar, definir, idear, prototipar y validar soluciones.",
                "Aprendizaje basado en proyectos: reto, producto, roles, cronograma, evidencias y sustentación.",
                "PRIMM para programación: predecir, ejecutar, investigar, modificar y crear.",
                "Aprendizaje basado en datos: formular preguntas, registrar, organizar, visualizar e interpretar.",
                "Laboratorio técnico: diagnóstico, prueba, cuidado, bitácora y protocolo.",
                "Producción editorial digital: documentos, presentaciones, infografías, mapas, campañas, prototipos y portafolios.",
            ],
            callout_title="Arquitectura de clase",
            callout="Cada sesión debe iniciar con sentido, avanzar con producción acompañada y cerrar con evidencia, reflexión y mejora.",
        ),
        PrelimSection(
            title="Evaluación, evidencias y promoción",
            color=PALETTE["vino"],
            paragraphs=[
                "El sistema de evaluación del área se fundamenta en una concepción formativa, criterial, integral y liberadora. Formativa, porque acompaña el proceso y ofrece retroalimentación para mejorar; criterial, porque valora desempeños observables y no impresiones subjetivas; integral, porque articula dimensiones cognitivas, procedimentales, actitudinales y socioemocionales; liberadora, porque ayuda al estudiante a comprender su propio aprendizaje y a transformar el error en conciencia, autonomía y mejora.",
                "En coherencia con MILC, la evaluación se organiza en cuatro momentos. En Escucha se valora la capacidad de observar, preguntar, escuchar al otro y reconocer el contexto. En Sistematización se valora la organización de información, la construcción de criterios, la explicación conceptual y la documentación del proceso. En Praxis se valora la producción, el prototipado, la aplicación técnica, la colaboración y la resolución situada de problemas. En Evaluación liberadora se valora la reflexión crítica, la mejora del producto, la sustentación, la autoevaluación, la coevaluación y el compromiso ético.",
                "El área adopta una lectura de desempeños coherente con el SIEE institucional, de modo que las valoraciones puedan traducirse a los niveles de desempeño definidos por la institución. No obstante, internamente se recomienda que cada periodo conserve una matriz de evidencias que permita rastrear guías, portafolio, proyecto integrador, prueba final y componente actitudinal.",
                "La promoción en Tecnología e Informática no debe depender de una única prueba ni de la entrega mecánica de tareas. Debe considerar la progresión del estudiante, la calidad de sus evidencias, su participación en procesos colaborativos, su capacidad de explicar decisiones, su cuidado de equipos e información y su disposición para mejorar productos a partir de retroalimentación.",
            ],
            bullets=[
                "Criterio cognitivo: comprende conceptos, sistemas, relaciones, riesgos, datos, algoritmos y criterios de decisión.",
                "Criterio procedimental: usa herramientas, documenta procesos, programa, analiza datos, prototipa y produce recursos digitales.",
                "Criterio actitudinal y socioemocional: cuida equipos, regula emociones, colabora, respeta acuerdos, reconoce al otro y decide responsablemente.",
                "Criterio ético-cultural: valida información, protege datos, reconoce contexto, evita plagio, considera impacto social y propone usos con sentido.",
                "Instrumentos: rúbricas, listas de cotejo, bitácoras, portafolios, guías, proyectos integradores, pruebas aplicadas, sustentaciones, autoevaluación y coevaluación.",
                "Ponderación orientativa por periodo: guías y portafolio 30%; proyecto o producto integrador 30%; prueba aplicada o examen final 20%; participación, cuidado, colaboración y reflexión MILC 20%. Esta ponderación puede ajustarse al SIEE institucional.",
                "Escala cualitativa sugerida: superior, alto, básico y bajo, definida por niveles de autonomía, calidad de evidencia, transferencia y responsabilidad ética.",
                "Plan de mejoramiento: revisión de evidencias, tutoría focalizada, reconstrucción del producto, sustentación breve y compromiso verificable.",
            ],
            callout_title="Pregunta evaluativa central",
            callout="¿Qué sabe el estudiante, qué sabe hacer con ese saber y qué decisiones éticas toma cuando usa tecnología?",
        ),
        PrelimSection(
            title="Referentes académicos que soportan la propuesta",
            color=PALETTE["verde"],
            paragraphs=[
                "La propuesta se apoya en referentes pedagógicos, filosóficos y tecnológicos que permiten defender el plan más allá de la organización temática. Freire fundamenta la crítica a la educación bancaria y la necesidad de una educación problematizadora. Dussel aporta la exigencia de escuchar al otro y orientar la producción tecnológica hacia la dignidad. Floridi permite comprender la vida en la infoesfera y la ética de la información. Wing ofrece el pensamiento computacional como habilidad transversal. CASEL aporta la dimensión socioemocional necesaria para aprender, convivir y decidir en comunidad. UNESCO actualiza la discusión sobre inteligencia artificial, ciudadanía y diseño humano-centrado. La Guía No. 30 MEN conserva el marco colombiano para competencias tecnológicas.",
                "MILC articula estos referentes en una metodología propia: Escucha, Sistematización, Praxis y Evaluación liberadora. Su aporte central es convertir la educación tecnológica en una práctica de investigación escolar situada, culturalmente consciente y éticamente responsable.",
            ],
            bullets=[
                "Freire, P. Pedagogía del oprimido: educación bancaria, diálogo, problematización y praxis.",
                "Dussel, E. Filosofía de la liberación: exterioridad, alteridad, ética y crítica de la dominación.",
                "Floridi, L. Filosofía de la información: infoesfera, vida onlife, identidad y ética informacional.",
                "Wing, J. M. Computational Thinking: formulación de problemas, abstracción y diseño de soluciones.",
                "CASEL: competencias socioemocionales para autoconciencia, autorregulación, conciencia social, relación y decisión responsable.",
                "UNESCO: competencias estudiantiles en inteligencia artificial con enfoque humano, ético, creativo e inclusivo.",
                "MEN, Guía No. 30: componentes para la educación en tecnología y alfabetización tecnológica en Colombia.",
                "Design Thinking, ABP, PRIMM y PDCA: metodologías activas para convertir conocimiento en producto, prototipo, mejora y transferencia.",
            ],
            callout_title="Lugar de MILC",
            callout="MILC no reemplaza los referentes: los integra en un lenguaje pedagógico operativo para el aula de Tecnología e Informática.",
        ),
        PrelimSection(
            title="Integración de recursos curriculares 2025",
            color=PALETTE["magenta"],
            paragraphs=[
                "El plan se operacionaliza mediante un ecosistema de recursos articulados: treinta guías por grado, tres proyectos integradores por grado y tres exámenes finales por grado. Esta organización garantiza continuidad semanal, cierre por periodo y trazabilidad de evidencias.",
                "Las guías no son documentos accesorios: son la unidad didáctica mínima del sistema. Cada una traduce la malla curricular en experiencia de aula mediante pregunta orientadora, conceptos, actividad, producción y reflexión. Los proyectos integradores elevan la complejidad y permiten aplicar Design Thinking, experimentación, prototipado y validación. Los exámenes finales verifican transferencia, argumentación y síntesis.",
            ],
            bullets=[
                "180 guías de aprendizaje: 30 por grado, distribuidas en 10 sesiones por periodo.",
                "18 proyectos integradores: uno por periodo para cada grado.",
                "18 exámenes finales: una síntesis aplicada por periodo y grado.",
                "Manifiestos y anexos: trazabilidad entre mallas, recursos, productos y evaluación.",
                "Formato editorial común: identidad visual MILC y continuidad de experiencia para estudiantes y docentes.",
            ],
            callout_title="Coherencia de sistema",
            callout="Plan, guía, proyecto y examen deben sentirse como partes de una misma arquitectura pedagógica.",
        ),
    ]


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for child in ("pdf", "tex", "docx", "anexos"):
        (OUT / child).mkdir(exist_ok=True)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_colored_heading(doc: Document, text: str, color: str = PALETTE["vino"]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("FFFFFF")


def add_small_table(doc: Document, rows: list[tuple[str, str]], header_color: str = PALETTE["turquesa"]) -> None:
    for left, right in rows:
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(4)
        label.paragraph_format.space_after = Pt(1)
        p_pr = label._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_color)
        p_pr.append(shd)
        run = label.add_run(left)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(8)
        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(3)
        body.paragraph_format.left_indent = Cm(0.25)
        for idx, part in enumerate(str(right).split("\n")):
            if idx:
                body.add_run().add_break()
            rr = body.add_run(part)
            rr.font.size = Pt(8)


def add_curricular_table_docx(doc: Document, pdata: PeriodData) -> None:
    headers = [
        "Estándares / orientaciones",
        "Competencias",
        "Aprendizajes",
        "Cognitivos",
        "Procedimentales",
        "Actitudinales",
    ]
    rows = component_rows(pdata)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, PALETTE["vino"])
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=7)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        values = [
            row.standard,
            row.competencias,
            row.aprendizajes,
            row.cognitivos,
            row.procedimentales,
            row.actitudinales,
        ]
        for col_idx, value in enumerate(values):
            if row_idx % 2 == 1:
                shade_cell(cells[col_idx], "F4F2F4")
            set_cell_text(cells[col_idx], value, bold=(col_idx == 0), size=6)
    doc.add_paragraph()


def add_compact_table_docx(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_color: str = PALETTE["vino"],
    font_size: int = 6,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, header_color)
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=font_size)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            if row_idx % 2 == 1:
                shade_cell(cells[col_idx], "F4F2F4")
            set_cell_text(cells[col_idx], value, size=font_size)
    doc.add_paragraph()


def add_annex_cover_docx(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("ANEXOS CURRICULARES\nPLAN DE ÁREA 2025")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(PALETTE["vino"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Uso pedagógico de guías, temarios, modelo MILC y trazabilidad de recursos")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(PALETTE["magenta"])
    add_small_table(
        doc,
        [
            ("Propósito", "Convertir los anexos en una herramienta de consulta para planear, orientar, estudiar, evaluar y actualizar el área."),
            ("Alcance", "Temario completo, uso docente, uso estudiantil, modelo MILC, criterios de evaluación y correspondencia de recursos."),
        ],
        PALETTE["vino"],
    )


def add_full_temario_docx(doc: Document, periods: list[PeriodData]) -> None:
    add_colored_heading(doc, "Anexo 1 · Temario completo", PALETTE["vino"])
    add_body_paragraph(
        doc,
        "Este anexo presenta el temario completo por grado en tres columnas, una por periodo. Cada fila corresponde a una sesión; así el docente puede leer la progresión horizontal del año y comparar cómo avanza cada periodo en diez encuentros.",
        size=8,
    )
    by_grade_period = {
        (pdata.grade, pdata.period): pdata.topics
        for pdata in periods
    }
    for grade in range(6, 12):
        add_colored_heading(doc, f"Temario · Grado {GRADE_LABELS[grade]}", PALETTE["vino"])
        rows: list[list[str]] = []
        for session in range(1, 11):
            rows.append(
                [
                    f"Sesión {session}: {by_grade_period[(grade, 1)][session - 1]}",
                    f"Sesión {session}: {by_grade_period[(grade, 2)][session - 1]}",
                    f"Sesión {session}: {by_grade_period[(grade, 3)][session - 1]}",
                ]
            )
        add_compact_table_docx(
            doc,
            ["Periodo 1", "Periodo 2", "Periodo 3"],
            rows,
            PALETTE["vino"],
            font_size=5,
        )


def add_annex_descriptions_docx(doc: Document) -> None:
    add_colored_heading(doc, "Anexo 2 · Uso de las guías para docentes", PALETTE["turquesa"])
    add_body_paragraph(
        doc,
        "Las guías no son talleres aislados: son la unidad didáctica mínima del plan. Cada guía debe orientar una sesión de trabajo autónomo acompañado, conectar con el temario del periodo y producir una evidencia que pueda entrar al portafolio.",
        size=8,
    )
    for item in [
        "Antes de clase: revisar el propósito, el producto esperado, los conceptos nuevos y los materiales necesarios; anticipar qué ejemplos o demostraciones necesitará el grupo.",
        "Inicio de clase: usar la Escucha para conectar el tema con una situación real del colegio, la familia, la comunidad, una empresa simulada o un problema digital cercano.",
        "Desarrollo conceptual: no pedir respuestas antes de explicar vocabulario, criterios, ejemplos, errores frecuentes y el procedimiento mínimo para avanzar.",
        "Acompañamiento: circular por el aula, pedir evidencias parciales, hacer preguntas de precisión y ayudar al estudiante a explicar qué intentó, qué encontró y qué debe corregir.",
        "Cierre: recoger una evidencia concreta, orientar autoevaluación breve, registrar dificultades comunes y dejar una mejora verificable para el portafolio.",
        "Evaluación: valorar comprensión, proceso, producto, cuidado de datos/equipos, colaboración y reflexión; no reducir la guía a entrega mecánica.",
        "Flexibilización: ajustar tiempo, cantidad de producto o apoyo visual para estudiantes que lo requieran, sin eliminar el propósito cognitivo de la sesión.",
    ]:
        add_bullet(doc, item, PALETTE["turquesa"])

    add_colored_heading(doc, "Anexo 3 · Uso de las guías para estudiantes", PALETTE["magenta"])
    add_body_paragraph(
        doc,
        "La guía está diseñada para que el estudiante aprenda con autonomía progresiva. No reemplaza al docente, pero sí le permite saber qué debe comprender, qué debe producir y cómo revisar su propio avance.",
        size=8,
    )
    for item in [
        "Leer el título, la pregunta orientadora y el producto esperado antes de empezar; si no se entiende la meta, se debe preguntar.",
        "Identificar conceptos nuevos y escribir una definición propia con un ejemplo; copiar sin comprender no cuenta como aprendizaje.",
        "Resolver primero la práctica guiada, compararla con el ejemplo y luego avanzar a la práctica autónoma.",
        "Guardar evidencias del proceso: borradores, capturas, tablas, enlaces, códigos, versiones corregidas, reflexiones y respuestas finales.",
        "Usar IA solo cuando la guía lo permita o el docente lo autorice; siempre verificar, reescribir con voz propia y declarar el apoyo recibido.",
        "Pedir ayuda con precisión: decir qué se intentó, dónde aparece el error, qué instrucción genera duda y qué evidencia se tiene.",
        "Cerrar cada guía con una revisión personal: qué aprendí, qué corregí, qué producto entrego y qué debo mejorar en la siguiente sesión.",
    ]:
        add_bullet(doc, item, PALETTE["magenta"])

    add_colored_heading(doc, "Anexo 4 · Modelo MILC: fundamentación, importancia y uso curricular", PALETTE["verde"])
    for paragraph in [
        "MILC significa Modelo de Investigación Liberadora y Científica. En este plan de área funciona como arquitectura pedagógica: organiza la clase para que el estudiante escuche el contexto, sistematice saberes, produzca una solución y evalúe el impacto de lo realizado.",
        "Su importancia está en superar la clase bancaria, entendida como copia de instrucciones o repetición de procedimientos. En Tecnología e Informática esto es decisivo: no basta aprender a usar una herramienta; el estudiante debe comprender sistemas, tomar decisiones, cuidar información, reconocer riesgos y crear productos útiles.",
        "El modelo permite que los cuatro estándares del área no queden como encabezados formales, sino como lentes permanentes de lectura. Cada periodo debe mostrar cómo la tecnología evoluciona, cómo se usa de forma apropiada, cómo ayuda a resolver problemas y qué consecuencias sociales, éticas, ambientales o económicas produce.",
        "En el plan de área, MILC no aparece como actividad adicional. Está presente en las guías, los proyectos integradores, los exámenes finales, los portafolios, las rúbricas y la lectura de los cuatro estándares: naturaleza y evolución, apropiación y uso, solución de problemas, y tecnología y sociedad.",
    ]:
        add_body_paragraph(doc, paragraph, size=8)
    add_compact_table_docx(
        doc,
        ["Referente", "Aporte central", "Uso dentro del plan"],
        [
            ["Paulo Freire", "Educación problematizadora, diálogo y lectura crítica de la realidad.", "La Escucha inicia desde preguntas del contexto y evita que la guía sea solo copia de procedimientos."],
            ["Enrique Dussel", "Ética de la alteridad: escuchar al otro, reconocer exclusiones y responder con responsabilidad.", "Los proyectos exigen pensar usuarios, beneficiarios, impactos y límites de cada solución tecnológica."],
            ["Guía No. 30 MEN", "Cuatro componentes de la educación en tecnología.", "Los estándares se trabajan como filas permanentes en cada periodo y no como temas aislados."],
            ["Wing, CSTA e ISTE", "Pensamiento computacional, ciudadanía digital, diseño, datos y solución de problemas.", "La Sistematización organiza conceptos, algoritmos, evidencias y criterios antes de producir."],
            ["UNESCO y Floridi", "Ciudadanía digital, IA responsable, infoesfera, datos, sesgos y bienestar.", "Orienta el uso crítico de inteligencia artificial, privacidad, autoría y verificación de información."],
            ["CASEL", "Autoconciencia, colaboración, toma responsable de decisiones y autorregulación.", "Aporta criterios para trabajo en equipo, portafolio, reflexión y evaluación liberadora."],
        ],
        PALETTE["verde"],
        font_size=7,
    )
    add_small_table(
        doc,
        [
            ("Contexto antes que herramienta", "La clase parte de una situación, necesidad, usuario o pregunta real; la herramienta se elige después de comprender el problema."),
            ("Concepto antes que producto", "La guía enseña vocabulario, criterios, ejemplos y procedimientos antes de pedir una entrega."),
            ("Evidencia antes que calificación", "El avance se verifica con borradores, capturas, datos, versiones, pruebas, decisiones y correcciones."),
            ("Ética antes que eficiencia", "Toda solución debe revisar privacidad, autoría, sesgos, inclusión, sostenibilidad y consecuencias para otros."),
            ("Mejora antes que cierre", "El producto final debe incluir revisión, explicación de cambios y aprendizaje transferible a otra situación."),
        ],
        PALETTE["verde"],
    )
    add_small_table(
        doc,
        [
            ("Escucha", "Momento de lectura del contexto. El estudiante observa, escucha, pregunta, identifica usuarios, reconoce necesidades y formula una situación problemática situada."),
            ("Sistematización", "Momento de organización del saber. El estudiante define conceptos, revisa fuentes, ordena datos, construye criterios y comprende procedimientos antes de producir."),
            ("Praxis", "Momento de acción transformadora. El estudiante diseña, programa, redacta, analiza, prototipa, documenta o comunica un producto verificable."),
            ("Evaluación liberadora", "Momento de conciencia crítica. El estudiante revisa evidencias, errores, impactos, sesgos, decisiones éticas, aprendizajes y mejoras posibles."),
        ],
        PALETTE["verde"],
    )
    add_compact_table_docx(
        doc,
        ["Elemento del plan", "Uso de MILC"],
        [
            ["Guías", "Cada sesión traduce MILC en pregunta, conceptos, práctica, producto y reflexión."],
            ["Proyectos", "Cada periodo cierra con escucha de una necesidad, diseño, validación, mejora y sustentación."],
            ["Exámenes", "Las preguntas evalúan aplicación, criterio, transferencia y justificación, no solo memoria."],
            ["Portafolio", "Reúne evidencias de proceso, versiones, errores, decisiones y reflexión final."],
            ["Evaluación", "Integra desempeño cognitivo, procedimental, actitudinal, ético y socioemocional."],
            ["Rol docente", "Diseña condiciones, modela procedimientos, pregunta, retroalimenta y ayuda a elevar el nivel del producto."],
            ["Rol estudiante", "Investiga, decide, produce evidencias, solicita ayuda con precisión, corrige y sustenta con argumentos."],
        ],
        PALETTE["verde"],
        font_size=8,
    )

    add_colored_heading(doc, "Anexo 5 · Criterios de evaluación y evidencias", PALETTE["mostaza"])
    add_compact_table_docx(
        doc,
        ["Evidencia", "Función pedagógica", "Criterios de calidad"],
        [
            ["Guías", "Acompañan la construcción semanal de conceptos y productos.", "Claridad, autonomía, aplicación, revisión y reflexión."],
            ["Portafolio", "Hace visible el proceso y la mejora.", "Orden, completitud, versiones, evidencias y metacognición."],
            ["Proyecto integrador", "Sintetiza el periodo en un producto situado.", "Problema real, investigación, prototipo, validación, impacto y sustentación."],
            ["Examen final", "Verifica transferencia, comprensión y argumentación.", "Lectura de contexto, uso de conceptos, toma de decisiones y justificación."],
            ["Actitud y cuidado", "Valora convivencia tecnológica y responsabilidad.", "Cuidado de equipos, datos, tiempos, colaboración y honestidad intelectual."],
        ],
        PALETTE["mostaza"],
        font_size=6,
    )


def add_correspondence_docx(doc: Document, periods: list[PeriodData]) -> None:
    add_colored_heading(doc, "Anexo 6 · Correspondencia de recursos", PALETTE["turquesa"])
    add_body_paragraph(
        doc,
        "Este anexo permite rastrear, por grado y periodo, las guías de aprendizaje, el proyecto integrador y el examen final asociados.",
        size=8,
    )
    rows = []
    for pdata in periods:
        guides, project, exam = resource_summary(pdata)
        rows.append([str(pdata.grade), PERIOD_NAMES[pdata.period], guides, project, exam])
    add_compact_table_docx(doc, ["Grado", "Periodo", "Guías", "Proyecto", "Examen"], rows, PALETTE["turquesa"], font_size=6)


def add_grade_brand_table_docx(doc: Document) -> None:
    add_body_paragraph(
        doc,
        "Código visual por grado. La paleta se toma de ConectaTE y se proyecta a portadas físicas, PDFs, DOCX, proyectos y exámenes para que el estudiante reconozca su ruta de aprendizaje con el mismo lenguaje visual en aula y plataforma.",
        size=8,
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Grado", "Color", "Código", "Contraste", "Foco visual"]
    for idx, header in enumerate(headers):
        shade_cell(table.cell(0, idx), PALETTE["vino"])
        set_cell_text(table.cell(0, idx), header, bold=True, color="FFFFFF", size=7)
    for grade in range(6, 12):
        item = GRADE_BRAND_COLORS[grade]
        cells = table.add_row().cells
        set_cell_text(cells[0], f"Grado {GRADE_LABELS[grade]}", bold=True, size=7)
        shade_cell(cells[1], item["hex"])
        set_cell_text(cells[1], item["name"], bold=True, color=item["text"], size=7)
        set_cell_text(cells[2], f"#{item['hex']}", size=7)
        contrast = "Texto blanco" if item["text"] == "FFFFFF" else "Texto oscuro"
        set_cell_text(cells[3], contrast, size=7)
        set_cell_text(cells[4], item["focus"], size=7)
    doc.add_paragraph()


def add_platform_annex_docx(doc: Document) -> None:
    add_colored_heading(doc, "Anexo 7 · Plataforma ConectaTE como mediación curricular", PALETTE["magenta"])
    for paragraph in [
        "ConectaTE es la plataforma educativa institucional que articula el Plan de Área de Tecnología e Informática con la experiencia concreta de aula. No se limita a alojar documentos: organiza el tránsito entre malla curricular, guías, proyectos, exámenes, evidencias y seguimiento del estudiante.",
        "Su puesta en escena permite que cada estudiante identifique su grado, navegue por periodos y sesiones, abra la guía completa, la descargue, la consulte en Drive, comparta el enlace y marque avances. De esta manera, la guía deja de ser un archivo aislado y se convierte en una ruta de aprendizaje accesible dentro y fuera del aula.",
        "La plataforma favorece la continuidad pedagógica porque permite retomar el trabajo fuera del horario de clase, consultar recursos desde celular, tableta o computador, y mantener una referencia común entre docente, estudiante y familia. También apoya la trazabilidad del proceso al ordenar los recursos por grado, periodo, sesión, proyecto y examen.",
    ]:
        add_body_paragraph(doc, paragraph, size=8)
    add_small_table(
        doc,
        [
            ("Cobertura curricular", "180 guías de aprendizaje, 18 proyectos integradores y 18 exámenes finales organizados para grados sexto a undécimo."),
            ("Interacción de aula", "Visor PDF embebido, descarga directa, apertura en Drive, navegación entre sesiones, búsqueda de recursos, opción de compartir y seguimiento de progreso."),
            ("Acceso extendido", "Diseño adaptable a móvil, tableta y computador, con comportamiento de aplicación web progresiva y soporte de consulta después de la primera carga cuando el recurso está en caché."),
            ("Vínculo institucional", PLATFORM_URL),
        ],
        PALETTE["magenta"],
    )
    add_grade_brand_table_docx(doc)
    add_compact_table_docx(
        doc,
        ["Momento MILC", "Función de ConectaTE"],
        [
            ["Escucha", "Sitúa al estudiante en grado, periodo, sesión y propósito; permite reconocer qué guía corresponde al proceso real de aula."],
            ["Sistematización", "Ordena contenidos, recursos y evidencias en una arquitectura navegable que evita dispersión documental."],
            ["Praxis", "Facilita el acceso al documento de trabajo para producir actividades, proyectos, portafolios y evidencias verificables."],
            ["Evaluación liberadora", "Apoya la revisión de avance, la continuidad fuera del aula, la socialización de enlaces y la responsabilidad del estudiante sobre su proceso."],
        ],
        PALETTE["magenta"],
        font_size=8,
    )
    for item in [
        "Para el docente, ConectaTE funciona como tablero de ruta: permite ubicar la sesión, proyectar la guía, orientar instrucciones comunes y dejar continuidad para casa.",
        "Para el estudiante, funciona como cuaderno digital de consulta: le muestra qué debe trabajar, dónde está el recurso, cómo descargarlo y cómo avanzar con autonomía.",
        "Para la institución, fortalece la coherencia entre plan de área, recursos curriculares, experiencia de aula, seguimiento y actualización permanente.",
    ]:
        add_bullet(doc, item, PALETTE["magenta"])


def add_bibliography_docx(doc: Document) -> None:
    doc.add_page_break()
    add_colored_heading(doc, "Bibliografía", PALETTE["vino"])
    add_body_paragraph(
        doc,
        "La bibliografía reúne los referentes pedagógicos, normativos, tecnológicos y curriculares que sostienen el enfoque MILC, la organización del plan de área, la ciudadanía digital, la inteligencia artificial responsable, la gestión de calidad y la mediación mediante plataforma educativa.",
        size=8,
    )
    for entry in BIBLIOGRAPHY:
        add_bullet(doc, entry, PALETTE["vino"])


def add_annexes_docx(doc: Document, periods: list[PeriodData]) -> None:
    add_annex_cover_docx(doc)
    add_full_temario_docx(doc, periods)
    add_annex_descriptions_docx(doc)
    add_correspondence_docx(doc, periods)
    add_platform_annex_docx(doc)
    add_bibliography_docx(doc)


def add_body_paragraph(doc: Document, text: str, size: int = 9) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(size)


def add_bullet(doc: Document, text: str, color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    bullet = p.add_run("• ")
    bullet.bold = True
    bullet.font.color.rgb = RGBColor.from_string(color)
    bullet.font.size = Pt(9)
    r = p.add_run(text)
    r.font.size = Pt(8)


def add_callout(doc: Document, title: str, text: str, color: str) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(5)
    label.paragraph_format.space_after = Pt(1)
    p_pr = label._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)
    run = label.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    run.font.size = Pt(9)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.35)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)


def add_preliminaries_docx(doc: Document) -> None:
    add_colored_heading(doc, "Preliminares académicos y metodológicos", PALETTE["magenta"])
    intro = (
        "Los apartados siguientes reconstruyen y amplían los elementos previos al plan de estudios. "
        "Su función es dar fundamento curricular, epistemológico, metodológico y evaluativo a las mallas, "
        "de modo que la metodología MILC no aparezca como un añadido, sino como el principio organizador "
        "de todo el documento."
    )
    add_callout(doc, "Clave de lectura", intro, PALETTE["vino"])
    for section in academic_preliminaries():
        add_colored_heading(doc, section.title, section.color)
        for paragraph in section.paragraphs:
            add_body_paragraph(doc, paragraph)
        if section.bullets:
            for item in section.bullets:
                add_bullet(doc, item, section.color)
        if section.callout:
            add_callout(doc, section.callout_title or "Síntesis", section.callout, section.color)
        # Evita saltos rígidos que dejan grandes espacios en blanco en la versión editable.


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor.from_string(PALETTE["vino"])


def build_docx(periods: list[PeriodData]) -> Path:
    doc = Document()
    configure_docx_styles(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN DE ÁREA 2025\nTECNOLOGÍA E INFORMÁTICA")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(PALETTE["vino"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Versión editorial MILC · Institución Educativa Sor María Juliana")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(PALETTE["magenta"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Autor: {AUTHOR}").bold = True
    doc.add_page_break()

    add_preliminaries_docx(doc)

    add_colored_heading(doc, "Criterio de articulación con las mallas")
    doc.add_paragraph(
        "Esta versión conserva el sentido curricular del plan de área 2025 y reorganiza las mallas "
        "con el formato observado en el plan 2017: estándares u orientaciones, competencias, "
        "aprendizajes y desempeños cognitivos, procedimentales y actitudinales. El ajuste explicita "
        "los cuatro componentes de las Orientaciones Generales para la Educación en Tecnología "
        "del MEN y articula cada periodo con MILC, estrategias activas y los recursos generados."
    )
    add_small_table(
        doc,
        [
            ("Componentes MEN", "; ".join(COMPONENTS)),
            ("Metodología", "MILC: Escucha, Sistematización, Praxis y Evaluación liberadora."),
            ("Recursos integrados", "180 guías de aprendizaje, 18 proyectos integradores y 18 exámenes finales."),
            ("Criterio editorial", "Paleta magenta/vino con acentos verde, turquesa y mostaza; tablas limpias y legibles."),
        ],
        PALETTE["magenta"],
    )

    add_colored_heading(doc, "Grados", PALETTE["vino"])
    add_body_paragraph(
        doc,
        "La sección de grados presenta las mallas curriculares de sexto a undécimo. Cada grado conserva sus tres periodos y cada periodo desarrolla los cuatro estándares como filas curriculares permanentes.",
        size=8,
    )
    for grade in range(6, 12):
        doc.add_page_break()
        add_colored_heading(doc, f"Mallas curriculares · Grado {GRADE_LABELS[grade]}")
        for pdata in [p for p in periods if p.grade == grade]:
            add_colored_heading(doc, f"Periodo {PERIOD_NAMES[pdata.period]} · Grado {pdata.grade}", PALETTE["turquesa"])
            guides, project, exam = resource_summary(pdata)
            add_small_table(
                doc,
                [
                    ("Asignación académica", pdata.asignacion),
                    ("Derecho básico / aprendizaje", compact(pdata.dba, 430)),
                    ("Referentes", compact(pdata.referentes, 430)),
                    ("Componentes MEN", "\n".join(f"• {c}" for c in COMPONENTS)),
                    ("Énfasis del periodo", infer_emphasis(pdata)),
                    ("Estrategias activas", active_strategy(pdata)),
                    ("Recursos", f"Guías: {guides}\nProyecto: {project}\nExamen: {exam}"),
                ],
                PALETTE["verde"],
            )
            add_curricular_table_docx(doc, pdata)
            doc.add_paragraph("Temario articulado: " + "; ".join(simplified_topics(pdata)), style=None)

    add_annexes_docx(doc, periods)

    path = OUT / "docx" / "plan_area_tecnologia_2025_editorial_milc.docx"
    doc.save(path)
    return path


def tex_curricular_rows(pdata: PeriodData) -> str:
    rendered: list[str] = []
    for idx, row in enumerate(component_rows(pdata)):
        prefix = "\\rowcolor{gris}\n" if idx % 2 else ""
        rendered.append(
            prefix
            + rf"\textbf{{{tex_escape(row.standard)}}} & "
            + f"{tex_escape(compact(row.competencias, 230))} & "
            + f"{tex_escape(compact(row.aprendizajes, 280))} & "
            + f"{tex_escape(compact(row.cognitivos, 240))} & "
            + f"{tex_escape(compact(row.procedimentales, 240))} & "
            + f"{tex_escape(compact(row.actitudinales, 240))}\\\\"
        )
    return "\n".join(rendered)


def tex_annex_divider() -> str:
    return rf"""
\clearpage
\phantomsection
\addcontentsline{{toc}}{{chapter}}{{Anexos}}
\pagecolor{{magentaMILC}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[magentaMILC] (current page.south west) rectangle (current page.north east);
  \foreach \x in {{-1,1,...,23}}{{
    \foreach \y in {{-1,1,...,31}}{{
      \draw[vino, opacity=.12, line width=.6pt] (\x,\y) .. controls +(1,.6) and +(-1,-.6) .. +(2,1.4);
      \draw[vino, opacity=.10, line width=.6pt] (\x+.6,\y+.4) arc (210:510:.45);
    }}
  }}
\end{{tikzpicture}}
\vspace*{{1.2cm}}
\begin{{center}}
\begin{{tcolorbox}}[
  enhanced,colback=vino,colframe=vino,arc=12mm,width=.86\textwidth,
  boxrule=0pt,left=12mm,right=12mm,top=12mm,bottom=12mm,center,
  overlay={{\draw[magentaMILC,line width=2pt,dash pattern=on 9pt off 8pt,rounded corners=10mm]
  ([xshift=7mm,yshift=-7mm]frame.north west) rectangle ([xshift=-7mm,yshift=7mm]frame.south east);}}]
  \centering
  {{\sffamily\bfseries\color{{white}}\fontsize{{15}}{{17}}\selectfont PLAN DE ÁREA · TECNOLOGÍA E INFORMÁTICA}}\\[9mm]
  {{\sffamily\bfseries\color{{white}}\fontsize{{31}}{{34}}\selectfont Anexos curriculares}}\\[3mm]
  {{\sffamily\bfseries\color{{mostaza}}\fontsize{{22}}{{25}}\selectfont Guías, temario, MILC y evaluación}}\\[8mm]
  {{\sffamily\color{{white}}\Large Grados 6 a 11 · Versión editorial MILC}}\\[8mm]
  {{\sffamily\bfseries\color{{white}}\large Institución Educativa Sor María Juliana}}\\[2mm]
  {{\sffamily\color{{white}}Autor: {tex_escape(AUTHOR)}}}
\end{{tcolorbox}}
\vfill
\begin{{tcolorbox}}[enhanced,colback=vino!72!black,colframe=vino!72!black,arc=5mm,boxrule=0pt,width=.86\textwidth,left=6mm,right=6mm,top=4mm,bottom=4mm,center]
{{\color{{white}}\sffamily\small Consulta docente · Aprendizaje autónomo · Trazabilidad · Evaluación}}
\end{{tcolorbox}}
\end{{center}}
\clearpage
\nopagecolor
"""


def tex_full_temario_annex(periods: list[PeriodData]) -> str:
    by_grade_period = {
        (pdata.grade, pdata.period): pdata.topics
        for pdata in periods
    }
    chunks = [
        r"""
\section*{Anexo 1 · Temario completo}\addcontentsline{toc}{section}{Anexo 1 · Temario completo}
\begin{tcolorbox}[breakable,colback=gris,colframe=vino,title=\textbf{Criterio de lectura},top=1mm,bottom=1mm]
Este anexo organiza el temario completo por grado. Cada grado aparece en una tabla de tres columnas, una por periodo, y diez filas, una por sesión. La correspondencia exacta con guías, proyectos y exámenes se consulta en el anexo de trazabilidad.
\end{tcolorbox}
"""
    ]
    for grade in range(6, 12):
        rows: list[str] = []
        for session in range(1, 11):
            prefix = "\\rowcolor{gris}\n" if session % 2 == 0 else ""
            cells = []
            for period in (1, 2, 3):
                topic = by_grade_period[(grade, period)][session - 1]
                cells.append(rf"\textbf{{Sesión {session}:}} {tex_escape(topic)}")
            rows.append(prefix + " & ".join(cells) + r"\\")
        chunks.append(
            rf"""
\subsection*{{Grado {tex_escape(GRADE_LABELS[grade])}}}
\begingroup
\footnotesize
\renewcommand{{\arraystretch}}{{1.05}}
\setlength{{\tabcolsep}}{{3pt}}
\setlength{{\LTleft}}{{0pt}}
\setlength{{\LTright}}{{0pt}}
\begin{{longtable}}{{L{{0.31\linewidth}}L{{0.31\linewidth}}L{{0.31\linewidth}}}}
\rowcolor{{vino}}\color{{white}}\textbf{{Periodo 1}} & \color{{white}}\textbf{{Periodo 2}} & \color{{white}}\textbf{{Periodo 3}}\\
{chr(10).join(rows)}
\end{{longtable}}
\endgroup
"""
        )
    return "\n".join(chunks)


def tex_usage_annexes() -> str:
    return r"""
\section*{Anexo 2 · Uso de las guías para docentes}\addcontentsline{toc}{section}{Anexo 2 · Uso de las guías para docentes}
\begin{tcolorbox}[breakable,colback=turquesa!7,colframe=turquesa,title=\textbf{Sentido pedagógico},top=1mm,bottom=1mm]
Las guías no son talleres aislados: son la unidad didáctica mínima del plan. Cada guía orienta una sesión de trabajo autónomo acompañado, conecta con el temario del periodo y produce una evidencia que puede entrar al portafolio.
\end{tcolorbox}
\begin{tcolorbox}[breakable,colback=white,colframe=turquesa,title=\textbf{Ruta de uso docente},top=1mm,bottom=1mm]
\begin{itemize}
\item \textbf{Antes de clase:} revisar propósito, producto esperado, conceptos nuevos y materiales; anticipar ejemplos, demostraciones o apoyos visuales.
\item \textbf{Inicio:} usar la Escucha para conectar el tema con una situación real del colegio, la familia, la comunidad, una empresa simulada o un problema digital cercano.
\item \textbf{Desarrollo conceptual:} explicar vocabulario, criterios, ejemplos, errores frecuentes y procedimiento mínimo antes de exigir respuestas.
\item \textbf{Acompañamiento:} pedir evidencias parciales, hacer preguntas de precisión y ayudar al estudiante a explicar qué intentó, qué encontró y qué debe corregir.
\item \textbf{Cierre:} recoger una evidencia concreta, orientar autoevaluación, registrar dificultades comunes y dejar una mejora verificable para el portafolio.
\item \textbf{Evaluación:} valorar comprensión, proceso, producto, cuidado de datos/equipos, colaboración y reflexión; no reducir la guía a entrega mecánica.
\item \textbf{Flexibilización:} ajustar tiempo, cantidad de producto o apoyos sin eliminar el propósito cognitivo de la sesión.
\end{itemize}
\end{tcolorbox}

\section*{Anexo 3 · Uso de las guías para estudiantes}\addcontentsline{toc}{section}{Anexo 3 · Uso de las guías para estudiantes}
\begin{tcolorbox}[breakable,colback=magentaMILC!7,colframe=magentaMILC,title=\textbf{Sentido para el estudiante},top=1mm,bottom=1mm]
La guía está diseñada para aprender con autonomía progresiva. No reemplaza al docente, pero permite saber qué se debe comprender, qué se debe producir y cómo revisar el avance propio.
\end{tcolorbox}
\begin{tcolorbox}[breakable,colback=white,colframe=magentaMILC,title=\textbf{Ruta de aprendizaje autónomo},top=1mm,bottom=1mm]
\begin{itemize}
\item Leer título, pregunta orientadora y producto esperado antes de empezar; si no se entiende la meta, se debe preguntar.
\item Identificar conceptos nuevos y escribir una definición propia con un ejemplo; copiar sin comprender no cuenta como aprendizaje.
\item Resolver primero la práctica guiada, compararla con el ejemplo y luego avanzar a la práctica autónoma.
\item Guardar evidencias del proceso: borradores, capturas, tablas, enlaces, códigos, versiones corregidas, reflexiones y respuestas finales.
\item Usar IA solo cuando la guía lo permita o el docente lo autorice; siempre verificar, reescribir con voz propia y declarar el apoyo recibido.
\item Pedir ayuda con precisión: decir qué se intentó, dónde aparece el error, qué instrucción genera duda y qué evidencia se tiene.
\item Cerrar cada guía con una revisión personal: qué aprendí, qué corregí, qué producto entrego y qué debo mejorar en la siguiente sesión.
\end{itemize}
\end{tcolorbox}

\section*{Anexo 4 · Modelo MILC: fundamentación, importancia y uso curricular}\addcontentsline{toc}{section}{Anexo 4 · Modelo MILC}
\begin{tcolorbox}[breakable,colback=verde!7,colframe=verde,title=\textbf{Fundamentación del modelo},top=1mm,bottom=1mm]
MILC significa Modelo de Investigación Liberadora y Científica. En este plan de área funciona como arquitectura pedagógica: organiza la clase para que el estudiante escuche el contexto, sistematice saberes, produzca una solución y evalúe el impacto de lo realizado.

Su importancia está en superar la clase bancaria, entendida como copia de instrucciones o repetición de procedimientos. En Tecnología e Informática no basta aprender a usar una herramienta; el estudiante debe comprender sistemas, tomar decisiones, cuidar información, reconocer riesgos y crear productos útiles.

El modelo permite que los cuatro estándares del área no queden como encabezados formales, sino como lentes permanentes de lectura. Cada periodo debe mostrar cómo la tecnología evoluciona, cómo se usa de forma apropiada, cómo ayuda a resolver problemas y qué consecuencias sociales, éticas, ambientales o económicas produce.
\end{tcolorbox}
\begin{tcolorbox}[breakable,colback=white,colframe=verde,title=\textbf{Referentes que dialogan con MILC},top=1mm,bottom=1mm]
MILC dialoga con Paulo Freire en la educación problematizadora; con Enrique Dussel en la exigencia ética de escuchar al otro y reconocer exclusiones; con Luciano Floridi en la comprensión de la infoesfera; con Jeannette Wing en el pensamiento computacional; con CASEL en la dimensión socioemocional; con UNESCO en ciudadanía digital e inteligencia artificial responsable; y con la Guía No. 30 del MEN en los cuatro componentes de la educación en tecnología.
\end{tcolorbox}
\begingroup
\footnotesize
\renewcommand{\arraystretch}{0.92}
\setlength{\tabcolsep}{2pt}
\begin{longtable}{L{0.17\linewidth}L{0.35\linewidth}L{0.40\linewidth}}
\rowcolor{vino}\color{white}\textbf{Referente} & \color{white}\textbf{Aporte central} & \color{white}\textbf{Uso dentro del plan}\\
Paulo Freire & Educación problematizadora, diálogo y lectura crítica de la realidad. & La Escucha inicia desde preguntas del contexto y evita que la guía sea solo copia de procedimientos.\\
Enrique Dussel & Ética de la alteridad: escuchar al otro, reconocer exclusiones y responder con responsabilidad. & Los proyectos exigen pensar usuarios, beneficiarios, impactos y límites de cada solución tecnológica.\\
Guía No. 30 MEN & Cuatro componentes de la educación en tecnología. & Los estándares se trabajan como filas permanentes en cada periodo y no como temas aislados.\\
Wing, CSTA e ISTE & Pensamiento computacional, ciudadanía digital, diseño, datos y solución de problemas. & La Sistematización organiza conceptos, algoritmos, evidencias y criterios antes de producir.\\
UNESCO y Floridi & Ciudadanía digital, IA responsable, infoesfera, datos, sesgos y bienestar. & Orienta el uso crítico de inteligencia artificial, privacidad, autoría y verificación de información.\\
CASEL & Autoconciencia, colaboración, toma responsable de decisiones y autorregulación. & Aporta criterios para trabajo en equipo, portafolio, reflexión y evaluación liberadora.\\
\end{longtable}
\endgroup
\begin{tcolorbox}[breakable,colback=white,colframe=verde,title=\textbf{Principios operativos},top=1mm,bottom=1mm]
\begin{itemize}
\item \textbf{Contexto antes que herramienta:} la clase parte de una situación, necesidad, usuario o pregunta real; la herramienta se elige después de comprender el problema.
\item \textbf{Concepto antes que producto:} la guía enseña vocabulario, criterios, ejemplos y procedimientos antes de pedir una entrega.
\item \textbf{Evidencia antes que calificación:} el avance se verifica con borradores, capturas, datos, versiones, pruebas, decisiones y correcciones.
\item \textbf{Ética antes que eficiencia:} toda solución revisa privacidad, autoría, sesgos, inclusión, sostenibilidad y consecuencias para otros.
\item \textbf{Mejora antes que cierre:} el producto final incluye revisión, explicación de cambios y aprendizaje transferible a otra situación.
\end{itemize}
\end{tcolorbox}
\begin{tcolorbox}[breakable,colback=verde!7,colframe=verde,title=\textbf{MILC en cuatro momentos},top=1mm,bottom=1mm]
\textbf{Escucha:} momento de lectura del contexto. El estudiante observa, escucha, pregunta, identifica usuarios, reconoce necesidades y formula una situación problemática situada.\\
\textbf{Sistematización:} momento de organización del saber. El estudiante define conceptos, revisa fuentes, ordena datos, construye criterios y comprende procedimientos antes de producir.\\
\textbf{Praxis:} momento de acción transformadora. El estudiante diseña, programa, redacta, analiza, prototipa, documenta o comunica un producto verificable.\\
\textbf{Evaluación liberadora:} momento de conciencia crítica. El estudiante revisa evidencias, errores, impactos, sesgos, decisiones éticas, aprendizajes y mejoras posibles.
\end{tcolorbox}
\begingroup
\small
\renewcommand{\arraystretch}{1.18}
\setlength{\tabcolsep}{5pt}
\setlength{\LTleft}{0pt}
\setlength{\LTright}{0pt}
\begin{longtable}{L{0.20\linewidth}L{0.72\linewidth}}
\rowcolor{vino}\color{white}\textbf{Elemento del plan} & \color{white}\textbf{Uso de MILC}\\
Guías & Cada sesión traduce MILC en pregunta, conceptos, práctica, producto y reflexión.\\
Proyectos & Cada periodo cierra con escucha de una necesidad, diseño, validación, mejora y sustentación.\\
Exámenes & Las preguntas evalúan aplicación, criterio, transferencia y justificación, no solo memoria.\\
Portafolio & Reúne evidencias de proceso, versiones, errores, decisiones y reflexión final.\\
Evaluación & Integra desempeño cognitivo, procedimental, actitudinal, ético y socioemocional.\\
Rol docente & Diseña condiciones, modela procedimientos, pregunta, retroalimenta y ayuda a elevar el nivel del producto.\\
Rol estudiante & Investiga, decide, produce evidencias, solicita ayuda con precisión, corrige y sustenta con argumentos.\\
\end{longtable}
\endgroup
\begin{tcolorbox}[breakable,colback=verde!7,colframe=verde,title=\textbf{Indicadores de que MILC sí se está aplicando},top=1mm,bottom=1mm]
Una sesión muestra aplicación real del modelo cuando el estudiante puede explicar qué problema o contexto atendió, qué conceptos necesitó, qué evidencia produjo, qué decisión tomó, qué riesgo reconoció y qué mejora haría. Una guía es débil cuando pide respuestas sin haber enseñado vocabulario, criterios o procedimiento; un proyecto es débil cuando solo exhibe un producto sin investigación, validación ni reflexión.
\end{tcolorbox}

\clearpage
\section*{Anexo 5 · Criterios de evaluación y evidencias}\addcontentsline{toc}{section}{Anexo 5 · Criterios de evaluación y evidencias}
\begingroup
\footnotesize
\setlength{\tabcolsep}{3pt}
\begin{longtable}{L{0.16\linewidth}L{0.32\linewidth}L{0.44\linewidth}}
\rowcolor{vino}\color{white}\textbf{Evidencia} & \color{white}\textbf{Función pedagógica} & \color{white}\textbf{Criterios de calidad}\\
Guías & Acompañan la construcción semanal de conceptos y productos. & Claridad, autonomía, aplicación, revisión y reflexión.\\
Portafolio & Hace visible el proceso y la mejora. & Orden, completitud, versiones, evidencias y metacognición.\\
Proyecto integrador & Sintetiza el periodo en un producto situado. & Problema real, investigación, prototipo, validación, impacto y sustentación.\\
Examen final & Verifica transferencia, comprensión y argumentación. & Lectura de contexto, uso de conceptos, toma de decisiones y justificación.\\
Actitud y cuidado & Valora convivencia tecnológica y responsabilidad. & Cuidado de equipos, datos, tiempos, colaboración y honestidad intelectual.\\
\end{longtable}
\endgroup
"""


def tex_correspondence_annex(correspondence: list[str]) -> str:
    return rf"""
\section*{{Anexo 6 · Correspondencia de recursos}}\addcontentsline{{toc}}{{section}}{{Anexo 6 · Correspondencia de recursos}}
\begin{{tcolorbox}}[colback=gris,colframe=turquesa,title=\textbf{{Criterio de lectura}},top=1mm,bottom=1mm]
Este anexo permite rastrear, por grado y periodo, las guías de aprendizaje, el proyecto integrador y el examen final asociados.
\end{{tcolorbox}}
\begingroup
\scriptsize
\renewcommand{{\arraystretch}}{{0.88}}
\setlength{{\tabcolsep}}{{2pt}}
\setlength{{\LTleft}}{{0pt}}
\setlength{{\LTright}}{{0pt}}
\begin{{longtable}}{{L{{0.06\linewidth}}L{{0.10\linewidth}}L{{0.34\linewidth}}L{{0.21\linewidth}}L{{0.21\linewidth}}}}
\rowcolor{{vino}}\color{{white}}\textbf{{Grado}} & \color{{white}}\textbf{{Periodo}} & \color{{white}}\textbf{{Guías}} & \color{{white}}\textbf{{Proyecto}} & \color{{white}}\textbf{{Examen}}\\
{chr(10).join(correspondence)}
\end{{longtable}}
\endgroup
"""


def tex_grade_brand_rows() -> str:
    rows = []
    for grade in range(6, 12):
        item = GRADE_BRAND_COLORS[grade]
        text_color = "white" if item["text"] == "FFFFFF" else "black"
        contrast = "Texto blanco" if item["text"] == "FFFFFF" else "Texto oscuro"
        rows.append(
            rf"Grado {tex_escape(GRADE_LABELS[grade])} & "
            rf"\cellcolor[HTML]{{{item['hex']}}}\textcolor{{{text_color}}}{{\textbf{{{tex_escape(item['name'])}}}}} & "
            rf"\#{item['hex']} & {tex_escape(contrast)} & {tex_escape(item['focus'])}\\"
        )
    return "\n".join(rows)


def tex_platform_annex() -> str:
    grade_brand_rows = tex_grade_brand_rows()
    return rf"""
\clearpage
\section*{{Anexo 7 · Plataforma ConectaTE como mediación curricular}}\addcontentsline{{toc}}{{section}}{{Anexo 7 · Plataforma ConectaTE}}
\begin{{tcolorbox}}[breakable,colback=magentaMILC!7,colframe=magentaMILC,title=\textbf{{Sentido institucional}},top=1mm,bottom=1mm]
ConectaTE es la plataforma educativa institucional que articula el Plan de Área de Tecnología e Informática con la experiencia concreta de aula. No se limita a alojar documentos: organiza el tránsito entre malla curricular, guías, proyectos, exámenes, evidencias y seguimiento del estudiante.

Su puesta en escena permite que cada estudiante identifique su grado, navegue por periodos y sesiones, abra la guía completa, la descargue, la consulte en Drive, comparta el enlace y marque avances. De esta manera, la guía deja de ser un archivo aislado y se convierte en una ruta de aprendizaje accesible dentro y fuera del aula.
\end{{tcolorbox}}
\begin{{tcolorbox}}[breakable,colback=white,colframe=magentaMILC,title=\textbf{{Características de la plataforma}},top=1mm,bottom=1mm]
\begin{{itemize}}
\item \textbf{{Cobertura curricular:}} 180 guías de aprendizaje, 18 proyectos integradores y 18 exámenes finales organizados para grados sexto a undécimo.
\item \textbf{{Navegación pedagógica:}} acceso por grado, periodo y sesión, con rutas claras entre guía anterior, guía siguiente, proyecto y examen.
\item \textbf{{Interacción con guías:}} visor PDF embebido, descarga directa, apertura en Drive, opción de compartir por WhatsApp y copia de enlace.
\item \textbf{{Seguimiento:}} registro de estudiante, selección de grado y grupo, tablero de progreso, marcación de guías completadas y experiencia de exploración.
\item \textbf{{Acceso extendido:}} diseño adaptable a móvil, tableta y computador; manifiesto de aplicación web progresiva y soporte de consulta después de la primera carga cuando el recurso queda en caché.
\item \textbf{{Enlace institucional:}} \url{{{PLATFORM_URL}}}
\end{{itemize}}
\end{{tcolorbox}}
\begin{{tcolorbox}}[breakable,colback=gris,colframe=vino,title=\textbf{{Código visual por grado}},top=1mm,bottom=1mm]
La paleta de ConectaTE se proyecta a portadas físicas, PDFs, DOCX, proyectos y exámenes para que cada grado conserve una identidad visual reconocible en aula y plataforma.
\end{{tcolorbox}}
\begingroup
\scriptsize
\renewcommand{{\arraystretch}}{{1.15}}
\setlength{{\tabcolsep}}{{3pt}}
\begin{{longtable}}{{L{{0.13\linewidth}}L{{0.20\linewidth}}L{{0.12\linewidth}}L{{0.14\linewidth}}L{{0.33\linewidth}}}}
\rowcolor{{vino}}\color{{white}}\textbf{{Grado}} & \color{{white}}\textbf{{Color}} & \color{{white}}\textbf{{Código}} & \color{{white}}\textbf{{Contraste}} & \color{{white}}\textbf{{Foco visual}}\\
{grade_brand_rows}
\end{{longtable}}
\endgroup
\begingroup
\small
\renewcommand{{\arraystretch}}{{1.18}}
\setlength{{\tabcolsep}}{{4pt}}
\begin{{longtable}}{{L{{0.19\linewidth}}L{{0.73\linewidth}}}}
\rowcolor{{vino}}\color{{white}}\textbf{{Momento MILC}} & \color{{white}}\textbf{{Función de ConectaTE}}\\
Escucha & Sitúa al estudiante en grado, periodo, sesión y propósito; permite reconocer qué guía corresponde al proceso real de aula.\\
Sistematización & Ordena contenidos, recursos y evidencias en una arquitectura navegable que evita dispersión documental.\\
Praxis & Facilita el acceso al documento de trabajo para producir actividades, proyectos, portafolios y evidencias verificables.\\
Evaluación liberadora & Apoya la revisión de avance, la continuidad fuera del aula, la socialización de enlaces y la responsabilidad del estudiante sobre su proceso.\\
\end{{longtable}}
\endgroup
\begin{{tcolorbox}}[breakable,colback=gris,colframe=turquesa,title=\textbf{{Uso pedagógico}},top=1mm,bottom=1mm]
Para el docente, ConectaTE funciona como tablero de ruta: permite ubicar la sesión, proyectar la guía, orientar instrucciones comunes y dejar continuidad para casa. Para el estudiante, funciona como cuaderno digital de consulta: le muestra qué debe trabajar, dónde está el recurso, cómo descargarlo y cómo avanzar con autonomía. Para la institución, fortalece la coherencia entre plan de área, recursos curriculares, experiencia de aula, seguimiento y actualización permanente.
\end{{tcolorbox}}
"""


def tex_bibliography() -> str:
    items = "\n".join(rf"\item {tex_escape(entry)}" for entry in BIBLIOGRAPHY)
    return rf"""
\clearpage
\section*{{Bibliografía}}\addcontentsline{{toc}}{{chapter}}{{Bibliografía}}
\begin{{tcolorbox}}[breakable,colback=gris,colframe=vino,title=\textbf{{Criterio de lectura}},top=1mm,bottom=1mm]
La bibliografía reúne los referentes pedagógicos, normativos, tecnológicos y curriculares que sostienen el enfoque MILC, la organización del plan de área, la ciudadanía digital, la inteligencia artificial responsable, la gestión de calidad y la mediación mediante plataforma educativa.
\end{{tcolorbox}}
\begin{{itemize}}
{items}
\end{{itemize}}
"""


def tex_annexes(periods: list[PeriodData], correspondence: list[str]) -> str:
    return "\n".join(
        [
            tex_annex_divider(),
            tex_full_temario_annex(periods),
            tex_usage_annexes(),
            tex_correspondence_annex(correspondence),
            tex_platform_annex(),
            tex_bibliography(),
        ]
    )


def tex_table(periods: list[PeriodData]) -> str:
    sections = [
        r"""
\section*{Grados}\addcontentsline{toc}{chapter}{Grados}
\begin{tcolorbox}[breakable,colback=gris,colframe=vino,title=\textbf{Criterio de lectura},top=1mm,bottom=1mm]
La sección de grados presenta las mallas curriculares de sexto a undécimo. Cada grado conserva sus tres periodos y cada periodo desarrolla los cuatro estándares como filas curriculares permanentes.
\end{tcolorbox}
"""
    ]
    for grade in range(6, 12):
        sections.append(rf"\clearpage\section*{{Grado {tex_escape(GRADE_LABELS[grade])}}}\addcontentsline{{toc}}{{section}}{{Grado {tex_escape(GRADE_LABELS[grade])}}}")
        for pdata in [p for p in periods if p.grade == grade]:
            guides, project, exam = resource_summary(pdata)
            topics = "; ".join(simplified_topics(pdata))
            sections.append(
                rf"""
\section*{{Periodo {tex_escape(PERIOD_NAMES[pdata.period])}}}
\begin{{tcolorbox}}[colback=verde!8,colframe=verde,title=\textbf{{Datos curriculares y recursos}}]
\textbf{{Asignación:}} {tex_escape(pdata.asignacion)}\\
\textbf{{DBA / aprendizaje:}} {tex_escape(compact(pdata.dba, 650))}\\
\textbf{{Referentes:}} {tex_escape(compact(pdata.referentes, 520))}\\
\textbf{{Estándares MEN:}} se desarrollan como cuatro filas curriculares en la tabla del periodo.\\
\textbf{{Énfasis:}} {tex_escape(infer_emphasis(pdata))}\\
\textbf{{Estrategia activa:}} {tex_escape(active_strategy(pdata))}\\
\textbf{{Recursos:}} Guías {tex_escape(guides)}; Proyecto {tex_escape(project)}; Examen {tex_escape(exam)}.
\end{{tcolorbox}}

\begingroup
\scriptsize
\setlength{{\tabcolsep}}{{2pt}}
\setlength{{\LTleft}}{{0pt}}
\setlength{{\LTright}}{{0pt}}
\begin{{longtable}}{{L{{0.14\linewidth}}L{{0.16\linewidth}}L{{0.18\linewidth}}L{{0.145\linewidth}}L{{0.145\linewidth}}L{{0.145\linewidth}}}}
\rowcolor{{vino}}\color{{white}}\textbf{{Estándares / orientaciones}} &
\color{{white}}\textbf{{Competencias}} &
\color{{white}}\textbf{{Aprendizajes}} &
\color{{white}}\textbf{{Cognitivos}} &
\color{{white}}\textbf{{Procedimentales}} &
\color{{white}}\textbf{{Actitudinales}}\\
{tex_curricular_rows(pdata)}
\end{{longtable}}
\endgroup

\textbf{{Temario articulado:}} {tex_escape(topics)}
\smallskip
"""
            )
    return "\n".join(sections)


def tex_preliminaries() -> str:
    chunks = [
        r"""\begin{tcolorbox}[colback=gris,colframe=vino,title=\textbf{Clave de lectura}]
Los apartados siguientes reconstruyen y amplían los elementos previos al plan de estudios. Su función es dar fundamento curricular, epistemológico, metodológico y evaluativo a las mallas, de modo que la metodología MILC no aparezca como un añadido, sino como el principio organizador de todo el documento.
\end{tcolorbox}"""
    ]
    for section in academic_preliminaries():
        color_name = {
            PALETTE["vino"]: "vino",
            PALETTE["turquesa"]: "turquesa",
            PALETTE["magenta"]: "magentaMILC",
            PALETTE["verde"]: "verde",
            PALETTE["mostaza"]: "mostaza",
        }.get(section.color, "vino")
        chunks.append(rf"\section*{{{tex_escape(section.title)}}}\addcontentsline{{toc}}{{section}}{{{tex_escape(section.title)}}}")
        chunks.append(rf"\begin{{tcolorbox}}[enhanced,breakable,colback={color_name}!7,colframe={color_name},arc=2mm,boxrule=.6pt,left=4mm,right=4mm,top=2mm,bottom=2mm,before skip=3pt,after skip=4pt]")
        for paragraph in section.paragraphs:
            chunks.append(tex_escape(paragraph) + "\n")
        if section.bullets:
            chunks.append(r"\begin{itemize}")
            for item in section.bullets:
                chunks.append(rf"\item {tex_escape(item)}")
            chunks.append(r"\end{itemize}")
        if section.callout:
            chunks.append(
                rf"\begin{{tcolorbox}}[breakable,colback=white,colframe={color_name},title=\textbf{{{tex_escape(section.callout_title or 'Síntesis')}}},top=1mm,bottom=1mm,before skip=2pt,after skip=2pt]"
            )
            chunks.append(tex_escape(section.callout))
            chunks.append(r"\end{tcolorbox}")
        chunks.append(r"\end{tcolorbox}")
    return "\n".join(chunks)


def build_tex(periods: list[PeriodData]) -> Path:
    body = tex_table(periods)
    correspondence = []
    for pdata in periods:
        guides, project, exam = resource_summary(pdata)
        correspondence.append(
            f"{pdata.grade} & {tex_escape(PERIOD_NAMES[pdata.period])} & {tex_escape(guides)} & {tex_escape(project)} & {tex_escape(exam)}\\\\"
        )
    preliminaries = tex_preliminaries()
    annexes = tex_annexes(periods, correspondence)
    tex = rf"""
\documentclass[10pt,letterpaper]{{report}}
\usepackage[margin=1.20cm]{{geometry}}
\usepackage{{fontspec}}
\usepackage{{xcolor}}
\usepackage{{tikz}}
\usepackage{{longtable}}
\usepackage{{array}}
\usepackage{{booktabs}}
\usepackage{{colortbl}}
\usepackage[most]{{tcolorbox}}
\usepackage{{hyperref}}
\setmainfont{{Helvetica Neue}}
\setsansfont{{Helvetica Neue}}
\definecolor{{magentaMILC}}{{HTML}}{{{PALETTE['magenta']}}}
\definecolor{{vino}}{{HTML}}{{{PALETTE['vino']}}}
\definecolor{{verde}}{{HTML}}{{{PALETTE['verde']}}}
\definecolor{{turquesa}}{{HTML}}{{{PALETTE['turquesa']}}}
\definecolor{{mostaza}}{{HTML}}{{{PALETTE['mostaza']}}}
\definecolor{{gris}}{{HTML}}{{{PALETTE['gris']}}}
\hypersetup{{colorlinks=true,linkcolor=vino,urlcolor=turquesa}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{4pt}}
\renewcommand{{\arraystretch}}{{1.22}}
\renewcommand{{\contentsname}}{{Tabla de contenido}}
\newcolumntype{{L}}[1]{{>{{\raggedright\arraybackslash}}p{{#1}}}}
\hyphenpenalty=10000
\exhyphenpenalty=10000
\pretolerance=10000
\emergencystretch=3em
\raggedbottom
\sloppy

\begin{{document}}
\begin{{titlepage}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[magentaMILC] (current page.south west) rectangle (current page.north east);
  \foreach \x in {{-1,1,...,23}}{{
    \foreach \y in {{-1,1,...,31}}{{
      \draw[vino, opacity=.12, line width=.6pt] (\x,\y) .. controls +(1,.6) and +(-1,-.6) .. +(2,1.4);
      \draw[vino, opacity=.10, line width=.6pt] (\x+.6,\y+.4) arc (210:510:.45);
    }}
  }}
\end{{tikzpicture}}
\vspace*{{.85cm}}
\begin{{center}}
\begin{{tcolorbox}}[
  enhanced,colback=vino,colframe=vino,arc=12mm,width=.86\textwidth,
  boxrule=0pt,left=12mm,right=12mm,top=11mm,bottom=11mm,center,
  overlay={{\draw[magentaMILC,line width=2pt,dash pattern=on 9pt off 8pt,rounded corners=10mm]
  ([xshift=7mm,yshift=-7mm]frame.north west) rectangle ([xshift=-7mm,yshift=7mm]frame.south east);}}]
  \centering
  {{\sffamily\bfseries\color{{white}}\fontsize{{15}}{{17}}\selectfont PLAN DE ÁREA · TECNOLOGÍA E INFORMÁTICA}}\\[1.0cm]
  {{\sffamily\bfseries\color{{white}}\fontsize{{31}}{{34}}\selectfont Tecnología e Informática}}\\[3mm]
  {{\sffamily\bfseries\color{{mostaza}}\fontsize{{24}}{{27}}\selectfont Plan de Área 2025}}\\[8mm]
  {{\sffamily\color{{white}}\Large Grados 6 a 11 · Versión editorial MILC}}\\[1.0cm]
  {{\sffamily\bfseries\color{{white}}\large Institución Educativa Sor María Juliana}}\\[2mm]
  {{\sffamily\color{{white}}Autor: {tex_escape(AUTHOR)}}}\\[4mm]
  {{\sffamily\itshape\color{{white}}Lo que necesitamos saber, cuidar y saber hacer}}
\end{{tcolorbox}}
\vfill
\begin{{tcolorbox}}[enhanced,colback=vino!72!black,colframe=vino!72!black,arc=5mm,boxrule=0pt,width=.86\textwidth,left=6mm,right=6mm,top=4mm,bottom=4mm,center]
{{\color{{white}}\sffamily\small Metodología MILC · Escucha · Sistematización · Praxis · Evaluación liberadora\hfill Recursos: 180 guías · 18 proyectos · 18 exámenes}}
\end{{tcolorbox}}
\end{{center}}
\end{{titlepage}}
\nopagecolor
\tableofcontents
\section*{{Preliminares institucionales}}\addcontentsline{{toc}}{{chapter}}{{Preliminares institucionales}}
{preliminaries}

\section*{{Criterio de actualización}}\addcontentsline{{toc}}{{chapter}}{{Criterio de actualización}}
Esta versión conserva el sentido curricular del plan de área 2025 y reorganiza las mallas con el formato observado en el plan 2017: estándares u orientaciones, competencias, aprendizajes y desempeños cognitivos, procedimentales y actitudinales. El ajuste explicita los cuatro componentes de las Orientaciones Generales para la Educación en Tecnología del MEN y articula cada periodo con MILC, estrategias activas y los recursos generados.

\begin{{tcolorbox}}[breakable,colback=gris,colframe=magentaMILC,title=\textbf{{Sistema metodológico y editorial}},top=1mm,bottom=1mm,before skip=2pt,after skip=4pt]
\textbf{{Componentes MEN:}} {tex_escape("; ".join(COMPONENTS))}.\\
\textbf{{MILC:}} Escucha, Sistematización, Praxis y Evaluación liberadora.\\
\textbf{{Estrategias activas:}} Design Thinking, aprendizaje basado en proyectos, experimentos escolares, laboratorio técnico, prototipado, validación y producción situada.\\
\textbf{{Estilo:}} títulos sobre formas sólidas, texto blanco, jerarquía limpia y paleta magenta/vino con acentos verde, turquesa y mostaza.
\end{{tcolorbox}}

{body}

{annexes}
\end{{document}}
"""
    path = OUT / "tex" / "plan_area_tecnologia_2025_editorial_milc.tex"
    path.write_text(tex, encoding="utf-8")
    return path


def build_manifest(periods: list[PeriodData]) -> Path:
    path = OUT / "anexos" / "correspondencia_plan_recursos.csv"
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "grado",
                "periodo",
                "guias",
                "proyecto",
                "examen",
                "componentes_men",
                "estrategia_activa",
            ],
        )
        writer.writeheader()
        for pdata in periods:
            guides, project, exam = resource_row(pdata)
            writer.writerow(
                {
                    "grado": pdata.grade,
                    "periodo": pdata.period,
                    "guias": guides,
                    "proyecto": project,
                    "examen": exam,
                    "componentes_men": " | ".join(COMPONENTS),
                    "estrategia_activa": active_strategy(pdata),
                }
            )
    return path


def compile_pdf(tex_path: Path) -> Path:
    for _ in range(2):
        subprocess.run(
            [XELATEX, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=tex_path.parent,
            check=True,
            capture_output=True,
            text=True,
        )
    pdf_src = tex_path.with_suffix(".pdf")
    pdf_dst = OUT / "pdf" / pdf_src.name
    shutil.copy2(pdf_src, pdf_dst)
    return pdf_dst


def main() -> None:
    ensure_dirs()
    periods = extract_periods()
    tex_path = build_tex(periods)
    pdf_path = compile_pdf(tex_path)
    docx_path = build_docx(periods)
    manifest_path = build_manifest(periods)
    print(json.dumps({
        "periodos": len(periods),
        "tex": str(tex_path),
        "pdf": str(pdf_path),
        "docx": str(docx_path),
        "anexo": str(manifest_path),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
